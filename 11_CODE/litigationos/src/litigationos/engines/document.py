"""Document generation engine — Jinja2 templates → DOCX / PDF.

Renders court filings from database-stored Jinja2 templates, converts to
DOCX with Michigan court formatting defaults, and produces PDFs via
*reportlab*.  Also supports merging multiple PDFs into one.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import TYPE_CHECKING, Optional

import jinja2
from docx import Document as DocxDocument  # python-docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

if TYPE_CHECKING:
    from litigationos.db.connection import DatabaseManager

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Court formatting defaults
# ---------------------------------------------------------------------------

_COURT_FMT = {
    "font_name": "Times New Roman",
    "font_size_pt": 12,
    "margin_inches": 1.0,
    "line_spacing": 2.0,       # double-spaced
    "page_numbers": True,
}


class DocumentEngine:
    """Generate court documents from Jinja2 templates.

    Workflow:
        1. ``generate_from_template()`` → rendered text/markdown
        2. ``to_docx()`` → formatted ``.docx``
        3. ``to_pdf()`` → ``.pdf`` via *reportlab*
        4. ``merge_pdfs()`` → combine exhibits / attachments

    Attributes:
        _db: :class:`DatabaseManager` for template and document storage.
        output_dir: Default directory for generated files.
    """

    def __init__(
        self,
        db: "DatabaseManager",
        output_dir: Optional[Path] = None,
    ) -> None:
        self._db = db
        self.output_dir = output_dir or Path.home() / "LitigationOS" / "documents"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._jinja_env = jinja2.Environment(
            undefined=jinja2.StrictUndefined,
            autoescape=False,
        )

    # ------------------------------------------------------------------
    # Template rendering
    # ------------------------------------------------------------------

    def generate_from_template(
        self, template_name: str, context_dict: dict
    ) -> str:
        """Load a Jinja2 template from the DB and render it.

        Args:
            template_name: ``templates.name`` value to look up.
            context_dict: Variables passed into ``Template.render()``.

        Returns:
            Rendered text (typically Markdown or plain text).

        Raises:
            ValueError: If the template is not found.
            jinja2.TemplateError: On rendering failures.
        """
        row = self._db.fetchone(
            "SELECT * FROM templates WHERE name = ?",
            (template_name,),
        )
        if row is None:
            raise ValueError(f"Template '{template_name}' not found in database.")

        tmpl = self._jinja_env.from_string(row["content"])
        rendered = tmpl.render(**context_dict)
        logger.info("Rendered template '%s' (%d chars).", template_name, len(rendered))
        return rendered

    # ------------------------------------------------------------------
    # DOCX generation
    # ------------------------------------------------------------------

    def to_docx(
        self,
        text_or_md: str,
        output_path: str | Path,
        formatting: Optional[dict] = None,
    ) -> Path:
        """Convert text/Markdown content to a ``.docx`` with court formatting.

        Args:
            text_or_md: Rendered text content.
            output_path: Destination file path.
            formatting: Optional overrides for ``_COURT_FMT`` keys.

        Returns:
            Resolved :class:`Path` to the written file.
        """
        fmt = {**_COURT_FMT, **(formatting or {})}
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = DocxDocument()

        # -- Page setup (margins) --
        for section in doc.sections:
            section.top_margin = Inches(fmt["margin_inches"])
            section.bottom_margin = Inches(fmt["margin_inches"])
            section.left_margin = Inches(fmt["margin_inches"])
            section.right_margin = Inches(fmt["margin_inches"])

        # -- Default paragraph style --
        style = doc.styles["Normal"]
        font = style.font
        font.name = fmt["font_name"]
        font.size = Pt(fmt["font_size_pt"])
        pf = style.paragraph_format
        pf.line_spacing = fmt["line_spacing"]

        # -- Page numbers (simple footer) --
        if fmt["page_numbers"]:
            for section in doc.sections:
                footer = section.footer
                footer.is_linked_to_previous = False
                fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add PAGE field via run XML
                run = fp.add_run()
                from docx.oxml.ns import qn
                fld_char_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
                run._r.append(fld_char_begin)
                run2 = fp.add_run()
                instr = run2._r.makeelement(qn("w:instrText"), {})
                instr.text = " PAGE "
                run2._r.append(instr)
                run3 = fp.add_run()
                fld_char_end = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
                run3._r.append(fld_char_end)

        # -- Content --
        for line in text_or_md.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            # Basic heading detection from Markdown
            if stripped.startswith("# "):
                p = doc.add_paragraph(stripped[2:])
                p.style = doc.styles["Heading 1"]
            elif stripped.startswith("## "):
                p = doc.add_paragraph(stripped[3:])
                p.style = doc.styles["Heading 2"]
            elif stripped.startswith("### "):
                p = doc.add_paragraph(stripped[4:])
                p.style = doc.styles["Heading 3"]
            else:
                doc.add_paragraph(stripped)

        doc.save(str(output_path))
        logger.info("Generated DOCX: %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # PDF generation
    # ------------------------------------------------------------------

    def to_pdf(
        self,
        docx_path: str | Path,
        output_path: str | Path,
    ) -> Path:
        """Convert a DOCX (or its text) to PDF via *reportlab*.

        Extracts text from the DOCX and re-renders it as a formatted PDF
        using court-standard formatting.

        Args:
            docx_path: Source ``.docx`` file.
            output_path: Destination ``.pdf`` file.

        Returns:
            Resolved :class:`Path` to the written PDF.
        """
        docx_path = Path(docx_path)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX not found: {docx_path}")

        # Extract paragraphs from the DOCX
        src = DocxDocument(str(docx_path))
        paragraphs_text: list[tuple[str, str]] = []  # (text, style_hint)
        for para in src.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            paragraphs_text.append((para.text, style_name))

        # Build PDF
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            leftMargin=1 * inch,
            rightMargin=1 * inch,
            topMargin=1 * inch,
            bottomMargin=1 * inch,
        )

        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            "CourtBody",
            parent=styles["Normal"],
            fontName="Times-Roman",
            fontSize=12,
            leading=24,  # double-spaced at 12pt
        )
        heading_style = ParagraphStyle(
            "CourtHeading",
            parent=styles["Heading1"],
            fontName="Times-Bold",
            fontSize=14,
            leading=28,
            spaceAfter=12,
        )

        flowables: list = []
        for text, style_hint in paragraphs_text:
            if not text.strip():
                flowables.append(Spacer(1, 12))
                continue
            if "Heading" in style_hint:
                flowables.append(Paragraph(text, heading_style))
            else:
                flowables.append(Paragraph(text, body_style))

        doc.build(flowables)
        logger.info("Generated PDF: %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # PDF merge
    # ------------------------------------------------------------------

    def merge_pdfs(
        self, pdf_paths: list[str | Path], output_path: str | Path
    ) -> Path:
        """Merge multiple PDF files into one.

        Uses *reportlab* to concatenate pages.  Falls back to a simple
        binary concatenation of individual PDFs re-rendered as single-
        page images if a richer library isn't available — but in practice
        this method uses *PyPDF2* / *pypdf* when present, otherwise
        rebuilds from extracted text via reportlab.

        Args:
            pdf_paths: Ordered list of PDF file paths.
            output_path: Destination merged PDF.

        Returns:
            Resolved :class:`Path` to the merged PDF.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Try pypdf first (commonly available)
        try:
            from pypdf import PdfReader, PdfWriter

            writer = PdfWriter()
            for p in pdf_paths:
                p = Path(p)
                if not p.exists():
                    logger.warning("Skipping missing PDF: %s", p)
                    continue
                reader = PdfReader(str(p))
                for page in reader.pages:
                    writer.add_page(page)
            with open(output_path, "wb") as f:
                writer.write(f)
            logger.info("Merged %d PDFs --> %s (pypdf)", len(pdf_paths), output_path)
            return output_path
        except ImportError:
            pass

        # Fallback: concatenate raw bytes with reportlab canvas
        try:
            from PyPDF2 import PdfReader as Pdf2Reader  # type: ignore[import-untyped]
            from PyPDF2 import PdfWriter as Pdf2Writer

            writer = Pdf2Writer()
            for p in pdf_paths:
                p = Path(p)
                if not p.exists():
                    continue
                reader = Pdf2Reader(str(p))
                for page in reader.pages:
                    writer.add_page(page)
            with open(output_path, "wb") as f:
                writer.write(f)
            logger.info("Merged %d PDFs --> %s (PyPDF2)", len(pdf_paths), output_path)
            return output_path
        except ImportError:
            pass

        # Last resort: re-render text from each PDF via reportlab
        logger.warning(
            "Neither pypdf nor PyPDF2 available; merging by text extraction."
        )
        all_text: list[str] = []
        for p in pdf_paths:
            p = Path(p)
            if not p.exists():
                continue
            all_text.append(f"--- {p.name} ---")
            all_text.append(p.read_text(errors="replace"))

        combined = "\n\n".join(all_text)
        # Build a simple PDF from the combined text
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "Merge", parent=styles["Normal"], fontName="Times-Roman", fontSize=12, leading=24,
        )
        doc = SimpleDocTemplate(
            str(output_path), pagesize=letter,
            leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
        )
        flowables = [Paragraph(line, body) for line in combined.split("\n") if line.strip()]
        doc.build(flowables)
        logger.info("Merged %d PDFs --> %s (text fallback)", len(pdf_paths), output_path)
        return output_path

    # ------------------------------------------------------------------
    # Template listing
    # ------------------------------------------------------------------

    def get_templates(
        self, category: Optional[str] = None
    ) -> list[dict]:
        """List available templates from the database.

        Args:
            category: Optional filter by ``templates.template_type``.

        Returns:
            List of template dicts.
        """
        if category:
            rows = self._db.fetchall(
                "SELECT * FROM templates WHERE template_type = ? ORDER BY name",
                (category,),
            )
        else:
            rows = self._db.fetchall(
                "SELECT * FROM templates ORDER BY name",
            )
        return [dict(r) for r in rows]
