#!/usr/bin/env python3
"""
court_packet_skeleton_docx_v47.py — NON-INTERPRETIVE DOCX SKELETON BUILDER (v47)

Creates a DOCX shell for a filing packet using *existing* structured inputs.
It does NOT generate arguments, does NOT invent facts, and does NOT summarize authority.

Inputs (all optional, but useful):
- outputs/facts.jsonl (operator-entered)
- outputs/intents.json
- outputs/authority_triples_candidates.jsonl
- outputs/deadlines.json
- docs/vehicle_registry_seed_v35.json or outputs/vehicle_registry_mined_v39.json

Output:
- outputs/packet_skeleton.docx

Requires: python-docx installed (offline).
"""
import argparse, json, os, sys

def load_json(p):
    if not (os.path.exists(p) and os.path.getsize(p)>0): return None
    import json
    return json.load(open(p,"r",encoding="utf-8"))

def iter_jsonl(p):
    if not (os.path.exists(p) and os.path.getsize(p)>0): return []
    out=[]
    import json
    with open(p,"r",encoding="utf-8",errors="replace") as f:
        for line in f:
            line=line.strip()
            if not line: continue
            out.append(json.loads(line))
    return out

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument("--bundle-root", default=".")
    ap.add_argument("--out-docx", default="outputs/packet_skeleton.docx")
    args=ap.parse_args()
    root=os.path.abspath(args.bundle_root)

    try:
        from docx import Document
    except Exception:
        print("ERROR: python-docx not installed.", file=sys.stderr)
        raise SystemExit(2)

    doc=Document()
    doc.add_heading("Filing Packet Skeleton (v47)", level=0)
    doc.add_paragraph("NON-INTERPRETIVE: This document is a structure shell only. Populate with verified content.")

    # Intents
    intents=load_json(os.path.join(root,"outputs","intents.json"))
    doc.add_heading("Intents / Relief Requested", level=1)
    if intents and intents.get("intents"):
        for it in intents["intents"]:
            doc.add_paragraph(f"{it.get('id')}: {it.get('text')}", style="List Bullet")
    else:
        doc.add_paragraph("(missing outputs/intents.json)")

    # Facts (no summarizing; list verbatim)
    facts=iter_jsonl(os.path.join(root,"outputs","facts.jsonl"))
    doc.add_heading("Facts (verbatim operator-entered)", level=1)
    if facts:
        for fx in facts[:300]:
            pin=fx.get("pinpoint") or {}
            loc=pin.get("locator") or ""
            doc.add_paragraph(f"{fx.get('fact_id')}: {fx.get('text')}", style="List Number")
            doc.add_paragraph(f"Pinpoint: {pin.get('source_path','')} | {loc}")
    else:
        doc.add_paragraph("(missing outputs/facts.jsonl)")

    # Authority triples candidates
    triples=iter_jsonl(os.path.join(root,"outputs","authority_triples_candidates.jsonl"))
    doc.add_heading("Authority Triples (candidates; do not file until verified)", level=1)
    if triples:
        for t in triples[:200]:
            doc.add_paragraph(f"{t.get('triple_id','')}: {t.get('proposition','')}", style="List Bullet")
            doc.add_paragraph(f"Authority_ref: {t.get('authority_ref','')}")
    else:
        doc.add_paragraph("(missing outputs/authority_triples_candidates.jsonl)")

    # Deadlines
    deadlines=load_json(os.path.join(root,"outputs","deadlines.json"))
    doc.add_heading("Deadlines (candidate)", level=1)
    if deadlines and deadlines.get("items"):
        for d in deadlines["items"][:200]:
            doc.add_paragraph(f"{d.get('deadline_id','')}: {d.get('due_date','')} — {d.get('label','')}", style="List Bullet")
            doc.add_paragraph(f"Basis: {d.get('basis','')}")
    else:
        doc.add_paragraph("(missing outputs/deadlines.json)")

    os.makedirs(os.path.dirname(os.path.join(root,args.out_docx)) or ".", exist_ok=True)
    out_path=os.path.join(root,args.out_docx)
    doc.save(out_path)
    print("OK wrote", args.out_docx)

if __name__=="__main__":
    main()
