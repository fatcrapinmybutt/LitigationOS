#!/usr/bin/env python3
"""
_export_filings.py — Convert OMEGA-generated filing templates to court-ready formats.

Reads all 0?_*.md files in this directory and produces:
  - court_ready/<name>.txt   (plain text, court-formatted)
  - court_ready/<name>.docx  (Word, Times New Roman 12pt, 1" margins, double-spaced)

The .md files remain the source of truth.
"""

import glob
import os
import re
import sys
import textwrap

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATTERN = os.path.join(SCRIPT_DIR, "0[0-9]_*.md")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "court_ready")

LINE_WIDTH = 65          # characters per line (simulates ~1" margins on letter paper)
LEFT_MARGIN = " " * 10   # 10-space left margin for plain text
PAGE_LINES = 50           # approximate lines per "page" for footer simulation


# ---------------------------------------------------------------------------
# Markdown → structured text helpers
# ---------------------------------------------------------------------------

def strip_metadata_header(lines: list[str]) -> list[str]:
    """Remove the OMEGA metadata block (lines starting with '# OMEGA',
    '# Action:', '# Forum:', '# Generated:', '# ⚠️'), HTML comments,
    and leading '---'."""
    start = 0
    in_html_comment = False
    for i, line in enumerate(lines):
        stripped = line.strip()
        # Track multi-line HTML comments
        if "<!--" in stripped:
            in_html_comment = True
        if in_html_comment:
            start = i + 1
            if "-->" in stripped:
                in_html_comment = False
            continue
        if stripped.startswith("# OMEGA") or stripped.startswith("# Action:") \
                or stripped.startswith("# Forum:") or stripped.startswith("# Generated:") \
                or stripped.startswith("# ⚠️") or stripped == "" or stripped == "---":
            start = i + 1
        else:
            break
    return lines[start:]


def extract_code_block(lines: list[str]) -> tuple[list[str], list[str]]:
    """Extract the first ``` ... ``` block (the court caption) and return
    (caption_lines, remaining_lines)."""
    caption = []
    rest = []
    in_block = False
    block_done = False
    for line in lines:
        if not block_done and line.strip() == "```":
            if in_block:
                block_done = True
            else:
                in_block = True
            continue
        if in_block and not block_done:
            caption.append(line.rstrip())
        else:
            if block_done:
                rest.append(line)
            elif not in_block:
                rest.append(line)
    return caption, rest


def md_bold_to_upper(text: str) -> str:
    """Convert **bold** to UPPERCASE."""
    def _upper(m):
        return m.group(1).upper()
    return re.sub(r"\*\*(.+?)\*\*", _upper, text)


def md_italic_strip(text: str) -> str:
    """Strip single * or _ italic markers, keep content."""
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)", r"\1", text)
    return text


def md_header_to_text(line: str) -> tuple[str, int] | None:
    """If line is a markdown header, return (text, level). Else None."""
    m = re.match(r"^(#{1,6})\s+(.*)", line.strip())
    if m:
        return m.group(2).strip(), len(m.group(1))
    return None


def convert_checkbox(text: str) -> str:
    """Convert - [ ] to ( ) and - [x] to (X)."""
    text = re.sub(r"^-\s*\[\s*\]", "( )", text)
    text = re.sub(r"^-\s*\[[xX]\]", "(X)", text)
    return text


def strip_html_comments(lines: list[str]) -> list[str]:
    """Remove HTML comment blocks from the body."""
    result = []
    in_comment = False
    for line in lines:
        if not in_comment and "<!--" in line:
            in_comment = True
        if in_comment:
            if "-->" in line:
                in_comment = False
            continue
        result.append(line)
    return result


def process_body_lines(lines: list[str]) -> list[str]:
    """Convert markdown body to plain-text court format lines."""
    lines = strip_html_comments(lines)
    output: list[str] = []
    wrapper = textwrap.TextWrapper(width=LINE_WIDTH, initial_indent="",
                                   subsequent_indent="    ")

    for raw_line in lines:
        line = raw_line.rstrip()

        # Horizontal rule → blank line
        if re.match(r"^\s*-{3,}\s*$", line) or re.match(r"^\s*\*{3,}\s*$", line):
            output.append("")
            continue

        # Blank line → double-space
        if line.strip() == "":
            output.append("")
            continue

        # Headers → UPPERCASE, centered
        hdr = md_header_to_text(line)
        if hdr:
            text, level = hdr
            text = md_bold_to_upper(text)
            text = md_italic_strip(text)
            text = text.upper()
            output.append("")
            output.append(text.center(LINE_WIDTH))
            output.append("")
            continue

        # Process inline formatting
        line = md_bold_to_upper(line)
        line = md_italic_strip(line)

        # Bullet list → lettered or plain
        bullet_match = re.match(r"^\s*[-*]\s+(.*)", line)
        if bullet_match:
            content = bullet_match.group(1)
            content = convert_checkbox(content)
            wrapped = wrapper.fill(content)
            output.append(wrapped)
            output.append("")
            continue

        # Numbered list → keep numbering
        num_match = re.match(r"^\s*(\d+)\.\s+(.*)", line)
        if num_match:
            num = num_match.group(1)
            content = num_match.group(2)
            prefix = f"{num}. "
            w = textwrap.TextWrapper(width=LINE_WIDTH,
                                     initial_indent=prefix,
                                     subsequent_indent=" " * len(prefix))
            output.append(w.fill(content))
            output.append("")
            continue

        # Underscore divider (caption closer)
        if re.match(r"^_{5,}", line):
            output.append(line)
            continue

        # Regular paragraph text
        wrapped = wrapper.fill(line)
        output.append(wrapped)

    return output


def add_page_footers(lines: list[str]) -> list[str]:
    """Insert 'Page X of Y' footers every PAGE_LINES lines."""
    # Count total pages first
    total_pages = max(1, (len(lines) + PAGE_LINES - 1) // PAGE_LINES)
    result: list[str] = []
    line_count = 0
    page = 1
    for line in lines:
        result.append(line)
        line_count += 1
        if line_count >= PAGE_LINES and page < total_pages:
            footer = f"Page {page} of {total_pages}".center(LINE_WIDTH)
            result.append("")
            result.append(footer)
            result.append("")
            result.append("-" * LINE_WIDTH)
            result.append("")
            page += 1
            line_count = 0
    # Final page footer
    footer = f"Page {page} of {total_pages}".center(LINE_WIDTH)
    result.append("")
    result.append(footer)
    return result


def add_left_margin(lines: list[str]) -> list[str]:
    """Prepend LEFT_MARGIN to every non-empty line to simulate 1" margin."""
    return [LEFT_MARGIN + ln if ln.strip() else "" for ln in lines]


# ---------------------------------------------------------------------------
# Markdown → plain text conversion (full pipeline)
# ---------------------------------------------------------------------------

def md_to_plain_text(md_content: str) -> str:
    """Convert a filing markdown file to court-formatted plain text."""
    lines = md_content.splitlines()
    lines = strip_metadata_header(lines)

    # Separate caption block
    caption, body = extract_code_block(lines)

    # Format caption (centered)
    caption_formatted = []
    for cl in caption:
        caption_formatted.append(cl.center(LINE_WIDTH))
    caption_formatted.append("")

    # Process body
    body_lines = process_body_lines(body)

    # Combine
    all_lines = caption_formatted + body_lines

    # Remove excessive blank lines (max 2 consecutive)
    cleaned: list[str] = []
    blank_count = 0
    for ln in all_lines:
        if ln.strip() == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned.append("")
        else:
            blank_count = 0
            cleaned.append(ln)

    # Add page footers
    with_footers = add_page_footers(cleaned)

    # Add left margin
    margined = add_left_margin(with_footers)

    return "\n".join(margined)


# ---------------------------------------------------------------------------
# Markdown → DOCX conversion
# ---------------------------------------------------------------------------

def md_to_docx(md_content: str, output_path: str) -> None:
    """Convert a filing markdown file to a court-ready .docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # -- Page setup: 1-inch margins, Letter size --
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Set double spacing via XML (line = 480 twentieths of a point)
    pf_element = pf._element
    spacing = pf_element.find(qn("w:spacing"))
    if spacing is None:
        spacing = pf_element.makeelement(qn("w:spacing"), {})
        pf_element.append(spacing)
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")

    def set_run_font(run):
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    def add_paragraph(text, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      font_size=12, space_after=0):
        p = doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(font_size)
        return p

    def add_mixed_paragraph(segments, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """Add a paragraph with mixed bold/normal/italic runs.
        segments: list of (text, bold, italic)"""
        p = doc.add_paragraph()
        p.alignment = alignment
        for text, is_bold, is_italic in segments:
            run = p.add_run(text)
            run.bold = is_bold
            run.italic = is_italic
            set_run_font(run)
        return p

    def parse_inline(text: str) -> list[tuple[str, bool, bool]]:
        """Parse markdown inline formatting into (text, bold, italic) segments."""
        segments = []
        # Handle **bold** and *italic*
        pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_)")
        pos = 0
        for m in pattern.finditer(text):
            if m.start() > pos:
                segments.append((text[pos:m.start()], False, False))
            if m.group(2) is not None:  # bold
                segments.append((m.group(2), True, False))
            elif m.group(3) is not None:  # italic with *
                segments.append((m.group(3), False, True))
            elif m.group(4) is not None:  # italic with _
                segments.append((m.group(4), False, True))
            pos = m.end()
        if pos < len(text):
            segments.append((text[pos:], False, False))
        if not segments:
            segments.append((text, False, False))
        return segments

    # -- Parse the markdown --
    lines = md_content.splitlines()
    lines = strip_metadata_header(lines)
    caption, body = extract_code_block(lines)
    body = strip_html_comments(body)

    # Write caption block (centered, single-spaced)
    for cl in caption:
        p = add_paragraph(cl.strip(), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # Single-space for caption
        sp = p.paragraph_format._element.find(qn("w:spacing"))
        if sp is None:
            sp = p.paragraph_format._element.makeelement(qn("w:spacing"), {})
            p.paragraph_format._element.append(sp)
        sp.set(qn("w:line"), "240")
        sp.set(qn("w:lineRule"), "auto")

    # Blank line after caption
    doc.add_paragraph()

    # Process body
    i = 0
    while i < len(body):
        line = body[i].rstrip()
        i += 1

        # Skip horizontal rules
        if re.match(r"^\s*-{3,}\s*$", line) or re.match(r"^\s*\*{3,}\s*$", line):
            continue

        # Skip blank lines (spacing handled by paragraph format)
        if line.strip() == "":
            continue

        # Headers
        hdr = md_header_to_text(line)
        if hdr:
            text, level = hdr
            # Strip any remaining markdown
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            if level <= 2:
                add_paragraph(text.upper(), bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              font_size=12, space_after=6)
            else:
                add_paragraph(text, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              font_size=12, space_after=6)
            continue

        # Underscore divider
        if re.match(r"^_{5,}", line):
            add_paragraph(line)
            continue

        # Checkbox lists
        cb = re.match(r"^\s*-\s*\[\s*\]\s*(.*)", line)
        if cb:
            add_paragraph(f"( ) {cb.group(1)}")
            continue
        cbx = re.match(r"^\s*-\s*\[[xX]\]\s*(.*)", line)
        if cbx:
            add_paragraph(f"(X) {cbx.group(1)}")
            continue

        # Bullet lists
        bullet = re.match(r"^\s*[-*]\s+(.*)", line)
        if bullet:
            content = bullet.group(1)
            segments = parse_inline(content)
            add_mixed_paragraph(segments)
            continue

        # Numbered lists
        num = re.match(r"^\s*(\d+)\.\s+(.*)", line)
        if num:
            number = num.group(1)
            content = num.group(2)
            segments = [(f"{number}. ", False, False)] + parse_inline(content)
            add_mixed_paragraph(segments)
            continue

        # Regular paragraph — may span multiple lines (gather continuation)
        para_text = line
        while i < len(body):
            next_line = body[i].rstrip()
            if next_line.strip() == "" or re.match(r"^#{1,6}\s", next_line) \
                    or re.match(r"^\s*-{3,}\s*$", next_line) \
                    or re.match(r"^\s*[-*]\s+", next_line) \
                    or re.match(r"^\s*\d+\.\s+", next_line):
                break
            para_text += " " + next_line.strip()
            i += 1

        segments = parse_inline(para_text)
        add_mixed_paragraph(segments)

    # Add page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        # Page number field code
        fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = run2._element.makeelement(qn("w:instrText"), {})
        instrText.text = " PAGE "
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
        run3._element.append(fldChar2)
        run4 = p.add_run("1")
        run5 = p.add_run()
        fldChar3 = run5._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run5._element.append(fldChar3)
        # " of "
        run6 = p.add_run(" of ")
        run6.font.name = "Times New Roman"
        run6.font.size = Pt(10)
        # Total pages field
        run7 = p.add_run()
        fldChar4 = run7._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run7._element.append(fldChar4)
        run8 = p.add_run()
        instrText2 = run8._element.makeelement(qn("w:instrText"), {})
        instrText2.text = " NUMPAGES "
        run8._element.append(instrText2)
        run9 = p.add_run()
        fldChar5 = run9._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
        run9._element.append(fldChar5)
        run10 = p.add_run("1")
        run11 = p.add_run()
        fldChar6 = run11._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run11._element.append(fldChar6)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    md_files = sorted(glob.glob(INPUT_PATTERN))
    if not md_files:
        print(f"No files matching {INPUT_PATTERN}")
        sys.exit(1)

    print(f"Found {len(md_files)} filing(s) to convert.\n")

    for md_path in md_files:
        basename = os.path.splitext(os.path.basename(md_path))[0]
        print(f"  Processing: {basename}")

        with open(md_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        # --- Plain text ---
        txt_path = os.path.join(OUTPUT_DIR, basename + ".txt")
        plain = md_to_plain_text(md_content)
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(plain)
        print(f"    ✓ {os.path.basename(txt_path)}")

        # --- DOCX ---
        try:
            docx_path = os.path.join(OUTPUT_DIR, basename + ".docx")
            md_to_docx(md_content, docx_path)
            print(f"    ✓ {os.path.basename(docx_path)}")
        except ImportError:
            print("    ⚠ python-docx not available — skipping DOCX generation")
        except Exception as e:
            print(f"    ✗ DOCX error: {e}")

    print(f"\nDone. Output in: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
