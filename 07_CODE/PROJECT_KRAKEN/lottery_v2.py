import sqlite3, random, json, os, sys, re, traceback
from pathlib import Path
from collections import defaultdict

print("=" * 70)
print("APEX LOTTERY HARVEST v2 — fd + DB inventory + content-first")
print("=" * 70)

# ── PHASE 1: Harvest file paths from MULTIPLE sources (zero os.walk) ──
all_files = []
EXTS = {'.pdf', '.txt', '.csv', '.html', '.json', '.docx', '.md'}

# Source 1: file_inventory DB table (611K+ files already indexed)
print("\n[SOURCE 1] file_inventory DB (pre-indexed, instant)...")
db_path = r"C:\Users\andre\LitigationOS\litigation_context.db"
try:
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA busy_timeout=30000")
    # Check what columns exist
    cols = {r[1] for r in conn.execute("PRAGMA table_info(file_inventory)")}
    
    if 'file_path' in cols:
        path_col = 'file_path'
    elif 'path' in cols:
        path_col = 'path'
    elif 'full_path' in cols:
        path_col = 'full_path'
    else:
        path_col = list(cols)[0]  # fallback
        
    print(f"  Using column: {path_col}")
    print(f"  Available columns: {sorted(cols)}")
    
    # Get files matching our extensions from DB
    rows = conn.execute(f"""
        SELECT {path_col} FROM file_inventory 
        WHERE {path_col} IS NOT NULL
        LIMIT 50000
    """).fetchall()
    
    db_files = []
    for (fp,) in rows:
        if fp and any(fp.lower().endswith(ext) for ext in EXTS):
            db_files.append(fp)
    
    print(f"  Found {len(db_files)} files with target extensions (from {len(rows)} total)")
    all_files.extend(db_files)
    conn.close()
except Exception as e:
    print(f"  DB error: {e}")

# Source 2: fd (Rust) for C:\Users\andre — blazing fast
print("\n[SOURCE 2] fd (Rust) scan of C:\\Users\\andre...")
import subprocess
try:
    ext_args = []
    for ext in ['pdf', 'txt', 'csv', 'html', 'json', 'docx', 'md']:
        ext_args.extend(['-e', ext])
    
    result = subprocess.run(
        ['fd', '--type', 'f', '--no-ignore', '--hidden'] + ext_args + 
        ['--exclude', 'node_modules', '--exclude', '.git', '--exclude', 'AppData',
         '--exclude', '__pycache__', '--exclude', '.cache',
         '.', r'C:\Users\andre'],
        capture_output=True, text=True, timeout=30, encoding='utf-8', errors='replace'
    )
    fd_files = [f.strip() for f in result.stdout.strip().split('\n') if f.strip()]
    print(f"  fd found {len(fd_files)} files in <30s")
    all_files.extend(fd_files)
except FileNotFoundError:
    print("  fd not found, using fallback...")
    # Fallback: just list key directories quickly
    for d in [r'C:\Users\andre\Desktop', r'C:\Users\andre\Documents', 
              r'C:\Users\andre\Downloads', r'C:\Users\andre\temp',
              r'C:\Users\andre\LitigationOS\04_ANALYSIS',
              r'C:\Users\andre\LitigationOS\05_FILINGS']:
        if os.path.isdir(d):
            for f in os.listdir(d):
                fp = os.path.join(d, f)
                if os.path.isfile(fp) and any(f.lower().endswith(e) for e in EXTS):
                    all_files.append(fp)
except subprocess.TimeoutExpired:
    print("  fd timed out at 30s, using partial results")
except Exception as e:
    print(f"  fd error: {e}")

# Source 3: Quick scan of D:\ (small drive, fast)
print("\n[SOURCE 3] Quick D:\\ scan...")
try:
    result = subprocess.run(
        ['fd', '--type', 'f', '-e', 'pdf', '-e', 'txt', '-e', 'csv', '-e', 'html',
         '-e', 'json', '-e', 'docx', '-e', 'md', '--no-ignore',
         '--exclude', 'node_modules', '--exclude', '.git',
         '.', 'D:\\'],
        capture_output=True, text=True, timeout=30, encoding='utf-8', errors='replace'
    )
    d_files = [f.strip() for f in result.stdout.strip().split('\n') if f.strip()]
    print(f"  fd found {len(d_files)} files on D:\\")
    all_files.extend(d_files)
except Exception as e:
    print(f"  D:\\ scan: {e}")

# Deduplicate
all_files = list(set(all_files))
print(f"\n[TOTAL] {len(all_files)} unique files across all sources")

# Filter to only files that actually exist (for DB entries from old scans)
print("\n[PHASE 2] Validating existence of candidate files...")
valid_files = []
# Sample a subset to validate (checking all 50K would be slow on USB)
sample_size = min(5000, len(all_files))
candidates = random.sample(all_files, sample_size) if len(all_files) > sample_size else all_files

for fp in candidates:
    try:
        if os.path.isfile(fp) and os.path.getsize(fp) > 100 and os.path.getsize(fp) < 50_000_000:
            valid_files.append(fp)
    except (OSError, PermissionError):
        pass

print(f"  {len(valid_files)} valid files (100B-50MB, accessible)")

# ── PHASE 3: THE LOTTERY — pick 10 random files ──
print("\n" + "=" * 70)
print("🎰 THE LOTTERY DRAW — 10 Random Files")
print("=" * 70)

lottery = random.sample(valid_files, min(10, len(valid_files)))

for i, fp in enumerate(lottery, 1):
    sz = os.path.getsize(fp)
    ext = Path(fp).suffix.lower()
    print(f"\n  [{i:02d}] {ext.upper()} | {sz:,} bytes")
    print(f"       {fp}")

# ── PHASE 4: READ AND ANALYZE EACH FILE ──
print("\n" + "=" * 70)
print("📖 CONTENT ANALYSIS — Reading actual file contents")
print("=" * 70)

# Adversary patterns
ADVERSARIES = {
    'Emily Watson': r'(?i)\bemily\b.*\bwatson\b|\bwatson\b.*\bemily\b|\bemily\s+a\.?\s+watson\b',
    'Judge McNeill': r'(?i)\bmcneill\b|\bmcneil\b',
    'Pamela Rusco': r'(?i)\brusco\b|\bpamela\b.*\brusco\b',
    'Albert Watson': r'(?i)\balbert\b.*\bwatson\b|\bwatson\b.*\balbert\b',
    'Lori Watson': r'(?i)\blori\b.*\bwatson\b',
    'Ronald Berry': r'(?i)\bronald\b.*\bberry\b|\bron\b.*\bberry\b',
    'Cavan Berry': r'(?i)\bcavan\b.*\bberry\b',
    'Jennifer Barnes': r'(?i)\bbarnes\b.*\bjennifer\b|\bjennifer\b.*\bbarnes\b|\bbarnes\b.*\bP55406\b',
    'Mandi Martini': r'(?i)\bmartini\b|\bmandi\b.*\bmartini\b',
    'Kenneth Hoopes': r'(?i)\bhoopes\b|\bkenneth\b.*\bhoopes\b',
    'Maria Ladas-Hoopes': r'(?i)\bladas\b.*\bhoopes\b|\bmaria\b.*\bladas\b',
    'Kim Davis': r'(?i)\bkim\b.*\bdavis\b',
    'Cassandra VanDam': r'(?i)\bvandam\b|\bvan\s*dam\b|\bcassandra\b.*\bvan',
    'Nicole Browley': r'(?i)\bbrowley\b|\bnicole\b.*\bbrowley\b',
    'Henry Brandell': r'(?i)\bbrandell\b|\bbrandel\b|\bhenry\b.*\bbrand',
    'Jeremy Brown': r'(?i)\bjeremy\b.*\bbrown\b',
    'Aaron Cox': r'(?i)\baaron\b.*\bcox\b|\bcox\b.*\bP69346\b',
    'DJ Hilson': r'(?i)\bhilson\b|\bd\.?j\.?\b.*\bhilson\b',
    'Lauren Duguid': r'(?i)\bduguid\b|\blauren\b.*\bduguid\b',
    'Shady Oaks': r'(?i)\bshady\s*oaks\b',
    'FOC': r'(?i)\bfriend\s+of\s+(the\s+)?court\b|\bFOC\b',
    'EGLE': r'(?i)\bEGLE\b|\benvironmental\s+quality\b',
}

# Legal authority patterns
LEGAL = {
    'MCR': r'\bMCR\s+\d+\.\d+',
    'MCL': r'\bMCL\s+\d+\.\d+',
    'MRE': r'\bMRE\s+\d+',
    'USC': r'\b\d+\s+U\.?S\.?C\.?\s+[§\s]*\d+',
    'FRCP': r'\bFRCP\s+\d+|\bFed\.?\s*R\.?\s*Civ\.?\s*P\.?\s*\d+',
    'Case Law': r'\b\d+\s+Mich\s+(App\s+)?\d+|\b\d+\s+F\.\d[a-z]*\s+\d+',
}

results = []

for i, fp in enumerate(lottery, 1):
    ext = Path(fp).suffix.lower()
    fname = Path(fp).name
    sz = os.path.getsize(fp)
    
    print(f"\n{'─' * 60}")
    print(f"[{i:02d}] {fname} ({sz:,} bytes)")
    print(f"     {fp}")
    
    content = ""
    try:
        if ext == '.pdf':
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(fp)
                pages = []
                for pg_idx in range(min(len(pdf), 20)):  # first 20 pages
                    page = pdf[pg_idx]
                    text = page.get_textpage().get_text_range()
                    pages.append(text)
                content = '\n'.join(pages)
                print(f"     PDF: {len(pdf)} pages, extracted {len(content):,} chars")
            except Exception as e:
                print(f"     PDF extraction failed: {e}")
                
        elif ext == '.docx':
            try:
                from docx import Document
                doc = Document(fp)
                content = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                print(f"     DOCX: {len(doc.paragraphs)} paragraphs, {len(content):,} chars")
            except Exception as e:
                print(f"     DOCX failed: {e}")
                
        elif ext in {'.txt', '.csv', '.html', '.json', '.md'}:
            for enc in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
                try:
                    with open(fp, 'r', encoding=enc, errors='replace') as f:
                        content = f.read(500_000)  # cap at 500KB
                    print(f"     TEXT: {len(content):,} chars ({enc})")
                    break
                except:
                    continue
    except Exception as e:
        print(f"     READ ERROR: {e}")
        continue
    
    if not content or len(content) < 50:
        print(f"     ⚠ No meaningful content extracted")
        results.append({'file': fp, 'ext': ext, 'size': sz, 'content_len': len(content),
                       'adversaries': {}, 'legal': {}, 'value': 'EMPTY'})
        continue
    
    # Analyze for adversaries
    adv_hits = {}
    for name, pattern in ADVERSARIES.items():
        matches = re.findall(pattern, content)
        if matches:
            adv_hits[name] = len(matches)
    
    # Analyze for legal authorities
    legal_hits = {}
    for cat, pattern in LEGAL.items():
        matches = re.findall(pattern, content)
        if matches:
            legal_hits[cat] = list(set(matches))[:10]
    
    # Extract key quotes (first 5 sentences with adversary mentions)
    key_quotes = []
    sentences = re.split(r'[.!?\n]+', content)
    for sent in sentences:
        sent = sent.strip()
        if len(sent) > 30 and len(sent) < 500:
            for name, pattern in ADVERSARIES.items():
                if re.search(pattern, sent):
                    key_quotes.append(sent[:200])
                    break
        if len(key_quotes) >= 5:
            break
    
    # Evidence categories
    categories = []
    cat_patterns = {
        'custody': r'(?i)\bcustody\b|\bparenting\s*time\b|\bvisitation\b',
        'PPO': r'(?i)\bprotection\s*order\b|\bPPO\b|\brestraining\b',
        'judicial': r'(?i)\bjudicial\b|\bbias\b|\bex\s*parte\b|\brecusal\b',
        'housing': r'(?i)\beviction\b|\btenant\b|\blandlord\b|\bhousing\b|\bmobile\s*home\b',
        'criminal': r'(?i)\bcontempt\b|\bjail\b|\bincarcerat\b|\barrest\b',
        'financial': r'(?i)\bchild\s*support\b|\bfiling\s*fee\b|\bdamages\b',
        'police': r'(?i)\bpolice\b|\bNSPD\b|\bofficer\b|\breport\b',
        'medical': r'(?i)\bhealthwest\b|\bmental\s*health\b|\bpsych\b|\bmedication\b',
    }
    for cat, pat in cat_patterns.items():
        if re.search(pat, content):
            categories.append(cat)
    
    # Score value
    value_score = len(adv_hits) * 3 + len(legal_hits) * 2 + len(categories) + len(key_quotes)
    value = 'HIGH' if value_score >= 10 else 'MEDIUM' if value_score >= 4 else 'LOW'
    
    print(f"     📊 Value: {value} (score={value_score})")
    if adv_hits:
        top3 = sorted(adv_hits.items(), key=lambda x: -x[1])[:3]
        print(f"     👤 Adversaries: {', '.join(f'{n}({c})' for n,c in top3)}")
    if legal_hits:
        print(f"     ⚖️  Legal: {', '.join(f'{k}:{len(v)}' for k,v in legal_hits.items())}")
    if categories:
        print(f"     📁 Categories: {', '.join(categories)}")
    if key_quotes:
        print(f"     💬 Key quote: \"{key_quotes[0][:120]}...\"")
    
    results.append({
        'file': fp, 'ext': ext, 'size': sz,
        'content_len': len(content),
        'adversaries': adv_hits,
        'legal': {k: v for k, v in legal_hits.items()},
        'categories': categories,
        'key_quotes': key_quotes[:5],
        'value': value,
        'value_score': value_score,
    })

# ── SUMMARY ──
print("\n" + "=" * 70)
print("📊 LOTTERY RESULTS SUMMARY")
print("=" * 70)

high = [r for r in results if r.get('value') == 'HIGH']
med = [r for r in results if r.get('value') == 'MEDIUM']
low = [r for r in results if r.get('value') in ('LOW', 'EMPTY')]

print(f"  🔴 HIGH value: {len(high)} files")
print(f"  🟡 MEDIUM value: {len(med)} files")
print(f"  ⚪ LOW/EMPTY: {len(low)} files")

all_advs = defaultdict(int)
all_legals = defaultdict(list)
for r in results:
    for k, v in r.get('adversaries', {}).items():
        all_advs[k] += v
    for k, v in r.get('legal', {}).items():
        all_legals[k].extend(v)

if all_advs:
    print(f"\n  Top adversary mentions across lottery:")
    for name, count in sorted(all_advs.items(), key=lambda x: -x[1])[:10]:
        print(f"    {name}: {count}")

if all_legals:
    print(f"\n  Legal authorities found:")
    for cat, refs in sorted(all_legals.items()):
        unique = list(set(refs))[:5]
        print(f"    {cat}: {', '.join(unique)}")

# Save results
out_path = r"D:\LitigationOS_tmp\lottery_results_v2.json"
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, indent=2, default=str)
print(f"\n  Results saved: {out_path}")

print("\n✅ APEX LOTTERY HARVEST v2 COMPLETE")
