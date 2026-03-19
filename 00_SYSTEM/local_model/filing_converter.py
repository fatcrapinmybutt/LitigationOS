#!/usr/bin/env python3
"""
Filing Converter — Markdown to DOCX for Michigan Court Filings
==============================================================
Converts court-ready .md filing documents into properly formatted .docx
Word files per MCR 2.113 formatting requirements.

Usage:
    from filing_converter import convert_filing, convert_all_filings
    output = convert_filing("path/to/motion.md")
    results = convert_all_filings("C:/Users/andre/LitigationOS/04_COURT_FILINGS/")

    # CLI
    python filing_converter.py path/to/filing.md
    python filing_converter.py --batch path/to/filings_dir/
"""

from __future__ import annotations

import logging
import os
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("[FilingConverter] python-docx not installed. Run: pip install python-docx",
          file=sys.stderr)
    raise

# ── Constants ──────────────────────────────────────────────────────
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_HEADING2 = Pt(14)
FONT_SIZE_HEADING3 = Pt(13)
FONT_SIZE_CAPTION = Pt(12)
MARGIN_INCHES = 1.0
LINE_SPACING = WD_LINE_SPACING.DOUBLE

DOC_TYPE_PATTERNS = {
    "motion": re.compile(r"\bMOTION\b", re.IGNORECASE),
    "brief": re.compile(r"\bBRIEF\b", re.IGNORECASE),
    "complaint": re.compile(r"\bCOMPLAINT\b", re.IGNORECASE),
    "response": re.compile(r"\bRESPONSE\b", re.IGNORECASE),
    "affidavit": re.compile(r"\bAFFIDAVIT\b", re.IGNORECASE),
    "order": re.compile(r"\bORDER\b", re.IGNORECASE),
    "petition": re.compile(r"\bPETITION\b", re.IGNORECASE),
    "appeal": re.compile(r"\bAPPEAL|CLAIM OF APPEAL\b", re.IGNORECASE),
}

# Regex patterns for markdown parsing
RE_HEADING2 = re.compile(r"^##\s+(.+)$")
RE_HEADING3 = re.compile(r"^###\s+(.+)$")
RE_HEADING4 = re.compile(r"^####\s+(.+)$")
RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
RE_ITALIC = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
RE_BOLD_ITALIC = re.compile(r"\*\*\*(.+?)\*\*\*")
RE_NUMBERED = re.compile(r"^(\d+)\.\s+(.+)$")
RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
RE_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")
RE_HORIZ_RULE = re.compile(r"^(-{3,}|_{3,}|\*{3,})$")
RE_PAGE_BREAK = re.compile(r"^<!--\s*pagebreak\s*-->$|^\\pagebreak$", re.IGNORECASE)
RE_CAPTION_LINE = re.compile(
    r"(CIRCUIT COURT|COURT OF APPEALS|SUPREME COURT|DISTRICT COURT|"
    r"Case No\.|Plaintiff|Defendant|Appellant|Appellee|Hon\.|v\.\s)",
    re.IGNORECASE,
)


def _detect_doc_type(text: str) -> str:
    """Auto-detect document type from content."""
    first_500 = text[:500].upper()
    for doc_type, pattern in DOC_TYPE_PATTERNS.items():
        if pattern.search(first_500):
            return doc_type
    return "filing"


def _set_document_margins(doc: Document) -> None:
    """Set 1-inch margins on all sides per MCR 2.113."""
    for section in doc.sections:
        section.top_margin = Inches(MARGIN_INCHES)
        section.bottom_margin = Inches(MARGIN_INCHES)
        section.left_margin = Inches(MARGIN_INCHES)
        section.right_margin = Inches(MARGIN_INCHES)


def _add_page_numbers(doc: Document) -> None:
    """Add centered page numbers in footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # Page number field
        fld_xml = (
            '<w:fldSimple %s w:instr=" PAGE \\* MERGEFORMAT ">'
            '<w:r><w:t>1</w:t></w:r>'
            '</w:fldSimple>' % nsdecls('w')
        )
        fld_elem = parse_xml(fld_xml)
        run._r.append(fld_elem)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)


def _set_paragraph_format(paragraph, spacing=LINE_SPACING, font_size=FONT_SIZE_BODY,
                          bold=False, alignment=None, keep_with_next=False):
    """Apply standard formatting to a paragraph."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = spacing
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    if alignment:
        paragraph.alignment = alignment
    if keep_with_next:
        pf.keep_with_next = True
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = font_size
        if bold:
            run.font.bold = True


def _add_formatted_text(paragraph, text: str, base_bold=False, base_size=FONT_SIZE_BODY):
    """Parse inline markdown (bold, italic) and add runs to paragraph."""
    # Process bold+italic first, then bold, then italic
    segments = _parse_inline_markdown(text)
    for seg_text, seg_bold, seg_italic in segments:
        run = paragraph.add_run(seg_text)
        run.font.name = FONT_NAME
        run.font.size = base_size
        run.font.bold = base_bold or seg_bold
        run.font.italic = seg_italic


def _parse_inline_markdown(text: str) -> List[Tuple[str, bool, bool]]:
    """Parse inline markdown into segments of (text, is_bold, is_italic)."""
    segments = []
    pos = 0
    combined = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|(?<!\*)\*([^*]+?)\*(?!\*)")

    for m in combined.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False, False))
        if m.group(1):  # bold+italic
            segments.append((m.group(1), True, True))
        elif m.group(2):  # bold
            segments.append((m.group(2), True, False))
        elif m.group(3):  # italic
            segments.append((m.group(3), False, True))
        pos = m.end()

    if pos < len(text):
        segments.append((text[pos:], False, False))

    return segments if segments else [(text, False, False)]


def _is_caption_block(lines: List[str], start: int) -> Tuple[bool, int]:
    """Detect if current position starts a caption block. Return (is_caption, end_index)."""
    if start >= len(lines):
        return False, start

    caption_count = 0
    end = start
    for i in range(start, min(start + 25, len(lines))):
        line = lines[i].strip()
        if not line:
            if caption_count >= 2:
                end = i
                break
            continue
        if RE_CAPTION_LINE.search(line):
            caption_count += 1
            end = i + 1
        elif caption_count == 0:
            break
        else:
            end = i + 1

    return caption_count >= 2, end


def _add_caption_block(doc: Document, lines: List[str], start: int, end: int) -> None:
    """Add caption block with single-spaced, centered formatting."""
    for i in range(start, end):
        line = lines[i].strip()
        if not line:
            continue
        p = doc.add_paragraph()
        _add_formatted_text(p, line, base_bold=False, base_size=FONT_SIZE_CAPTION)
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        # Center court name and case number lines
        upper = line.upper()
        if any(kw in upper for kw in ["CIRCUIT COURT", "COURT OF APPEALS", "SUPREME COURT",
                                       "DISTRICT COURT", "CASE NO", "HON."]):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True


def _add_table(doc: Document, header_cells: List[str], rows: List[List[str]]) -> None:
    """Add a formatted Word table from markdown table data."""
    if not header_cells:
        return
    n_cols = len(header_cells)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_BODY
        run.font.bold = True
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data[:n_cols]):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_formatted_text(p, cell_text.strip(), base_size=FONT_SIZE_BODY)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def convert_filing(md_path: str, output_path: str = None) -> str:
    """
    Convert a markdown court filing to a properly formatted .docx file.

    Args:
        md_path: Path to the input .md file.
        output_path: Optional path for the output .docx file.
                     Defaults to same directory/name with .docx extension.

    Returns:
        Path to the generated .docx file.
    """
    md_path = os.path.abspath(md_path)
    if not os.path.isfile(md_path):
        raise FileNotFoundError(f"Filing not found: {md_path}")

    with open(md_path, "r", encoding="utf-8-sig") as f:
        content = f.read()

    if output_path is None:
        output_path = os.path.splitext(md_path)[0] + ".docx"
    output_path = os.path.abspath(output_path)

    doc_type = _detect_doc_type(content)
    logger.info(f"Converting {os.path.basename(md_path)} as '{doc_type}' document")

    doc = Document()
    _set_document_margins(doc)
    _add_page_numbers(doc)

    # Set default font for Normal style
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.line_spacing_rule = LINE_SPACING
    style.paragraph_format.space_after = Pt(0)

    # Configure heading styles
    for level, size in [(2, FONT_SIZE_HEADING2), (3, FONT_SIZE_HEADING3)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = FONT_NAME
        h_style.font.size = size
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(6)
        h_style.paragraph_format.keep_with_next = True

    lines = content.split("\n")
    i = 0
    in_table = False
    table_header: List[str] = []
    table_rows: List[List[str]] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Page break markers
        if RE_PAGE_BREAK.match(stripped):
            doc.add_page_break()
            i += 1
            continue

        # Horizontal rules → page break between major sections
        if RE_HORIZ_RULE.match(stripped):
            doc.add_page_break()
            i += 1
            continue

        # Caption block detection (at document start)
        if i < 5:
            is_caption, cap_end = _is_caption_block(lines, i)
            if is_caption:
                _add_caption_block(doc, lines, i, cap_end)
                i = cap_end
                continue

        # Flush pending table
        if in_table and not RE_TABLE_ROW.match(stripped):
            _add_table(doc, table_header, table_rows)
            table_header = []
            table_rows = []
            in_table = False

        # Table rows
        if RE_TABLE_ROW.match(stripped):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if RE_TABLE_SEP.match(stripped):
                # Separator row — skip
                i += 1
                continue
            if not in_table:
                table_header = cells
                in_table = True
            else:
                table_rows.append(cells)
            i += 1
            continue

        # Heading 2
        m = RE_HEADING2.match(stripped)
        if m:
            p = doc.add_heading(m.group(1), level=2)
            for run in p.runs:
                run.font.name = FONT_NAME
            i += 1
            continue

        # Heading 3
        m = RE_HEADING3.match(stripped)
        if m:
            p = doc.add_heading(m.group(1), level=3)
            for run in p.runs:
                run.font.name = FONT_NAME
            i += 1
            continue

        # Heading 4 → bold paragraph
        m = RE_HEADING4.match(stripped)
        if m:
            p = doc.add_paragraph()
            _add_formatted_text(p, m.group(1), base_bold=True, base_size=FONT_SIZE_BODY)
            _set_paragraph_format(p, keep_with_next=True)
            i += 1
            continue

        # Top-level heading (# Title) → centered bold title
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_formatted_text(p, title_text, base_bold=True, base_size=Pt(14))
            _set_paragraph_format(p)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Numbered paragraph
        m = RE_NUMBERED.match(stripped)
        if m:
            num = m.group(1)
            text = m.group(2)
            p = doc.add_paragraph()
            run_num = p.add_run(f"{num}. ")
            run_num.font.name = FONT_NAME
            run_num.font.size = FONT_SIZE_BODY
            run_num.font.bold = True
            _add_formatted_text(p, text, base_size=FONT_SIZE_BODY)
            _set_paragraph_format(p)
            i += 1
            continue

        # Regular paragraph — may span multiple lines until blank line
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line:
                break
            # Stop if next line is a heading, table, rule, or numbered para
            if (next_line.startswith("#") or RE_TABLE_ROW.match(next_line) or
                    RE_HORIZ_RULE.match(next_line) or RE_PAGE_BREAK.match(next_line) or
                    RE_NUMBERED.match(next_line)):
                break
            para_lines.append(next_line)
            i += 1

        full_text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_formatted_text(p, full_text, base_size=FONT_SIZE_BODY)
        _set_paragraph_format(p)

    # Flush any remaining table
    if in_table:
        _add_table(doc, table_header, table_rows)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info(f"Saved: {output_path}")

    return output_path


def convert_all_filings(filing_dir: str) -> list:
    """
    Batch-convert all .md files in a directory to .docx.

    Args:
        filing_dir: Directory containing .md filing documents.

    Returns:
        List of dicts with 'input', 'output', 'status', and 'error' keys.
    """
    filing_dir = os.path.abspath(filing_dir)
    if not os.path.isdir(filing_dir):
        raise NotADirectoryError(f"Directory not found: {filing_dir}")

    results = []
    md_files = sorted(Path(filing_dir).rglob("*.md"))

    for md_file in md_files:
        md_path = str(md_file)
        entry = {"input": md_path, "output": None, "status": "pending", "error": None}
        try:
            output = convert_filing(md_path)
            entry["output"] = output
            entry["status"] = "success"
        except Exception as e:
            entry["status"] = "error"
            entry["error"] = str(e)[:200]
            logger.error(f"Failed to convert {md_path}: {e}")
        results.append(entry)

    success = sum(1 for r in results if r["status"] == "success")
    logger.info(f"Batch complete: {success}/{len(results)} files converted")
    return results


# ── CLI ────────────────────────────────────────────────────────────
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(message)s")

    if len(sys.argv) < 2:
        print("Usage:")
        print("  python filing_converter.py <file.md>              # Convert single file")
        print("  python filing_converter.py --batch <directory>    # Convert all .md in dir")
        sys.exit(1)

    if sys.argv[1] == "--batch":
        if len(sys.argv) < 3:
            print("Error: --batch requires a directory path")
            sys.exit(1)
        results = convert_all_filings(sys.argv[2])
        for r in results:
            status = "✓" if r["status"] == "success" else "✗"
            print(f"  {status} {os.path.basename(r['input'])}", end="")
            if r["error"]:
                print(f"  — {r['error']}", end="")
            print()
        success = sum(1 for r in results if r["status"] == "success")
        print(f"\n{success}/{len(results)} files converted.")
    else:
        md_path = sys.argv[1]
        output = sys.argv[2] if len(sys.argv) > 2 else None
        try:
            result = convert_filing(md_path, output)
            print(f"✓ Converted: {result}")
        except Exception as e:
            print(f"✗ Error: {e}", file=sys.stderr)
            sys.exit(1)
