"""
Extract key SHADY OAKS ledger evidence. Fixed encoding.
"""
import os, sys

OUT = r"C:\Users\andre\temp\ledger_extraction.txt"
lines = []

# 1. Read ledger CSV
csv_path = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\shady_oaks_park_mhp_llc_LEDGER.csv"
if os.path.exists(csv_path):
    with open(csv_path, 'rb') as f:
        raw = f.read()
    null_count = raw.count(b'\x00')
    lines.append(f"=== LEDGER CSV: {len(raw)} bytes, {null_count} null bytes ===")
    if null_count > 100:
        lines.append("SPOLIATION: CSV is null-padded")
        non_null = bytes(b for b in raw if b != 0)
        sample = non_null[:1000].decode('utf-8', errors='replace')
        lines.append(f"Non-null content: {repr(sample[:500])}")
    else:
        sample = raw[:3000].decode('utf-8', errors='replace')
        lines.append(sample)

# 2. Read COOKING THE BOOKS PDF
ctb_pdf = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\ledger COOKING THE BOOKS shady oaks.pdf"
if os.path.exists(ctb_pdf):
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(ctb_pdf)
        lines.append(f"\n=== COOKING THE BOOKS PDF ({len(doc)} pages) ===")
        for i in range(min(len(doc), 8)):
            page = doc[i]
            tp = page.get_textpage()
            text = tp.get_text_bounded()
            text = text.encode('ascii', errors='replace').decode('ascii')
            if text.strip():
                lines.append(f"--- Page {i+1} ---")
                lines.append(text[:2000])
        doc.close()
    except Exception as e:
        lines.append(f"CTB PDF error: {e}")

# 3. Motion for Sanctions DOCX
ms_docx = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\04_MOTION_FOR_SANCTIONS_LEDGER_FRAUD.docx"
if os.path.exists(ms_docx):
    try:
        import docx
        doc = docx.Document(ms_docx)
        lines.append(f"\n=== MOTION FOR SANCTIONS DOCX ===")
        for p in doc.paragraphs[:100]:
            t = p.text.encode('ascii', errors='replace').decode('ascii')
            if t.strip():
                lines.append(t)
    except Exception as e:
        lines.append(f"DOCX error: {e}")

# 4. Gmail false ledger email
gmail_pdf = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\Gmail -SHADY OAKS CASSANDRA UNPROFESSIONAL FALSE LEDGER Andrew Pigors Current Ledger.pdf"
if os.path.exists(gmail_pdf):
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(gmail_pdf)
        lines.append(f"\n=== GMAIL FALSE LEDGER EMAIL ({len(doc)} pages) ===")
        for i in range(min(len(doc), 5)):
            page = doc[i]
            tp = page.get_textpage()
            text = tp.get_text_bounded().encode('ascii', errors='replace').decode('ascii')
            if text.strip():
                lines.append(f"--- Page {i+1} ---")
                lines.append(text[:2000])
        doc.close()
    except Exception as e:
        lines.append(f"Gmail PDF error: {e}")

# 5. Andrew Pigors Ledger 8-22
ap_pdf = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\Andrew Pigors Ledger 8-22.pdf"
if os.path.exists(ap_pdf):
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(ap_pdf)
        lines.append(f"\n=== ANDREW PIGORS LEDGER 8-22 ({len(doc)} pages) ===")
        for i in range(min(len(doc), 5)):
            page = doc[i]
            tp = page.get_textpage()
            text = tp.get_text_bounded().encode('ascii', errors='replace').decode('ascii')
            if text.strip():
                lines.append(f"--- Page {i+1} ---")
                lines.append(text[:2000])
        doc.close()
    except Exception as e:
        lines.append(f"AP Ledger error: {e}")

# Write output
with open(OUT, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print('\n'.join(lines))
print(f"\n\nSaved to: {OUT}")
