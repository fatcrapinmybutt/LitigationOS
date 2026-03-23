#!/usr/bin/env python3
"""
LitigationOS DOCX Converter — Michigan Court Filing Formatter
Converts markdown filing documents to properly formatted DOCX for Michigan courts.

Usage:
  python docx_converter.py input.md output.docx --case-number "2024-001507-DC" --court "14th Circuit"
  python docx_converter.py --batch 01_FILINGS/ --output 01_FILINGS/DOCX/
"""

import argparse
import logging
import os
import re
import sys
from pathlib import Path
from datetime import datetime
from xml.sax.saxutils import escape as xml_escape

sys.stdout = open(sys.stdout.fileno(), mode="w", encoding="utf-8", errors="replace")

logger = logging.getLogger("docx_converter")
if not logger.handlers:
    _handler = logging.StreamHandler(sys.stderr)
    _handler.setFormatter(logging.Formatter("%(levelname)s: %(message)s"))
    logger.addHandler(_handler)
    logger.setLevel(logging.INFO)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── Constants ──────────────────────────────────────────────────────────────────

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
HEADING_SIZE = Pt(14)
CAPTION_SIZE = Pt(12)
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)

PLAINTIFF = "ANDREW JAMES PIGORS"
DEFENDANT = "EMILY A. WATSON"
JUDGE = "Hon. Jenny L. McNeill"
DEFAULT_COURT = "14TH CIRCUIT COURT, FAMILY DIVISION"
DEFAULT_CASE = "2024-001507-DC"
COUNTY = "MUSKEGON"

SIGNATURE_BLOCK = (
    "Respectfully submitted,\n\n"
    "Date: {date}\n\n\n"
    "____________________________________\n"
    "Andrew James Pigors, Pro Se\n"
    "1977 Whitehall Road, Lot 17\n"
    "North Muskegon, MI 49445\n"
    "(231) 903-5690\n"
    "andrewjpigors@gmail.com"
)

CERTIFICATE_OF_SERVICE = (
    "CERTIFICATE OF SERVICE\n\n"
    "I hereby certify that on {date}, I served a true and correct copy of the "
    "foregoing document upon the following party(ies) by [U.S. Mail / E-Filing / Personal Service]:\n\n"
    "Emily A. Watson\n"
    "2160 Garland Drive\n"
    "Norton Shores, MI 49441\n\n\n"
    "____________________________________\n"
    "Andrew James Pigors"
)

# Legal citation patterns for italic formatting
CITATION_PATTERNS = [
    r"\b(\d+\s+Mich(?:\s+App)?\s+\d+)",      # Michigan cases: 123 Mich 456, 123 Mich App 456
    r"\b(\d+\s+NW(?:2d)?\s+\d+)",             # NW reporters
    r"\b(\d+\s+F(?:\.\d+[a-z]*)?\s+\d+)",     # Federal reporters
    r"\b(\d+\s+US\s+\d+)",                     # US reports
    r"\b(\d+\s+S\.?\s*Ct\.?\s+\d+)",           # Supreme Court reporter
]

# Statute reference patterns for bold
STATUTE_PATTERNS = [
    r"(MCL\s+\d+[\.\d]*\w*)",                  # MCL 722.23
    r"(MCR\s+\d+[\.\d]*\w*)",                  # MCR 2.003
    r"(42\s+U\.?S\.?C\.?\s+§?\s*\d+)",         # 42 USC § 1983
    r"(MRE\s+\d+[\.\d]*)",                     # MRE 801
]

# Placeholder patterns — rendered in red highlight so they stand out for review
PLACEHOLDER_RE = re.compile(
    r'(\[(?:VERIFY|ANDREW_REQUIRED|INSERT|ATTACH|UNKNOWN)[^\]]*\])'
)


# ── Document Setup ─────────────────────────────────────────────────────────────

def setup_document():
    """Create a new Document with Michigan court formatting."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.orientation = WD_ORIENT.PORTRAIT

    # Add page number footer
    _add_page_number_footer(section)

    # Define styles
    _setup_styles(doc)

    return doc


def _add_page_number_footer(section):
    """Add centered page numbers to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)

    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)

    for r in [run, run2, run3]:
        r.font.name = FONT_NAME
        r.font.size = Pt(10)


def _setup_styles(doc):
    """Configure document styles."""
    styles = doc.styles

    # Normal style — double-spaced body
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # Heading 1 — 14pt bold centered
    h1 = styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = HEADING_SIZE
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Heading 2 — 12pt bold left-aligned
    h2 = styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = BODY_SIZE
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Block quote — single-spaced, indented
    try:
        bq = styles.add_style("BlockQuote", 1)  # 1 = paragraph style
    except ValueError:
        bq = styles["BlockQuote"]
    bq.font.name = FONT_NAME
    bq.font.size = BODY_SIZE
    bq.paragraph_format.left_indent = Inches(0.5)
    bq.paragraph_format.right_indent = Inches(0.5)
    bq.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    bq.paragraph_format.space_before = Pt(6)
    bq.paragraph_format.space_after = Pt(6)

    # Heading 3 — 12pt bold+underline left-aligned subsections
    h3 = styles["Heading 3"]
    h3.font.name = FONT_NAME
    h3.font.size = BODY_SIZE
    h3.font.bold = True
    h3.font.underline = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Caption style — single-spaced for case caption
    try:
        cap = styles.add_style("CaseCaption", 1)
    except ValueError:
        cap = styles["CaseCaption"]
    cap.font.name = FONT_NAME
    cap.font.size = CAPTION_SIZE
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cap.paragraph_format.space_after = Pt(0)
    cap.paragraph_format.space_before = Pt(0)


# ── Case Caption ───────────────────────────────────────────────────────────────

def add_case_caption(doc, case_number=None, court=None, title=None):
    """Add Michigan court case caption block."""
    case_number = case_number or DEFAULT_CASE
    court = court or DEFAULT_COURT
    court_upper = court.upper()
    if "CIRCUIT" not in court_upper and "COURT" not in court_upper:
        court_upper = f"{court_upper} COURT"

    # STATE OF MICHIGAN
    p = doc.add_paragraph(style="CaseCaption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STATE OF MICHIGAN")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = CAPTION_SIZE

    # IN THE [COURT]
    p = doc.add_paragraph(style="CaseCaption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"IN THE {court_upper}")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = CAPTION_SIZE

    # COUNTY OF MUSKEGON
    p = doc.add_paragraph(style="CaseCaption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"COUNTY OF {COUNTY}")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = CAPTION_SIZE

    # Blank line
    doc.add_paragraph(style="CaseCaption")

    # Party block with case number — use tab stops for alignment
    lines = [
        (f"{PLAINTIFF},", f"Case No. {case_number}"),
        ("    Plaintiff/Petitioner,", f"{JUDGE}"),
        ("", ""),
        ("    v.", ""),
        ("", ""),
        (f"{DEFENDANT},", ""),
        ("    Defendant/Respondent.", ""),
        ("________________________________/", ""),
    ]

    for left_text, right_text in lines:
        p = doc.add_paragraph(style="CaseCaption")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(3.25))

        run_left = p.add_run(left_text)
        run_left.font.name = FONT_NAME
        run_left.font.size = CAPTION_SIZE

        if right_text:
            run_tab = p.add_run("\t")
            run_tab.font.name = FONT_NAME
            run_right = p.add_run(right_text)
            run_right.font.name = FONT_NAME
            run_right.font.size = CAPTION_SIZE

    # Blank line after caption
    doc.add_paragraph(style="CaseCaption")

    # Document title if provided
    if title:
        p = doc.add_paragraph(style="Heading 1")
        run = p.add_run(title.upper())
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = HEADING_SIZE
        run.underline = True


# ── Markdown Parsing ───────────────────────────────────────────────────────────

def parse_markdown(md_text):
    """Parse markdown into structured blocks."""
    blocks = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip YAML frontmatter
        if i == 0 and line.strip() == "---":
            i += 1
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings (h1-h6 — handle all levels to prevent infinite loop)
        if line.startswith("# ") and not line.startswith("## "):
            blocks.append({"type": "h1", "text": line[2:].strip()})
            i += 1
            continue

        if line.startswith("## ") and not line.startswith("### "):
            blocks.append({"type": "h2", "text": line[3:].strip()})
            i += 1
            continue

        if line.startswith("### ") and not line.startswith("#### "):
            blocks.append({"type": "h3", "text": line[4:].strip()})
            i += 1
            continue

        # h4-h6: render as bold text (same as h3 style)
        heading_match = re.match(r"^(#{4,6})\s+(.*)", line)
        if heading_match:
            blocks.append({"type": "h3", "text": heading_match.group(2).strip()})
            i += 1
            continue

        # Horizontal rules
        if re.match(r"^[-*_]{3,}\s*$", line.strip()):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Block quotes — collect consecutive > lines
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            blocks.append({"type": "blockquote", "text": "\n".join(quote_lines).strip()})
            continue

        # Table — lines starting with |
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        # Numbered list
        if re.match(r"^\s*\d+[\.\)]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+[\.\)]\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+[\.\)]\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "numbered_list", "items": items})
            continue

        # Bullet list
        if re.match(r"^\s*[-*+]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*+]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*+]\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "bullet_list", "items": items})
            continue

        # Regular paragraph — collect consecutive non-empty, non-special lines
        para_lines = []
        while i < len(lines):
            l = lines[i]
            if not l.strip():
                break
            if l.startswith("#") or l.strip().startswith(">") or l.strip().startswith("|"):
                break
            if re.match(r"^\s*[-*+]\s+", l) or re.match(r"^\s*\d+[\.\)]\s+", l):
                break
            if re.match(r"^[-*_]{3,}\s*$", l.strip()):
                break
            para_lines.append(l)
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": " ".join(para_lines).strip()})

    return blocks


def _add_formatted_text(paragraph, text):
    """Add text with inline markdown formatting (bold, italic, underline, citations, placeholders)."""
    # Split on placeholder patterns first, render those specially
    segments = PLACEHOLDER_RE.split(text)

    for segment in segments:
        if not segment:
            continue
        if PLACEHOLDER_RE.match(segment):
            # Placeholder — red text on yellow highlight
            run = paragraph.add_run(segment)
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            # Normal text — process inline formatting
            parts = _split_inline(segment)
            for part_text, fmt in parts:
                run = paragraph.add_run(part_text)
                run.font.name = FONT_NAME
                run.font.size = BODY_SIZE
                if "bold" in fmt:
                    run.bold = True
                if "italic" in fmt:
                    run.italic = True
                if "underline" in fmt:
                    run.underline = True


def _split_inline(text):
    """Split text into (text, format_set) tuples for inline formatting."""
    # For very long texts, skip statute/citation auto-detection to avoid performance issues
    if len(text) < 5000:
        # First apply statute patterns — wrap in bold markers (skip already-wrapped)
        for pat in STATUTE_PATTERNS:
            text = re.sub(r'(?<!\*)' + pat + r'(?!\*)', r"**\1**", text)

        # Apply citation patterns — wrap in italic markers (skip already-wrapped)
        for pat in CITATION_PATTERNS:
            text = re.sub(r'(?<!\*)' + pat + r'(?!\*)', r"*\1*", text)

    # Remove duplicate markers from overlapping patterns
    text = text.replace("****", "**")
    text = text.replace("***", "**")

    result = []
    # Parse bold (**text**), italic (*text* or _text_), underline (++text++ or <u>text</u>)
    pattern = re.compile(
        r"\*\*([^*]+)\*\*"         # **bold**
        r"|\*([^*]+)\*"            # *italic*
        r"|_([^_]+)_"              # _italic_
        r"|\+\+([^+]+)\+\+"       # ++underline++
        r"|<u>([^<]+)</u>"         # <u>underline</u>
        r"|([^*_+<]+)"            # plain text
        r"|([*_+<])"              # stray markers
    )

    for m in pattern.finditer(text):
        if m.group(1):
            result.append((m.group(1), {"bold"}))
        elif m.group(2):
            result.append((m.group(2), {"italic"}))
        elif m.group(3):
            result.append((m.group(3), {"italic"}))
        elif m.group(4):
            result.append((m.group(4), {"underline"}))
        elif m.group(5):
            result.append((m.group(5), {"underline"}))
        elif m.group(6):
            result.append((m.group(6), set()))
        elif m.group(7):
            result.append((m.group(7), set()))

    return result if result else [(text, set())]


# ── Block Renderers ────────────────────────────────────────────────────────────

def render_blocks(doc, blocks):
    """Render parsed markdown blocks into the DOCX document."""
    for block in blocks:
        btype = block["type"]

        if btype == "h1":
            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(block["text"].upper())
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = HEADING_SIZE

        elif btype == "h2":
            p = doc.add_paragraph(style="Heading 2")
            run = p.add_run(block["text"])
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE

        elif btype == "h3":
            p = doc.add_paragraph(style="Heading 3")
            run = p.add_run(block["text"])
            run.bold = True
            run.underline = True
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE

        elif btype == "paragraph":
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Inches(0.5)
            _add_formatted_text(p, block["text"])

        elif btype == "blockquote":
            p = doc.add_paragraph(style="BlockQuote")
            _add_formatted_text(p, block["text"])

        elif btype == "bullet_list":
            for item in block["items"]:
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run("•  ")
                run.font.name = FONT_NAME
                run.font.size = BODY_SIZE
                _add_formatted_text(p, item)

        elif btype == "numbered_list":
            for idx, item in enumerate(block["items"], 1):
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run(f"{idx}.  ")
                run.font.name = FONT_NAME
                run.font.size = BODY_SIZE
                _add_formatted_text(p, item)

        elif btype == "table":
            _render_table(doc, block["lines"])

        elif btype == "hr":
            p = doc.add_paragraph(style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE


def _render_table(doc, table_lines):
    """Render a markdown table into a DOCX table."""
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip separator rows (---|---|---)
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = p.add_run(cell_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(10)
                if r_idx == 0:
                    run.bold = True

    # Table spacing
    doc.add_paragraph(style="Normal")


# ── Signature & Certificate ────────────────────────────────────────────────────

def add_signature_block(doc):
    """Add the pro se signature block."""
    doc.add_paragraph(style="Normal")  # spacing
    date_str = datetime.now().strftime("%B %d, %Y")
    text = SIGNATURE_BLOCK.format(date=date_str)

    for line in text.split("\n"):
        p = doc.add_paragraph(style="CaseCaption")
        if line.strip():
            run = p.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE


def add_certificate_of_service(doc):
    """Add certificate of service template."""
    # Page break before certificate
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(break_type=WD_BREAK.PAGE)

    date_str = datetime.now().strftime("%B %d, %Y")
    text = CERTIFICATE_OF_SERVICE.format(date=date_str)

    for line in text.split("\n"):
        if line == "CERTIFICATE OF SERVICE":
            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(line)
            run.bold = True
            run.underline = True
            run.font.name = FONT_NAME
            run.font.size = HEADING_SIZE
        else:
            p = doc.add_paragraph(style="CaseCaption")
            if line.strip():
                run = p.add_run(line)
                run.font.name = FONT_NAME
                run.font.size = BODY_SIZE


# ── Watermark & TOC ────────────────────────────────────────────────────────────

def _add_watermark(doc, text="DRAFT \u2014 NOT FOR FILING"):
    """Add diagonal text watermark to all pages via VML shape in header."""
    safe_text = xml_escape(text)
    # Explicit namespace declarations — python-docx nsdecls() doesn't include VML/Office
    ns = (
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office"'
    )
    try:
        watermark_xml = (
            f'<w:r {ns}>'
            '<w:rPr><w:noProof/></w:rPr>'
            '<w:pict>'
            '<v:shapetype id="_x0000_t136" coordsize="21600,21600" '
            'o:spt="136" adj="10800" '
            'path="m@7,l@8,m@5,21600l@6,21600e">'
            '<v:formulas>'
            '<v:f eqn="sum #0 0 10800"/>'
            '<v:f eqn="prod #0 2 1"/>'
            '<v:f eqn="sum 21600 0 @1"/>'
            '<v:f eqn="sum 0 0 @2"/>'
            '<v:f eqn="sum 21600 0 @3"/>'
            '<v:f eqn="if @0 @3 0"/>'
            '<v:f eqn="if @0 21600 @1"/>'
            '<v:f eqn="if @0 0 @2"/>'
            '<v:f eqn="if @0 @4 21600"/>'
            '<v:f eqn="mid @5 @6"/>'
            '<v:f eqn="mid @8 @5"/>'
            '<v:f eqn="mid @7 @8"/>'
            '<v:f eqn="mid @6 @7"/>'
            '<v:f eqn="sum @6 0 @5"/>'
            '</v:formulas>'
            '<v:path textpathok="t" o:connecttype="custom" '
            'o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" '
            'o:connectangles="270,180,90,0"/>'
            '<v:textpath on="t" fitshape="t"/>'
            '<v:handles><v:h position="#0,bottomRight" xrange="6629,14971"/></v:handles>'
            '<o:lock v:ext="edit" text="t" shapetype="t"/>'
            '</v:shapetype>'
            '<v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s2049" '
            'type="#_x0000_t136" '
            'style="position:absolute;margin-left:0;margin-top:0;'
            'width:527.85pt;height:131.95pt;rotation:315;'
            'z-index:-251658752;'
            'mso-position-horizontal:center;'
            'mso-position-horizontal-relative:margin;'
            'mso-position-vertical:center;'
            'mso-position-vertical-relative:margin" '
            'o:allowincell="f" fillcolor="silver" stroked="f">'
            '<v:fill opacity=".5"/>'
            '<v:textpath style="font-family:&quot;Times New Roman&quot;;font-size:1pt" '
            f'string="{safe_text}"/>'
            '</v:shape>'
            '</w:pict>'
            '</w:r>'
        )
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p._p.append(parse_xml(watermark_xml))
        logger.info("VML watermark applied: %s", text)
    except Exception as exc:
        # Fallback: simple large gray text in header
        logger.warning("VML watermark failed (%s), using header text fallback", exc)
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = Pt(36)
            run.font.color.rgb = RGBColor(192, 192, 192)
            run.bold = True


def _add_toc(doc):
    """Insert a Table of Contents field that Word will populate on open."""
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = HEADING_SIZE

    p = doc.add_paragraph()
    # Begin field
    run = p.add_run()
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    # Field instruction — TOC with heading levels 1-3, hyperlinks, hide page numbers in web
    run2 = p.add_run()
    run2._r.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    ))
    # Separate
    run3 = p.add_run()
    run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    # Placeholder text until Word updates the field
    run4 = p.add_run("[Right-click and select 'Update Field' to generate Table of Contents]")
    run4.font.name = FONT_NAME
    run4.font.size = BODY_SIZE
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.italic = True
    # End field
    run5 = p.add_run()
    run5._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    # Page break after TOC
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(break_type=WD_BREAK.PAGE)
    logger.debug("Table of contents field inserted")


def _estimate_pages(blocks):
    """Rough page count estimate (double-spaced, 12pt Times, 1" margins ≈ 23 lines/page)."""
    lines = 0
    for block in blocks:
        btype = block.get("type", "")
        if btype in ("h1", "h2", "h3", "hr"):
            lines += 3
        elif btype == "paragraph":
            lines += max(1, len(block.get("text", "")) // 80) + 1
        elif btype == "blockquote":
            lines += max(1, len(block.get("text", "")) // 70) + 2
        elif btype in ("bullet_list", "numbered_list"):
            lines += len(block.get("items", [])) + 1
        elif btype == "table":
            lines += len(block.get("lines", [])) + 2
    return max(1, lines // 23)


# ── Metadata Extraction ───────────────────────────────────────────────────────

def extract_metadata(md_text):
    """Extract YAML frontmatter metadata if present."""
    meta = {}
    if md_text.startswith("---"):
        end = md_text.find("---", 3)
        if end != -1:
            yaml_block = md_text[3:end]
            for line in yaml_block.strip().split("\n"):
                if ":" in line:
                    key, val = line.split(":", 1)
                    meta[key.strip().lower().replace("-", "_")] = val.strip().strip('"').strip("'")
    return meta


def extract_title_from_md(md_text):
    """Extract the first H1 heading as the document title."""
    for line in md_text.split("\n"):
        line = line.strip()
        if line.startswith("# ") and not line.startswith("## "):
            return line[2:].strip()
    return None


# ── Main Conversion ───────────────────────────────────────────────────────────

def convert_md_to_docx(md_path, docx_path, case_number=None, court=None,
                       no_caption=False, no_signature=False, no_service=False,
                       draft=False):
    """Convert a markdown file to a formatted DOCX court filing.

    Args:
        md_path: Path to source markdown file.
        docx_path: Path for output DOCX file.
        case_number: Override case number (default from frontmatter or 2024-001507-DC).
        court: Override court name.
        no_caption: Skip the case caption block.
        no_signature: Skip the signature block.
        no_service: Skip the certificate of service.
        draft: Add "DRAFT — NOT FOR FILING" watermark.

    Returns:
        bool: True on success, False on failure.
    """
    md_path = Path(md_path)
    if not md_path.exists():
        logger.error("Input file not found: %s", md_path)
        return False

    try:
        md_text = md_path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        logger.error("Cannot read %s: %s", md_path, exc)
        return False

    # Strip XML-incompatible control characters (NULL bytes, etc.) that crash python-docx
    md_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', md_text)

    # Extract metadata from frontmatter
    meta = extract_metadata(md_text)
    case_number = case_number or meta.get("case_number") or meta.get("case") or DEFAULT_CASE
    court = court or meta.get("court") or DEFAULT_COURT
    title = meta.get("title") or extract_title_from_md(md_text)

    # Create document
    doc = setup_document()

    # Draft watermark
    if draft:
        _add_watermark(doc)

    # Case caption
    if not no_caption:
        add_case_caption(doc, case_number=case_number, court=court, title=title)

    # Parse and render markdown body
    blocks = parse_markdown(md_text)
    # Skip the first h1 block if we already used it as caption title
    if title and blocks and blocks[0]["type"] == "h1":
        blocks = blocks[1:]

    # Auto-generate TOC for long documents (>10 estimated pages)
    if _estimate_pages(blocks) > 10:
        _add_toc(doc)
        logger.info("Auto-inserted TOC (estimated %d pages)", _estimate_pages(blocks))

    render_blocks(doc, blocks)

    # Signature block
    if not no_signature:
        add_signature_block(doc)

    # Certificate of service
    if not no_service:
        add_certificate_of_service(doc)

    # Save
    docx_path = Path(docx_path)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(docx_path))
    except OSError as exc:
        logger.error("Cannot save %s: %s", docx_path, exc)
        return False

    print(f"OK: {md_path.name} \u2192 {docx_path}")
    return True


def batch_convert(input_dir, output_dir, case_number=None, court=None, **kwargs):
    """Convert all markdown files in a directory to DOCX."""
    input_dir = Path(input_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    md_files = sorted(input_dir.glob("*.md"))
    if not md_files:
        logger.warning("No .md files found in %s", input_dir)
        return 0

    count = 0
    errors = []
    for md_file in md_files:
        docx_name = md_file.stem + ".docx"
        docx_path = output_dir / docx_name
        try:
            if convert_md_to_docx(md_file, docx_path, case_number=case_number,
                                  court=court, **kwargs):
                count += 1
            else:
                errors.append(md_file.name)
        except Exception as exc:
            logger.error("Failed to convert %s: %s", md_file.name, exc)
            errors.append(md_file.name)

    print(f"\nBatch complete: {count}/{len(md_files)} files converted to {output_dir}")
    if errors:
        logger.warning("Failed files: %s", ", ".join(errors))
    return count


# ── Public API ─────────────────────────────────────────────────────────────────

def convert_single(input_md, output_docx, **kwargs):
    """Public API: Convert a single markdown file to a court-ready DOCX.

    Args:
        input_md: Path to input markdown file.
        output_docx: Path to output DOCX file.
        **kwargs: Options forwarded to convert_md_to_docx
            (case_number, court, draft, no_caption, no_signature, no_service).

    Returns:
        bool: True if conversion succeeded.
    """
    return convert_md_to_docx(input_md, output_docx, **kwargs)


def convert_directory(input_dir, output_dir, **kwargs):
    """Public API: Batch convert all .md files in a directory to DOCX.

    Args:
        input_dir: Directory containing .md files.
        output_dir: Directory for output .docx files (created if needed).
        **kwargs: Options forwarded to convert_md_to_docx
            (case_number, court, draft, no_caption, no_signature, no_service).

    Returns:
        int: Number of files successfully converted.
    """
    return batch_convert(input_dir, output_dir, **kwargs)


# ── CLI ────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="LitigationOS DOCX Converter — Michigan Court Filing Formatter",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  %(prog)s input.md output.docx\n"
            '  %(prog)s input.md output.docx --case-number "2024-001507-DC"\n'
            "  %(prog)s --batch 01_FILINGS/ --output 01_FILINGS/DOCX/\n"
        ),
    )
    parser.add_argument("input", nargs="?", help="Input markdown file")
    parser.add_argument("output", nargs="?", help="Output DOCX file")
    parser.add_argument("--case-number", "-c", default=None, help="Case number (e.g. 2024-001507-DC)")
    parser.add_argument("--court", default=None, help='Court name (e.g. "14th Circuit")')
    parser.add_argument("--batch", "-b", default=None, help="Batch mode: input directory of .md files")
    parser.add_argument("--output-dir", "-o", default=None, help="Batch mode: output directory for .docx files")
    parser.add_argument("--no-caption", action="store_true", help="Omit case caption")
    parser.add_argument("--no-signature", action="store_true", help="Omit signature block")
    parser.add_argument("--no-service", action="store_true", help="Omit certificate of service")
    parser.add_argument("--draft", action="store_true",
                        help='Add "DRAFT — NOT FOR FILING" watermark to every page')
    parser.add_argument("--verbose", "-v", action="store_true", help="Verbose logging output")

    args = parser.parse_args()

    if args.verbose:
        logger.setLevel(logging.DEBUG)

    opts = dict(
        no_caption=args.no_caption,
        no_signature=args.no_signature,
        no_service=args.no_service,
        draft=args.draft,
    )

    # Batch mode
    if args.batch:
        out_dir = args.output_dir or os.path.join(args.batch, "DOCX")
        count = batch_convert(args.batch, out_dir, case_number=args.case_number,
                              court=args.court, **opts)
        sys.exit(0 if count > 0 else 1)

    # Single file mode
    if not args.input:
        parser.print_help()
        sys.exit(1)

    if not args.output:
        args.output = Path(args.input).with_suffix(".docx")

    ok = convert_md_to_docx(args.input, args.output, case_number=args.case_number,
                            court=args.court, **opts)
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
