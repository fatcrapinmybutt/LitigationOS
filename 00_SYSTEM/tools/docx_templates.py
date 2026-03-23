#!/usr/bin/env python3
"""
LitigationOS DOCX Templates — Michigan Court Document Templates
================================================================

Pre-configured document templates for Michigan courts with proper
margins, fonts, spacing, captions, and page numbering.

Templates:
    MichiganCircuitTemplate  — 14th Circuit Court (Family + Civil)
    MichiganDistrictTemplate — Michigan District Court (61st District)
    MichiganCOATemplate      — Michigan Court of Appeals

Usage:
    from docx_templates import get_template

    # Get a pre-configured Document
    template = get_template("circuit_family")
    doc = template.create_document()
    template.add_caption(doc, title="MOTION TO COMPEL DISCOVERY")
    # ... add content ...
    template.add_signature(doc)
    template.add_certificate_of_service(doc)
    doc.save("output.docx")

    # Or use from CLI:
    python docx_templates.py --list
    python docx_templates.py circuit_family --title "MOTION" --output motion.docx
"""

import sys
import os
from pathlib import Path
from datetime import datetime

sys.stdout = open(sys.stdout.fileno(), mode="w", encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── Base Template ──────────────────────────────────────────────────────────────

class _BaseTemplate:
    """Base class for Michigan court document templates."""

    FONT_NAME = "Times New Roman"
    BODY_SIZE = Pt(12)
    HEADING_SIZE = Pt(14)
    CAPTION_SIZE = Pt(12)
    PAGE_WIDTH = Inches(8.5)
    PAGE_HEIGHT = Inches(11)
    MARGIN_TOP = Inches(1)
    MARGIN_BOTTOM = Inches(1)
    MARGIN_LEFT = Inches(1)
    MARGIN_RIGHT = Inches(1)

    # Subclasses override these
    COURT_NAME = ""
    COUNTY = ""
    CASE_NUMBER = ""
    JUDGE = ""
    PLAINTIFF = "ANDREW JAMES PIGORS"
    DEFENDANT = "EMILY A. WATSON"
    PLAINTIFF_ROLE = "Plaintiff/Petitioner"
    DEFENDANT_ROLE = "Defendant/Respondent"
    FILING_METHOD = "MiFILE"
    PAGE_LIMIT = 20
    MCR_BRIEF = ""
    ADDRESS = ""

    SIGNATURE_NAME = "Andrew James Pigors, Pro Se"
    SIGNATURE_ADDRESS = (
        "1977 Whitehall Road, Lot 17\n"
        "North Muskegon, MI 49445\n"
        "(231) 903-5690\n"
        "andrewjpigors@gmail.com"
    )

    SERVICE_PARTIES = [
        {
            "name": "Emily A. Watson",
            "address": "2160 Garland Drive\nNorton Shores, MI 49441",
            "method": "[U.S. Mail / E-Filing / Personal Service]",
        },
    ]

    def create_document(self):
        """Create a new Document pre-configured with court formatting."""
        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.page_width = self.PAGE_WIDTH
        section.page_height = self.PAGE_HEIGHT
        section.top_margin = self.MARGIN_TOP
        section.bottom_margin = self.MARGIN_BOTTOM
        section.left_margin = self.MARGIN_LEFT
        section.right_margin = self.MARGIN_RIGHT
        section.orientation = WD_ORIENT.PORTRAIT

        # Page numbers
        self._add_page_numbers(section)

        # Styles
        self._setup_styles(doc)

        return doc

    def _add_page_numbers(self, section):
        """Add centered page numbers to footer."""
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        run2 = p.add_run()
        run2._r.append(parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        ))
        run3 = p.add_run()
        run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

        for r in [run, run2, run3]:
            r.font.name = self.FONT_NAME
            r.font.size = Pt(10)

    def _setup_styles(self, doc):
        """Configure standard document styles."""
        styles = doc.styles

        # Normal — double-spaced body
        normal = styles["Normal"]
        normal.font.name = self.FONT_NAME
        normal.font.size = self.BODY_SIZE
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.space_before = Pt(0)

        # Heading 1 — centered bold title
        h1 = styles["Heading 1"]
        h1.font.name = self.FONT_NAME
        h1.font.size = self.HEADING_SIZE
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)
        h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        # Heading 2 — bold left-aligned
        h2 = styles["Heading 2"]
        h2.font.name = self.FONT_NAME
        h2.font.size = self.BODY_SIZE
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        # Heading 3 — bold underline subsection
        h3 = styles["Heading 3"]
        h3.font.name = self.FONT_NAME
        h3.font.size = self.BODY_SIZE
        h3.font.bold = True
        h3.font.underline = True
        h3.font.color.rgb = RGBColor(0, 0, 0)
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)
        h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        # Block quote — single-spaced, indented
        try:
            bq = styles.add_style("BlockQuote", 1)
        except ValueError:
            bq = styles["BlockQuote"]
        bq.font.name = self.FONT_NAME
        bq.font.size = self.BODY_SIZE
        bq.paragraph_format.left_indent = Inches(0.5)
        bq.paragraph_format.right_indent = Inches(0.5)
        bq.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        bq.paragraph_format.space_before = Pt(6)
        bq.paragraph_format.space_after = Pt(6)

        # Caption style — single-spaced for case caption block
        try:
            cap = styles.add_style("CaseCaption", 1)
        except ValueError:
            cap = styles["CaseCaption"]
        cap.font.name = self.FONT_NAME
        cap.font.size = self.CAPTION_SIZE
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cap.paragraph_format.space_after = Pt(0)
        cap.paragraph_format.space_before = Pt(0)

    def add_caption(self, doc, title=None, case_number=None, judge=None):
        """Add court caption block with party names and case number.

        Args:
            doc: Document to add caption to.
            title: Document title (e.g., "MOTION TO COMPEL DISCOVERY").
            case_number: Override the template's default case number.
            judge: Override the template's default judge.
        """
        case_number = case_number or self.CASE_NUMBER
        judge = judge or self.JUDGE

        # STATE OF MICHIGAN
        p = doc.add_paragraph(style="CaseCaption")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("STATE OF MICHIGAN")
        run.bold = True
        run.font.name = self.FONT_NAME
        run.font.size = self.CAPTION_SIZE

        # IN THE [COURT]
        p = doc.add_paragraph(style="CaseCaption")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"IN THE {self.COURT_NAME}")
        run.bold = True
        run.font.name = self.FONT_NAME
        run.font.size = self.CAPTION_SIZE

        # COUNTY (if applicable)
        if self.COUNTY:
            p = doc.add_paragraph(style="CaseCaption")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"COUNTY OF {self.COUNTY}")
            run.bold = True
            run.font.name = self.FONT_NAME
            run.font.size = self.CAPTION_SIZE

        doc.add_paragraph(style="CaseCaption")  # blank line

        # Party block with case number
        lines = [
            (f"{self.PLAINTIFF},", f"Case No. {case_number}"),
            (f"    {self.PLAINTIFF_ROLE},", f"{judge}" if judge else ""),
            ("", ""),
            ("    v.", ""),
            ("", ""),
            (f"{self.DEFENDANT},", ""),
            (f"    {self.DEFENDANT_ROLE}.", ""),
            ("________________________________/", ""),
        ]

        for left_text, right_text in lines:
            p = doc.add_paragraph(style="CaseCaption")
            p.paragraph_format.tab_stops.add_tab_stop(Inches(3.25))
            run_left = p.add_run(left_text)
            run_left.font.name = self.FONT_NAME
            run_left.font.size = self.CAPTION_SIZE
            if right_text:
                p.add_run("\t").font.name = self.FONT_NAME
                run_right = p.add_run(right_text)
                run_right.font.name = self.FONT_NAME
                run_right.font.size = self.CAPTION_SIZE

        doc.add_paragraph(style="CaseCaption")  # blank line

        # Document title
        if title:
            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(title.upper())
            run.bold = True
            run.underline = True
            run.font.name = self.FONT_NAME
            run.font.size = self.HEADING_SIZE

    def add_signature(self, doc, date_str=None):
        """Add pro se signature block with date."""
        date_str = date_str or datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(style="Normal")

        sig_lines = [
            "Respectfully submitted,",
            "",
            f"Date: {date_str}",
            "",
            "",
            "____________________________________",
            self.SIGNATURE_NAME,
        ] + self.SIGNATURE_ADDRESS.split("\n")

        for line in sig_lines:
            p = doc.add_paragraph(style="CaseCaption")
            if line.strip():
                run = p.add_run(line)
                run.font.name = self.FONT_NAME
                run.font.size = self.BODY_SIZE

    def add_certificate_of_service(self, doc, date_str=None, method=None):
        """Add certificate of service with party addresses.

        Args:
            doc: Document to append to.
            date_str: Date of service (default: today).
            method: Service method override for all parties.
        """
        date_str = date_str or datetime.now().strftime("%B %d, %Y")

        # Page break
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(break_type=WD_BREAK.PAGE)

        # Title
        p = doc.add_paragraph(style="Heading 1")
        run = p.add_run("CERTIFICATE OF SERVICE")
        run.bold = True
        run.underline = True
        run.font.name = self.FONT_NAME
        run.font.size = self.HEADING_SIZE

        # Body
        p = doc.add_paragraph(style="CaseCaption")
        run = p.add_run(
            f"I hereby certify that on {date_str}, I served a true and correct copy "
            "of the foregoing document upon the following party(ies) by "
        )
        run.font.name = self.FONT_NAME
        run.font.size = self.BODY_SIZE

        # List each party
        for party in self.SERVICE_PARTIES:
            svc_method = method or party.get("method", "[U.S. Mail / E-Filing / Personal Service]")
            p = doc.add_paragraph(style="CaseCaption")
            run = p.add_run(svc_method + ":")
            run.font.name = self.FONT_NAME
            run.font.size = self.BODY_SIZE

            doc.add_paragraph(style="CaseCaption")
            for line in [party["name"]] + party["address"].split("\n"):
                p = doc.add_paragraph(style="CaseCaption")
                run = p.add_run(line)
                run.font.name = self.FONT_NAME
                run.font.size = self.BODY_SIZE

        # Signature
        doc.add_paragraph(style="CaseCaption")
        doc.add_paragraph(style="CaseCaption")
        p = doc.add_paragraph(style="CaseCaption")
        run = p.add_run("____________________________________")
        run.font.name = self.FONT_NAME
        run.font.size = self.BODY_SIZE
        p = doc.add_paragraph(style="CaseCaption")
        run = p.add_run(self.SIGNATURE_NAME.split(",")[0])
        run.font.name = self.FONT_NAME
        run.font.size = self.BODY_SIZE

    def info(self):
        """Return a dict describing this template's configuration."""
        return {
            "court": self.COURT_NAME,
            "county": self.COUNTY,
            "case_number": self.CASE_NUMBER,
            "judge": self.JUDGE,
            "page_limit": self.PAGE_LIMIT,
            "mcr_brief": self.MCR_BRIEF,
            "filing_method": self.FILING_METHOD,
            "margins": f"{self.MARGIN_TOP}, {self.MARGIN_BOTTOM}, {self.MARGIN_LEFT}, {self.MARGIN_RIGHT}",
        }


# ── Michigan Circuit Court ─────────────────────────────────────────────────────

class MichiganCircuitTemplate(_BaseTemplate):
    """14th Circuit Court — Family Division (Pigors v. Watson).

    Pre-configured for:
      - Case No. 2024-001507-DC
      - Hon. Jenny L. McNeill
      - MCR 2.119(A)(2) brief rules (20-page limit)
      - MiFILE e-filing
    """

    COURT_NAME = "14TH JUDICIAL CIRCUIT COURT \u2014 FAMILY DIVISION"
    COUNTY = "MUSKEGON"
    CASE_NUMBER = "2024-001507-DC"
    JUDGE = "Hon. Jenny L. McNeill"
    PAGE_LIMIT = 20
    MCR_BRIEF = "MCR 2.119(A)(2)"
    FILING_METHOD = "MiFILE (mifile.courts.michigan.gov)"
    ADDRESS = "990 Terrace Street, Muskegon, MI 49442"


class MichiganCircuitCivilTemplate(_BaseTemplate):
    """14th Circuit Court — Civil Division (Shady Oaks housing case).

    Pre-configured for:
      - Case No. 2025-002760-CZ
      - MCR 2.119(A)(2) brief rules
    """

    COURT_NAME = "14TH JUDICIAL CIRCUIT COURT"
    COUNTY = "MUSKEGON"
    CASE_NUMBER = "2025-002760-CZ"
    JUDGE = "[ASSIGNED JUDGE]"
    PAGE_LIMIT = 20
    MCR_BRIEF = "MCR 2.119(A)(2)"
    FILING_METHOD = "MiFILE"
    ADDRESS = "990 Terrace Street, Muskegon, MI 49442"


# ── Michigan District Court ────────────────────────────────────────────────────

class MichiganDistrictTemplate(_BaseTemplate):
    """Michigan 61st District Court.

    Pre-configured for:
      - General district court filings
      - Smaller margins (0.75" sides) per local practice
      - 10-page brief limit for district motions
    """

    COURT_NAME = "61ST DISTRICT COURT"
    COUNTY = "MUSKEGON"
    CASE_NUMBER = "[TO BE ASSIGNED]"
    JUDGE = ""
    PAGE_LIMIT = 10
    MCR_BRIEF = "MCR 2.119(A)(2)"
    FILING_METHOD = "MiFILE"
    ADDRESS = "990 Terrace Street, Muskegon, MI 49442"


# ── Michigan Court of Appeals ──────────────────────────────────────────────────

class MichiganCOATemplate(_BaseTemplate):
    """Michigan Court of Appeals.

    Pre-configured for:
      - COA Case No. 366810
      - MCR 7.212(B) brief rules (50-page limit)
      - Larger margins per COA requirements
      - No county in caption (statewide court)
    """

    COURT_NAME = "MICHIGAN COURT OF APPEALS"
    COUNTY = ""
    CASE_NUMBER = "COA 366810"
    JUDGE = ""
    MARGIN_TOP = Inches(1)
    MARGIN_BOTTOM = Inches(1)
    MARGIN_LEFT = Inches(1.25)
    MARGIN_RIGHT = Inches(1)
    PAGE_LIMIT = 50
    MCR_BRIEF = "MCR 7.212(B)"
    FILING_METHOD = "MiFILE"
    ADDRESS = "Hall of Justice, 925 W Ottawa St, Lansing, MI 48915"

    PLAINTIFF_ROLE = "Plaintiff-Appellant"
    DEFENDANT_ROLE = "Defendant-Appellee"


# ── Template Registry ──────────────────────────────────────────────────────────

TEMPLATES = {
    "circuit_family": MichiganCircuitTemplate,
    "circuit_civil": MichiganCircuitCivilTemplate,
    "district": MichiganDistrictTemplate,
    "coa": MichiganCOATemplate,
}


def get_template(court_type):
    """Factory function to get a court template instance.

    Args:
        court_type: One of "circuit_family", "circuit_civil", "district", "coa".

    Returns:
        Template instance with pre-configured court settings.

    Raises:
        ValueError: If court_type is not recognized.
    """
    cls = TEMPLATES.get(court_type)
    if cls is None:
        valid = ", ".join(sorted(TEMPLATES.keys()))
        raise ValueError(f"Unknown court type '{court_type}'. Valid types: {valid}")
    return cls()


def list_templates():
    """Return info for all available templates."""
    return {name: cls().info() for name, cls in TEMPLATES.items()}


# ── CLI ────────────────────────────────────────────────────────────────────────

def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="LitigationOS DOCX Templates — Michigan Court Document Templates",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Available templates:\n"
            "  circuit_family  — 14th Circuit Court, Family Division\n"
            "  circuit_civil   — 14th Circuit Court, Civil Division\n"
            "  district        — 61st District Court\n"
            "  coa             — Michigan Court of Appeals\n"
            "\nExamples:\n"
            '  %(prog)s circuit_family --title "MOTION TO COMPEL" -o motion.docx\n'
            "  %(prog)s coa --title \"APPLICATION FOR LEAVE\" -o app_leave.docx\n"
            "  %(prog)s --list\n"
        ),
    )
    parser.add_argument("template", nargs="?", help="Template name (see --list)")
    parser.add_argument("--list", action="store_true", help="List available templates")
    parser.add_argument("--title", "-t", default=None, help="Document title")
    parser.add_argument("--output", "-o", default=None, help="Output DOCX path")
    parser.add_argument("--case-number", "-c", default=None, help="Override case number")
    parser.add_argument("--no-caption", action="store_true", help="Omit case caption")
    parser.add_argument("--no-signature", action="store_true", help="Omit signature block")
    parser.add_argument("--no-service", action="store_true", help="Omit certificate of service")

    args = parser.parse_args()

    if args.list:
        print("Available Michigan Court Templates:")
        print("=" * 60)
        for name, info in list_templates().items():
            print(f"\n  {name}:")
            for k, v in info.items():
                if v:
                    print(f"    {k}: {v}")
        sys.exit(0)

    if not args.template:
        parser.print_help()
        sys.exit(1)

    try:
        template = get_template(args.template)
    except ValueError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)

    doc = template.create_document()

    if not args.no_caption:
        template.add_caption(doc, title=args.title, case_number=args.case_number)

    # Add placeholder body text
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run("[Document body goes here — use docx_converter.py to convert markdown content]")
    run.font.name = template.FONT_NAME
    run.font.size = template.BODY_SIZE
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    if not args.no_signature:
        template.add_signature(doc)
    if not args.no_service:
        template.add_certificate_of_service(doc)

    output = args.output or f"{args.template}_template.docx"
    Path(output).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"OK: Template '{args.template}' saved to {output}")


if __name__ == "__main__":
    main()
