#!/usr/bin/env python3
"""
SCAO form overlay engine (DOCX). Fills fields into DOCX templates.
Usage:
  python scao_overlay.py --form MC20 --data data.json --template templates/MC20_template.docx --out OUT/MC20_filled.docx
Deps (optional): pip install python-docx
Maps file: forms/maps/scao_maps.json
"""
import json, argparse
from pathlib import Path
try:
    from docx import Document
except Exception:
    Document=None

def apply_text(doc, mapping, payload):
    # Very simple text replace in all paragraphs and runs for keys like {{FIELD}}
    repl = {f"{{{{{k}}}}}":str(payload.get(v,"")) for k,v in mapping.items()}
    for p in doc.paragraphs:
        for k,v in repl.items():
            if k in p.text:
                for r in p.runs:
                    if k in r.text: r.text=r.text.replace(k, v)
    # also in tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for k,v in repl.items():
                    if k in cell.text:
                        for pr in cell.paragraphs:
                            for r in pr.runs:
                                if k in r.text: r.text=r.text.replace(k, v)

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument('--form', required=True, help='e.g., MC20, MC230, FOC50')
    ap.add_argument('--data', required=True, help='JSON with field:value')
    ap.add_argument('--template', required=True, help='DOCX template with {{FIELD}} tokens')
    ap.add_argument('--out', required=True, help='Output DOCX')
    ap.add_argument('--maps', default='forms/maps/scao_maps.json')
    a=ap.parse_args()
    maps=json.loads(Path(a.maps).read_text(encoding='utf-8'))
    if a.form not in maps: print(f"[!] No map for {a.form}"); raise SystemExit(1)
    mapping=maps[a.form]
    payload=json.loads(Path(a.data).read_text(encoding='utf-8'))
    if Document is None: print("[!] Install python-docx to fill DOCX: pip install python-docx"); raise SystemExit(1)
    doc=Document(a.template)
    apply_text(doc, mapping, payload)
    Path(a.out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(a.out)
    print("[OK]", a.out)
if __name__=='__main__': main()
