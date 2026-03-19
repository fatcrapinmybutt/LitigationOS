#!/usr/bin/env python3
"""
Document Assembly Engine v1.0 — LitigationOS
Template-to-DOCX-to-PDF pipeline with variable substitution from litigation_context.db.

Usage:
  python doc_assembly_engine.py --input filing.md --output filing.pdf --case COA_366810
  python doc_assembly_engine.py --input filing.md --format docx --output filing.docx
  python doc_assembly_engine.py --batch-dir 02_TRIAL_14TH\\WATSON_TORT\\ --output-dir output\\
"""

import os
import sys
import re
import json
import glob
import sqlite3
import argparse
from datetime import datetime

sys.stdout.reconfigure(encoding="utf-8")

# ─── python-docx imports ─────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── reportlab imports ────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, KeepTogether,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ─── Constants ────────────────────────────────────────────────────────────────

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DB_PATH = os.path.join(BASE_DIR, "litigation_context.db")

# Court-format presets
COURT_FORMAT = {
    "font": "Times New Roman",
    "size": 12,
    "line_spacing": 2.0,   # double-spaced
    "margin": 1.0,         # inches
}

# Case-identifier mapping: prefix -> court_address_book entity_id
CASE_PREFIX_MAP = {
    "COA":     "mi_court_of_appeals",
    "TRIAL":   "14th_circuit_court",
    "14TH":    "14th_circuit_court",
    "FED":     "us_district_court_wdmi",
    "WDMI":    "us_district_court_wdmi",
    "MSC":     "mi_supreme_court",
    "JTC":     "jtc",
    "BAR":     "attorney_grievance_commission",
}

# ─── Database helpers ─────────────────────────────────────────────────────────

def get_db(db_path=None):
    """Return a WAL-mode connection with generous timeouts."""
    path = db_path or DB_PATH
    conn = sqlite3.connect(path, timeout=120)
    conn.execute("PRAGMA busy_timeout=60000")
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA cache_size=-32000")
    conn.row_factory = sqlite3.Row
    return conn


def load_case_variables(conn, case_id):
    """
    Build a variable dict from the DB for a given case identifier.
    case_id examples: COA_366810, TRIAL_001507, 14TH_001507
    """
    variables = {}
    variables["DATE"] = datetime.now().strftime("%B %d, %Y")

    # Parse prefix
    parts = case_id.split("_", 1)
    prefix = parts[0].upper() if parts else ""
    case_num_hint = parts[1] if len(parts) > 1 else ""

    # ── Parties from party_contacts ──
    try:
        rows = conn.execute(
            "SELECT party_name, role FROM party_contacts ORDER BY rowid"
        ).fetchall()
        for r in rows:
            role_upper = (r["role"] or "").upper()
            if "PLAINTIFF" in role_upper:
                variables["PLAINTIFF"] = r["party_name"]
            elif "DEFENDANT" in role_upper:
                variables["DEFENDANT"] = r["party_name"]
        # Fallbacks
        variables.setdefault("PLAINTIFF", "PLAINTIFF")
        variables.setdefault("DEFENDANT", "DEFENDANT")
    except Exception:
        variables.setdefault("PLAINTIFF", "PLAINTIFF")
        variables.setdefault("DEFENDANT", "DEFENDANT")

    # ── Court & judge from court_address_book ──
    entity_id = CASE_PREFIX_MAP.get(prefix, "14th_circuit_court")
    try:
        row = conn.execute(
            "SELECT * FROM court_address_book WHERE entity_id = ?", (entity_id,)
        ).fetchone()
        if row:
            variables["COURT_NAME"] = row["name"]
            extra = json.loads(row["extra_json"]) if row["extra_json"] else {}
            variables["JUDGE"] = extra.get("judge", "")
            variables["CASE_NUMBER"] = extra.get("case_number", case_num_hint)
        else:
            variables["COURT_NAME"] = ""
            variables["JUDGE"] = ""
            variables["CASE_NUMBER"] = case_num_hint
    except Exception:
        variables.setdefault("COURT_NAME", "")
        variables.setdefault("JUDGE", "")
        variables.setdefault("CASE_NUMBER", case_num_hint)

    # ── Additional lookup for [PLACEHOLDER:xxx] via case_intelligence_hub ──
    # Stored as a callable for deferred resolution
    variables["_conn"] = conn
    return variables


def resolve_placeholder(conn, tag):
    """Resolve a [PLACEHOLDER:xxx] tag by searching common DB tables."""
    tag_lower = tag.lower().strip()

    # Try case_intelligence_hub first
    try:
        row = conn.execute(
            "SELECT summary FROM case_intelligence_hub WHERE entity_id = ? LIMIT 1",
            (tag,),
        ).fetchone()
        if row and row["summary"]:
            return row["summary"]
    except Exception:
        pass

    # Try court_address_book
    try:
        row = conn.execute(
            "SELECT name FROM court_address_book WHERE entity_id = ? LIMIT 1",
            (tag_lower,),
        ).fetchone()
        if row and row["name"]:
            return row["name"]
    except Exception:
        pass

    # Try party_contacts
    try:
        row = conn.execute(
            "SELECT party_name FROM party_contacts WHERE party_name LIKE ? LIMIT 1",
            (f"%{tag}%",),
        ).fetchone()
        if row and row["party_name"]:
            return row["party_name"]
    except Exception:
        pass

    return f"[{tag}]"


# ─── Variable substitution ───────────────────────────────────────────────────

def substitute(text, variables):
    """
    Replace {{VAR}} and [PLACEHOLDER:xxx] tags with values from *variables* dict.
    Unknown {{VAR}} are left as-is; [PLACEHOLDER:xxx] are resolved via DB.
    """
    conn = variables.get("_conn")

    # {{VAR}} replacement
    def _replace_var(m):
        key = m.group(1).strip()
        return variables.get(key, m.group(0))

    text = re.sub(r"\{\{(\w+)\}\}", _replace_var, text)

    # [PLACEHOLDER:xxx] replacement
    def _replace_ph(m):
        tag = m.group(1).strip()
        if conn:
            return resolve_placeholder(conn, tag)
        return f"[{tag}]"

    text = re.sub(r"\[PLACEHOLDER:(.*?)\]", _replace_ph, text)
    return text


# ─── Markdown parsing ────────────────────────────────────────────────────────

def parse_markdown(text):
    """
    Parse markdown into structured blocks.
    Returns list of dicts with keys: type, text, level (for headings), num (for numbered).
    """
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    lines = text.split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Page break
        if stripped in ("---", "***", "___") or re.match(r"^[-*_]{3,}$", stripped):
            blocks.append({"type": "separator"})
            i += 1
            continue

        # Headings
        hm = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if hm:
            blocks.append({"type": "heading", "level": len(hm.group(1)), "text": hm.group(2).strip()})
            i += 1
            continue

        # Table formatting rows — skip
        if re.match(r"^\|[-\s|:]+\|$", stripped):
            i += 1
            continue
        # Table data rows — keep as paragraph
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            blocks.append({"type": "table_row", "text": "  |  ".join(cells)})
            i += 1
            continue

        # Numbered list
        nm = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if nm:
            blocks.append({"type": "numbered", "num": nm.group(1), "text": nm.group(2)})
            i += 1
            continue

        # Bullet
        if stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append({"type": "bullet", "text": stripped[2:].strip()})
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            blocks.append({"type": "blockquote", "text": stripped[2:].strip()})
            i += 1
            continue

        # Regular paragraph
        blocks.append({"type": "paragraph", "text": stripped})
        i += 1

    return blocks


def clean_md(text):
    """Strip bold/italic markdown markers."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    return text.strip()


# ─── DOCX generation ─────────────────────────────────────────────────────────

def _add_rich_run(paragraph, text, font_name, font_size):
    """Add runs to a paragraph, handling **bold** and *italic* markdown."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(font_size)


def _set_paragraph_spacing(paragraph, line_spacing=2.0):
    """Set double-spacing on a paragraph."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = Pt(12 * line_spacing)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def _add_page_number(section):
    """Insert page number in the footer center."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page number field
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)

    run.font.name = COURT_FORMAT["font"]
    run.font.size = Pt(COURT_FORMAT["size"])


def build_caption_block(doc, variables):
    """
    Add a standard court caption block at the top of the document.
    """
    font = COURT_FORMAT["font"]
    sz = COURT_FORMAT["size"]

    # Court name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(variables.get("COURT_NAME", "").upper())
    run.bold = True
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    # Blank line
    doc.add_paragraph()

    # Caption table: Plaintiff v. Defendant | Case No.
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)

    # Remove table borders — use borderless style
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                borders.append(el)
            tcPr.append(borders)

    # Row 0: plaintiff
    c = table.cell(0, 0)
    p = c.paragraphs[0]
    run = p.add_run(variables.get("PLAINTIFF", "").upper() + ",")
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    c = table.cell(0, 1)
    p = c.paragraphs[0]
    run = p.add_run("Case No. " + variables.get("CASE_NUMBER", ""))
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    # Row 1: role labels and judge
    c = table.cell(1, 0)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plaintiff,")
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    c = table.cell(1, 1)
    p = c.paragraphs[0]
    judge = variables.get("JUDGE", "")
    if judge:
        run = p.add_run(judge)
        run.font.name = font
        run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    # Row 2: v. + defendant
    c = table.cell(2, 0)
    p = c.paragraphs[0]
    run = p.add_run("v.\n" + variables.get("DEFENDANT", "").upper() + ",")
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    c = table.cell(2, 1)
    # empty
    _set_paragraph_spacing(c.paragraphs[0], 1.0)

    # Defendant role label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("        Defendant.")
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    # Separator line
    p = doc.add_paragraph()
    run = p.add_run("_" * 65)
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    # Blank line before body
    doc.add_paragraph()


def build_cos_block(doc, variables):
    """Append a Certificate of Service block at the end."""
    font = COURT_FORMAT["font"]
    sz = COURT_FORMAT["size"]

    p = doc.add_paragraph()
    _set_paragraph_spacing(p, 2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CERTIFICATE OF SERVICE")
    run.bold = True
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 2.0)

    date_str = variables.get("DATE", datetime.now().strftime("%B %d, %Y"))
    plaintiff = variables.get("PLAINTIFF", "Plaintiff")
    defendant = variables.get("DEFENDANT", "Defendant")

    cos_text = (
        f"I, {plaintiff}, hereby certify that on {date_str}, "
        f"I served a true and correct copy of the foregoing document upon "
        f"{defendant} and/or their counsel of record by electronic filing "
        f"via MiFILE and/or by first-class U.S. Mail, postage prepaid, "
        f"to the last known address on file."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_rich_run(p, cos_text, font, sz)
    _set_paragraph_spacing(p, 2.0)

    # Signature line
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("_" * 40)
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    p = doc.add_paragraph()
    run = p.add_run(plaintiff)
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    p = doc.add_paragraph()
    run = p.add_run("Pro Se Litigant")
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)

    p = doc.add_paragraph()
    run = p.add_run(f"Dated: {date_str}")
    run.font.name = font
    run.font.size = Pt(sz)
    _set_paragraph_spacing(p, 1.0)


def markdown_to_docx(md_text, variables, output_path, include_caption=True, include_cos=True):
    """Convert substituted markdown to a court-formatted DOCX file."""
    font = COURT_FORMAT["font"]
    sz = COURT_FORMAT["size"]
    margin = COURT_FORMAT["margin"]

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    _add_page_number(section)

    # ── Caption block ──
    if include_caption:
        build_caption_block(doc, variables)

    # ── Body ──
    blocks = parse_markdown(md_text)
    for blk in blocks:
        btype = blk["type"]

        if btype == "separator":
            doc.add_page_break()
            continue

        if btype == "heading":
            level = blk.get("level", 1)
            p = doc.add_paragraph()
            if level <= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_md(blk["text"]).upper() if level == 1 else clean_md(blk["text"]))
            run.bold = True
            run.font.name = font
            run.font.size = Pt(sz + (2 if level == 1 else 0))
            if level == 1:
                run.underline = True
            _set_paragraph_spacing(p, 2.0)
            continue

        if btype == "numbered":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(f"{blk['num']}. ")
            run.bold = True
            run.font.name = font
            run.font.size = Pt(sz)
            _add_rich_run(p, blk["text"], font, sz)
            _set_paragraph_spacing(p, 2.0)
            continue

        if btype == "bullet":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run("\u2022 ")
            run.font.name = font
            run.font.size = Pt(sz)
            _add_rich_run(p, blk["text"], font, sz)
            _set_paragraph_spacing(p, 2.0)
            continue

        if btype == "blockquote":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.75)
            p.paragraph_format.right_indent = Inches(0.75)
            _add_rich_run(p, blk["text"], font, sz)
            run = p.runs[0] if p.runs else p.add_run("")
            run.italic = True
            _set_paragraph_spacing(p, 1.5)
            continue

        # Default: paragraph / table_row
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.5)
        _add_rich_run(p, blk.get("text", ""), font, sz)
        _set_paragraph_spacing(p, 2.0)

    # ── Certificate of Service ──
    if include_cos:
        build_cos_block(doc, variables)

    doc.save(output_path)
    return output_path


# ─── PDF generation (reportlab) ──────────────────────────────────────────────

def _register_fonts():
    """Register Times New Roman if available on system, else fall back to Times-Roman."""
    tnr_paths = [
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\Times New Roman.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
    ]
    for p in tnr_paths:
        if os.path.isfile(p):
            try:
                pdfmetrics.registerFont(TTFont("TimesNewRoman", p))
                return "TimesNewRoman"
            except Exception:
                continue
    # Fallback to built-in
    return "Times-Roman"


def _build_pdf_styles(font_name):
    """Create reportlab paragraph styles for court documents."""
    styles = getSampleStyleSheet()

    base = ParagraphStyle(
        "CourtBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=12,
        leading=24,          # double-spaced (12pt * 2)
        alignment=TA_JUSTIFY,
        firstLineIndent=36,  # 0.5 inch
        spaceAfter=0,
        spaceBefore=0,
    )
    styles.add(base)

    styles.add(ParagraphStyle(
        "CourtHeading1",
        parent=base,
        fontSize=14,
        leading=28,
        alignment=TA_CENTER,
        firstLineIndent=0,
        spaceAfter=12,
        spaceBefore=12,
        underline=True,
    ))

    styles.add(ParagraphStyle(
        "CourtHeading2",
        parent=base,
        fontSize=12,
        leading=24,
        alignment=TA_CENTER,
        firstLineIndent=0,
        spaceAfter=6,
        spaceBefore=6,
    ))

    styles.add(ParagraphStyle(
        "CourtCaption",
        parent=base,
        alignment=TA_CENTER,
        firstLineIndent=0,
        leading=18,
    ))

    styles.add(ParagraphStyle(
        "CourtCaptionLeft",
        parent=base,
        alignment=TA_LEFT,
        firstLineIndent=0,
        leading=18,
    ))

    styles.add(ParagraphStyle(
        "CourtBlockquote",
        parent=base,
        leftIndent=54,
        rightIndent=54,
        firstLineIndent=0,
        leading=18,
    ))

    styles.add(ParagraphStyle(
        "CourtBullet",
        parent=base,
        leftIndent=36,
        firstLineIndent=0,
    ))

    styles.add(ParagraphStyle(
        "CourtCOS",
        parent=base,
        firstLineIndent=0,
    ))

    return styles


def _md_to_rl(text):
    """Convert markdown bold/italic to reportlab XML tags."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    text = re.sub(r"__(.+?)__", r"<b>\1</b>", text)
    text = re.sub(r"_(.+?)_", r"<i>\1</i>", text)
    return text


def _page_footer(canvas, doc_template):
    """Draw page number in footer."""
    canvas.saveState()
    canvas.setFont("Times-Roman", 12)
    canvas.drawCentredString(
        letter[0] / 2, 0.5 * inch,
        str(canvas.getPageNumber()),
    )
    canvas.restoreState()


def markdown_to_pdf(md_text, variables, output_path, include_caption=True, include_cos=True):
    """Convert substituted markdown directly to a court-formatted PDF."""
    font_name = _register_fonts()
    styles = _build_pdf_styles(font_name)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
    )

    story = []

    # ── Caption block ──
    if include_caption:
        court = variables.get("COURT_NAME", "").upper()
        plaintiff = variables.get("PLAINTIFF", "").upper()
        defendant = variables.get("DEFENDANT", "").upper()
        case_no = variables.get("CASE_NUMBER", "")
        judge = variables.get("JUDGE", "")

        story.append(Paragraph(f"<b>{court}</b>", styles["CourtCaption"]))
        story.append(Spacer(1, 18))
        story.append(Paragraph(f"{plaintiff},", styles["CourtCaptionLeft"]))
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;Plaintiff,", styles["CourtCaptionLeft"]))
        story.append(Spacer(1, 6))

        right_block = f"Case No. {case_no}"
        if judge:
            right_block += f"<br/>{judge}"
        story.append(Paragraph(right_block, styles["CourtCaptionLeft"]))
        story.append(Spacer(1, 6))

        story.append(Paragraph("v.", styles["CourtCaptionLeft"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"{defendant},", styles["CourtCaptionLeft"]))
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;Defendant.", styles["CourtCaptionLeft"]))
        story.append(Spacer(1, 6))
        story.append(Paragraph("_" * 80, styles["CourtCaptionLeft"]))
        story.append(Spacer(1, 18))

    # ── Body ──
    blocks = parse_markdown(md_text)
    for blk in blocks:
        btype = blk["type"]

        if btype == "separator":
            story.append(PageBreak())
            continue

        if btype == "heading":
            level = blk.get("level", 1)
            txt = clean_md(blk["text"])
            if level == 1:
                story.append(Paragraph(f"<b><u>{txt.upper()}</u></b>", styles["CourtHeading1"]))
            else:
                story.append(Paragraph(f"<b>{txt}</b>", styles["CourtHeading2"]))
            continue

        if btype == "numbered":
            txt = _md_to_rl(blk["text"])
            story.append(Paragraph(f"<b>{blk['num']}.</b> {txt}", styles["CourtBody"]))
            continue

        if btype == "bullet":
            txt = _md_to_rl(blk["text"])
            story.append(Paragraph(f"\u2022 {txt}", styles["CourtBullet"]))
            continue

        if btype == "blockquote":
            txt = _md_to_rl(blk["text"])
            story.append(Paragraph(f"<i>{txt}</i>", styles["CourtBlockquote"]))
            continue

        # paragraph / table_row
        txt = _md_to_rl(blk.get("text", ""))
        if txt.strip():
            story.append(Paragraph(txt, styles["CourtBody"]))

    # ── Certificate of Service ──
    if include_cos:
        story.append(Spacer(1, 24))
        story.append(Paragraph("<b>CERTIFICATE OF SERVICE</b>", styles["CourtHeading2"]))
        story.append(Spacer(1, 12))

        date_str = variables.get("DATE", datetime.now().strftime("%B %d, %Y"))
        plaintiff = variables.get("PLAINTIFF", "Plaintiff")
        defendant = variables.get("DEFENDANT", "Defendant")

        cos_text = (
            f"I, {plaintiff}, hereby certify that on {date_str}, "
            f"I served a true and correct copy of the foregoing document upon "
            f"{defendant} and/or their counsel of record by electronic filing "
            f"via MiFILE and/or by first-class U.S. Mail, postage prepaid, "
            f"to the last known address on file."
        )
        story.append(Paragraph(cos_text, styles["CourtCOS"]))
        story.append(Spacer(1, 48))
        story.append(Paragraph("_" * 50, styles["CourtCOS"]))
        story.append(Paragraph(plaintiff, styles["CourtCOS"]))
        story.append(Paragraph("Pro Se Litigant", styles["CourtCOS"]))
        story.append(Paragraph(f"Dated: {date_str}", styles["CourtCOS"]))

    doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)
    return output_path


# ─── Pipeline orchestrator ────────────────────────────────────────────────────

def assemble(input_path, output_path, case_id=None, fmt="pdf",
             include_caption=True, include_cos=True, db_path=None):
    """
    Full pipeline: read markdown → substitute → generate DOCX or PDF.
    Returns the output file path.
    """
    # Read template
    with open(input_path, "r", encoding="utf-8", errors="replace") as f:
        md_text = f.read()

    # Load variables
    conn = get_db(db_path)
    if case_id:
        variables = load_case_variables(conn, case_id)
    else:
        # Default variables without specific case
        variables = {
            "DATE": datetime.now().strftime("%B %d, %Y"),
            "PLAINTIFF": "PLAINTIFF",
            "DEFENDANT": "DEFENDANT",
            "COURT_NAME": "",
            "JUDGE": "",
            "CASE_NUMBER": "",
            "_conn": conn,
        }
        # Try loading from default court entry
        try:
            row = conn.execute(
                "SELECT * FROM court_address_book WHERE entity_id = '14th_circuit_court'"
            ).fetchone()
            if row:
                variables["COURT_NAME"] = row["name"]
                extra = json.loads(row["extra_json"]) if row["extra_json"] else {}
                variables["JUDGE"] = extra.get("judge", "")
                variables["CASE_NUMBER"] = extra.get("case_number", "")
        except Exception:
            pass
        # Load parties
        try:
            rows = conn.execute("SELECT party_name, role FROM party_contacts ORDER BY rowid").fetchall()
            for r in rows:
                role_upper = (r["role"] or "").upper()
                if "PLAINTIFF" in role_upper:
                    variables["PLAINTIFF"] = r["party_name"]
                elif "DEFENDANT" in role_upper:
                    variables["DEFENDANT"] = r["party_name"]
        except Exception:
            pass

    # Substitute
    md_text = substitute(md_text, variables)

    # Generate
    fmt = fmt.lower()
    if fmt == "docx":
        result = markdown_to_docx(md_text, variables, output_path, include_caption, include_cos)
    elif fmt == "both":
        docx_path = output_path.rsplit(".", 1)[0] + ".docx"
        markdown_to_docx(md_text, variables, docx_path, include_caption, include_cos)
        result = markdown_to_pdf(md_text, variables, output_path, include_caption, include_cos)
    else:
        result = markdown_to_pdf(md_text, variables, output_path, include_caption, include_cos)

    conn.close()
    return result


def batch_assemble(batch_dir, output_dir, case_id=None, fmt="pdf", db_path=None):
    """Process all .md files in batch_dir, writing output to output_dir."""
    os.makedirs(output_dir, exist_ok=True)
    md_files = sorted(glob.glob(os.path.join(batch_dir, "*.md")))
    if not md_files:
        print(f"No .md files found in {batch_dir}")
        return []

    results = []
    for md_path in md_files:
        base = os.path.splitext(os.path.basename(md_path))[0]
        ext = "docx" if fmt == "docx" else "pdf"
        out_path = os.path.join(output_dir, f"{base}.{ext}")
        try:
            assemble(md_path, out_path, case_id=case_id, fmt=fmt, db_path=db_path)
            print(f"  OK: {md_path} -> {out_path}")
            results.append(out_path)
        except Exception as e:
            print(f"  FAIL: {md_path}: {e}")
    return results


# ─── Skill registration ──────────────────────────────────────────────────────

def register_skill(db_path=None):
    """Register this engine as a skill in the LitigationOS skill registry."""
    conn = get_db(db_path)
    conn.execute("""
        INSERT OR REPLACE INTO skill_registry (id, name, file_path, version, db_tables_used, methods)
        VALUES (
            'doc_assembly',
            'doc_assembly',
            'engines/doc_assembly_engine.py',
            '1.0',
            'party_contacts,court_address_book,case_intelligence_hub',
            'assemble,batch_assemble,markdown_to_docx,markdown_to_pdf,substitute'
        )
    """, )
    conn.commit()
    conn.close()
    print("Skill 'doc_assembly' registered in skill_registry.")


# ─── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="LitigationOS Document Assembly Engine — template-to-DOCX/PDF pipeline",
    )
    parser.add_argument("--input", "-i", help="Path to input markdown template")
    parser.add_argument("--output", "-o", help="Output file path (.pdf or .docx)")
    parser.add_argument("--format", "-f", default="pdf", choices=["pdf", "docx", "both"],
                        help="Output format (default: pdf)")
    parser.add_argument("--case", "-c", help="Case ID (e.g. COA_366810, TRIAL_001507)")
    parser.add_argument("--batch-dir", help="Directory of .md files for batch processing")
    parser.add_argument("--output-dir", help="Output directory for batch mode")
    parser.add_argument("--no-caption", action="store_true", help="Omit caption block")
    parser.add_argument("--no-cos", action="store_true", help="Omit certificate of service")
    parser.add_argument("--db", help="Override database path")
    parser.add_argument("--register", action="store_true", help="Register skill in DB and exit")
    parser.add_argument("--test", action="store_true", help="Run built-in test")

    args = parser.parse_args()

    if args.register:
        register_skill(args.db)
        return

    if args.test:
        run_test(args.db)
        return

    if args.batch_dir:
        if not args.output_dir:
            print("ERROR: --output-dir required for batch mode")
            sys.exit(1)
        results = batch_assemble(
            args.batch_dir, args.output_dir,
            case_id=args.case, fmt=args.format, db_path=args.db,
        )
        print(f"Batch complete: {len(results)} files generated.")
        return

    if not args.input:
        parser.print_help()
        sys.exit(1)

    if not args.output:
        ext = "docx" if args.format == "docx" else "pdf"
        base = os.path.splitext(args.input)[0]
        args.output = f"{base}.{ext}"

    result = assemble(
        args.input, args.output,
        case_id=args.case,
        fmt=args.format,
        include_caption=not args.no_caption,
        include_cos=not args.no_cos,
        db_path=args.db,
    )
    print(f"Generated: {result}")


# ─── Built-in test ────────────────────────────────────────────────────────────

def run_test(db_path=None):
    """Generate a test document to verify end-to-end pipeline."""
    test_md = """# Motion for Summary Disposition

## Statement of Facts

1. On or about **January 15, 2024**, {{PLAINTIFF}} filed an action against {{DEFENDANT}} in the {{COURT_NAME}}, Case No. {{CASE_NUMBER}}, assigned to {{JUDGE}}.

2. The Defendant has failed to comply with the court's prior orders regarding parenting time and custody arrangements.

3. Discovery has revealed that the Defendant made material misrepresentations to the court in prior proceedings.

## Legal Standard

Summary disposition is proper under *MCR 2.116(C)(10)* when there is no genuine issue of material fact and the moving party is entitled to judgment as a matter of law. *Maiden v Rozwood*, 461 Mich 109, 120 (1999).

> The court must consider the affidavits, pleadings, depositions, admissions, and other documentary evidence submitted by the parties in the light most favorable to the nonmoving party.

## Argument

The undisputed facts establish that:

- Defendant violated the custody order on multiple documented occasions
- Defendant made false statements under oath regarding her compliance
- The children's best interests require modification of the current arrangement

## Relief Requested

**WHEREFORE**, {{PLAINTIFF}} respectfully requests that this Honorable Court:

1. Grant summary disposition in favor of Plaintiff
2. Modify the existing custody order
3. Award such other relief as the Court deems just and equitable

Respectfully submitted,

{{PLAINTIFF}}, Pro Se
Dated: {{DATE}}
"""
    test_dir = os.path.join(BASE_DIR, "00_SYSTEM")
    os.makedirs(test_dir, exist_ok=True)

    # Write temp markdown
    tmp_md = os.path.join(test_dir, "_test_template.md")
    with open(tmp_md, "w", encoding="utf-8") as f:
        f.write(test_md)

    # Generate PDF
    pdf_out = os.path.join(test_dir, "test_doc_output.pdf")
    assemble(tmp_md, pdf_out, case_id="TRIAL_001507", fmt="pdf", db_path=db_path)
    print(f"Test PDF: {pdf_out} ({os.path.getsize(pdf_out):,} bytes)")

    # Generate DOCX too
    docx_out = os.path.join(test_dir, "test_doc_output.docx")
    assemble(tmp_md, docx_out, case_id="TRIAL_001507", fmt="docx", db_path=db_path)
    print(f"Test DOCX: {docx_out} ({os.path.getsize(docx_out):,} bytes)")

    # Cleanup temp
    try:
        os.remove(tmp_md)
    except OSError:
        pass

    # Register skill
    register_skill(db_path)
    print("End-to-end test PASSED.")


if __name__ == "__main__":
    main()
