#!/usr/bin/env python3
"""
Filing Assembly Pipeline Engine
Converts markdown filing stacks to court-ready DOCX format.
Supports: State (14th Circuit), COA, MSC, Federal (WDMI), Administrative
"""

import os
import sys
import re
import glob
import sqlite3
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Inches, Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ─── Constants ────────────────────────────────────────────────────────────────

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DB_PATH = os.path.join(BASE_DIR, "litigation_context.db")

CASE_INFO = {
    "plaintiff": "ANDREW J. PIGORS",
    "case_name": "Pigors v. Watson",
    "trial_case_no": "2024-001507-DC",
    "coa_case_no": "366810",
    "court_14th": "14th Judicial Circuit Court, Muskegon County",
    "court_coa": "Michigan Court of Appeals",
    "court_wdmi": "United States District Court, Western District of Michigan, Southern Division",
    "judge": "Hon. Jenny L. McNeill",
}

COURT_FORMATS = {
    "state": {"rule": "MCR 2.119", "font": "Times New Roman", "size": 12, "spacing": 2, "margins": 1.0},
    "coa":   {"rule": "MCR 7.212", "font": "Times New Roman", "size": 12, "spacing": 2, "margins": 1.0},
    "msc":   {"rule": "MCR 7.306", "font": "Times New Roman", "size": 12, "spacing": 2, "margins": 1.0},
    "federal": {"rule": "FRCP 8",  "font": "Times New Roman", "size": 12, "spacing": 2, "margins": 1.0},
    "admin": {"rule": "Agency",    "font": "Times New Roman", "size": 12, "spacing": 2, "margins": 1.0},
}


# ─── Pipeline Functions ──────────────────────────────────────────────────────

def read_stack(stack_path):
    """Read all .md files in a stack directory, sorted by filename."""
    if not os.path.isdir(stack_path):
        print(f"  WARNING: Stack path not found: {stack_path}")
        return []
    md_files = sorted(glob.glob(os.path.join(stack_path, "*.md")))
    results = []
    for fp in md_files:
        basename = os.path.basename(fp)
        # Skip readiness/instruction/score files - not court content
        skip_keywords = ["READINESS", "FILING_INSTRUCTIONS", "FILING_READINESS"]
        if any(kw in basename.upper() for kw in skip_keywords):
            continue
        try:
            with open(fp, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            results.append({"filename": basename, "path": fp, "content": content})
        except Exception as e:
            print(f"  WARNING: Could not read {fp}: {e}")
    print(f"  Read {len(results)} files from {stack_path}")
    return results


def merge_documents(files):
    """Combine multiple markdown files into a single document string."""
    merged = []
    for f in files:
        merged.append(f["content"])
    return "\n\n---\n\n".join(merged)


def _strip_html_comments(text):
    """Remove HTML comments from markdown."""
    return re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)


def _md_to_paragraphs(text):
    """
    Parse markdown into a list of dicts with type and content.
    Types: heading1, heading2, heading3, paragraph, numbered, bullet, separator, table_row
    """
    text = _strip_html_comments(text)
    lines = text.split('\n')
    paragraphs = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            i += 1
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            paragraphs.append({"type": "heading1", "text": stripped[2:].strip()})
            i += 1
            continue
        if stripped.startswith('## ') and not stripped.startswith('### '):
            paragraphs.append({"type": "heading2", "text": stripped[3:].strip()})
            i += 1
            continue
        if stripped.startswith('### '):
            paragraphs.append({"type": "heading3", "text": stripped[4:].strip()})
            i += 1
            continue

        # Skip markdown table formatting rows
        if re.match(r'^\|[-\s|]+\|$', stripped):
            i += 1
            continue

        # Table rows (caption tables) - skip them, caption handled separately
        if stripped.startswith('|') and stripped.endswith('|'):
            i += 1
            continue

        # Numbered paragraphs
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            paragraphs.append({"type": "numbered", "num": num_match.group(1), "text": num_match.group(2)})
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            paragraphs.append({"type": "bullet", "text": stripped[2:].strip()})
            i += 1
            continue

        # Block quotes
        if stripped.startswith('> '):
            paragraphs.append({"type": "paragraph", "text": stripped[2:].strip()})
            i += 1
            continue

        # Regular paragraph
        if stripped and not stripped.startswith('[') or (stripped.startswith('[') and not stripped.startswith('[TABLE')):
            paragraphs.append({"type": "paragraph", "text": stripped})

        i += 1

    return paragraphs


def _clean_md_formatting(text):
    """Remove markdown bold/italic markers for clean text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    return text.strip()


def _add_bold_italic_run(paragraph, text, font_name="Times New Roman", font_size=12):
    """Add a run handling **bold** and *italic* markdown markers."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(font_size)


def format_for_court(document_text, court_type):
    """Apply court-specific formatting rules. Returns formatting metadata."""
    fmt = COURT_FORMATS.get(court_type, COURT_FORMATS["state"])
    return {
        "court_type": court_type,
        "rule": fmt["rule"],
        "font": fmt["font"],
        "size": fmt["size"],
        "spacing": fmt["spacing"],
        "margins": fmt["margins"],
        "content": document_text,
    }


def _set_cell_border(cell, **kwargs):
    """Set cell border for table cells."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), val.get('space', '0'))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_page_numbers(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _add_caption_state(doc, case_name, case_no, court_name, judge=None, doc_title=None):
    """Add state court caption block."""
    # Court name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STATE OF MICHIGAN")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(court_name.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # Caption table
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True

    # Plaintiff
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run("ANDREW J. PIGORS,\nPlaintiff,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    run = p.add_run(f"Case No. {case_no}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # v.
    cell = table.cell(1, 0)
    p = cell.paragraphs[0]
    run = p.add_run("v.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    if judge:
        cell = table.cell(1, 1)
        p = cell.paragraphs[0]
        run = p.add_run(judge)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Defendant
    cell = table.cell(2, 0)
    p = cell.paragraphs[0]
    defendant = case_name.split(" v. ")[-1] if " v. " in case_name else "DEFENDANT"
    run = p.add_run(f"{defendant.upper()},\nDefendant.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Remove table borders for caption
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                left={"val": "none", "sz": "0", "color": "FFFFFF"},
                right={"val": "none", "sz": "0", "color": "FFFFFF"})

    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Document title
    if doc_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(doc_title.upper())
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    doc.add_paragraph()  # spacer


def _add_caption_coa(doc, case_no, lower_case_no, doc_title=None):
    """Add Court of Appeals caption block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STATE OF MICHIGAN")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IN THE COURT OF APPEALS")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Caption table
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True

    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run("ANDREW J. PIGORS,\nPlaintiff-Appellant,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    run = p.add_run(f"Court of Appeals No. {case_no}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(1, 0)
    p = cell.paragraphs[0]
    run = p.add_run("v.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(1, 1)
    p = cell.paragraphs[0]
    run = p.add_run(f"Lower Court Case No. {lower_case_no}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(2, 0)
    p = cell.paragraphs[0]
    run = p.add_run("TIFFANY A. WATSON,\nDefendant-Appellee.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(2, 1)
    p = cell.paragraphs[0]
    run = p.add_run(CASE_INFO["judge"])
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                left={"val": "none", "sz": "0", "color": "FFFFFF"},
                right={"val": "none", "sz": "0", "color": "FFFFFF"})

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.name = "Times New Roman"

    if doc_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(doc_title.upper())
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    doc.add_paragraph()


def _add_caption_federal(doc, case_name, doc_title=None):
    """Add federal court caption block."""
    for line in [
        "UNITED STATES DISTRICT COURT",
        "WESTERN DISTRICT OF MICHIGAN",
        "SOUTHERN DIVISION",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.autofit = True

    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run("ANDREW J. PIGORS,\nPlaintiff,")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    run = p.add_run("Case No. [TO BE ASSIGNED]")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(1, 0)
    p = cell.paragraphs[0]
    run = p.add_run("v.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    cell = table.cell(1, 1)
    p = cell.paragraphs[0]
    run = p.add_run("JURY TRIAL DEMANDED")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    defendants = case_name.split(" v. ")[-1] if " v. " in case_name else "DEFENDANTS"
    cell = table.cell(2, 0)
    p = cell.paragraphs[0]
    run = p.add_run(f"{defendants.upper()},\nDefendants.")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                left={"val": "none", "sz": "0", "color": "FFFFFF"},
                right={"val": "none", "sz": "0", "color": "FFFFFF"})

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.name = "Times New Roman"

    if doc_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(doc_title.upper())
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    doc.add_paragraph()


def generate_certificate_of_service(doc, court_type="state"):
    """Append MCR 2.107 / FRCP 5 compliant certificate of service."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CERTIFICATE OF SERVICE")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    doc.add_paragraph()

    if court_type == "federal":
        rule_ref = "Federal Rule of Civil Procedure 5(b)(2)"
        method = "CM/ECF electronic filing system"
    else:
        rule_ref = "MCR 2.107(C)(3)"
        method = "first-class U.S. Mail, postage prepaid, and/or MiFILE electronic service"

    text = (
        f"I, Andrew J. Pigors, hereby certify that on {datetime.now().strftime('%B %d, %Y')}, "
        f"I served a true and correct copy of the foregoing document upon all parties and/or their "
        f"counsel of record by {method}, in compliance with {rule_ref}, addressed as follows:"
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    _add_bold_italic_run(p, text)

    # Service list placeholder
    p = doc.add_paragraph()
    run = p.add_run("[OPPOSING COUNSEL / PARTY NAME]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("[ADDRESS LINE 1]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("[CITY, STATE ZIP]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    p = doc.add_paragraph()
    run = p.add_run("_" * 40)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Andrew J. Pigors, Pro Se")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(f"Dated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def generate_cover_sheet(doc, stack_info, court_type="state"):
    """Generate court-specific cover sheet as first page."""
    if court_type == "coa":
        title = "MICHIGAN COURT OF APPEALS\nFILING COVER SHEET"
    elif court_type == "federal":
        title = "CIVIL COVER SHEET (JS-44)"
    else:
        title = "CIRCUIT COURT FILING COVER SHEET\nSCAO APPROVED"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    doc.add_paragraph()

    fields = [
        ("Case Name:", stack_info.get("case_name", CASE_INFO["case_name"])),
        ("Case Number:", stack_info.get("case_no", "[TO BE ASSIGNED]")),
        ("Court:", stack_info.get("court", "")),
        ("Filing Party:", "Andrew J. Pigors, Plaintiff, Pro Se"),
        ("Document Type:", stack_info.get("doc_type", "")),
        ("Date Filed:", datetime.now().strftime("%B %d, %Y")),
    ]

    for label, value in fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}  ")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run = p.add_run(value)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_page_break()


def generate_fee_waiver(doc):
    """Generate MC 20 fee waiver template."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FEE WAIVER REQUEST")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(MC 20 / Application to Waive Fees)")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_paragraph()

    text = (
        "Plaintiff Andrew J. Pigors respectfully requests that the Court waive all filing fees "
        "and costs associated with this action pursuant to MCR 2.002 and MCL 600.2529. "
        "Plaintiff is unable to pay fees and costs because of indigency."
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    _add_bold_italic_run(p, text)

    doc.add_paragraph()

    sections = [
        "1. I am unable to pay the fees and costs of this action.",
        "2. I receive or am eligible for public assistance: [  ] Yes  [  ] No",
        "3. My gross income for the past 12 months: $__________",
        "4. Number of dependents: __________",
        "5. I have attached a completed MC 20 Affidavit of Indigency.",
    ]
    for s in sections:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(s)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("_" * 40)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Andrew J. Pigors, Pro Se")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(f"Dated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def generate_docx(formatted_doc, output_path, caption_type="state", doc_title=None,
                  case_name=None, case_no=None, court_name=None, judge=None,
                  include_cover=True, include_fee_waiver=False, stack_info=None):
    """Create Word document with court formatting from markdown content."""
    doc = Document()
    fmt = COURT_FORMATS.get(caption_type, COURT_FORMATS["state"])

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(fmt["margins"])
        section.bottom_margin = Inches(fmt["margins"])
        section.left_margin = Inches(fmt["margins"])
        section.right_margin = Inches(fmt["margins"])

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = fmt["font"]
    font.size = Pt(fmt["size"])
    pf = style.paragraph_format
    pf.line_spacing = fmt["spacing"]

    # Cover sheet
    if include_cover and stack_info:
        generate_cover_sheet(doc, stack_info, caption_type)

    # Caption
    if caption_type == "coa":
        _add_caption_coa(doc, 
                         case_no or CASE_INFO["coa_case_no"],
                         CASE_INFO["trial_case_no"],
                         doc_title)
    elif caption_type == "federal":
        _add_caption_federal(doc, case_name or CASE_INFO["case_name"], doc_title)
    else:
        _add_caption_state(doc,
                           case_name or CASE_INFO["case_name"],
                           case_no or CASE_INFO["trial_case_no"],
                           court_name or CASE_INFO["court_14th"],
                           judge or CASE_INFO["judge"],
                           doc_title)

    # Parse and add content
    content = formatted_doc.get("content", "") if isinstance(formatted_doc, dict) else formatted_doc
    paragraphs = _md_to_paragraphs(content)

    for para in paragraphs:
        ptype = para["type"]
        text = para.get("text", "")

        if ptype == "heading1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(_clean_md_formatting(text))
            run.bold = True
            run.font.name = fmt["font"]
            run.font.size = Pt(fmt["size"])

        elif ptype == "heading2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(_clean_md_formatting(text))
            run.bold = True
            run.font.name = fmt["font"]
            run.font.size = Pt(fmt["size"])

        elif ptype == "heading3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(_clean_md_formatting(text))
            run.bold = True
            run.italic = True
            run.font.name = fmt["font"]
            run.font.size = Pt(fmt["size"])

        elif ptype == "numbered":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Inches(0.5)
            num = para.get("num", "")
            _add_bold_italic_run(p, f"{num}. {text}", fmt["font"], fmt["size"])

        elif ptype == "bullet":
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 2.0
            _add_bold_italic_run(p, text, fmt["font"], fmt["size"])

        else:  # paragraph
            if text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.first_line_indent = Inches(0.5)
                _add_bold_italic_run(p, text, fmt["font"], fmt["size"])

    # Certificate of service
    generate_certificate_of_service(doc, caption_type)

    # Fee waiver
    if include_fee_waiver:
        generate_fee_waiver(doc)

    # Page numbers
    _add_page_numbers(doc)

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"  SAVED: {output_path}")
    return output_path


def run_pipeline(stack_path, court_type, output_dir, output_filename,
                 doc_title=None, case_name=None, case_no=None,
                 court_name=None, judge=None, include_fee_waiver=False):
    """Full pipeline: read → merge → format → generate DOCX."""
    print(f"\n{'='*60}")
    print(f"PIPELINE: {doc_title or output_filename}")
    print(f"  Stack: {stack_path}")
    print(f"  Court: {court_type} ({COURT_FORMATS.get(court_type, {}).get('rule', 'N/A')})")
    print(f"{'='*60}")

    # Step 1: Read stack
    files = read_stack(stack_path)
    if not files:
        print("  ERROR: No files found in stack. Skipping.")
        return None

    # Step 2: Merge
    merged = merge_documents(files)
    print(f"  Merged: {len(merged):,} characters from {len(files)} files")

    # Step 3: Format
    formatted = format_for_court(merged, court_type)

    # Step 4: Generate DOCX
    output_path = os.path.join(output_dir, output_filename)
    stack_info = {
        "case_name": case_name or CASE_INFO["case_name"],
        "case_no": case_no or CASE_INFO["trial_case_no"],
        "court": court_name or CASE_INFO["court_14th"],
        "doc_type": doc_title or output_filename,
    }

    result = generate_docx(
        formatted, output_path,
        caption_type=court_type,
        doc_title=doc_title,
        case_name=case_name,
        case_no=case_no,
        court_name=court_name,
        judge=judge,
        include_cover=True,
        include_fee_waiver=include_fee_waiver,
        stack_info=stack_info,
    )

    print(f"  COMPLETE: {output_filename}")
    return result


# ─── Priority Stack Configurations ───────────────────────────────────────────

def run_all_priority_stacks():
    """Generate court-ready DOCX for all 5 priority stacks."""
    print("\n" + "=" * 70)
    print("FILING ASSEMBLY PIPELINE — PRIORITY STACKS")
    print(f"Timestamp: {datetime.now().isoformat()}")
    print("=" * 70)

    results = []

    # 1. COA 366810 — Appellant's Brief
    r = run_pipeline(
        stack_path=os.path.join(BASE_DIR, "01_COA_366810", "FINAL_BRIEF_STACK"),
        court_type="coa",
        output_dir=os.path.join(BASE_DIR, "01_COA_366810", "COURT_READY"),
        output_filename="Appellants_Brief.docx",
        doc_title="Appellant's Brief on Appeal",
        case_name="Pigors v. Watson",
        case_no="366810",
        court_name=CASE_INFO["court_coa"],
        judge=CASE_INFO["judge"],
    )
    results.append(("COA 366810 — Appellant's Brief", r))

    # 2. Watson Tort — Civil Complaint
    r = run_pipeline(
        stack_path=os.path.join(BASE_DIR, "02_TRIAL_14TH", "WATSON_TORT_STACK"),
        court_type="state",
        output_dir=os.path.join(BASE_DIR, "02_TRIAL_14TH", "WATSON_TORT_STACK", "COURT_READY"),
        output_filename="Civil_Complaint.docx",
        doc_title="Verified Complaint for Damages",
        case_name="Pigors v. Watson et al.",
        case_no="[TO BE ASSIGNED]",
        court_name=CASE_INFO["court_14th"],
        judge=None,
    )
    results.append(("Watson Tort — Civil Complaint", r))

    # 3. Shady Oaks — Civil Complaint
    r = run_pipeline(
        stack_path=os.path.join(BASE_DIR, "02_TRIAL_14TH", "SHADY_OAKS_EXPANDED_STACK"),
        court_type="state",
        output_dir=os.path.join(BASE_DIR, "02_TRIAL_14TH", "SHADY_OAKS_EXPANDED_STACK", "COURT_READY"),
        output_filename="Civil_Complaint.docx",
        doc_title="Complaint and Demand for Jury Trial",
        case_name="Pigors v. Shady Oaks MHC, LLC et al.",
        case_no="[TO BE ASSIGNED]",
        court_name=CASE_INFO["court_14th"],
        judge=None,
    )
    results.append(("Shady Oaks — Civil Complaint", r))

    # 4. Emergency Motions — Emergency Motion Packet
    r = run_pipeline(
        stack_path=os.path.join(BASE_DIR, "06_EMERGENCY", "FINAL_EMERGENCY_STACK"),
        court_type="state",
        output_dir=os.path.join(BASE_DIR, "06_EMERGENCY", "FINAL_EMERGENCY_STACK", "COURT_READY"),
        output_filename="Emergency_Motion_Packet.docx",
        doc_title="Emergency Motion Packet",
        case_name="Pigors v. Watson",
        case_no=CASE_INFO["trial_case_no"],
        court_name=CASE_INFO["court_14th"],
        judge=CASE_INFO["judge"],
    )
    results.append(("Emergency Motions — Motion Packet", r))

    # 5. WDMI 1983 — Section 1983 Complaint
    r = run_pipeline(
        stack_path=os.path.join(BASE_DIR, "03_FEDERAL_1983", "WDMI_FULL_STACK"),
        court_type="federal",
        output_dir=os.path.join(BASE_DIR, "03_FEDERAL_1983", "WDMI_FULL_STACK", "COURT_READY"),
        output_filename="Section_1983_Complaint.docx",
        doc_title="Complaint for Deprivation of Civil Rights",
        case_name="Pigors v. McNeill et al.",
        case_no="[TO BE ASSIGNED]",
        court_name=CASE_INFO["court_wdmi"],
        judge=None,
        include_fee_waiver=True,
    )
    results.append(("WDMI 1983 — Section 1983 Complaint", r))

    # Summary
    print("\n" + "=" * 70)
    print("PIPELINE RESULTS SUMMARY")
    print("=" * 70)
    for name, path in results:
        status = "✓ GENERATED" if path else "✗ FAILED"
        print(f"  {status}  {name}")
        if path:
            size = os.path.getsize(path)
            print(f"           → {path} ({size:,} bytes)")
    print("=" * 70)

    return results


# ─── Entry Point ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    run_all_priority_stacks()
