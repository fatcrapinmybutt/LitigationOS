"""Harvest Engine — File content extraction.

Extracts text from PDF (pypdfium2), DOCX (python-docx), TXT, CSV, JSON.
Graceful fallback if optional libraries are unavailable.
"""

import csv
import hashlib
import io
import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ── Lazy imports for optional libraries ──────────────────────────────────────
_pypdfium2 = None
_docx = None
_orjson = None

def _get_pypdfium2():
    global _pypdfium2
    if _pypdfium2 is None:
        try:
            import pypdfium2
            _pypdfium2 = pypdfium2
        except ImportError:
            logger.warning("pypdfium2 not available — PDF extraction disabled")
            _pypdfium2 = False
    return _pypdfium2 if _pypdfium2 is not False else None

def _get_docx():
    global _docx
    if _docx is None:
        try:
            import docx
            _docx = docx
        except ImportError:
            logger.warning("python-docx not available — DOCX extraction disabled")
            _docx = False
    return _docx if _docx is not False else None

def _get_orjson():
    global _orjson
    if _orjson is None:
        try:
            import orjson
            _orjson = orjson
        except ImportError:
            _orjson = False
    return _orjson if _orjson is not False else None


class ExtractionResult:
    """Result of extracting text from a file."""
    __slots__ = (
        "file_path", "file_name", "file_size", "sha256",
        "pages", "full_text", "page_count", "extraction_method",
        "error",
    )

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.file_size = 0
        self.sha256 = ""
        self.pages = []       # list of (page_num, text)
        self.full_text = ""
        self.page_count = 0
        self.extraction_method = ""
        self.error = None

    @property
    def success(self) -> bool:
        return self.error is None and len(self.full_text) > 0


def compute_sha256(file_path: str) -> str:
    """Stream-hash a file with SHA-256."""
    sha = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            while chunk := f.read(1024 * 1024):
                sha.update(chunk)
    except (OSError, IOError) as e:
        logger.error("SHA-256 failed for %s: %s", file_path, e)
        return ""
    return sha.hexdigest()


def extract_pdf(file_path: str) -> ExtractionResult:
    """Extract text from PDF using pypdfium2."""
    result = ExtractionResult(file_path)
    pdfium = _get_pypdfium2()
    if pdfium is None:
        result.error = "pypdfium2 not installed"
        return result

    try:
        result.file_size = os.path.getsize(file_path)
        result.sha256 = compute_sha256(file_path)
        result.extraction_method = "pypdfium2"

        pdf = pdfium.PdfDocument(file_path)
        result.page_count = len(pdf)
        all_text = []

        for i in range(len(pdf)):
            page = pdf[i]
            try:
                textpage = page.get_textpage()
                text = textpage.get_text_bounded()
                if text and len(text.strip()) >= 10:
                    result.pages.append((i + 1, text.strip()))
                    all_text.append(text.strip())
                textpage.close()
            except Exception as e:
                logger.debug("Page %d extraction failed in %s: %s", i + 1, file_path, e)
            finally:
                page.close()

        pdf.close()
        result.full_text = "\n\n".join(all_text)

    except Exception as e:
        result.error = f"PDF extraction failed: {e}"
        logger.error("PDF extraction failed for %s: %s", file_path, e)

    return result


def extract_docx(file_path: str) -> ExtractionResult:
    """Extract text from DOCX using python-docx."""
    result = ExtractionResult(file_path)
    docx_lib = _get_docx()
    if docx_lib is None:
        result.error = "python-docx not installed"
        return result

    try:
        result.file_size = os.path.getsize(file_path)
        result.sha256 = compute_sha256(file_path)
        result.extraction_method = "python-docx"

        doc = docx_lib.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Treat each ~50 lines as a "page" for consistency
        page_size = 50
        for i in range(0, len(paragraphs), page_size):
            chunk = paragraphs[i:i + page_size]
            text = "\n".join(chunk)
            if len(text.strip()) >= 10:
                result.pages.append((i // page_size + 1, text.strip()))

        result.full_text = "\n".join(paragraphs)
        result.page_count = max(1, len(result.pages))

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    result.full_text += "\n" + " | ".join(cells)

    except Exception as e:
        result.error = f"DOCX extraction failed: {e}"
        logger.error("DOCX extraction failed for %s: %s", file_path, e)

    return result


def extract_text_file(file_path: str) -> ExtractionResult:
    """Extract text from plain text files (TXT, MD, CSV, HTML)."""
    result = ExtractionResult(file_path)
    try:
        result.file_size = os.path.getsize(file_path)
        result.sha256 = compute_sha256(file_path)
        result.extraction_method = "text_read"

        # Try UTF-8 first, then latin-1 as fallback
        text = None
        for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                with open(file_path, "r", encoding=enc, errors="replace") as f:
                    text = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if text is None:
            result.error = "Could not decode file with any encoding"
            return result

        result.full_text = text.strip()
        # Split into pages of ~100 lines
        lines = text.split("\n")
        page_size = 100
        for i in range(0, len(lines), page_size):
            chunk = "\n".join(lines[i:i + page_size]).strip()
            if len(chunk) >= 10:
                result.pages.append((i // page_size + 1, chunk))
        result.page_count = max(1, len(result.pages))

    except Exception as e:
        result.error = f"Text extraction failed: {e}"
        logger.error("Text extraction failed for %s: %s", file_path, e)

    return result


def extract_csv_file(file_path: str) -> ExtractionResult:
    """Extract text from CSV files."""
    result = ExtractionResult(file_path)
    try:
        result.file_size = os.path.getsize(file_path)
        result.sha256 = compute_sha256(file_path)
        result.extraction_method = "csv_reader"

        rows_text = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i > 10000:
                    break
                rows_text.append(" | ".join(row))

        result.full_text = "\n".join(rows_text)
        result.page_count = 1
        if result.full_text.strip():
            result.pages.append((1, result.full_text.strip()))

    except Exception as e:
        result.error = f"CSV extraction failed: {e}"
        logger.error("CSV extraction failed for %s: %s", file_path, e)

    return result


def extract_json_file(file_path: str) -> ExtractionResult:
    """Extract text content from JSON files."""
    result = ExtractionResult(file_path)
    try:
        result.file_size = os.path.getsize(file_path)
        result.sha256 = compute_sha256(file_path)
        result.extraction_method = "json_read"

        import json
        orjson_lib = _get_orjson()

        with open(file_path, "rb") as f:
            raw = f.read()

        if orjson_lib:
            data = orjson_lib.loads(raw)
        else:
            data = json.loads(raw)

        # Flatten JSON to text
        text = _flatten_json(data)
        result.full_text = text
        result.page_count = 1
        if text.strip():
            result.pages.append((1, text.strip()[:50000]))

    except Exception as e:
        result.error = f"JSON extraction failed: {e}"
        logger.error("JSON extraction failed for %s: %s", file_path, e)

    return result


def _flatten_json(data, prefix: str = "", depth: int = 0) -> str:
    """Recursively flatten JSON to searchable text."""
    if depth > 10:
        return ""
    parts = []
    if isinstance(data, dict):
        for k, v in data.items():
            parts.append(_flatten_json(v, f"{prefix}{k}: ", depth + 1))
    elif isinstance(data, list):
        for i, item in enumerate(data[:500]):
            parts.append(_flatten_json(item, f"{prefix}[{i}] ", depth + 1))
    elif isinstance(data, str):
        if len(data) >= 5:
            parts.append(f"{prefix}{data}")
    elif data is not None:
        parts.append(f"{prefix}{data}")
    return "\n".join(p for p in parts if p)


def extract_file(file_path: str) -> ExtractionResult:
    """Auto-detect file type and extract text content."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_docx(file_path)
    elif ext in (".txt", ".md", ".html", ".htm", ".xml", ".rtf", ".eml"):
        return extract_text_file(file_path)
    elif ext == ".csv":
        return extract_csv_file(file_path)
    elif ext in (".json", ".jsonl"):
        return extract_json_file(file_path)
    else:
        result = ExtractionResult(file_path)
        result.error = f"Unsupported file type: {ext}"
        return result
