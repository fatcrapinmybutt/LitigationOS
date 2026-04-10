"""
Harvest Engine — Content-Based Deduplicator v1.0
User mandate: peek inside documents, NOT hash-only.
SHA-256 clustering → first 500 chars content comparison → flag (never delete).
"""

import sqlite3
import hashlib
import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

logger = logging.getLogger(__name__)

# How many characters to peek for content comparison
PEEK_CHARS = 500


class DedupResult:
    """Result of a deduplication scan."""
    __slots__ = (
        'files_scanned', 'hash_clusters', 'content_matches',
        'unique_files', 'duplicate_pairs', 'errors'
    )

    def __init__(self):
        self.files_scanned = 0
        self.hash_clusters = 0
        self.content_matches = 0
        self.unique_files = 0
        self.duplicate_pairs = []
        self.errors = []


def _table_exists(conn: sqlite3.Connection, table: str) -> bool:
    """Check if table exists."""
    row = conn.execute(
        "SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name=?",
        (table,)
    ).fetchone()
    return row[0] > 0


def init_dedup_tables(conn: sqlite3.Connection) -> None:
    """Create dedup tracking tables if they don't exist."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS dedup_clusters (
            cluster_id TEXT,
            file_path TEXT,
            sha256 TEXT,
            file_size INTEGER,
            content_peek TEXT,
            is_canonical INTEGER DEFAULT 0,
            match_confidence REAL DEFAULT 0.0,
            detected_at TEXT,
            PRIMARY KEY (cluster_id, file_path)
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS dedup_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            action TEXT,
            source_path TEXT,
            canonical_path TEXT,
            sha256 TEXT,
            match_type TEXT,
            confidence REAL,
            logged_at TEXT
        )
    """)
    conn.commit()


def compute_sha256(file_path: str) -> Optional[str]:
    """Compute SHA-256 hash of a file (streaming for large files)."""
    try:
        sha = hashlib.sha256()
        with open(file_path, 'rb') as f:
            while True:
                chunk = f.read(65536)
                if not chunk:
                    break
                sha.update(chunk)
        return sha.hexdigest()
    except (OSError, PermissionError) as e:
        logger.warning("Cannot hash %s: %s", file_path, e)
        return None


def peek_content(file_path: str, chars: int = PEEK_CHARS) -> Optional[str]:
    """
    Peek inside a file to get first N characters of content.
    User mandate: content-based dedup, not hash-only.
    """
    ext = os.path.splitext(file_path)[1].lower()

    # Text-based files — read directly
    if ext in ('.txt', '.md', '.csv', '.json', '.jsonl', '.html', '.xml',
               '.log', '.cfg', '.ini', '.yml', '.yaml'):
        for enc in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
            try:
                with open(file_path, 'r', encoding=enc, errors='replace') as f:
                    return f.read(chars)
            except (OSError, UnicodeDecodeError):
                continue
        return None

    # PDF files — extract first page text
    if ext == '.pdf':
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(file_path)
            if len(pdf) > 0:
                page = pdf[0]
                tp = page.get_textpage()
                text = tp.get_text_bounded()
                tp.close()
                page.close()
                pdf.close()
                return text[:chars] if text else None
            pdf.close()
        except Exception as e:
            logger.debug("PDF peek failed for %s: %s", file_path, e)
        return None

    # DOCX files — extract first paragraph
    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            text = '\n'.join(p.text for p in doc.paragraphs[:10])
            return text[:chars] if text else None
        except Exception as e:
            logger.debug("DOCX peek failed for %s: %s", file_path, e)
        return None

    # Binary/unknown — read raw bytes as hex preview
    try:
        with open(file_path, 'rb') as f:
            raw = f.read(min(chars, 256))
            return raw.hex()[:chars]
    except (OSError, PermissionError):
        return None


def content_similarity(text_a: str, text_b: str) -> float:
    """
    Compute content similarity between two text peeks.
    Returns 0.0 (completely different) to 1.0 (identical).
    Uses character-level comparison for speed.
    """
    if not text_a or not text_b:
        return 0.0

    if text_a == text_b:
        return 1.0

    # Normalize whitespace
    a = ' '.join(text_a.split())
    b = ' '.join(text_b.split())

    if a == b:
        return 0.99

    # Character overlap ratio (bidirectional)
    set_a = set(a)
    set_b = set(b)
    if not set_a or not set_b:
        return 0.0

    intersection = len(set_a & set_b)
    union = len(set_a | set_b)
    char_sim = intersection / union if union > 0 else 0.0

    # Prefix match — same beginning strongly suggests same document
    prefix_len = 0
    min_len = min(len(a), len(b))
    for i in range(min(min_len, 200)):
        if a[i] == b[i]:
            prefix_len += 1
        else:
            break
    prefix_ratio = prefix_len / min(min_len, 200) if min_len > 0 else 0.0

    # Combined score: 60% prefix match + 40% character overlap
    return 0.6 * prefix_ratio + 0.4 * char_sim


def find_duplicates_in_directory(
    directory: str,
    conn: sqlite3.Connection,
    extensions: tuple = ('.pdf', '.docx', '.txt', '.md', '.csv', '.json'),
    min_size: int = 100,
    similarity_threshold: float = 0.85
) -> DedupResult:
    """
    Scan a directory for duplicate files using two-stage dedup:
    1. SHA-256 hash clustering (exact duplicates)
    2. Content peek comparison (near-duplicates)

    Never deletes — only flags and records in dedup_clusters table.
    """
    result = DedupResult()
    init_dedup_tables(conn)

    # Stage 1: Collect file info and hash
    file_info = []
    for root, dirs, files in os.walk(directory):
        for fname in files:
            ext = os.path.splitext(fname)[1].lower()
            if ext not in extensions:
                continue
            fpath = os.path.join(root, fname)
            try:
                fsize = os.path.getsize(fpath)
                if fsize < min_size:
                    continue
                file_info.append({
                    'path': fpath,
                    'name': fname,
                    'size': fsize,
                    'ext': ext
                })
                result.files_scanned += 1
            except OSError:
                continue

    if not file_info:
        return result

    # Stage 2: Hash clustering
    hash_map = {}
    for fi in file_info:
        sha = compute_sha256(fi['path'])
        if sha:
            fi['sha256'] = sha
            hash_map.setdefault(sha, []).append(fi)

    # Find exact-hash clusters (2+ files with same hash)
    for sha, cluster in hash_map.items():
        if len(cluster) < 2:
            result.unique_files += 1
            continue

        result.hash_clusters += 1
        # Mark first file as canonical, rest as duplicates
        canonical = cluster[0]
        canonical_peek = peek_content(canonical['path'])

        cluster_id = f"hash_{sha[:16]}"

        conn.execute(
            "INSERT OR REPLACE INTO dedup_clusters "
            "(cluster_id, file_path, sha256, file_size, content_peek, "
            "is_canonical, match_confidence, detected_at) "
            "VALUES (?,?,?,?,?,1,1.0,?)",
            (cluster_id, canonical['path'], sha, canonical['size'],
             canonical_peek or '', datetime.now().isoformat())
        )

        for dup in cluster[1:]:
            dup_peek = peek_content(dup['path'])
            result.duplicate_pairs.append((canonical['path'], dup['path']))
            result.content_matches += 1

            conn.execute(
                "INSERT OR REPLACE INTO dedup_clusters "
                "(cluster_id, file_path, sha256, file_size, content_peek, "
                "is_canonical, match_confidence, detected_at) "
                "VALUES (?,?,?,?,?,0,1.0,?)",
                (cluster_id, dup['path'], sha, dup['size'],
                 dup_peek or '', datetime.now().isoformat())
            )

            conn.execute(
                "INSERT INTO dedup_log "
                "(action, source_path, canonical_path, sha256, "
                "match_type, confidence, logged_at) "
                "VALUES (?,?,?,?,?,?,?)",
                ('duplicate_found', dup['path'], canonical['path'],
                 sha, 'exact_hash', 1.0, datetime.now().isoformat())
            )

    # Stage 3: Content-based near-duplicate detection
    # Group by similar size (within 10%) for efficiency
    unique_files = [fi for sha, cluster in hash_map.items()
                    if len(cluster) == 1 for fi in cluster]

    # Sort by size for windowed comparison
    unique_files.sort(key=lambda x: x.get('size', 0))

    for i, fi_a in enumerate(unique_files):
        if i >= len(unique_files) - 1:
            break

        peek_a = peek_content(fi_a['path'])
        if not peek_a:
            continue

        # Compare with nearby-sized files (within 20% size difference)
        for j in range(i + 1, min(i + 20, len(unique_files))):
            fi_b = unique_files[j]

            # Size ratio check — skip if too different
            size_ratio = fi_a.get('size', 1) / max(fi_b.get('size', 1), 1)
            if size_ratio < 0.8 or size_ratio > 1.25:
                continue

            peek_b = peek_content(fi_b['path'])
            if not peek_b:
                continue

            sim = content_similarity(peek_a, peek_b)
            if sim >= similarity_threshold:
                cluster_id = f"content_{hashlib.md5((fi_a['path'] + fi_b['path']).encode()).hexdigest()[:16]}"
                result.content_matches += 1
                result.duplicate_pairs.append((fi_a['path'], fi_b['path']))

                # Record both files in cluster
                conn.execute(
                    "INSERT OR REPLACE INTO dedup_clusters "
                    "(cluster_id, file_path, sha256, file_size, content_peek, "
                    "is_canonical, match_confidence, detected_at) "
                    "VALUES (?,?,?,?,?,1,?,?)",
                    (cluster_id, fi_a['path'], fi_a.get('sha256', ''),
                     fi_a.get('size', 0), peek_a[:200],
                     sim, datetime.now().isoformat())
                )
                conn.execute(
                    "INSERT OR REPLACE INTO dedup_clusters "
                    "(cluster_id, file_path, sha256, file_size, content_peek, "
                    "is_canonical, match_confidence, detected_at) "
                    "VALUES (?,?,?,?,?,0,?,?)",
                    (cluster_id, fi_b['path'], fi_b.get('sha256', ''),
                     fi_b.get('size', 0), peek_b[:200],
                     sim, datetime.now().isoformat())
                )

                conn.execute(
                    "INSERT INTO dedup_log "
                    "(action, source_path, canonical_path, sha256, "
                    "match_type, confidence, logged_at) "
                    "VALUES (?,?,?,?,?,?,?)",
                    ('near_duplicate', fi_b['path'], fi_a['path'],
                     fi_b.get('sha256', ''), 'content_peek',
                     sim, datetime.now().isoformat())
                )

    conn.commit()
    result.unique_files = result.files_scanned - result.content_matches
    return result


def check_evidence_duplicates(
    conn: sqlite3.Connection,
    quote_text: str,
    source_file: str
) -> Dict[str, Any]:
    """
    Check if a specific evidence quote is a duplicate of existing data.
    Returns match info or None if unique.
    """
    results = {'is_duplicate': False, 'matches': []}

    # Check evidence_quotes
    try:
        rows = conn.execute(
            "SELECT quote_text, source_file, lane "
            "FROM evidence_quotes "
            "WHERE quote_text = ? LIMIT 5",
            (quote_text,)
        ).fetchall()
        if rows:
            results['is_duplicate'] = True
            results['matches'].extend([
                {'table': 'evidence_quotes', 'source': r[1], 'lane': r[2]}
                for r in rows
            ])
    except sqlite3.OperationalError:
        pass

    # Check harvest_evidence
    try:
        rows = conn.execute(
            "SELECT quote_text, source_file, lane "
            "FROM harvest_evidence "
            "WHERE quote_text = ? LIMIT 5",
            (quote_text,)
        ).fetchall()
        if rows:
            results['is_duplicate'] = True
            results['matches'].extend([
                {'table': 'harvest_evidence', 'source': r[1], 'lane': r[2]}
                for r in rows
            ])
    except sqlite3.OperationalError:
        pass

    return results


def get_dedup_stats(conn: sqlite3.Connection) -> Dict[str, Any]:
    """Get dedup statistics from the dedup_clusters table."""
    stats = {
        'total_clusters': 0,
        'total_files_in_clusters': 0,
        'exact_hash_dupes': 0,
        'content_near_dupes': 0,
        'canonical_files': 0,
        'duplicate_files': 0,
    }

    if not _table_exists(conn, 'dedup_clusters'):
        return stats

    try:
        row = conn.execute(
            "SELECT "
            "  COUNT(DISTINCT cluster_id) AS clusters, "
            "  COUNT(*) AS total_files, "
            "  SUM(CASE WHEN is_canonical = 1 THEN 1 ELSE 0 END) AS canonical, "
            "  SUM(CASE WHEN is_canonical = 0 THEN 1 ELSE 0 END) AS duplicates "
            "FROM dedup_clusters"
        ).fetchone()
        if row:
            stats['total_clusters'] = row[0]
            stats['total_files_in_clusters'] = row[1]
            stats['canonical_files'] = row[2]
            stats['duplicate_files'] = row[3]
    except sqlite3.OperationalError:
        pass

    if _table_exists(conn, 'dedup_log'):
        try:
            row = conn.execute(
                "SELECT "
                "  SUM(CASE WHEN match_type = 'exact_hash' THEN 1 ELSE 0 END), "
                "  SUM(CASE WHEN match_type = 'content_peek' THEN 1 ELSE 0 END) "
                "FROM dedup_log"
            ).fetchone()
            if row:
                stats['exact_hash_dupes'] = row[0] or 0
                stats['content_near_dupes'] = row[1] or 0
        except sqlite3.OperationalError:
            pass

    return stats
