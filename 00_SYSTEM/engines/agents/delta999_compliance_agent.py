#!/usr/bin/env python3
"""
Delta999 Compliance Agent — Filing compliance checker.

Checks document format, deadlines, service requirements, fee waiver status,
and generates comprehensive compliance reports.

CLI:
    python delta999_compliance_agent.py --action check_format --doc-path "path" --court-rule "MCR 2.107"
    python delta999_compliance_agent.py --action check_deadlines --filing-stack "MEEK1"
    python delta999_compliance_agent.py --action check_service --filing-stack "MEEK1"
    python delta999_compliance_agent.py --action check_fee_waiver --filing-stack "MEEK1"
    python delta999_compliance_agent.py --action full_report --filing-stack "MEEK1"
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

import argparse
import json
import os
import re
import sqlite3
from datetime import datetime, timedelta
from pathlib import Path

# ── paths ────────────────────────────────────────────────────────────────────
AGENT_DIR = Path(__file__).parent
ENGINE_DIR = AGENT_DIR.parent
sys.path.insert(0, str(ENGINE_DIR))

DB = r'C:\Users\andre\LitigationOS\litigation_context.db'
AGENT_NAME = 'delta999_compliance_agent'

from llm_bridge import llm_ask

# Michigan court format requirements
FORMAT_RULES = {
    'MCR 7.212': {
        'font': 'Times New Roman 12pt or equivalent',
        'margins': '1 inch all sides',
        'spacing': 'Double-spaced (except block quotes)',
        'page_limit': None,
        'word_limit': 16000,
        'required_sections': [
            'Cover Page', 'Table of Contents', 'Table of Authorities',
            'Statement of Jurisdiction', 'Statement of Questions Presented',
            'Statement of Facts', 'Argument', 'Relief Requested',
        ],
    },
    'MCR 2.113': {
        'font': 'Standard readable font',
        'margins': '1 inch all sides',
        'spacing': 'Double-spaced',
        'caption': 'Required: case number, parties, court name',
    },
    'MCR 2.107': {
        'service_methods': ['Personal service', 'First-class mail', 'Electronic service (if agreed)'],
        'proof_of_service': 'Required for all filings',
    },
}


# ── DB helpers ───────────────────────────────────────────────────────────────

def get_conn():
    conn = sqlite3.connect(DB, timeout=120)
    conn.execute('PRAGMA busy_timeout=60000')
    conn.row_factory = sqlite3.Row
    return conn


def log_activity(action, result):
    conn = get_conn()
    conn.execute(
        'INSERT INTO agent_activity_log (agent_name, action, result) VALUES (?,?,?)',
        (AGENT_NAME, action, str(result)[:2000])
    )
    conn.commit()
    conn.close()


# ── Core functions ───────────────────────────────────────────────────────────

def check_format(document_path: str, court_rule: str = 'MCR 7.212') -> dict:
    """Check document format compliance against a specific court rule."""
    p = Path(document_path)
    issues = []
    checks = {}

    if not p.exists():
        return {'error': f'File not found: {document_path}'}

    # Read text
    if p.suffix.lower() in ('.txt', '.md'):
        text = p.read_text(encoding='utf-8')
    elif p.suffix.lower() == '.docx':
        try:
            from docx import Document
            doc = Document(str(p))
            text = '\n'.join(para.text for para in doc.paragraphs)
            # Check font/spacing from docx metadata
            for para in doc.paragraphs[:5]:
                if para.runs:
                    run = para.runs[0]
                    if run.font.size:
                        pt = run.font.size.pt
                        checks['font_size'] = f'{pt}pt'
                        if pt < 11:
                            issues.append(f'Font size {pt}pt may be too small (12pt recommended)')
        except ImportError:
            text = p.read_text(encoding='utf-8', errors='replace')
    else:
        text = p.read_text(encoding='utf-8', errors='replace')

    # Word count
    word_count = len(text.split())
    checks['word_count'] = word_count
    rule_info = FORMAT_RULES.get(court_rule, {})
    if rule_info.get('word_limit') and word_count > rule_info['word_limit']:
        issues.append(f'Word count {word_count} exceeds limit of {rule_info["word_limit"]}')

    # Required sections check
    text_lower = text.lower()
    if 'required_sections' in rule_info:
        for section in rule_info['required_sections']:
            found = section.lower() in text_lower
            checks[f'section_{section}'] = found
            if not found:
                issues.append(f'Missing required section: {section}')

    # Proof of service
    has_cos = 'certificate of service' in text_lower or 'proof of service' in text_lower
    checks['certificate_of_service'] = has_cos
    if not has_cos:
        issues.append('Missing Certificate/Proof of Service')

    result = {
        'document_path': document_path,
        'court_rule': court_rule,
        'rule_requirements': rule_info,
        'checks': checks,
        'issues': issues,
        'compliant': len(issues) == 0,
    }
    log_activity(f'check_format:{court_rule}', json.dumps(result, default=str)[:2000])
    return result


def check_deadlines(filing_stack: str) -> dict:
    """Check deadline compliance for a filing stack."""
    conn = get_conn()

    # Get filings and their deadlines
    filings = []
    try:
        rows = conn.execute(
            "SELECT * FROM apex_filing_stack_index WHERE stack_label LIKE ? OR meek_track LIKE ?",
            (f'%{filing_stack}%', f'%{filing_stack}%')
        ).fetchall()
        filings = [dict(r) for r in rows]
    except Exception:
        pass

    # Check timeline for deadline events
    deadlines = []
    try:
        rows = conn.execute(
            "SELECT * FROM master_timeline WHERE event_type LIKE '%deadline%' "
            "OR event_type LIKE '%due%' OR description LIKE '%deadline%' "
            "ORDER BY event_date DESC LIMIT 20"
        ).fetchall()
        deadlines = [dict(r) for r in rows]
    except Exception:
        pass

    # MCR deadline rules
    try:
        rows = conn.execute(
            "SELECT * FROM mcr_authority_library WHERE rule_text LIKE '%days%' "
            "AND (rule_text LIKE '%file%' OR rule_text LIKE '%serve%') LIMIT 15"
        ).fetchall()
        deadline_rules = [dict(r) for r in rows]
    except Exception:
        deadline_rules = []

    conn.close()

    # LLM analysis
    try:
        analysis = llm_ask(
            f"Analyze deadline compliance for filing stack: '{filing_stack}'\n"
            f"Filings found: {len(filings)}\n"
            f"Known deadlines: {len(deadlines)}\n"
            f"Deadline-related rules: {len(deadline_rules)}\n\n"
            f"Filing info: {json.dumps(filings[:5], default=str)[:600]}\n"
            f"Deadlines: {json.dumps(deadlines[:5], default=str)[:600]}\n\n"
            f"Identify: (1) upcoming deadlines, (2) missed deadlines, "
            f"(3) tolling arguments if deadlines passed.",
            system_prompt="You are a Michigan court deadline compliance specialist."
        )
    except Exception as e:
        analysis = f"LLM unavailable: {e}"

    result = {
        'filing_stack': filing_stack,
        'filings_found': len(filings),
        'deadlines_found': len(deadlines),
        'deadline_rules': len(deadline_rules),
        'analysis': analysis,
    }
    log_activity(f'check_deadlines:{filing_stack}', json.dumps(result, default=str)[:2000])
    return result


def check_service(filing_stack: str) -> dict:
    """Check service requirements for a filing stack."""
    conn = get_conn()

    # MCR 2.107 service requirements
    service_rules = []
    try:
        rows = conn.execute(
            "SELECT * FROM mcr_authority_library WHERE rule_number LIKE '2.107%' "
            "OR rule_number LIKE '2.105%' OR rule_number LIKE '2.106%'"
        ).fetchall()
        service_rules = [dict(r) for r in rows]
    except Exception:
        pass

    # Check filing inventory for service status
    service_records = []
    try:
        rows = conn.execute(
            "SELECT * FROM filing_inventory WHERE stack LIKE ? OR description LIKE ?",
            (f'%{filing_stack}%', f'%{filing_stack}%')
        ).fetchall()
        service_records = [dict(r) for r in rows]
    except Exception:
        pass

    conn.close()

    result = {
        'filing_stack': filing_stack,
        'service_rules_found': len(service_rules),
        'service_records': len(service_records),
        'mcr_2107_requirements': {
            'personal_service': 'MCR 2.105 — personal delivery',
            'mail_service': 'MCR 2.107(C)(1) — first-class mail',
            'electronic_service': 'MCR 2.107(C)(4) — if parties agreed',
            'proof_required': 'MCR 2.104 — proof of service with all filings',
        },
        'service_rule_details': service_rules[:5],
    }
    log_activity(f'check_service:{filing_stack}', json.dumps(result, default=str)[:2000])
    return result


def check_fee_waiver(filing_stack: str) -> dict:
    """Check MC 20 / IFP (In Forma Pauperis) fee waiver status."""
    conn = get_conn()

    # Search for fee waiver records
    fee_records = []
    try:
        rows = conn.execute(
            "SELECT * FROM filing_inventory WHERE description LIKE '%fee%' "
            "OR description LIKE '%waiver%' OR description LIKE '%IFP%' "
            "OR description LIKE '%forma pauperis%' OR description LIKE '%MC 20%'"
        ).fetchall()
        fee_records = [dict(r) for r in rows]
    except Exception:
        pass

    # MCL fee waiver statute
    fee_statutes = []
    try:
        rows = conn.execute(
            "SELECT * FROM mcl_authority_library WHERE statute_text LIKE '%fee%waiv%' "
            "OR statute_number LIKE '600.2529%' LIMIT 5"
        ).fetchall()
        fee_statutes = [dict(r) for r in rows]
    except Exception:
        pass

    conn.close()

    result = {
        'filing_stack': filing_stack,
        'fee_waiver_records': len(fee_records),
        'fee_waiver_details': fee_records[:5],
        'applicable_statutes': fee_statutes,
        'mc20_requirements': {
            'form': 'MC 20 — Affidavit and Order, Suspension of Fees/Costs',
            'criteria': 'Inability to pay fees without depriving necessities',
            'mcl': 'MCL 600.2529 — Filing fees',
            'note': 'Must be filed with or before the filing requiring the fee',
        },
    }
    log_activity(f'check_fee_waiver:{filing_stack}', json.dumps(result, default=str)[:2000])
    return result


def full_compliance_report(filing_stack: str) -> dict:
    """Generate comprehensive compliance report for a filing stack."""
    report = {
        'filing_stack': filing_stack,
        'generated': datetime.now().isoformat(),
        'sections': {},
    }

    print(f"  ▸ Checking deadlines...")
    report['sections']['deadlines'] = check_deadlines(filing_stack)

    print(f"  ▸ Checking service requirements...")
    report['sections']['service'] = check_service(filing_stack)

    print(f"  ▸ Checking fee waiver status...")
    report['sections']['fee_waiver'] = check_fee_waiver(filing_stack)

    # Overall compliance summary
    issues = []
    for section_name, section_data in report['sections'].items():
        if isinstance(section_data, dict):
            section_issues = section_data.get('issues', [])
            if section_issues:
                issues.extend([f"[{section_name}] {i}" for i in section_issues])

    report['total_issues'] = len(issues)
    report['all_issues'] = issues
    report['overall_status'] = 'COMPLIANT' if not issues else 'NEEDS_ATTENTION'

    log_activity(f'full_report:{filing_stack}', json.dumps(report, default=str)[:2000])
    return report


# ── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Delta999 Compliance Agent')
    parser.add_argument('--action', required=True,
                        choices=['check_format', 'check_deadlines', 'check_service',
                                 'check_fee_waiver', 'full_report'],
                        help='Action to perform')
    parser.add_argument('--doc-path', type=str, help='Document path')
    parser.add_argument('--court-rule', type=str, default='MCR 7.212', help='Court rule to check against')
    parser.add_argument('--filing-stack', type=str, help='Filing stack identifier')
    args = parser.parse_args()

    if args.action == 'check_format':
        if not args.doc_path:
            parser.error('--doc-path required')
        result = check_format(args.doc_path, args.court_rule)
    elif args.action == 'check_deadlines':
        if not args.filing_stack:
            parser.error('--filing-stack required')
        result = check_deadlines(args.filing_stack)
    elif args.action == 'check_service':
        if not args.filing_stack:
            parser.error('--filing-stack required')
        result = check_service(args.filing_stack)
    elif args.action == 'check_fee_waiver':
        if not args.filing_stack:
            parser.error('--filing-stack required')
        result = check_fee_waiver(args.filing_stack)
    elif args.action == 'full_report':
        if not args.filing_stack:
            parser.error('--filing-stack required')
        result = full_compliance_report(args.filing_stack)
    else:
        parser.print_help()
        return

    print(json.dumps(result, indent=2, default=str))


if __name__ == '__main__':
    main()
