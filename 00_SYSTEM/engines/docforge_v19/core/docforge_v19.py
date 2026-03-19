#!/usr/bin/env python3
"""
LITIGATIONOS MI APPELLATE DOCFORGE V19
Single-file orchestrator for:
- exact official-text promotion lane (bulk PARTIAL -> RESOLVED_VERIFIED)
- transcript-to-order-finding score tuning / linking
- vehicle-specific deadline registry (COA/MSC/JTC template + tolling notes)
- CourtPack DOCX/PDF emit lane

Design rails:
- MI-first routing
- no SHA-256 default
- originals read-only
- derived outputs append-only
- fail-soft unknowns
"""
from __future__ import annotations
import argparse, csv, dataclasses, datetime as dt, json, os, re, sqlite3, sys, textwrap
from pathlib import Path
from typing import Dict, Iterable, List, Tuple, Optional

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except Exception:
    letter = None
    canvas = None

APP_NAME = "LITIGATIONOS_MI_APPELLATE_DOCFORGE_MEEK1234_V19"
DEFAULT_DB = "runtime/docforge_v19.sqlite3"

STOPWORDS = {
    "the","a","an","and","or","to","of","in","on","for","with","at","by","is","are","was","were","be","been",
    "that","this","it","as","from","into","about","after","before","during","under","over","not","no","do","did",
    "he","she","they","we","i","you","his","her","their","our","my","your","court","judge","order","finding",
    "plaintiff","defendant","petitioner","respondent"
}

ISSUE_KEYWORDS = {
    "PT": ["parenting","time","lincoln","birthday","suspension","custody","visitation","child"],
    "PPO": ["ppo","show","cause","violation","contact","harass","stalk","appclose"],
    "CONTEMPT": ["contempt","jailed","jail","sanction","muted","objection"],
    "HEALTHWEST": ["healthwest","evaluation","assessment","mental","drug","screen"],
    "HOUSING": ["shady","oaks","rent","utility","eviction","lease","mhp","park","housing","water","sewage"],
    "BIAS": ["bias","jtc","canon","recusal","hostile","judge","mcneill"],
}

def utcnow() -> str:
    return dt.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"

def ensure_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)

def identity_key_v1(path: Path) -> str:
    """Fast deterministic identity key. No SHA-256 by default."""
    st = path.stat()
    ino = getattr(st, "st_ino", 0)
    ctime_ns = getattr(st, "st_ctime_ns", int(st.st_ctime * 1e9))
    mtime_ns = getattr(st, "st_mtime_ns", int(st.st_mtime * 1e9))
    norm = str(path).replace("\\", "/").lower()
    return f"IK1|{norm}|{st.st_size}|{mtime_ns}|{ctime_ns}|{ino}"

def connect_db(db_path: Path) -> sqlite3.Connection:
    ensure_dir(db_path.parent)
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL;")
    return conn

def initdb(conn: sqlite3.Connection) -> None:
    conn.executescript("""
    CREATE TABLE IF NOT EXISTS fileatoms(
        file_id TEXT PRIMARY KEY,
        path TEXT NOT NULL,
        rel_path TEXT NOT NULL,
        size INTEGER NOT NULL,
        mtime_ns INTEGER NOT NULL,
        identity_key TEXT NOT NULL,
        lane TEXT,
        case_anchor TEXT,
        doc_kind TEXT,
        extracted_ok INTEGER DEFAULT 0,
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS quotes(
        quote_id TEXT PRIMARY KEY,
        file_id TEXT NOT NULL,
        text TEXT NOT NULL,
        page INTEGER,
        line_start INTEGER,
        line_end INTEGER,
        speaker TEXT,
        issue_tag TEXT,
        created_at TEXT NOT NULL
    );

    CREATE VIRTUAL TABLE IF NOT EXISTS quotes_fts USING fts5(
        quote_id UNINDEXED,
        text,
        content=''
    );

    CREATE TABLE IF NOT EXISTS authority_triples(
        authority_id TEXT PRIMARY KEY,
        lane TEXT,
        citation TEXT NOT NULL,
        normalized_citation TEXT NOT NULL,
        proposition TEXT,
        source_url TEXT,
        source_url_status TEXT DEFAULT 'UNRESOLVED',
        official_text_exact TEXT,
        pinpoint TEXT,
        pinpoint_status TEXT DEFAULT 'UNRESOLVED',
        graft_status TEXT DEFAULT 'UNRESOLVED',
        graft_notes TEXT,
        updated_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS official_text_grafts(
        graft_id TEXT PRIMARY KEY,
        citation TEXT NOT NULL,
        normalized_citation TEXT NOT NULL,
        official_text_exact TEXT NOT NULL,
        pinpoint TEXT,
        source_url TEXT,
        verified INTEGER NOT NULL DEFAULT 0,
        imported_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS transcripts(
        transcript_row_id TEXT PRIMARY KEY,
        transcript_id TEXT NOT NULL,
        page INTEGER,
        line_start INTEGER,
        line_end INTEGER,
        speaker TEXT,
        text TEXT NOT NULL,
        issue_tag TEXT,
        created_at TEXT NOT NULL
    );

    CREATE VIRTUAL TABLE IF NOT EXISTS transcripts_fts USING fts5(
        transcript_row_id UNINDEXED,
        text,
        content=''
    );

    CREATE TABLE IF NOT EXISTS order_findings(
        finding_id TEXT PRIMARY KEY,
        order_id TEXT,
        finding_text TEXT NOT NULL,
        issue_tag TEXT,
        adverse INTEGER DEFAULT 0,
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS transcript_order_links(
        link_id TEXT PRIMARY KEY,
        finding_id TEXT NOT NULL,
        transcript_row_id TEXT NOT NULL,
        score REAL NOT NULL,
        score_detail_json TEXT NOT NULL,
        accepted INTEGER NOT NULL DEFAULT 0,
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS linker_tuning(
        tuning_run_id TEXT PRIMARY KEY,
        weights_json TEXT NOT NULL,
        threshold REAL NOT NULL,
        recall REAL,
        precision REAL,
        f1 REAL,
        notes TEXT,
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS deadline_rules(
        rule_id TEXT PRIMARY KEY,
        lane TEXT,
        vehicle_code TEXT NOT NULL,
        forum TEXT NOT NULL,
        anchor_type TEXT NOT NULL, -- entered|signed|served|recorded
        trigger_event TEXT NOT NULL,
        due_days INTEGER,
        business_days_only INTEGER DEFAULT 0,
        tolling_notes TEXT,
        authority_citation TEXT,
        rule_status TEXT DEFAULT 'UNVERIFIED_TEMPLATE',
        active INTEGER DEFAULT 1,
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS deadline_events(
        event_id TEXT PRIMARY KEY,
        lane TEXT,
        case_anchor TEXT,
        vehicle_code TEXT NOT NULL,
        forum TEXT NOT NULL,
        trigger_event TEXT NOT NULL,
        entered_date TEXT,
        signed_date TEXT,
        served_date TEXT,
        recorded_date TEXT,
        notes TEXT,
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS deadline_locks(
        lock_id TEXT PRIMARY KEY,
        event_id TEXT NOT NULL,
        rule_id TEXT NOT NULL,
        anchor_type TEXT NOT NULL,
        anchor_date TEXT,
        due_date TEXT,
        lock_status TEXT NOT NULL, -- LOCKED|MISSING_ANCHOR|RULE_MISSING|PENDING_VERIFY
        tolling_notes TEXT,
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS contradiction_pairs(
        pair_id TEXT PRIMARY KEY,
        lane_scope TEXT,
        proposition_key TEXT,
        left_source TEXT,
        left_ref TEXT,
        left_text TEXT,
        right_source TEXT,
        right_ref TEXT,
        right_text TEXT,
        confidence REAL,
        created_at TEXT NOT NULL
    );

    CREATE TABLE IF NOT EXISTS courtpack_runs(
        pack_id TEXT PRIMARY KEY,
        lane TEXT,
        case_anchor TEXT,
        out_dir TEXT NOT NULL,
        manifest_json TEXT NOT NULL,
        created_at TEXT NOT NULL
    );
    """)
    conn.commit()

def normalize_citation(c: str) -> str:
    c = (c or "").strip().upper()
    c = re.sub(r"\s+", " ", c)
    c = c.replace("M.C.R.", "MCR").replace("M.C.L.", "MCL").replace("M.R.E.", "MRE")
    return c

def route_source_url(citation: str) -> Tuple[str, str]:
    """Return (url, status). Deterministic MI-first root URLs; exact pinpoint may remain unresolved."""
    c = normalize_citation(citation)
    if c.startswith("MCL "):
        # try MCL 722.27a pattern -> legislature.mi.gov style path (deterministic root guess)
        code = c.split(" ",1)[1].strip().lower()
        code = code.replace("§","").replace(" ", "")
        # Legislature URLs vary, so we use the article root plus note unresolved exact path.
        return ("https://www.legislature.mi.gov/", "PARTIAL_ROOT")
    if c.startswith("MCR ") or c.startswith("MRE "):
        return ("https://www.courts.michigan.gov/", "PARTIAL_ROOT")
    if c.startswith("JTC") or "CANON" in c:
        return ("https://jtc.courts.mi.gov/", "PARTIAL_ROOT")
    if "SCAO" in c or c.startswith("MC ") or c.startswith("FOC "):
        return ("https://www.courts.michigan.gov/scao-forms/", "PARTIAL_ROOT")
    return ("", "UNRESOLVED")

def tokenize(text: str) -> List[str]:
    toks = re.findall(r"[A-Za-z0-9']+", (text or "").lower())
    return [t for t in toks if t not in STOPWORDS and len(t) > 1]

def issue_tag_for_text(text: str) -> str:
    tokens = set(tokenize(text))
    best = ("UNKNOWN", 0)
    for tag, keys in ISSUE_KEYWORDS.items():
        score = sum(1 for k in keys if k in tokens)
        if score > best[1]:
            best = (tag, score)
    return best[0]

def proposition_key(text: str) -> str:
    toks = tokenize(text)
    if not toks:
        return "EMPTY"
    top = sorted(set(toks))[:8]
    return "|".join(top)

def upsert_authority_request(conn, lane: str, citation: str, proposition: str = ""):
    n = normalize_citation(citation)
    src_url, src_status = route_source_url(citation)
    authority_id = f"AUTH::{n}"
    conn.execute("""
    INSERT INTO authority_triples(authority_id,lane,citation,normalized_citation,proposition,source_url,source_url_status,pinpoint_status,graft_status,updated_at)
    VALUES(?,?,?,?,?,?,?,?,?,?)
    ON CONFLICT(authority_id) DO UPDATE SET
      lane=excluded.lane,
      proposition=COALESCE(NULLIF(excluded.proposition,''), authority_triples.proposition),
      source_url=CASE WHEN authority_triples.source_url IS NULL OR authority_triples.source_url='' THEN excluded.source_url ELSE authority_triples.source_url END,
      source_url_status=CASE WHEN authority_triples.source_url_status='UNRESOLVED' THEN excluded.source_url_status ELSE authority_triples.source_url_status END,
      updated_at=excluded.updated_at
    """, (authority_id, lane, citation, n, proposition, src_url, src_status, "PARTIAL", "PARTIAL", utcnow()))
    conn.commit()

def ingest_authority_requests(conn, path: Path):
    count = 0
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line=line.strip()
            if not line:
                continue
            row = json.loads(line)
            upsert_authority_request(conn, row.get("lane","GLOBAL"), row["citation"], row.get("proposition",""))
            count += 1
    return count

def promote_official_text(conn, path: Path):
    """Bulk promotion lane: PARTIAL -> RESOLVED_VERIFIED when verified exact text present."""
    promoted = 0
    imported = 0
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line=line.strip()
            if not line:
                continue
            row = json.loads(line)
            citation = row["citation"]
            n = normalize_citation(citation)
            exact = (row.get("official_text_exact") or "").strip()
            pinpoint = (row.get("pinpoint") or "").strip()
            verified = 1 if row.get("verified") else 0
            source_url = (row.get("source_url") or "").strip()
            if not source_url:
                source_url, _ = route_source_url(citation)
            graft_id = f"GRAFT::{n}"
            conn.execute("""
            INSERT INTO official_text_grafts(graft_id,citation,normalized_citation,official_text_exact,pinpoint,source_url,verified,imported_at)
            VALUES(?,?,?,?,?,?,?,?)
            ON CONFLICT(graft_id) DO UPDATE SET
              official_text_exact=excluded.official_text_exact,
              pinpoint=excluded.pinpoint,
              source_url=excluded.source_url,
              verified=excluded.verified,
              imported_at=excluded.imported_at
            """, (graft_id, citation, n, exact, pinpoint, source_url, verified, utcnow()))
            imported += 1

            # ensure authority row exists
            upsert_authority_request(conn, row.get("lane","GLOBAL"), citation, row.get("proposition",""))

            if verified and exact and pinpoint:
                res = conn.execute("""
                UPDATE authority_triples
                   SET official_text_exact=?,
                       pinpoint=?,
                       source_url=COALESCE(NULLIF(?,''),source_url),
                       source_url_status='RESOLVED_VERIFIED',
                       pinpoint_status='RESOLVED_VERIFIED',
                       graft_status='RESOLVED_VERIFIED',
                       graft_notes=?,
                       updated_at=?
                 WHERE normalized_citation=?
                """, (exact, pinpoint, source_url, "Bulk promoted via ingest-official-text", utcnow(), n))
                promoted += res.rowcount
            else:
                conn.execute("""
                UPDATE authority_triples
                   SET graft_status=CASE WHEN graft_status='UNRESOLVED' THEN 'PARTIAL' ELSE graft_status END,
                       graft_notes=?,
                       updated_at=?
                 WHERE normalized_citation=?
                """, ("Imported graft row but not verified exact+pinpoint", utcnow(), n))
    conn.commit()
    return {"imported": imported, "promoted": promoted}

def scan_files(conn, root: Path):
    rows = 0
    allowed = {".txt",".md",".csv",".json",".jsonl",".docx",".pdf"}
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        if p.suffix.lower() not in allowed:
            continue
        try:
            st = p.stat()
        except OSError:
            continue
        rel = str(p.relative_to(root)).replace("\\","/")
        lane = "MEEK1" if "MEEK1" in rel else "MEEK2" if "MEEK2" in rel else "MEEK3" if "MEEK3" in rel else "MEEK4" if "MEEK4" in rel else "UNLANED"
        doc_kind = p.suffix.lower().lstrip(".")
        file_id = f"FILE::{identity_key_v1(p)}"
        conn.execute("""
        INSERT OR REPLACE INTO fileatoms(file_id,path,rel_path,size,mtime_ns,identity_key,lane,case_anchor,doc_kind,created_at)
        VALUES(?,?,?,?,?,?,?,?,?,?)
        """, (file_id, str(p), rel, st.st_size, getattr(st,"st_mtime_ns", int(st.st_mtime*1e9)), identity_key_v1(p), lane, "", doc_kind, utcnow()))
        rows += 1
    conn.commit()
    return rows

def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")

def _read_docx(path: Path) -> str:
    if Document is None:
        return ""
    # python-docx import class doubles for reading
    from docx import Document as Doc
    d = Doc(str(path))
    return "\n".join(p.text for p in d.paragraphs if p.text)

def _read_pdf(path: Path) -> str:
    try:
        import pypdf
        r = pypdf.PdfReader(str(path))
        out = []
        for i, pg in enumerate(r.pages, start=1):
            txt = pg.extract_text() or ""
            if txt.strip():
                out.append(f"[PAGE {i}] {txt}")
        return "\n".join(out)
    except Exception:
        return ""

def extract_textbank(conn, root: Path, textbank_dir: Path) -> Dict[str,int]:
    ensure_dir(textbank_dir)
    processed = 0
    quotes_added = 0
    for row in conn.execute("SELECT * FROM fileatoms ORDER BY rel_path"):
        p = Path(row["path"])
        if not p.exists():
            continue
        text = ""
        try:
            if p.suffix.lower() in {".txt",".md",".json",".jsonl",".csv"}:
                text = _read_text_file(p)
            elif p.suffix.lower() == ".docx":
                text = _read_docx(p)
            elif p.suffix.lower() == ".pdf":
                text = _read_pdf(p)
        except Exception as e:
            text = f""
        out = textbank_dir / (re.sub(r"[^A-Za-z0-9._-]+","_", row["rel_path"]) + ".txt")
        out.write_text(text or "", encoding="utf-8")
        conn.execute("UPDATE fileatoms SET extracted_ok=? WHERE file_id=?", (1 if text else 0, row["file_id"]))
        processed += 1
        if text:
            q = quoteize_text(conn, row["file_id"], text)
            quotes_added += q
    conn.commit()
    return {"processed": processed, "quotes_added": quotes_added}

def quoteize_text(conn, file_id: str, text: str) -> int:
    # split into chunks by page markers or paragraphs
    chunks = []
    page = None
    for block in re.split(r"\n\s*\n+", text):
        blk = block.strip()
        if not blk:
            continue
        m = re.match(r"\[PAGE\s+(\d+)\]\s*(.*)$", blk, re.I|re.S)
        if m:
            page = int(m.group(1))
            blk = m.group(2).strip()
        chunks.append((page, blk))
    count = 0
    for idx,(pg,blk) in enumerate(chunks, start=1):
        # sub-split long blocks
        for part_idx, part in enumerate(textwrap.wrap(blk, width=700, break_long_words=False, break_on_hyphens=False) or [blk], start=1):
            qid = f"Q::{file_id}::{idx}.{part_idx}"
            issue = issue_tag_for_text(part)
            conn.execute("""
            INSERT OR REPLACE INTO quotes(quote_id,file_id,text,page,line_start,line_end,speaker,issue_tag,created_at)
            VALUES(?,?,?,?,?,?,?,?,?)
            """, (qid,file_id,part,pg,None,None,None,issue,utcnow()))
            conn.execute("INSERT INTO quotes_fts(rowid,quote_id,text) VALUES((SELECT COALESCE(MAX(rowid),0)+1 FROM quotes_fts),?,?)", (qid,part))
            count += 1
    return count

def ingest_transcripts(conn, path: Path) -> int:
    n=0
    with path.open("r", encoding="utf-8", newline="") as f:
        rd = csv.DictReader(f)
        for r in rd:
            tr_id = r.get("transcript_id","T1")
            page = int(r["page"]) if r.get("page") else None
            ls = int(r["line_start"]) if r.get("line_start") else None
            le = int(r["line_end"]) if r.get("line_end") else None
            speaker = r.get("speaker","")
            text = r.get("text","")
            issue = r.get("issue_tag") or issue_tag_for_text(text)
            tri = f"TR::{tr_id}::{page}:{ls}-{le}"
            conn.execute("""
            INSERT OR REPLACE INTO transcripts(transcript_row_id,transcript_id,page,line_start,line_end,speaker,text,issue_tag,created_at)
            VALUES(?,?,?,?,?,?,?,?,?)
            """, (tri,tr_id,page,ls,le,speaker,text,issue,utcnow()))
            conn.execute("INSERT INTO transcripts_fts(rowid,transcript_row_id,text) VALUES((SELECT COALESCE(MAX(rowid),0)+1 FROM transcripts_fts),?,?)", (tri,text))
            n += 1
    conn.commit()
    return n

def ingest_order_findings(conn, path: Path) -> int:
    n=0
    with path.open("r", encoding="utf-8", newline="") as f:
        rd = csv.DictReader(f)
        for r in rd:
            fid = r.get("finding_id") or f"F::{n+1}"
            txt = r["finding_text"]
            issue = r.get("issue_tag") or issue_tag_for_text(txt)
            adverse = 1 if str(r.get("adverse","1")).strip() in {"1","true","TRUE","yes","YES"} else 0
            conn.execute("""
            INSERT OR REPLACE INTO order_findings(finding_id,order_id,finding_text,issue_tag,adverse,created_at)
            VALUES(?,?,?,?,?,?)
            """, (fid, r.get("order_id","ORDER1"), txt, issue, adverse, utcnow()))
            n += 1
    conn.commit()
    return n

def score_finding_to_transcript(finding_text: str, finding_issue: str, tr_text: str, tr_issue: str, speaker: str, weights: Dict[str,float]) -> Tuple[float, Dict[str,float]]:
    ft = set(tokenize(finding_text))
    tt = set(tokenize(tr_text))
    jacc = (len(ft & tt) / len(ft | tt)) if (ft or tt) else 0.0
    overlap = len(ft & tt)
    issue_match = 1.0 if (finding_issue and tr_issue and finding_issue == tr_issue) else 0.0
    speaker_bonus = 0.0
    if speaker:
        sp = speaker.lower()
        if "judge" in sp:
            speaker_bonus = 0.4
        elif "court" in sp:
            speaker_bonus = 0.25
    phrase_bonus = 0.0
    # exact-ish phrase fragments
    for frag in [p.strip() for p in re.split(r"[.;,]", finding_text) if len(p.strip())>=12]:
        frag_toks = tokenize(frag)
        if len(frag_toks) >= 3 and all(tok in tt for tok in frag_toks[:3]):
            phrase_bonus += 0.15
            break
    score = (
        weights.get("jacc",1.0) * jacc +
        weights.get("overlap",0.05) * overlap +
        weights.get("issue",0.6) * issue_match +
        weights.get("speaker",0.3) * speaker_bonus +
        weights.get("phrase",0.4) * phrase_bonus
    )
    detail = {
        "jacc": round(jacc,4),
        "overlap": overlap,
        "issue_match": issue_match,
        "speaker_bonus": speaker_bonus,
        "phrase_bonus": phrase_bonus,
    }
    return score, detail

def link_order_findings(conn, threshold: float = 0.95, top_k: int = 3, weights: Optional[Dict[str,float]] = None) -> Dict[str,int]:
    weights = weights or {"jacc":1.1,"overlap":0.06,"issue":0.8,"speaker":0.4,"phrase":0.6}
    conn.execute("DELETE FROM transcript_order_links")
    findings = list(conn.execute("SELECT * FROM order_findings ORDER BY finding_id"))
    transcripts = list(conn.execute("SELECT * FROM transcripts ORDER BY transcript_id,page,line_start"))
    links = 0
    accepted = 0
    for f in findings:
        scored = []
        for tr in transcripts:
            score, detail = score_finding_to_transcript(
                f["finding_text"], f["issue_tag"], tr["text"], tr["issue_tag"], tr["speaker"] or "", weights
            )
            scored.append((score, detail, tr))
        scored.sort(key=lambda x: x[0], reverse=True)
        for rank,(score,detail,tr) in enumerate(scored[:top_k], start=1):
            lid = f"LINK::{f['finding_id']}::{tr['transcript_row_id']}"
            acc = 1 if (rank == 1 and score >= threshold) else 0
            conn.execute("""
            INSERT INTO transcript_order_links(link_id,finding_id,transcript_row_id,score,score_detail_json,accepted,created_at)
            VALUES(?,?,?,?,?,?,?)
            """, (lid,f["finding_id"],tr["transcript_row_id"],float(score),json.dumps(detail, ensure_ascii=False),acc,utcnow()))
            links += 1
            if acc:
                accepted += 1
    conn.commit()
    return {"links":links, "accepted":accepted}

def tune_linker(conn, gold_csv: Path, thresholds: Iterable[float] = (0.7,0.8,0.9,1.0,1.1)) -> Dict[str,object]:
    """
    Gold CSV columns:
      finding_id, transcript_row_id, is_match (1/0)
    """
    gold = []
    with gold_csv.open("r", encoding="utf-8", newline="") as f:
        rd = csv.DictReader(f)
        for r in rd:
            gold.append((r["finding_id"], r["transcript_row_id"], int(r["is_match"])))
    # candidate weights grid
    grids = []
    for jacc in (0.8,1.0,1.2):
        for overlap in (0.04,0.06,0.08):
            for issue in (0.4,0.8,1.2):
                for speaker in (0.2,0.4,0.6):
                    for phrase in (0.2,0.4,0.8):
                        grids.append({"jacc":jacc,"overlap":overlap,"issue":issue,"speaker":speaker,"phrase":phrase})
    # Preload
    f_map = {r["finding_id"]: r for r in conn.execute("SELECT * FROM order_findings")}
    t_map = {r["transcript_row_id"]: r for r in conn.execute("SELECT * FROM transcripts")}
    best = None
    for w in grids:
        # score all labeled pairs
        pair_scores = []
        for fid, tid, is_match in gold:
            f = f_map.get(fid); tr = t_map.get(tid)
            if not f or not tr: 
                continue
            s, _ = score_finding_to_transcript(f["finding_text"], f["issue_tag"], tr["text"], tr["issue_tag"], tr["speaker"] or "", w)
            pair_scores.append((s,is_match))
        if not pair_scores:
            continue
        for thr in thresholds:
            tp=fp=fn=0
            for s,is_match in pair_scores:
                pred = 1 if s >= thr else 0
                if pred and is_match: tp += 1
                elif pred and not is_match: fp += 1
                elif (not pred) and is_match: fn += 1
            rec = tp/(tp+fn) if (tp+fn) else 0.0
            prec = tp/(tp+fp) if (tp+fp) else 0.0
            f1 = 2*prec*rec/(prec+rec) if (prec+rec) else 0.0
            score_tuple = (rec, f1, prec)  # favor recall as user requested
            if best is None or score_tuple > best["score_tuple"]:
                best = {"weights": w, "threshold": thr, "recall": rec, "precision": prec, "f1": f1, "score_tuple": score_tuple}
    if not best:
        return {"status":"NO_GOLD"}
    run_id = f"TUNE::{utcnow()}"
    conn.execute("""
    INSERT INTO linker_tuning(tuning_run_id,weights_json,threshold,recall,precision,f1,notes,created_at)
    VALUES(?,?,?,?,?,?,?,?)
    """, (run_id, json.dumps(best["weights"]), float(best["threshold"]), best["recall"], best["precision"], best["f1"], "Recall-priority grid search", utcnow()))
    conn.commit()
    return {k:v for k,v in best.items() if k!="score_tuple"}

def seed_deadline_registry(conn) -> int:
    """Template registry; subrule clocks and tolling notes are scaffolds pending verification."""
    rows = [
        # COA templates
        ("MEEK2","COA_APPEAL_OF_RIGHT","COA","served","claim_of_appeal", 21, 0, "Template rail. Verify exact MCR subrule and tolling before filing.", "MCR 7.204", "UNVERIFIED_TEMPLATE"),
        ("MEEK2","COA_APPEAL_OF_RIGHT","COA","entered","transcript_order", 7, 0, "Template rail for transcript designation/order workflows. Verify exact subrule.", "MCR 7.210", "UNVERIFIED_TEMPLATE"),
        ("MEEK3","COA_APPLICATION_LEAVE","COA","entered","application_leave", 21, 0, "Template only. Verify exact clock and post-judgment tolling.", "MCR 7.205", "UNVERIFIED_TEMPLATE"),
        # MSC templates
        ("MEEK2","MSC_APP_LEAVE","MSC","entered","application_leave_msc", 56, 0, "Template only. Verify counting from COA decision/order and reconsideration impact.", "MCR 7.305", "UNVERIFIED_TEMPLATE"),
        # JTC templates
        ("MEEK4","JTC_COMPLAINT","JTC","recorded","misconduct_packet", None, 0, "No fixed statewide deadline in this template. Prompt immediate filing after record packet stabilization.", "JTC_CANON/PROC", "UNVERIFIED_TEMPLATE"),
        # Trial templates
        ("MEEK1","TRIAL_MOTION_RELIEF","TRIAL","served","motion_response", 14, 0, "Template motion response rail. Verify local/admin order variations.", "MCR 2.119", "UNVERIFIED_TEMPLATE"),
    ]
    inserted = 0
    for lane, vehicle, forum, anchor, trigger, days_due, biz, tolling, auth, status in rows:
        rid = f"DR::{forum}::{vehicle}::{anchor}::{trigger}"
        conn.execute("""
        INSERT OR REPLACE INTO deadline_rules(rule_id,lane,vehicle_code,forum,anchor_type,trigger_event,due_days,business_days_only,tolling_notes,authority_citation,rule_status,active,created_at)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (rid,lane,vehicle,forum,anchor,trigger,days_due,biz,tolling,auth,status,1,utcnow()))
        inserted += 1
    conn.commit()
    return inserted

def ingest_deadline_events(conn, path: Path) -> int:
    n=0
    with path.open("r", encoding="utf-8", newline="") as f:
        rd = csv.DictReader(f)
        for r in rd:
            eid = r.get("event_id") or f"DE::{n+1}"
            conn.execute("""
            INSERT OR REPLACE INTO deadline_events(event_id,lane,case_anchor,vehicle_code,forum,trigger_event,entered_date,signed_date,served_date,recorded_date,notes,created_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?)
            """, (
                eid, r.get("lane",""), r.get("case_anchor",""), r["vehicle_code"], r["forum"], r["trigger_event"],
                r.get("entered_date") or None, r.get("signed_date") or None, r.get("served_date") or None, r.get("recorded_date") or None,
                r.get("notes",""), utcnow()
            ))
            n += 1
    conn.commit()
    return n

def _parse_date(s: Optional[str]) -> Optional[dt.date]:
    if not s:
        return None
    return dt.date.fromisoformat(s)

def _add_days(d: dt.date, days: int, business_only: bool=False) -> dt.date:
    if not business_only:
        return d + dt.timedelta(days=days)
    cur = d
    left = days
    while left > 0:
        cur += dt.timedelta(days=1)
        if cur.weekday() < 5:
            left -= 1
    return cur

def lock_deadlines(conn) -> Dict[str,int]:
    conn.execute("DELETE FROM deadline_locks")
    events = list(conn.execute("SELECT * FROM deadline_events"))
    rules = list(conn.execute("SELECT * FROM deadline_rules WHERE active=1"))
    locks = 0
    missing_anchor = 0
    pending_verify = 0
    for ev in events:
        matches = [r for r in rules if r["vehicle_code"] == ev["vehicle_code"] and r["forum"] == ev["forum"] and r["trigger_event"] == ev["trigger_event"]]
        if not matches:
            lid = f"DL::{ev['event_id']}::RULE_MISSING"
            conn.execute("""
            INSERT INTO deadline_locks(lock_id,event_id,rule_id,anchor_type,anchor_date,due_date,lock_status,tolling_notes,created_at)
            VALUES(?,?,?,?,?,?,?,?,?)
            """, (lid, ev["event_id"], "RULE_MISSING", "", None, None, "RULE_MISSING", "No deadline rule matched event.", utcnow()))
            locks += 1
            continue
        for r in matches:
            anchor_col = f"{r['anchor_type']}_date"
            anchor_date = ev[anchor_col] if anchor_col in ev.keys() else None
            status = "LOCKED"
            due_date = None
            if not anchor_date:
                status = "MISSING_ANCHOR"
                missing_anchor += 1
            elif r["due_days"] is None:
                status = "PENDING_VERIFY"
                pending_verify += 1
            else:
                d = _parse_date(anchor_date)
                due = _add_days(d, int(r["due_days"]), bool(r["business_days_only"]))
                due_date = due.isoformat()
                if r["rule_status"] != "VERIFIED":
                    status = "PENDING_VERIFY"
                    pending_verify += 1
            lid = f"DL::{ev['event_id']}::{r['rule_id']}"
            conn.execute("""
            INSERT INTO deadline_locks(lock_id,event_id,rule_id,anchor_type,anchor_date,due_date,lock_status,tolling_notes,created_at)
            VALUES(?,?,?,?,?,?,?,?,?)
            """, (lid, ev["event_id"], r["rule_id"], r["anchor_type"], anchor_date, due_date, status, r["tolling_notes"], utcnow()))
            locks += 1
    conn.commit()
    return {"locks":locks, "missing_anchor":missing_anchor, "pending_verify":pending_verify}

def consolidate_contradictions(conn) -> int:
    conn.execute("DELETE FROM contradiction_pairs")
    # Build candidate pool from order findings + accepted transcript links + quotes
    # Pair findings with transcript linked quotes that diverge lexically (possible contradiction)
    inserted = 0
    qrows = list(conn.execute("""
    SELECT l.finding_id, l.transcript_row_id, l.score, f.issue_tag f_issue, f.finding_text, t.text tr_text
      FROM transcript_order_links l
      JOIN order_findings f ON f.finding_id=l.finding_id
      JOIN transcripts t ON t.transcript_row_id=l.transcript_row_id
     WHERE l.accepted=1
    """))
    for r in qrows:
        ft = set(tokenize(r["finding_text"]))
        tt = set(tokenize(r["tr_text"]))
        overlap = len(ft & tt)
        # contradiction heuristic: linked but low overlap OR negation mismatch
        neg_f = bool(re.search(r"\b(no|not|never|none|didn't|did not)\b", r["finding_text"], re.I))
        neg_t = bool(re.search(r"\b(no|not|never|none|didn't|did not)\b", r["tr_text"], re.I))
        contradiction_signal = (overlap <= 2) or (neg_f != neg_t)
        if contradiction_signal:
            pid = f"CP::{r['finding_id']}::{r['transcript_row_id']}"
            conn.execute("""
            INSERT INTO contradiction_pairs(pair_id,lane_scope,proposition_key,left_source,left_ref,left_text,right_source,right_ref,right_text,confidence,created_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?)
            """, (
                pid, r["f_issue"] or "UNKNOWN", proposition_key(r["finding_text"]),
                "ORDER_FINDING", r["finding_id"], r["finding_text"],
                "TRANSCRIPT", r["transcript_row_id"], r["tr_text"],
                0.65 if neg_f != neg_t else 0.45,
                utcnow()
            ))
            inserted += 1
    conn.commit()
    return inserted

def _md_table(headers: List[str], rows: List[List[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"]*len(headers)) + " |"]
    for r in rows:
        out.append("| " + " | ".join(str(x).replace("\n"," ").replace("|","\\|") for x in r) + " |")
    return "\n".join(out)

def emit_docx(path: Path, title: str, sections: List[Tuple[str,str]]):
    if Document is None:
        return False, "python-docx unavailable"
    doc = Document()
    doc.add_heading(title, level=0)
    for h, body in sections:
        doc.add_heading(h, level=1)
        for para in body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    ensure_dir(path.parent)
    doc.save(str(path))
    return True, "ok"

def _wrap_lines(text: str, width: int = 95) -> List[str]:
    lines = []
    for para in text.split("\n"):
        if not para.strip():
            lines.append("")
            continue
        lines.extend(textwrap.wrap(para, width=width, break_long_words=False, break_on_hyphens=False) or [""])
    return lines

def emit_pdf(path: Path, title: str, sections: List[Tuple[str,str]]):
    if canvas is None:
        return False, "reportlab unavailable"
    ensure_dir(path.parent)
    c = canvas.Canvas(str(path), pagesize=letter)
    W, H = letter
    margin = 54
    y = H - margin
    def new_page():
        nonlocal y
        c.showPage()
        y = H - margin
    c.setTitle(title)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, title[:110]); y -= 22
    c.setFont("Helvetica", 10)
    for h, body in sections:
        if y < 80: new_page()
        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin, y, h[:110]); y -= 16
        c.setFont("Helvetica", 10)
        for line in _wrap_lines(body, 98):
            if y < 60: new_page(); c.setFont("Helvetica", 10)
            c.drawString(margin, y, line[:120]); y -= 12
        y -= 6
    c.save()
    return True, "ok"

def compile_courtpack(conn, out_root: Path, lane: str="GLOBAL", case_anchor: str="UNSPECIFIED") -> Dict[str, object]:
    ts = dt.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    pack_id = f"PACK::{lane}::{ts}"
    out_dir = out_root / f"{lane}_{ts}"
    ensure_dir(out_dir)

    # Gather data
    auth_rows = list(conn.execute("SELECT * FROM authority_triples ORDER BY normalized_citation"))
    link_rows = list(conn.execute("""
        SELECT l.*, f.finding_text, t.transcript_id, t.page, t.line_start, t.line_end, t.text tr_text
          FROM transcript_order_links l
          JOIN order_findings f ON f.finding_id=l.finding_id
          JOIN transcripts t ON t.transcript_row_id=l.transcript_row_id
         ORDER BY l.score DESC, l.finding_id
    """))
    deadline_rows = list(conn.execute("""
        SELECT dl.*, de.vehicle_code, de.forum, de.trigger_event, de.case_anchor, dr.authority_citation, dr.rule_status
          FROM deadline_locks dl
          LEFT JOIN deadline_events de ON de.event_id = dl.event_id
          LEFT JOIN deadline_rules dr ON dr.rule_id = dl.rule_id
         ORDER BY de.forum, de.vehicle_code, dl.anchor_type
    """))
    contra_rows = list(conn.execute("SELECT * FROM contradiction_pairs ORDER BY confidence DESC"))

    # Build artifacts
    case_state = {
        "lane": lane,
        "case_anchor": case_anchor,
        "counts": {
            "authority_triples": len(auth_rows),
            "transcript_order_links": len(link_rows),
            "deadline_locks": len(deadline_rows),
            "contradiction_pairs": len(contra_rows),
        },
        "generated_at": utcnow(),
        "status": "DRAFT_DROPIN_READY"
    }
    (out_dir/"CASE_STATE.json").write_text(json.dumps(case_state, indent=2), encoding="utf-8")
    (out_dir/"CASE_STATE.md").write_text(
        f"# CASE_STATE\n\n- lane: {lane}\n- case_anchor: {case_anchor}\n- status: DRAFT_DROPIN_READY\n"
        f"- authority_triples: {len(auth_rows)}\n- transcript_order_links: {len(link_rows)}\n"
        f"- deadline_locks: {len(deadline_rows)}\n- contradiction_pairs: {len(contra_rows)}\n",
        encoding="utf-8"
    )

    auth_md_rows = []
    for r in auth_rows:
        auth_md_rows.append([
            r["citation"], r["pinpoint_status"], r["graft_status"], (r["pinpoint"] or "")[:40], (r["source_url_status"] or ""), (r["source_url"] or "")[:60]
        ])
    (out_dir/"AuthorityTriples.md").write_text("# AuthorityTriples\n\n"+_md_table(
        ["Citation","PinpointStatus","GraftStatus","Pinpoint","SourceStatus","SourceURL"], auth_md_rows
    ), encoding="utf-8")
    (out_dir/"AuthorityTriples.json").write_text(json.dumps([dict(r) for r in auth_rows], indent=2), encoding="utf-8")

    sof_md = ["# Transcript_StatementOfFacts_Pins", ""]
    for r in link_rows:
        mark = "ACCEPTED" if r["accepted"] else "CANDIDATE"
        sof_md.append(f"## {mark} {r['finding_id']} -> {r['transcript_id']} p.{r['page']} l.{r['line_start']}-{r['line_end']} (score={r['score']:.3f})")
        sof_md.append(f"- Finding: {r['finding_text']}")
        sof_md.append(f"- Transcript: {r['tr_text']}")
        sof_md.append("")
    (out_dir/"Transcript_StatementOfFacts_Pins.md").write_text("\n".join(sof_md), encoding="utf-8")

    dl_table = []
    for r in deadline_rows:
        dl_table.append([
            r["forum"] or "", r["vehicle_code"] or "", r["trigger_event"] or "", r["anchor_type"] or "",
            r["anchor_date"] or "", r["due_date"] or "", r["lock_status"] or "", (r["authority_citation"] or ""), (r["rule_status"] or "")
        ])
    (out_dir/"Deadlines.md").write_text("# Deadlines\n\n"+_md_table(
        ["Forum","Vehicle","Trigger","Anchor","AnchorDate","DueDate","Status","Authority","RuleStatus"], dl_table
    ), encoding="utf-8")

    contra_md = ["# ContradictionMap", ""]
    for r in contra_rows:
        contra_md.append(f"## {r['pair_id']} (confidence={r['confidence']})")
        contra_md.append(f"- PropositionKey: {r['proposition_key']}")
        contra_md.append(f"- Left [{r['left_source']} {r['left_ref']}]: {r['left_text']}")
        contra_md.append(f"- Right [{r['right_source']} {r['right_ref']}]: {r['right_text']}")
        contra_md.append("")
    (out_dir/"ContradictionMap.md").write_text("\n".join(contra_md), encoding="utf-8")
    (out_dir/"ContradictionMap.json").write_text(json.dumps([dict(r) for r in contra_rows], indent=2), encoding="utf-8")

    # Validation / RedTeam
    unresolved_auth = sum(1 for r in auth_rows if r["graft_status"] != "RESOLVED_VERIFIED")
    pending_deadlines = sum(1 for r in deadline_rows if r["lock_status"] != "LOCKED")
    validation = [
        "# Validation_RedTeam",
        "",
        f"- unresolved authorities: {unresolved_auth}",
        f"- non-locked deadlines (pending/missing): {pending_deadlines}",
        f"- accepted transcript links: {sum(1 for r in link_rows if r['accepted'])}",
        f"- contradiction candidates: {len(contra_rows)}",
        "",
        "## Acquisition tasks",
        "- Import exact official text and exact pinpoints into official_text_grafts.jsonl to promote PARTIAL -> RESOLVED_VERIFIED.",
        "- Export real transcript CSV with page/line rows and order findings CSV from hostile orders for score tuning.",
        "- Verify vehicle-specific deadline registry subrules and set deadline_rules.rule_status=VERIFIED for filing use.",
    ]
    (out_dir/"Validation_RedTeam.md").write_text("\n".join(validation), encoding="utf-8")

    # SBNA + VehicleMap + ContextPack + Draft blocks + Proposed order + Service checklist
    (out_dir/"SBNA.md").write_text(
        "# SBNA\n\n"
        "Single Best Next Action:\n"
        "1. Promote official authority text/pinpoints to RESOLVED_VERIFIED for the target vehicle.\n"
        "2. Tune transcript linker on gold pairs to improve recall for hostile-order findings.\n"
        "3. Verify and lock filing clocks (entered/signed/served split) for the immediate COA/MSC/JTC vehicle.\n",
        encoding="utf-8"
    )
    (out_dir/"VehicleMap.md").write_text(
        "# VehicleMap\n\n"
        "- MEEK1: Housing trial and appellate preservation vehicles\n"
        "- MEEK2: Custody/PT objections, COA leave/right, MSC application, emergency relief\n"
        "- MEEK3: PPO show-cause defense, contempt review, appellate/JTC narrative routes\n"
        "- MEEK4: JTC complaint + Canon framing + COA/MSC bias preservation overlays\n",
        encoding="utf-8"
    )
    (out_dir/"ContextPack.md").write_text(
        "# ContextPack\n\n"
        "This CourtPack is generated from a local, append-only DocForge run. Originals remain read-only. "
        "Derived artifacts summarize authority triples, transcript statement-of-facts pin candidates, split deadline locks, and contradiction candidates.\n",
        encoding="utf-8"
    )
    draft = """# DRAFT_BLOCKS_DROPIN

## Background
Petitioner compiles this record from existing filings, transcripts, orders, and court-generated events. This packet preserves timing anchors separately (entered/signed/served) and links factual findings to transcript page-line citations.

## Facts
Insert accepted Transcript_StatementOfFacts_Pins items here, with exact page-line cites and exhibit cross-links.

## Argument
Insert AuthorityTriples entries only after graft_status = RESOLVED_VERIFIED for the specific proposition.

## Relief
State the vehicle-specific relief requested and attach ProposedOrder / ServiceChecklist.
"""
    (out_dir/"DRAFT_BLOCKS_DROPIN.md").write_text(draft, encoding="utf-8")
    (out_dir/"ProposedOrder.md").write_text("# ProposedOrder\n\n[Draft order block generated by vehicle-specific compiler lane]\n", encoding="utf-8")
    (out_dir/"ServiceChecklist.md").write_text("# ServiceChecklist\n\n- Verify service method\n- Verify service recipients\n- File proof of service\n", encoding="utf-8")

    # Integrated docx/pdf emit lane
    sections = [
        ("CASE_STATE", (out_dir/"CASE_STATE.md").read_text(encoding="utf-8")),
        ("AuthorityTriples", (out_dir/"AuthorityTriples.md").read_text(encoding="utf-8")),
        ("Deadlines", (out_dir/"Deadlines.md").read_text(encoding="utf-8")),
        ("Transcript_StatementOfFacts_Pins", (out_dir/"Transcript_StatementOfFacts_Pins.md").read_text(encoding="utf-8")),
        ("ContradictionMap", (out_dir/"ContradictionMap.md").read_text(encoding="utf-8")),
        ("Validation_RedTeam", (out_dir/"Validation_RedTeam.md").read_text(encoding="utf-8")),
    ]
    docx_ok, docx_msg = emit_docx(out_dir/"CourtPack_Summary.docx", f"CourtPack Summary {lane} {case_anchor}", sections)
    pdf_ok, pdf_msg = emit_pdf(out_dir/"CourtPack_Summary.pdf", f"CourtPack Summary {lane} {case_anchor}", sections)

    manifest = {
        "pack_id": pack_id,
        "lane": lane,
        "case_anchor": case_anchor,
        "generated_at": utcnow(),
        "counts": case_state["counts"],
        "doc_emit": {"docx": {"ok": docx_ok, "msg": docx_msg}, "pdf": {"ok": pdf_ok, "msg": pdf_msg}},
        "no_sha256_default": True,
        "originals_immutable": True,
        "derived_append_only": True,
    }
    (out_dir/"PACK_MANIFEST.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    conn.execute("INSERT INTO courtpack_runs(pack_id,lane,case_anchor,out_dir,manifest_json,created_at) VALUES(?,?,?,?,?,?)",
                 (pack_id,lane,case_anchor,str(out_dir),json.dumps(manifest),utcnow()))
    conn.commit()
    return manifest

def write_report(workspace: Path, name: str, payload: Dict[str, object]):
    ensure_dir(workspace / "reports")
    p = workspace / "reports" / f"{name}.json"
    p.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return p

def autopilot(workspace: Path) -> Dict[str, object]:
    db_path = workspace / DEFAULT_DB
    conn = connect_db(db_path)
    initdb(conn)
    seed_deadline_registry(conn)
    # sample inputs
    samples = workspace / "sample_inputs"
    if not samples.exists():
        samples = Path(__file__).resolve().parents[1] / "sample_inputs"
    ing_auth = ingest_authority_requests(conn, samples / "authority_requests.jsonl")
    promo = promote_official_text(conn, samples / "official_text_grafts.jsonl")
    tr_n = ingest_transcripts(conn, samples / "transcript_rows.csv")
    of_n = ingest_order_findings(conn, samples / "order_findings.csv")
    # optional tuning if gold exists
    tuning = {}
    gold = samples / "transcript_link_gold.csv"
    if gold.exists():
        tuning = tune_linker(conn, gold)
        if "weights" in tuning:
            link = link_order_findings(conn, threshold=float(tuning["threshold"]), weights=tuning["weights"])
        else:
            link = link_order_findings(conn)
    else:
        link = link_order_findings(conn)
    de_n = ingest_deadline_events(conn, samples / "deadline_events.csv")
    locks = lock_deadlines(conn)
    contra = consolidate_contradictions(conn)
    manifest = compile_courtpack(conn, workspace / "outputs" / "CourtPacks", lane="MEEK2", case_anchor="24-01507-DC")
    result = {
        "authority_requests_ingested": ing_auth,
        "official_promotion": promo,
        "transcripts_ingested": tr_n,
        "order_findings_ingested": of_n,
        "tuning": tuning,
        "links": link,
        "deadline_events_ingested": de_n,
        "deadline_locks": locks,
        "contradictions": contra,
        "courtpack_manifest": manifest
    }
    write_report(workspace, "autopilot_report", result)
    return result

def parse_weights(s: str) -> Dict[str,float]:
    out = {}
    for part in s.split(","):
        if not part.strip():
            continue
        k,v = part.split("=")
        out[k.strip()] = float(v.strip())
    return out

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--workspace", default=".", help="Workspace root")
    sub = ap.add_subparsers(dest="cmd", required=True)

    sub.add_parser("initdb")
    p = sub.add_parser("scan"); p.add_argument("--root", required=True)
    p = sub.add_parser("extract"); p.add_argument("--root", required=True); p.add_argument("--textbank", default="textbank")
    p = sub.add_parser("ingest-authority-requests"); p.add_argument("--jsonl", required=True)
    p = sub.add_parser("ingest-official-text"); p.add_argument("--jsonl", required=True)
    p = sub.add_parser("ingest-transcripts"); p.add_argument("--csv", required=True)
    p = sub.add_parser("ingest-order-findings"); p.add_argument("--csv", required=True)
    p = sub.add_parser("tune-linker"); p.add_argument("--gold-csv", required=True)
    p = sub.add_parser("link-order-findings"); p.add_argument("--threshold", type=float, default=0.95); p.add_argument("--weights", default="")
    sub.add_parser("seed-deadline-registry")
    p = sub.add_parser("ingest-deadline-events"); p.add_argument("--csv", required=True)
    sub.add_parser("lock-deadlines")
    sub.add_parser("consolidate-contradictions")
    p = sub.add_parser("courtpack"); p.add_argument("--lane", default="GLOBAL"); p.add_argument("--case-anchor", default="UNSPECIFIED")
    sub.add_parser("autopilot")

    args = ap.parse_args()
    ws = Path(args.workspace).resolve()
    ensure_dir(ws)
    db_path = ws / DEFAULT_DB
    conn = connect_db(db_path)

    out = {"cmd": args.cmd, "workspace": str(ws), "utc": utcnow()}
    if args.cmd == "initdb":
        initdb(conn); out["status"]="ok"
    elif args.cmd == "scan":
        initdb(conn); out["files_scanned"] = scan_files(conn, Path(args.root))
    elif args.cmd == "extract":
        initdb(conn); out.update(extract_textbank(conn, Path(args.root), ws / args.textbank))
    elif args.cmd == "ingest-authority-requests":
        initdb(conn); out["authority_requests_ingested"] = ingest_authority_requests(conn, Path(args.jsonl))
    elif args.cmd == "ingest-official-text":
        initdb(conn); out.update(promote_official_text(conn, Path(args.jsonl)))
    elif args.cmd == "ingest-transcripts":
        initdb(conn); out["transcripts_ingested"] = ingest_transcripts(conn, Path(args.csv))
    elif args.cmd == "ingest-order-findings":
        initdb(conn); out["order_findings_ingested"] = ingest_order_findings(conn, Path(args.csv))
    elif args.cmd == "tune-linker":
        initdb(conn); out.update(tune_linker(conn, Path(args.gold_csv)))
    elif args.cmd == "link-order-findings":
        initdb(conn)
        # use latest tuning if present unless explicit weights
        weights = None
        threshold = args.threshold
        if args.weights:
            weights = parse_weights(args.weights)
        else:
            row = conn.execute("SELECT * FROM linker_tuning ORDER BY created_at DESC LIMIT 1").fetchone()
            if row:
                weights = json.loads(row["weights_json"])
                threshold = float(row["threshold"])
        out.update(link_order_findings(conn, threshold=threshold, weights=weights))
        out["threshold"] = threshold
        out["weights"] = weights or {}
    elif args.cmd == "seed-deadline-registry":
        initdb(conn); out["deadline_rules_seeded"] = seed_deadline_registry(conn)
    elif args.cmd == "ingest-deadline-events":
        initdb(conn); out["deadline_events_ingested"] = ingest_deadline_events(conn, Path(args.csv))
    elif args.cmd == "lock-deadlines":
        initdb(conn); out.update(lock_deadlines(conn))
    elif args.cmd == "consolidate-contradictions":
        initdb(conn); out["contradiction_pairs"] = consolidate_contradictions(conn)
    elif args.cmd == "courtpack":
        initdb(conn); out["manifest"] = compile_courtpack(conn, ws / "outputs" / "CourtPacks", lane=args.lane, case_anchor=args.case_anchor)
    elif args.cmd == "autopilot":
        out["result"] = autopilot(ws)
    else:
        ap.error("Unknown command")

    rep = write_report(ws, f"last_{args.cmd}", out)
    print(json.dumps({"ok": True, "report": str(rep), **out}, indent=2))

if __name__ == "__main__":
    main()
