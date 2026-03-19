"""
OMEGA Phase 4B: DOCX Extraction
Extract text, tables, headers/footers from .docx/.doc files using python-docx.
"""
import json
import sqlite3
import sys
import time
from pathlib import Path

from config import get_cyclepack_dir, long_path, report_progress
from safety import write_phase_checkpoint, is_phase_done


def extract_docx(docx_path: str, output_dir: Path) -> dict | None:
    try:
        from docx import Document
    except ImportError:
        print("[PHASE4B] python-docx not available", file=sys.stderr)
        return None

    out_name = Path(docx_path).stem
    txt_path = output_dir / f"{out_name}.txt"
    if txt_path.exists():
        return {"method": "cached", "path": str(txt_path)}

    try:
        doc = Document(long_path(docx_path))
        parts: list[str] = []

        # Headers from all sections
        for i, section in enumerate(doc.sections):
            header_text = section.header.text.strip() if section.header else ""
            if header_text:
                parts.append(f"--- HEADER (Section {i + 1}) ---\n{header_text}")

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables
        for t_idx, table in enumerate(doc.tables):
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            if rows_text:
                parts.append(f"--- TABLE {t_idx + 1} ---")
                parts.extend(rows_text)

        # Footers from all sections
        for i, section in enumerate(doc.sections):
            footer_text = section.footer.text.strip() if section.footer else ""
            if footer_text:
                parts.append(f"--- FOOTER (Section {i + 1}) ---\n{footer_text}")

        full_text = "\n".join(parts)
        txt_path.write_text(full_text, encoding="utf-8")

        return {
            "method": "python-docx",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "chars": len(full_text),
            "path": str(txt_path),
        }
    except Exception as e:
        return {"method": "error", "error": str(e)[:200]}


def run_docx_extract(cycle_dir: Path, dry_run: bool = False):
    db_path = cycle_dir / "inventory.db"
    if not db_path.exists():
        print("[PHASE4B] inventory.db not found", file=sys.stderr)
        if dry_run:
            print("[PHASE4B] DRY RUN: would extract DOCX files from inventory.db", file=sys.stderr)
            return
        sys.exit(1)

    if is_phase_done(cycle_dir, "phase4b"):
        print("[PHASE4B] Already complete, skipping", file=sys.stderr)
        return

    extracts_dir = cycle_dir / "extracts" / "docx"
    extracts_dir.mkdir(parents=True, exist_ok=True)

    conn = sqlite3.connect(str(db_path))
    try:
        rows = conn.execute("""
            SELECT id, file_path, sha256, size_bytes FROM files
            WHERE is_canonical = 1 AND extension IN ('.docx', '.doc')
                  AND priority IN ('HIGH', 'MEDIUM')
            ORDER BY content_score DESC, size_bytes ASC
        """).fetchall()

        start = time.time()
        extracted = 0
        errors = 0
        total_chars = 0
        log_path = cycle_dir / "extraction_log_docx.jsonl"

        print(f"[PHASE4B] Extracting {len(rows):,} DOCX files...", file=sys.stderr)

        with open(log_path, "a", encoding="utf-8") as log:
            for i, (fid, fpath, sha, size) in enumerate(rows):
                if dry_run:
                    extracted += 1
                    continue

                result = extract_docx(fpath, extracts_dir)
                if result and result["method"] != "error":
                    total_chars += result.get("chars", 0)
                    extracted += 1
                else:
                    errors += 1

                entry = {
                    "sha256": sha,
                    "source_path": fpath,
                    "result": result,
                    "extraction_ts": time.strftime("%Y-%m-%dT%H:%M:%S"),
                }
                log.write(json.dumps(entry) + "\n")

                if (i + 1) % 1000 == 0:
                    report_progress("phase4b", i + 1, len(rows))

        elapsed = time.time() - start
        print(f"[PHASE4B] Done: {extracted:,} extracted, {errors} errors, {total_chars:,} chars in {elapsed:.0f}s", file=sys.stderr)

        if not dry_run:
            stats = {"extracted": extracted, "errors": errors, "total_chars": total_chars, "elapsed": round(elapsed, 1)}
            (cycle_dir / "docx_extract_stats.json").write_text(json.dumps(stats, indent=2), encoding="utf-8")
            write_phase_checkpoint(cycle_dir, "phase4b", {"status": "done", "extracted": extracted, "elapsed": f"{elapsed:.0f}s"})
    finally:
        conn.close()


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Phase 4B: DOCX Extract")
    parser.add_argument("--cycle-ts", default=None)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    from config import CYCLE_TS
    run_docx_extract(get_cyclepack_dir(args.cycle_ts or CYCLE_TS), args.dry_run)
