"""
context_harvester.py — Multi-drive, multi-format context harvester for LitigationOS.

Orchestrates document harvesting from 6 local drives + Google Drive into
the LitigationOS context system.  100 % local-first, zero required external
deps (PyMuPDF, python-docx, pydrive2 are optional with graceful fallbacks).

Classes
-------
DriveScanner        Scans local drives for litigation-relevant files.
TextExtractor       Extracts text from PDF, DOCX, TXT, MD, HTML, CSV, JSON.
GoogleDriveHarvester  Harvests documents from Google Drive (optional).
ContentIndexer      Indexes extracted content to SQLite + FTS5.
HarvestScheduler    Manages incremental harvesting with checkpoints.
ContextHarvester    Unified façade that ties everything together.

Database
--------
Tables: harvested_content, harvested_content_fts, harvest_progress,
harvest_errors.  All live in ``litigation_context.db`` next to the repo root.
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import logging
import os
import pathlib
import re
import sqlite3
import sys
import threading
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import (
    Any,
    Callable,
    Dict,
    Generator,
    List,
    Optional,
    Set,
    Tuple,
)

# ---------------------------------------------------------------------------
# UTF-8 stdout safety (mandatory per LitigationOS rules)
# ---------------------------------------------------------------------------
if hasattr(sys.stdout, "fileno"):
    try:
        sys.stdout = open(
            sys.stdout.fileno(), mode="w", encoding="utf-8",
            errors="replace", closefd=False,
        )
        sys.stderr = open(
            sys.stderr.fileno(), mode="w", encoding="utf-8",
            errors="replace", closefd=False,
        )
    except Exception:
        pass

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logger = logging.getLogger("litigationos.context_harvester")
if not logger.handlers:
    _handler = logging.StreamHandler(sys.stderr)
    _handler.setFormatter(
        logging.Formatter("[%(levelname)s] %(name)s | %(message)s")
    )
    logger.addHandler(_handler)
    logger.setLevel(logging.INFO)

# ---------------------------------------------------------------------------
# Path constants
# ---------------------------------------------------------------------------
_REPO = pathlib.Path(__file__).resolve().parent.parent.parent
_DB_PATH = _REPO / "litigation_context.db"
_CHECKPOINT_DIR = _REPO / "temp"

# ---------------------------------------------------------------------------
# Lazy optional imports (zero required external deps)
# ---------------------------------------------------------------------------
_fitz = None
_docx_mod = None
_pydrive2_auth = None
_pydrive2_drive = None
_bs4 = None


def _get_fitz():
    global _fitz
    if _fitz is None:
        try:
            import fitz as _f  # PyMuPDF
            _fitz = _f
        except ImportError:
            try:
                import pymupdf as _f
                _fitz = _f
            except ImportError:
                pass
    return _fitz


def _get_docx():
    global _docx_mod
    if _docx_mod is None:
        try:
            from docx import Document as _D  # noqa: N811
            _docx_mod = _D
        except ImportError:
            pass
    return _docx_mod


def _get_bs4():
    global _bs4
    if _bs4 is None:
        try:
            from bs4 import BeautifulSoup as _B
            _bs4 = _B
        except ImportError:
            pass
    return _bs4


def _get_pydrive2():
    global _pydrive2_auth, _pydrive2_drive
    if _pydrive2_auth is None:
        try:
            from pydrive2.auth import GoogleAuth as _GA
            from pydrive2.drive import GoogleDrive as _GD
            _pydrive2_auth = _GA
            _pydrive2_drive = _GD
        except ImportError:
            pass
    return _pydrive2_auth, _pydrive2_drive

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
LEGAL_EXTENSIONS: Set[str] = {
    ".pdf", ".txt", ".md", ".docx", ".doc",
    ".html", ".htm", ".csv", ".json",
}

SKIP_DIRS: Set[str] = {
    "node_modules", ".git", "__pycache__", "$Recycle.Bin", "$RECYCLE.BIN",
    "Windows", "Program Files", "Program Files (x86)", "pip_cache",
    ".tox", ".venv", "venv", "env", ".mypy_cache", ".pytest_cache",
    "AppData", "ProgramData", "Recovery", "System Volume Information",
}

DEFAULT_DRIVES: List[Tuple[str, bool]] = [
    ("C:\\", True),
    ("D:\\", True),
    ("F:\\", True),
    ("G:\\", True),
    ("H:\\", True),
    ("I:\\", False),  # slow USB — disabled by default
]

MAX_FILE_SIZE: int = 100 * 1024 * 1024  # 100 MB
MAX_EXTRACT_TIMEOUT: int = 120  # seconds
DEFAULT_MAX_DEPTH: int = 8
DEFAULT_BATCH_SIZE: int = 100
SHA256_CHUNK: int = 65_536  # 64 KB

# ---------------------------------------------------------------------------
# Lane detection patterns
# ---------------------------------------------------------------------------
LANE_PATTERNS: Dict[str, List[re.Pattern]] = {
    "E": [
        re.compile(r"misconduct|JTC|judicial\s+misconduct", re.I),
        re.compile(r"McNeill.*bias|bias.*McNeill", re.I),
    ],
    "D": [
        re.compile(r"\bPPO\b|protection\s+order|5907[\-\s]?PP", re.I),
    ],
    "F": [
        re.compile(r"appellate|COA|366810|\bappeal\b|supreme\s+court", re.I),
    ],
    "C": [
        re.compile(r"convergence|multi[\-\s]?lane|cross[\-\s]?case", re.I),
    ],
    "A": [
        re.compile(r"custody|parenting|001507[\-\s]?DC|McNeill", re.I),
    ],
    "B": [
        re.compile(r"housing|shady\s+oaks|002760[\-\s]?CZ|habitability", re.I),
    ],
}

# Detection priority: E → D → F → C → A → B
LANE_PRIORITY = ["E", "D", "F", "C", "A", "B"]


def detect_lane(filepath: str, preview: str = "") -> Optional[str]:
    """Return the highest-priority lane matching *filepath* + *preview*."""
    combined = f"{filepath}\n{preview}"
    for lane_id in LANE_PRIORITY:
        for pat in LANE_PATTERNS[lane_id]:
            if pat.search(combined):
                return lane_id
    return None

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _long_path(p: str) -> str:
    """Convert to Windows extended-length path for >260 chars."""
    if sys.platform == "win32" and not p.startswith("\\\\?\\"):
        return "\\\\?\\" + os.path.abspath(p)
    return p


def compute_sha256(filepath: str, max_bytes: int = 50 * 1024 * 1024) -> str:
    """Stream SHA-256 of *filepath*, reading at most *max_bytes*."""
    h = hashlib.sha256()
    read = 0
    try:
        with open(_long_path(filepath), "rb") as fh:
            while True:
                chunk = fh.read(SHA256_CHUNK)
                if not chunk:
                    break
                h.update(chunk)
                read += len(chunk)
                if read >= max_bytes:
                    break
    except (OSError, PermissionError):
        return ""
    return h.hexdigest()


def _iso_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _file_modified_iso(filepath: str) -> str:
    try:
        ts = os.path.getmtime(_long_path(filepath))
        return datetime.fromtimestamp(ts, tz=timezone.utc).isoformat(
            timespec="seconds"
        )
    except OSError:
        return ""


# ===================================================================
#  DataClasses
# ===================================================================

@dataclass
class FileRecord:
    """Metadata about a file discovered during a drive scan."""
    path: str
    size: int
    modified: str
    ext: str
    sha256: str
    drive: str


@dataclass
class ExtractedContent:
    """Result of text extraction from a single file."""
    text: str
    page_count: int = 0
    char_count: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_method: str = "unknown"


# ===================================================================
#  DriveScanner
# ===================================================================

class DriveScanner:
    """Scans local drives for litigation-relevant files.

    Thread-safe: uses a ThreadPoolExecutor to scan multiple drives in
    parallel.  Tracks ``last_scan_time`` per drive for incremental scans
    and de-duplicates by SHA-256.
    """

    def __init__(
        self,
        drives: Optional[List[Tuple[str, bool]]] = None,
        extensions: Optional[Set[str]] = None,
        skip_dirs: Optional[Set[str]] = None,
        max_depth: int = DEFAULT_MAX_DEPTH,
        max_file_size: int = MAX_FILE_SIZE,
        max_workers: int = 4,
        progress_callback: Optional[Callable[[str, int], None]] = None,
    ) -> None:
        self.drives = drives or list(DEFAULT_DRIVES)
        self.extensions = extensions or set(LEGAL_EXTENSIONS)
        self.skip_dirs = skip_dirs or set(SKIP_DIRS)
        self.max_depth = max_depth
        self.max_file_size = max_file_size
        self.max_workers = max_workers
        self.progress_callback = progress_callback

        self._lock = threading.Lock()
        self._known_hashes: Set[str] = set()
        self._last_scan_times: Dict[str, float] = {}
        self._stats: Dict[str, int] = {
            "total_scanned": 0,
            "total_matched": 0,
            "total_skipped_ext": 0,
            "total_skipped_size": 0,
            "total_skipped_dedup": 0,
            "total_errors": 0,
        }

    # -- public --

    def load_known_hashes(self, hashes: Set[str]) -> None:
        """Pre-seed known SHA-256 hashes (from the index) for dedup."""
        with self._lock:
            self._known_hashes.update(hashes)

    def set_last_scan_time(self, drive: str, ts: float) -> None:
        with self._lock:
            self._last_scan_times[drive] = ts

    def scan_all(self) -> List[FileRecord]:
        """Scan every *enabled* drive, return combined results."""
        records: List[FileRecord] = []
        enabled = [(d, e) for d, e in self.drives if e]
        if not enabled:
            logger.warning("No drives enabled for scanning.")
            return records

        with ThreadPoolExecutor(max_workers=min(self.max_workers, len(enabled))) as pool:
            futures = {
                pool.submit(self._scan_drive, drv): drv
                for drv, _ in enabled
            }
            for fut in as_completed(futures):
                drv = futures[fut]
                try:
                    recs = fut.result()
                    records.extend(recs)
                except Exception as exc:
                    logger.error("Drive %s scan failed: %s", drv, exc)
                    with self._lock:
                        self._stats["total_errors"] += 1
        return records

    def scan_drive(self, drive_path: str) -> List[FileRecord]:
        """Scan a single drive root."""
        return self._scan_drive(drive_path)

    @property
    def stats(self) -> Dict[str, int]:
        with self._lock:
            return dict(self._stats)

    # -- internal --

    def _scan_drive(self, root: str) -> List[FileRecord]:
        records: List[FileRecord] = []
        drive_letter = pathlib.Path(root).anchor.rstrip("\\:/")
        last_ts = self._last_scan_times.get(root, 0.0)
        scanned = 0
        matched = 0
        skipped_ext = 0
        skipped_size = 0
        skipped_dup = 0
        errors = 0

        logger.info("Scanning drive %s (incremental since %.0f)", root, last_ts)

        try:
            for dirpath, dirnames, filenames in os.walk(
                _long_path(root), topdown=True, onerror=lambda e: None
            ):
                # Depth check
                rel = os.path.relpath(dirpath, root)
                depth = 0 if rel == "." else rel.count(os.sep) + 1
                if depth >= self.max_depth:
                    dirnames.clear()
                    continue

                # Skip forbidden directories
                dirnames[:] = [
                    d for d in dirnames
                    if d not in self.skip_dirs
                    and not d.startswith(".")
                ]

                for fname in filenames:
                    scanned += 1
                    ext = os.path.splitext(fname)[1].lower()
                    if ext not in self.extensions:
                        skipped_ext += 1
                        continue

                    fpath = os.path.join(dirpath, fname)
                    try:
                        stat = os.stat(_long_path(fpath))
                    except (OSError, PermissionError):
                        errors += 1
                        continue

                    if stat.st_size > self.max_file_size:
                        skipped_size += 1
                        continue

                    if stat.st_size == 0:
                        continue

                    # Incremental: skip files not modified since last scan
                    if last_ts and stat.st_mtime <= last_ts:
                        continue

                    sha = compute_sha256(fpath)
                    if not sha:
                        errors += 1
                        continue

                    with self._lock:
                        if sha in self._known_hashes:
                            skipped_dup += 1
                            continue
                        self._known_hashes.add(sha)

                    mod_iso = datetime.fromtimestamp(
                        stat.st_mtime, tz=timezone.utc
                    ).isoformat(timespec="seconds")
                    records.append(FileRecord(
                        path=fpath,
                        size=stat.st_size,
                        modified=mod_iso,
                        ext=ext,
                        sha256=sha,
                        drive=drive_letter or root,
                    ))
                    matched += 1

                    if self.progress_callback and matched % 100 == 0:
                        self.progress_callback(root, matched)

        except Exception as exc:
            logger.error("Walk error on %s: %s", root, exc)
            errors += 1

        # Roll up stats
        with self._lock:
            self._stats["total_scanned"] += scanned
            self._stats["total_matched"] += matched
            self._stats["total_skipped_ext"] += skipped_ext
            self._stats["total_skipped_size"] += skipped_size
            self._stats["total_skipped_dedup"] += skipped_dup
            self._stats["total_errors"] += errors

        logger.info(
            "Drive %s done — scanned=%d matched=%d dup=%d err=%d",
            root, scanned, matched, skipped_dup, errors,
        )
        return records


# ===================================================================
#  TextExtractor
# ===================================================================

class TextExtractor:
    """Extracts text from multiple file formats.

    Uses try/except imports for every optional library so the module
    works even with zero external packages installed.
    """

    def __init__(
        self,
        max_file_size: int = MAX_FILE_SIZE,
        timeout: int = MAX_EXTRACT_TIMEOUT,
    ) -> None:
        self.max_file_size = max_file_size
        self.timeout = timeout
        self._lock = threading.Lock()
        self._stats: Dict[str, Dict[str, int]] = {}

    # -- public --

    def extract(self, filepath: str) -> ExtractedContent:
        """Dispatch to the correct extractor based on file extension."""
        ext = os.path.splitext(filepath)[1].lower()
        dispatch = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".doc": self._extract_docx,
            ".txt": self._extract_text,
            ".md": self._extract_text,
            ".html": self._extract_html,
            ".htm": self._extract_html,
            ".csv": self._extract_csv,
            ".json": self._extract_json,
        }
        extractor = dispatch.get(ext, self._extract_text)
        try:
            result = self._with_timeout(extractor, filepath)
        except Exception as exc:
            self._record_stat(ext, "fail")
            logger.debug("Extraction failed for %s: %s", filepath, exc)
            return ExtractedContent(
                text="", extraction_method=f"error:{type(exc).__name__}",
            )
        self._record_stat(ext, "success")
        return result

    @property
    def stats(self) -> Dict[str, Dict[str, int]]:
        with self._lock:
            return {k: dict(v) for k, v in self._stats.items()}

    # -- timeout wrapper (Windows-safe, threaded) --

    def _with_timeout(
        self, func: Callable, filepath: str,
    ) -> ExtractedContent:
        result_box: List[Optional[ExtractedContent]] = [None]
        error_box: List[Optional[Exception]] = [None]

        def _worker():
            try:
                result_box[0] = func(filepath)
            except Exception as exc:
                error_box[0] = exc

        t = threading.Thread(target=_worker, daemon=True)
        t.start()
        t.join(timeout=self.timeout)

        if t.is_alive():
            raise TimeoutError(
                f"Extraction timed out after {self.timeout}s: {filepath}"
            )
        if error_box[0] is not None:
            raise error_box[0]
        return result_box[0] or ExtractedContent(text="", extraction_method="empty")

    # -- extractors --

    def _extract_pdf(self, filepath: str) -> ExtractedContent:
        fitz = _get_fitz()
        if fitz is not None:
            return self._extract_pdf_fitz(filepath, fitz)
        # Fallback: try PyPDF2
        try:
            from PyPDF2 import PdfReader  # type: ignore[import-untyped]
            return self._extract_pdf_pypdf2(filepath, PdfReader)
        except ImportError:
            pass
        # Last resort: raw bytes heuristic
        return self._extract_pdf_raw(filepath)

    def _extract_pdf_fitz(self, filepath: str, fitz) -> ExtractedContent:
        pages_text: List[str] = []
        errors: List[str] = []
        try:
            doc = fitz.open(_long_path(filepath))
        except Exception as exc:
            return ExtractedContent(
                text="", extraction_method=f"fitz_open_error:{exc}",
            )
        try:
            total_pages = len(doc)
            for i in range(total_pages):
                try:
                    page = doc[i]
                    text = page.get_text("text") or ""
                    pages_text.append(text)
                except Exception as page_exc:
                    errors.append(f"page {i}: {page_exc}")
                    pages_text.append("")
        finally:
            doc.close()

        full = "\n".join(pages_text)
        return ExtractedContent(
            text=full,
            page_count=total_pages,
            char_count=len(full),
            metadata={"errors": errors} if errors else {},
            extraction_method="fitz",
        )

    def _extract_pdf_pypdf2(self, filepath: str, PdfReader) -> ExtractedContent:
        pages_text: List[str] = []
        try:
            reader = PdfReader(_long_path(filepath))
            for page in reader.pages:
                try:
                    text = page.extract_text() or ""
                    pages_text.append(text)
                except Exception:
                    pages_text.append("")
        except Exception as exc:
            return ExtractedContent(
                text="", extraction_method=f"pypdf2_error:{exc}",
            )
        full = "\n".join(pages_text)
        return ExtractedContent(
            text=full,
            page_count=len(pages_text),
            char_count=len(full),
            extraction_method="pypdf2",
        )

    def _extract_pdf_raw(self, filepath: str) -> ExtractedContent:
        """Heuristic byte-level text extraction for PDFs when no library available."""
        try:
            with open(_long_path(filepath), "rb") as fh:
                raw = fh.read(self.max_file_size)
        except (OSError, PermissionError) as exc:
            return ExtractedContent(
                text="", extraction_method=f"raw_error:{exc}",
            )
        # Extract ASCII runs between stream markers
        parts: List[str] = []
        for match in re.finditer(rb"BT\s(.*?)\sET", raw, re.DOTALL):
            segment = match.group(1)
            # Extract parenthesized strings
            for txt_match in re.finditer(rb"\(([^)]*)\)", segment):
                decoded = txt_match.group(1).decode("latin-1", errors="replace")
                if decoded.strip():
                    parts.append(decoded.strip())
        full = " ".join(parts)
        return ExtractedContent(
            text=full, char_count=len(full), extraction_method="raw_bytes",
        )

    def _extract_docx(self, filepath: str) -> ExtractedContent:
        DocxDocument = _get_docx()
        if DocxDocument is not None:
            return self._extract_docx_lib(filepath, DocxDocument)
        # Fallback: zipfile XML extraction
        return self._extract_docx_zip(filepath)

    def _extract_docx_lib(self, filepath: str, DocxDocument) -> ExtractedContent:
        parts: List[str] = []
        table_count = 0
        try:
            doc = DocxDocument(_long_path(filepath))
        except Exception as exc:
            return ExtractedContent(
                text="", extraction_method=f"docx_open_error:{exc}",
            )
        # Headers
        for i, section in enumerate(doc.sections, 1):
            try:
                header = section.header
                if header and header.paragraphs:
                    htext = "\n".join(p.text for p in header.paragraphs if p.text.strip())
                    if htext.strip():
                        parts.append(f"[--- HEADER (Section {i}) ---]\n{htext}")
            except Exception:
                pass
        # Body
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Tables
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        # Footers
        for i, section in enumerate(doc.sections, 1):
            try:
                footer = section.footer
                if footer and footer.paragraphs:
                    ftext = "\n".join(p.text for p in footer.paragraphs if p.text.strip())
                    if ftext.strip():
                        parts.append(f"[--- FOOTER (Section {i}) ---]\n{ftext}")
            except Exception:
                pass
        full = "\n".join(parts)
        return ExtractedContent(
            text=full,
            page_count=0,
            char_count=len(full),
            metadata={"tables": table_count, "paragraphs": len(doc.paragraphs)},
            extraction_method="python-docx",
        )

    def _extract_docx_zip(self, filepath: str) -> ExtractedContent:
        """Fallback DOCX extractor using zipfile + XML parsing."""
        try:
            with zipfile.ZipFile(_long_path(filepath), "r") as zf:
                if "word/document.xml" not in zf.namelist():
                    return ExtractedContent(text="", extraction_method="zip_no_doc_xml")
                xml_bytes = zf.read("word/document.xml")
        except Exception as exc:
            return ExtractedContent(
                text="", extraction_method=f"zip_error:{exc}",
            )
        # Strip XML tags, keep text
        text = re.sub(rb"<[^>]+>", b" ", xml_bytes).decode("utf-8", errors="replace")
        text = re.sub(r"\s+", " ", text).strip()
        return ExtractedContent(
            text=text, char_count=len(text), extraction_method="zip_xml",
        )

    def _extract_text(self, filepath: str) -> ExtractedContent:
        """UTF-8 plain text read with fallback to latin-1."""
        try:
            with open(_long_path(filepath), "r", encoding="utf-8", errors="replace") as fh:
                text = fh.read(self.max_file_size)
        except UnicodeDecodeError:
            try:
                with open(_long_path(filepath), "r", encoding="latin-1") as fh:
                    text = fh.read(self.max_file_size)
            except (OSError, PermissionError) as exc:
                return ExtractedContent(
                    text="", extraction_method=f"text_error:{exc}",
                )
        except (OSError, PermissionError) as exc:
            return ExtractedContent(
                text="", extraction_method=f"text_error:{exc}",
            )
        return ExtractedContent(
            text=text, char_count=len(text), extraction_method="utf8_read",
        )

    def _extract_html(self, filepath: str) -> ExtractedContent:
        raw_result = self._extract_text(filepath)
        if not raw_result.text:
            return raw_result
        BeautifulSoup = _get_bs4()
        if BeautifulSoup is not None:
            try:
                soup = BeautifulSoup(raw_result.text, "html.parser")
                text = soup.get_text(separator="\n", strip=True)
                return ExtractedContent(
                    text=text, char_count=len(text), extraction_method="bs4_html",
                )
            except Exception:
                pass
        # Fallback: strip tags with regex
        text = re.sub(r"<[^>]+>", " ", raw_result.text)
        text = re.sub(r"\s+", " ", text).strip()
        return ExtractedContent(
            text=text, char_count=len(text), extraction_method="regex_html",
        )

    def _extract_csv(self, filepath: str) -> ExtractedContent:
        rows: List[str] = []
        try:
            with open(
                _long_path(filepath), "r", encoding="utf-8", errors="replace",
                newline="",
            ) as fh:
                reader = csv.reader(io.StringIO(fh.read(self.max_file_size)))
                for row in reader:
                    rows.append(" | ".join(row))
                    if len(rows) >= 50_000:
                        break
        except (OSError, csv.Error) as exc:
            return ExtractedContent(
                text="", extraction_method=f"csv_error:{exc}",
            )
        full = "\n".join(rows)
        return ExtractedContent(
            text=full,
            char_count=len(full),
            metadata={"rows": len(rows)},
            extraction_method="csv_reader",
        )

    def _extract_json(self, filepath: str) -> ExtractedContent:
        try:
            with open(
                _long_path(filepath), "r", encoding="utf-8", errors="replace",
            ) as fh:
                raw = fh.read(self.max_file_size)
        except (OSError, PermissionError) as exc:
            return ExtractedContent(
                text="", extraction_method=f"json_error:{exc}",
            )
        try:
            data = json.loads(raw)
            text = json.dumps(data, indent=2, ensure_ascii=False, default=str)
        except (json.JSONDecodeError, ValueError):
            text = raw
        return ExtractedContent(
            text=text, char_count=len(text), extraction_method="json_load",
        )

    # -- stats helpers --

    def _record_stat(self, ext: str, outcome: str) -> None:
        with self._lock:
            if ext not in self._stats:
                self._stats[ext] = {"success": 0, "fail": 0, "skip": 0}
            self._stats[ext][outcome] = self._stats[ext].get(outcome, 0) + 1


# ===================================================================
#  GoogleDriveHarvester
# ===================================================================

class GoogleDriveHarvester:
    """Harvests documents from Google Drive.

    Uses pydrive2 if available, else google-api-python-client, else
    gracefully skips.  Credentials are read from environment variables
    — never hardcoded.
    """

    SUPPORTED_MIMES: Dict[str, str] = {
        "application/pdf": ".pdf",
        "text/plain": ".txt",
        "text/markdown": ".md",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/html": ".html",
        "text/csv": ".csv",
        "application/json": ".json",
    }

    EXPORT_MIMES: Dict[str, Tuple[str, str]] = {
        "application/vnd.google-apps.document": ("text/plain", ".txt"),
        "application/vnd.google-apps.spreadsheet": ("text/csv", ".csv"),
    }

    def __init__(
        self,
        download_dir: Optional[str] = None,
        rate_limit_delay: float = 0.5,
        last_sync_ts: Optional[str] = None,
    ) -> None:
        self.download_dir = download_dir or str(_REPO / "temp" / "gdrive_mirror")
        self.rate_limit_delay = rate_limit_delay
        self.last_sync_ts = last_sync_ts
        self._lock = threading.Lock()
        self._stats: Dict[str, int] = {
            "listed": 0, "downloaded": 0, "exported": 0,
            "skipped": 0, "errors": 0,
        }
        self._available = False
        self._drive = None

    def authenticate(self) -> bool:
        """Authenticate to Google Drive.  Returns True on success."""
        client_id = os.environ.get("GOOGLE_DRIVE_CLIENT_ID", "")
        if not client_id:
            logger.warning(
                "GOOGLE_DRIVE_CLIENT_ID not set — Google Drive harvesting disabled."
            )
            return False

        GA, GD = _get_pydrive2()
        if GA is not None and GD is not None:
            return self._auth_pydrive2(GA, GD, client_id)

        # Fallback: google-api-python-client
        try:
            return self._auth_google_api_client(client_id)
        except ImportError:
            logger.warning(
                "Neither pydrive2 nor google-api-python-client installed — "
                "Google Drive harvesting disabled."
            )
            return False

    def list_files(self) -> List[Dict[str, Any]]:
        """List harvestable files from Google Drive."""
        if not self._available or self._drive is None:
            return []
        results: List[Dict[str, Any]] = []
        try:
            mime_filter = " or ".join(
                f"mimeType='{m}'" for m in
                list(self.SUPPORTED_MIMES.keys()) + list(self.EXPORT_MIMES.keys())
            )
            query = f"trashed=false and ({mime_filter})"
            if self.last_sync_ts:
                query += f" and modifiedDate > '{self.last_sync_ts}'"

            file_list = self._drive.ListFile({"q": query}).GetList()
            for f in file_list:
                results.append({
                    "id": f["id"],
                    "title": f["title"],
                    "mimeType": f.get("mimeType", ""),
                    "modifiedDate": f.get("modifiedDate", ""),
                    "fileSize": int(f.get("fileSize", 0) or 0),
                })
            with self._lock:
                self._stats["listed"] = len(results)
        except Exception as exc:
            logger.error("Google Drive list failed: %s", exc)
            with self._lock:
                self._stats["errors"] += 1
        return results

    def download_files(
        self, file_list: Optional[List[Dict[str, Any]]] = None,
    ) -> List[str]:
        """Download (or export) files to local temp directory.

        Returns list of local file paths.
        """
        if file_list is None:
            file_list = self.list_files()
        if not self._available or self._drive is None:
            return []

        os.makedirs(self.download_dir, exist_ok=True)
        local_paths: List[str] = []

        for fmeta in file_list:
            mime = fmeta.get("mimeType", "")
            title = fmeta.get("title", "untitled")
            fid = fmeta["id"]

            try:
                gf = self._drive.CreateFile({"id": fid})

                if mime in self.EXPORT_MIMES:
                    export_mime, ext = self.EXPORT_MIMES[mime]
                    safe_name = re.sub(r'[<>:"/\\|?*]', "_", title) + ext
                    dest = os.path.join(self.download_dir, safe_name)
                    gf.GetContentFile(dest, mimetype=export_mime)
                    with self._lock:
                        self._stats["exported"] += 1
                elif mime in self.SUPPORTED_MIMES:
                    ext = self.SUPPORTED_MIMES[mime]
                    safe_name = re.sub(r'[<>:"/\\|?*]', "_", title)
                    if not safe_name.lower().endswith(ext):
                        safe_name += ext
                    dest = os.path.join(self.download_dir, safe_name)
                    gf.GetContentFile(dest)
                    with self._lock:
                        self._stats["downloaded"] += 1
                else:
                    with self._lock:
                        self._stats["skipped"] += 1
                    continue

                local_paths.append(dest)

            except Exception as exc:
                logger.warning("Download failed for %s (%s): %s", title, fid, exc)
                with self._lock:
                    self._stats["errors"] += 1

            time.sleep(self.rate_limit_delay)

        return local_paths

    @property
    def stats(self) -> Dict[str, int]:
        with self._lock:
            return dict(self._stats)

    # -- auth implementations --

    def _auth_pydrive2(self, GA, GD, client_id: str) -> bool:
        try:
            settings = {
                "client_config_backend": "settings",
                "client_config": {
                    "client_id": client_id,
                    "client_secret": os.environ.get(
                        "GOOGLE_DRIVE_CLIENT_SECRET", ""
                    ),
                },
                "save_credentials": True,
                "save_credentials_backend": "file",
                "save_credentials_file": str(
                    _REPO / "temp" / "gdrive_credentials.json"
                ),
            }
            gauth = GA(settings=settings)
            gauth.LocalWebserverAuth()
            self._drive = GD(gauth)
            self._available = True
            logger.info("Google Drive authenticated via pydrive2.")
            return True
        except Exception as exc:
            logger.warning("pydrive2 auth failed: %s", exc)
            return False

    def _auth_google_api_client(self, client_id: str) -> bool:
        """Attempt auth via google-api-python-client + google-auth-oauthlib."""
        try:
            from google.oauth2.credentials import Credentials  # type: ignore
            from googleapiclient.discovery import build  # type: ignore
        except ImportError:
            raise ImportError("google-api-python-client not installed")

        creds_path = _REPO / "temp" / "gdrive_token.json"
        creds = None
        if creds_path.exists():
            try:
                creds = Credentials.from_authorized_user_file(str(creds_path))
            except Exception:
                pass

        if creds is None or not creds.valid:
            logger.warning(
                "Google API client credentials not found or expired at %s. "
                "Cannot authenticate non-interactively.", creds_path,
            )
            return False

        try:
            service = build("drive", "v3", credentials=creds)
            # Wrap service in a pydrive-like interface (limited)
            self._drive = _GoogleApiClientWrapper(service)
            self._available = True
            logger.info("Google Drive authenticated via google-api-python-client.")
            return True
        except Exception as exc:
            logger.warning("google-api-python-client auth failed: %s", exc)
            return False


class _GoogleApiClientWrapper:
    """Minimal wrapper around google-api-python-client to mimic pydrive2 API."""

    def __init__(self, service) -> None:
        self._svc = service

    def ListFile(self, params: Dict[str, Any]):  # noqa: N802
        return _GFileList(self._svc, params.get("q", ""))

    def CreateFile(self, params: Dict[str, Any]):  # noqa: N802
        return _GFile(self._svc, params.get("id", ""))


class _GFileList:
    def __init__(self, svc, query: str) -> None:
        self._svc = svc
        self._query = query

    def GetList(self) -> List[Dict[str, Any]]:  # noqa: N802
        results: List[Dict[str, Any]] = []
        page_token = None
        while True:
            resp = (
                self._svc.files()
                .list(
                    q=self._query,
                    pageSize=100,
                    fields="nextPageToken, files(id, name, mimeType, modifiedTime, size)",
                    pageToken=page_token,
                )
                .execute()
            )
            for f in resp.get("files", []):
                results.append({
                    "id": f["id"],
                    "title": f.get("name", ""),
                    "mimeType": f.get("mimeType", ""),
                    "modifiedDate": f.get("modifiedTime", ""),
                    "fileSize": int(f.get("size", 0) or 0),
                })
            page_token = resp.get("nextPageToken")
            if not page_token:
                break
        return results


class _GFile:
    def __init__(self, svc, file_id: str) -> None:
        self._svc = svc
        self._id = file_id

    def GetContentFile(self, dest: str, mimetype: Optional[str] = None) -> None:  # noqa: N802
        if mimetype:
            request = self._svc.files().export_media(
                fileId=self._id, mimeType=mimetype,
            )
        else:
            request = self._svc.files().get_media(fileId=self._id)
        content = request.execute()
        if isinstance(content, bytes):
            with open(dest, "wb") as fh:
                fh.write(content)
        else:
            with open(dest, "w", encoding="utf-8") as fh:
                fh.write(str(content))


# ===================================================================
#  ContentIndexer
# ===================================================================

_SCHEMA_DDL = """
CREATE TABLE IF NOT EXISTS harvested_content (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    file_path TEXT NOT NULL,
    drive TEXT NOT NULL,
    sha256 TEXT NOT NULL,
    content_text TEXT,
    content_preview TEXT,
    file_type TEXT NOT NULL,
    file_size INTEGER,
    page_count INTEGER DEFAULT 0,
    char_count INTEGER DEFAULT 0,
    lane_detected TEXT,
    extraction_method TEXT,
    source_type TEXT DEFAULT 'local',
    indexed_at TEXT DEFAULT (datetime('now')),
    last_modified TEXT,
    UNIQUE(sha256)
);

CREATE VIRTUAL TABLE IF NOT EXISTS harvested_content_fts USING fts5(
    content_text,
    content_preview,
    file_path,
    content='harvested_content',
    content_rowid='id',
    tokenize='porter unicode61'
);

CREATE TRIGGER IF NOT EXISTS harvested_content_ai AFTER INSERT ON harvested_content BEGIN
    INSERT INTO harvested_content_fts(rowid, content_text, content_preview, file_path)
    VALUES (new.id, new.content_text, new.content_preview, new.file_path);
END;

CREATE TRIGGER IF NOT EXISTS harvested_content_ad AFTER DELETE ON harvested_content BEGIN
    INSERT INTO harvested_content_fts(harvested_content_fts, rowid, content_text, content_preview, file_path)
    VALUES ('delete', old.id, old.content_text, old.content_preview, old.file_path);
END;

CREATE TRIGGER IF NOT EXISTS harvested_content_au AFTER UPDATE ON harvested_content BEGIN
    INSERT INTO harvested_content_fts(harvested_content_fts, rowid, content_text, content_preview, file_path)
    VALUES ('delete', old.id, old.content_text, old.content_preview, old.file_path);
    INSERT INTO harvested_content_fts(rowid, content_text, content_preview, file_path)
    VALUES (new.id, new.content_text, new.content_preview, new.file_path);
END;

CREATE TABLE IF NOT EXISTS harvest_progress (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    drive TEXT NOT NULL,
    last_scan_path TEXT,
    files_scanned INTEGER DEFAULT 0,
    files_indexed INTEGER DEFAULT 0,
    files_skipped INTEGER DEFAULT 0,
    files_errored INTEGER DEFAULT 0,
    started_at TEXT DEFAULT (datetime('now')),
    completed_at TEXT,
    status TEXT DEFAULT 'in_progress'
);

CREATE TABLE IF NOT EXISTS harvest_errors (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    file_path TEXT NOT NULL,
    error_type TEXT,
    error_message TEXT,
    occurred_at TEXT DEFAULT (datetime('now'))
);

CREATE INDEX IF NOT EXISTS idx_harvested_sha256
    ON harvested_content(sha256);
CREATE INDEX IF NOT EXISTS idx_harvested_drive_type
    ON harvested_content(drive, file_type);
CREATE INDEX IF NOT EXISTS idx_harvested_lane
    ON harvested_content(lane_detected, file_type);
CREATE INDEX IF NOT EXISTS idx_harvest_progress_drive
    ON harvest_progress(drive, status);
"""


class ContentIndexer:
    """Indexes extracted content to SQLite + FTS5 for fast retrieval.

    Uses WAL mode, batch inserts via ``executemany``, composite indexes,
    and consolidated ``COUNT(*)`` queries per LitigationOS SQL conventions.
    """

    def __init__(self, db_path: Optional[str] = None) -> None:
        self.db_path = db_path or str(_DB_PATH)
        self._lock = threading.Lock()
        self._local = threading.local()
        self._ensure_schema()

    # -- connection --

    def _get_conn(self) -> sqlite3.Connection:
        conn = getattr(self._local, "conn", None)
        if conn is None:
            conn = sqlite3.connect(self.db_path, timeout=120)
            conn.execute("PRAGMA journal_mode=WAL")
            conn.execute("PRAGMA busy_timeout=60000")
            conn.execute("PRAGMA cache_size=-32000")
            conn.execute("PRAGMA temp_store=MEMORY")
            conn.execute("PRAGMA synchronous=NORMAL")
            conn.row_factory = sqlite3.Row
            self._local.conn = conn
        return conn

    def _ensure_schema(self) -> None:
        conn = self._get_conn()
        conn.executescript(_SCHEMA_DDL)
        conn.commit()

    # -- public --

    def get_known_hashes(self) -> Set[str]:
        """Return all SHA-256 hashes already in the index."""
        conn = self._get_conn()
        rows = conn.execute("SELECT sha256 FROM harvested_content").fetchall()
        return {r[0] for r in rows}

    def index_batch(
        self,
        records: List[FileRecord],
        contents: List[ExtractedContent],
        source_type: str = "local",
    ) -> int:
        """Insert a batch of file records + extracted contents.

        Returns the number of rows actually inserted (skips duplicates).
        """
        if not records:
            return 0
        conn = self._get_conn()
        rows: List[Tuple] = []
        for rec, content in zip(records, contents):
            preview = content.text[:500] if content.text else ""
            lane = detect_lane(rec.path, preview)
            rows.append((
                rec.path,
                rec.drive,
                rec.sha256,
                content.text,
                preview,
                rec.ext,
                rec.size,
                content.page_count,
                content.char_count,
                lane,
                content.extraction_method,
                source_type,
                _iso_now(),
                rec.modified,
            ))

        inserted = 0
        with self._lock:
            try:
                conn.executemany(
                    """INSERT OR IGNORE INTO harvested_content
                       (file_path, drive, sha256, content_text, content_preview,
                        file_type, file_size, page_count, char_count,
                        lane_detected, extraction_method, source_type,
                        indexed_at, last_modified)
                       VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    rows,
                )
                inserted = conn.total_changes
                conn.commit()
            except sqlite3.Error as exc:
                logger.error("Batch insert error: %s", exc)
                conn.rollback()
        return inserted

    def index_single(
        self,
        rec: FileRecord,
        content: ExtractedContent,
        source_type: str = "local",
    ) -> bool:
        """Insert a single record. Returns True if inserted, False if dup."""
        conn = self._get_conn()
        preview = content.text[:500] if content.text else ""
        lane = detect_lane(rec.path, preview)
        with self._lock:
            try:
                conn.execute(
                    """INSERT OR IGNORE INTO harvested_content
                       (file_path, drive, sha256, content_text, content_preview,
                        file_type, file_size, page_count, char_count,
                        lane_detected, extraction_method, source_type,
                        indexed_at, last_modified)
                       VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (
                        rec.path, rec.drive, rec.sha256, content.text, preview,
                        rec.ext, rec.size, content.page_count, content.char_count,
                        lane, content.extraction_method, source_type,
                        _iso_now(), rec.modified,
                    ),
                )
                conn.commit()
                return conn.total_changes > 0
            except sqlite3.Error as exc:
                logger.error("Insert error for %s: %s", rec.path, exc)
                conn.rollback()
                return False

    def search(
        self,
        query: str,
        lane: Optional[str] = None,
        file_type: Optional[str] = None,
        drive: Optional[str] = None,
        limit: int = 50,
    ) -> List[Dict[str, Any]]:
        """FTS5 search with optional lane/type/drive filters."""
        conn = self._get_conn()
        # Build WHERE clauses for the joined content table
        filters: List[str] = []
        params: List[Any] = []

        if lane:
            filters.append("hc.lane_detected = ?")
            params.append(lane)
        if file_type:
            if not file_type.startswith("."):
                file_type = "." + file_type
            filters.append("hc.file_type = ?")
            params.append(file_type)
        if drive:
            filters.append("hc.drive = ?")
            params.append(drive)

        where_clause = ""
        if filters:
            where_clause = "AND " + " AND ".join(filters)

        sql = f"""
            SELECT hc.id, hc.file_path, hc.drive, hc.sha256,
                   hc.content_preview, hc.file_type, hc.file_size,
                   hc.page_count, hc.char_count, hc.lane_detected,
                   hc.extraction_method, hc.source_type,
                   hc.indexed_at, hc.last_modified,
                   rank
            FROM harvested_content_fts fts
            JOIN harvested_content hc ON hc.id = fts.rowid
            WHERE harvested_content_fts MATCH ?
            {where_clause}
            ORDER BY rank
            LIMIT ?
        """
        params_final = [query] + params + [limit]
        try:
            rows = conn.execute(sql, params_final).fetchall()
            return [dict(r) for r in rows]
        except sqlite3.Error as exc:
            logger.error("Search error: %s", exc)
            return []

    def get_stats(self) -> Dict[str, Any]:
        """Consolidated stats via subquery pattern."""
        conn = self._get_conn()
        try:
            row = conn.execute("""
                SELECT
                    (SELECT COUNT(*) FROM harvested_content) AS total,
                    (SELECT COUNT(DISTINCT drive) FROM harvested_content) AS drives,
                    (SELECT COUNT(DISTINCT lane_detected) FROM harvested_content
                     WHERE lane_detected IS NOT NULL) AS lanes,
                    (SELECT SUM(file_size) FROM harvested_content) AS total_bytes,
                    (SELECT SUM(char_count) FROM harvested_content) AS total_chars
            """).fetchone()
            per_drive = conn.execute("""
                SELECT drive, COUNT(*) AS cnt, SUM(file_size) AS bytes
                FROM harvested_content GROUP BY drive
            """).fetchall()
            per_type = conn.execute("""
                SELECT file_type, COUNT(*) AS cnt
                FROM harvested_content GROUP BY file_type
            """).fetchall()
            per_lane = conn.execute("""
                SELECT lane_detected, COUNT(*) AS cnt
                FROM harvested_content
                WHERE lane_detected IS NOT NULL
                GROUP BY lane_detected
            """).fetchall()
        except sqlite3.Error as exc:
            logger.error("Stats query error: %s", exc)
            return {"error": str(exc)}

        return {
            "total_indexed": row["total"] or 0,
            "unique_drives": row["drives"] or 0,
            "unique_lanes": row["lanes"] or 0,
            "total_bytes": row["total_bytes"] or 0,
            "total_chars": row["total_chars"] or 0,
            "by_drive": {r["drive"]: {"count": r["cnt"], "bytes": r["bytes"] or 0} for r in per_drive},
            "by_type": {r["file_type"]: r["cnt"] for r in per_type},
            "by_lane": {r["lane_detected"]: r["cnt"] for r in per_lane},
        }

    def log_progress(
        self,
        drive: str,
        last_path: str,
        scanned: int,
        indexed: int,
        skipped: int,
        errored: int,
        status: str = "in_progress",
    ) -> int:
        """Insert or update a harvest_progress record. Returns row id."""
        conn = self._get_conn()
        with self._lock:
            completed = _iso_now() if status == "completed" else None
            cur = conn.execute(
                """INSERT INTO harvest_progress
                   (drive, last_scan_path, files_scanned, files_indexed,
                    files_skipped, files_errored, completed_at, status)
                   VALUES (?,?,?,?,?,?,?,?)""",
                (drive, last_path, scanned, indexed, skipped, errored,
                 completed, status),
            )
            conn.commit()
            return cur.lastrowid or 0

    def log_error(self, file_path: str, error_type: str, error_msg: str) -> None:
        """Record an extraction/indexing error."""
        conn = self._get_conn()
        with self._lock:
            try:
                conn.execute(
                    "INSERT INTO harvest_errors (file_path, error_type, error_message) "
                    "VALUES (?,?,?)",
                    (file_path, error_type, error_msg[:2000]),
                )
                conn.commit()
            except sqlite3.Error:
                pass

    def close(self) -> None:
        conn = getattr(self._local, "conn", None)
        if conn is not None:
            try:
                conn.close()
            except Exception:
                pass
            self._local.conn = None


# ===================================================================
#  HarvestScheduler
# ===================================================================

class HarvestScheduler:
    """Manages incremental harvesting with checkpoints and scheduling.

    Drives are processed in priority order (C first, I last).
    Checkpoints are saved after every batch so the harvester can resume
    after a crash.
    """

    DRIVE_PRIORITY = ["C", "D", "F", "G", "H", "I"]

    def __init__(
        self,
        scanner: DriveScanner,
        extractor: TextExtractor,
        indexer: ContentIndexer,
        batch_size: int = DEFAULT_BATCH_SIZE,
        rate_limit: float = 0.0,
        gdrive_sync_hours: int = 24,
        checkpoint_path: Optional[str] = None,
        progress_callback: Optional[Callable[[str, int, int], None]] = None,
    ) -> None:
        self.scanner = scanner
        self.extractor = extractor
        self.indexer = indexer
        self.batch_size = batch_size
        self.rate_limit = rate_limit
        self.gdrive_sync_hours = gdrive_sync_hours
        self.checkpoint_path = checkpoint_path or str(
            _CHECKPOINT_DIR / "harvest_checkpoint.json"
        )
        self.progress_callback = progress_callback

        self._lock = threading.Lock()
        self._checkpoint: Dict[str, Any] = self._load_checkpoint()
        self._running = False
        self._stop_event = threading.Event()

    # -- checkpoint persistence --

    def _load_checkpoint(self) -> Dict[str, Any]:
        try:
            with open(self.checkpoint_path, "r", encoding="utf-8") as fh:
                return json.load(fh)
        except (OSError, json.JSONDecodeError):
            return {
                "last_scan_times": {},
                "last_gdrive_sync": None,
                "drive_stats": {},
            }

    def _save_checkpoint(self) -> None:
        os.makedirs(os.path.dirname(self.checkpoint_path), exist_ok=True)
        try:
            with open(self.checkpoint_path, "w", encoding="utf-8") as fh:
                json.dump(self._checkpoint, fh, indent=2, default=str)
        except OSError as exc:
            logger.error("Checkpoint save failed: %s", exc)

    # -- public --

    def run_incremental(self) -> Dict[str, Any]:
        """Run one incremental harvest cycle across all enabled drives.

        Returns summary dict with per-drive stats.
        """
        self._running = True
        self._stop_event.clear()
        summary: Dict[str, Any] = {"drives": {}, "total_indexed": 0, "errors": 0}

        # Load known hashes for dedup
        try:
            known = self.indexer.get_known_hashes()
            self.scanner.load_known_hashes(known)
        except Exception as exc:
            logger.error("Failed to load known hashes: %s", exc)

        # Restore last scan times from checkpoint
        for drv, ts in self._checkpoint.get("last_scan_times", {}).items():
            self.scanner.set_last_scan_time(drv, float(ts))

        # Process drives in priority order
        for drive_letter in self.DRIVE_PRIORITY:
            if self._stop_event.is_set():
                break

            drive_path = f"{drive_letter}:\\"
            # Check if drive is enabled in scanner config
            enabled = any(
                d.upper().startswith(drive_letter.upper()) and e
                for d, e in self.scanner.drives
            )
            if not enabled:
                continue

            if not os.path.isdir(drive_path):
                logger.info("Drive %s not accessible, skipping.", drive_path)
                continue

            drive_summary = self._harvest_drive(drive_path, drive_letter)
            summary["drives"][drive_letter] = drive_summary
            summary["total_indexed"] += drive_summary.get("indexed", 0)
            summary["errors"] += drive_summary.get("errors", 0)

            # Checkpoint after each drive
            self._checkpoint["last_scan_times"][drive_path] = time.time()
            self._save_checkpoint()

        self._running = False
        return summary

    def stop(self) -> None:
        """Signal the scheduler to stop after the current batch."""
        self._stop_event.set()

    @property
    def is_running(self) -> bool:
        return self._running

    def should_sync_gdrive(self) -> bool:
        """Check if enough time has elapsed since last Google Drive sync."""
        last = self._checkpoint.get("last_gdrive_sync")
        if last is None:
            return True
        try:
            last_dt = datetime.fromisoformat(str(last))
            elapsed = (datetime.now(timezone.utc) - last_dt).total_seconds()
            return elapsed >= self.gdrive_sync_hours * 3600
        except (ValueError, TypeError):
            return True

    def mark_gdrive_synced(self) -> None:
        self._checkpoint["last_gdrive_sync"] = _iso_now()
        self._save_checkpoint()

    # -- internal --

    def _harvest_drive(
        self, drive_path: str, drive_letter: str,
    ) -> Dict[str, int]:
        """Scan one drive, extract, index, with batch checkpoints."""
        logger.info("Harvesting drive %s ...", drive_path)
        stats = {"scanned": 0, "indexed": 0, "skipped": 0, "errors": 0}

        try:
            records = self.scanner.scan_drive(drive_path)
        except Exception as exc:
            logger.error("Scan failed for %s: %s", drive_path, exc)
            stats["errors"] += 1
            return stats

        stats["scanned"] = len(records)
        logger.info("Drive %s: %d new files to process.", drive_letter, len(records))

        # Process in batches
        for batch_start in range(0, len(records), self.batch_size):
            if self._stop_event.is_set():
                break

            batch = records[batch_start: batch_start + self.batch_size]
            contents: List[ExtractedContent] = []
            valid_records: List[FileRecord] = []

            for rec in batch:
                if self._stop_event.is_set():
                    break
                try:
                    content = self.extractor.extract(rec.path)
                    if content.text:
                        contents.append(content)
                        valid_records.append(rec)
                    else:
                        stats["skipped"] += 1
                except Exception as exc:
                    stats["errors"] += 1
                    self.indexer.log_error(
                        rec.path, type(exc).__name__, str(exc)[:500],
                    )

                if self.rate_limit > 0:
                    time.sleep(1.0 / self.rate_limit)

            # Batch index
            if valid_records:
                try:
                    self.indexer.index_batch(valid_records, contents)
                    stats["indexed"] += len(valid_records)
                except Exception as exc:
                    logger.error("Batch index error: %s", exc)
                    stats["errors"] += len(valid_records)

            # Progress checkpoint
            last_path = batch[-1].path if batch else ""
            self.indexer.log_progress(
                drive_letter, last_path,
                stats["scanned"], stats["indexed"],
                stats["skipped"], stats["errors"],
            )

            if self.progress_callback:
                self.progress_callback(
                    drive_letter, stats["indexed"], stats["scanned"],
                )

        # Final progress entry
        self.indexer.log_progress(
            drive_letter, "",
            stats["scanned"], stats["indexed"],
            stats["skipped"], stats["errors"],
            status="completed",
        )

        self._checkpoint.setdefault("drive_stats", {})[drive_letter] = stats
        self._save_checkpoint()

        logger.info(
            "Drive %s complete — indexed=%d skipped=%d errors=%d",
            drive_letter, stats["indexed"], stats["skipped"], stats["errors"],
        )
        return stats


# ===================================================================
#  ContextHarvester  (main façade)
# ===================================================================

class ContextHarvester:
    """Unified multi-drive, multi-format context harvester.

    Wraps DriveScanner, TextExtractor, GoogleDriveHarvester,
    ContentIndexer, and HarvestScheduler behind a single API.

    Usage
    -----
    >>> harvester = ContextHarvester()
    >>> harvester.harvest_all()          # full scan
    >>> results = harvester.search("custody motion MCR 2.003")
    >>> stats = harvester.get_stats()
    """

    def __init__(
        self,
        db_path: Optional[str] = None,
        drives: Optional[List[Tuple[str, bool]]] = None,
        extensions: Optional[Set[str]] = None,
        max_depth: int = DEFAULT_MAX_DEPTH,
        max_file_size: int = MAX_FILE_SIZE,
        extract_timeout: int = MAX_EXTRACT_TIMEOUT,
        batch_size: int = DEFAULT_BATCH_SIZE,
        rate_limit: float = 0.0,
        max_workers: int = 4,
        gdrive_sync_hours: int = 24,
        progress_callback: Optional[Callable] = None,
    ) -> None:
        self.indexer = ContentIndexer(db_path=db_path)
        self.scanner = DriveScanner(
            drives=drives,
            extensions=extensions,
            max_depth=max_depth,
            max_file_size=max_file_size,
            max_workers=max_workers,
            progress_callback=self._scan_progress if progress_callback is None else progress_callback,
        )
        self.extractor = TextExtractor(
            max_file_size=max_file_size,
            timeout=extract_timeout,
        )
        self.gdrive = GoogleDriveHarvester()
        self.scheduler = HarvestScheduler(
            scanner=self.scanner,
            extractor=self.extractor,
            indexer=self.indexer,
            batch_size=batch_size,
            rate_limit=rate_limit,
            gdrive_sync_hours=gdrive_sync_hours,
            progress_callback=progress_callback,
        )
        self._lock = threading.Lock()

    # -- main operations --

    def harvest_all(self) -> Dict[str, Any]:
        """Full incremental harvest of all enabled drives + Google Drive."""
        result: Dict[str, Any] = {"local": {}, "gdrive": {}, "stats": {}}

        # Local drives
        result["local"] = self.scheduler.run_incremental()

        # Google Drive (if configured and due)
        if self.scheduler.should_sync_gdrive():
            result["gdrive"] = self.harvest_google_drive()
            self.scheduler.mark_gdrive_synced()

        result["stats"] = self.get_stats()
        return result

    def harvest_drive(self, drive_path: str) -> Dict[str, int]:
        """Harvest a single local drive."""
        drive_letter = pathlib.Path(drive_path).anchor.rstrip("\\:/")

        # Load known hashes
        try:
            known = self.indexer.get_known_hashes()
            self.scanner.load_known_hashes(known)
        except Exception:
            pass

        records = self.scanner.scan_drive(drive_path)
        stats = {"scanned": len(records), "indexed": 0, "skipped": 0, "errors": 0}

        for batch_start in range(0, len(records), self.scheduler.batch_size):
            batch = records[batch_start: batch_start + self.scheduler.batch_size]
            contents: List[ExtractedContent] = []
            valid: List[FileRecord] = []

            for rec in batch:
                try:
                    c = self.extractor.extract(rec.path)
                    if c.text:
                        contents.append(c)
                        valid.append(rec)
                    else:
                        stats["skipped"] += 1
                except Exception as exc:
                    stats["errors"] += 1
                    self.indexer.log_error(rec.path, type(exc).__name__, str(exc)[:500])

            if valid:
                self.indexer.index_batch(valid, contents)
                stats["indexed"] += len(valid)

            self.indexer.log_progress(
                drive_letter, batch[-1].path if batch else "",
                stats["scanned"], stats["indexed"],
                stats["skipped"], stats["errors"],
            )

        self.indexer.log_progress(
            drive_letter, "", stats["scanned"], stats["indexed"],
            stats["skipped"], stats["errors"], status="completed",
        )
        return stats

    def harvest_google_drive(self) -> Dict[str, Any]:
        """Sync and index files from Google Drive."""
        result: Dict[str, Any] = {"authenticated": False, "indexed": 0, "errors": 0}

        if not self.gdrive.authenticate():
            result["error"] = "Authentication failed or not configured."
            logger.warning("Google Drive harvest skipped — auth failed.")
            return result

        result["authenticated"] = True
        file_list = self.gdrive.list_files()
        if not file_list:
            result["message"] = "No files found or listing failed."
            return result

        local_paths = self.gdrive.download_files(file_list)
        for lp in local_paths:
            try:
                sha = compute_sha256(lp)
                if not sha:
                    result["errors"] += 1
                    continue
                stat = os.stat(lp)
                ext = os.path.splitext(lp)[1].lower()
                rec = FileRecord(
                    path=lp,
                    size=stat.st_size,
                    modified=_file_modified_iso(lp),
                    ext=ext,
                    sha256=sha,
                    drive="gdrive",
                )
                content = self.extractor.extract(lp)
                if content.text:
                    if self.indexer.index_single(rec, content, source_type="gdrive"):
                        result["indexed"] += 1
            except Exception as exc:
                result["errors"] += 1
                self.indexer.log_error(lp, type(exc).__name__, str(exc)[:500])

        result["gdrive_stats"] = self.gdrive.stats
        return result

    # -- search & retrieval --

    def search(
        self,
        query: str,
        lane: Optional[str] = None,
        file_type: Optional[str] = None,
        drive: Optional[str] = None,
        limit: int = 50,
    ) -> List[Dict[str, Any]]:
        """FTS5 search with optional filters."""
        return self.indexer.search(
            query, lane=lane, file_type=file_type, drive=drive, limit=limit,
        )

    def get_context_for_query(
        self,
        query: str,
        max_tokens: int = 4096,
        lane: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        """Return token-budgeted context snippets for an LLM query.

        Estimates ~4 chars per token and fills the budget from
        highest-ranked FTS5 results.
        """
        max_chars = max_tokens * 4
        results = self.search(query, lane=lane, limit=100)
        snippets: List[Dict[str, Any]] = []
        chars_used = 0

        for row in results:
            preview = row.get("content_preview", "") or ""
            if not preview:
                continue
            snippet_len = len(preview)
            if chars_used + snippet_len > max_chars:
                remaining = max_chars - chars_used
                if remaining > 100:
                    snippets.append({
                        "file_path": row.get("file_path", ""),
                        "drive": row.get("drive", ""),
                        "lane": row.get("lane_detected", ""),
                        "snippet": preview[:remaining],
                        "truncated": True,
                    })
                break
            snippets.append({
                "file_path": row.get("file_path", ""),
                "drive": row.get("drive", ""),
                "lane": row.get("lane_detected", ""),
                "snippet": preview,
                "truncated": False,
            })
            chars_used += snippet_len

        return snippets

    # -- stats & health --

    def get_stats(self) -> Dict[str, Any]:
        """Comprehensive stats across all subsystems."""
        idx_stats = self.indexer.get_stats()
        return {
            "index": idx_stats,
            "scanner": self.scanner.stats,
            "extractor": self.extractor.stats,
            "gdrive": self.gdrive.stats,
            "scheduler_running": self.scheduler.is_running,
        }

    def get_health(self) -> Dict[str, Any]:
        """Health check: drive accessibility, index freshness, error counts."""
        health: Dict[str, Any] = {
            "drives": {},
            "index_ok": False,
            "gdrive_configured": bool(os.environ.get("GOOGLE_DRIVE_CLIENT_ID")),
            "errors": [],
        }

        # Check each drive
        for drv, enabled in self.scanner.drives:
            drive_letter = pathlib.Path(drv).anchor.rstrip("\\:/")
            accessible = os.path.isdir(drv)
            health["drives"][drive_letter] = {
                "enabled": enabled,
                "accessible": accessible,
                "path": drv,
            }
            if enabled and not accessible:
                health["errors"].append(f"Drive {drive_letter} enabled but not accessible.")

        # Check index
        try:
            stats = self.indexer.get_stats()
            health["index_ok"] = stats.get("total_indexed", 0) >= 0
            health["total_indexed"] = stats.get("total_indexed", 0)
        except Exception as exc:
            health["errors"].append(f"Index check failed: {exc}")

        # Check recent errors
        try:
            conn = self.indexer._get_conn()
            recent_errors = conn.execute(
                "SELECT COUNT(*) AS cnt FROM harvest_errors "
                "WHERE occurred_at > datetime('now', '-24 hours')"
            ).fetchone()
            health["recent_errors_24h"] = recent_errors["cnt"] if recent_errors else 0
        except Exception:
            health["recent_errors_24h"] = -1

        return health

    def export_to_context_orchestrator(
        self, query: str, max_tokens: int = 4096,
    ) -> Dict[str, Any]:
        """Produce a context payload compatible with ContextOrchestrator.

        Returns a dict ready to be consumed by
        ``legal_ai.context_orchestrator.ContextOrchestrator.add_context()``.
        """
        snippets = self.get_context_for_query(query, max_tokens=max_tokens)
        return {
            "source": "context_harvester",
            "query": query,
            "max_tokens": max_tokens,
            "snippet_count": len(snippets),
            "snippets": snippets,
            "generated_at": _iso_now(),
        }

    # -- lifecycle --

    def stop(self) -> None:
        """Stop any running harvest."""
        self.scheduler.stop()

    def close(self) -> None:
        """Release resources."""
        self.scheduler.stop()
        self.indexer.close()

    # -- internal --

    @staticmethod
    def _scan_progress(drive: str, count: int) -> None:
        logger.info("Scanning %s — %d files found so far", drive, count)


# ===================================================================
#  CLI entry point
# ===================================================================

def _cli_main() -> None:
    """Minimal CLI for testing the harvester."""
    import argparse

    parser = argparse.ArgumentParser(
        description="LitigationOS Context Harvester — multi-drive document harvester",
    )
    sub = parser.add_subparsers(dest="command")

    # harvest
    harvest_p = sub.add_parser("harvest", help="Run incremental harvest")
    harvest_p.add_argument("--drive", help="Single drive letter (e.g. C)")
    harvest_p.add_argument("--batch-size", type=int, default=DEFAULT_BATCH_SIZE)
    harvest_p.add_argument("--include-gdrive", action="store_true")

    # search
    search_p = sub.add_parser("search", help="Search indexed content")
    search_p.add_argument("query", help="FTS5 search query")
    search_p.add_argument("--lane", help="Filter by lane (A-F)")
    search_p.add_argument("--type", dest="file_type", help="File extension filter")
    search_p.add_argument("--limit", type=int, default=20)

    # stats
    sub.add_parser("stats", help="Show index statistics")

    # health
    sub.add_parser("health", help="Run health check")

    args = parser.parse_args()

    if args.command == "harvest":
        harvester = ContextHarvester(batch_size=args.batch_size)
        if args.drive:
            drive_path = f"{args.drive.upper()}:\\"
            result = harvester.harvest_drive(drive_path)
        else:
            drives_cfg = list(DEFAULT_DRIVES)
            if not args.include_gdrive:
                harvester.scheduler.gdrive_sync_hours = 999_999
            result = harvester.harvest_all()
        print(json.dumps(result, indent=2, default=str))
        harvester.close()

    elif args.command == "search":
        harvester = ContextHarvester()
        results = harvester.search(
            args.query, lane=args.lane, file_type=args.file_type,
            limit=args.limit,
        )
        for r in results:
            print(f"[{r.get('lane_detected', '?')}] {r.get('file_path', '?')}")
            preview = r.get("content_preview", "")
            if preview:
                print(f"  {preview[:200]}")
            print()
        print(f"--- {len(results)} results ---")
        harvester.close()

    elif args.command == "stats":
        harvester = ContextHarvester()
        print(json.dumps(harvester.get_stats(), indent=2, default=str))
        harvester.close()

    elif args.command == "health":
        harvester = ContextHarvester()
        print(json.dumps(harvester.get_health(), indent=2, default=str))
        harvester.close()

    else:
        parser.print_help()


if __name__ == "__main__":
    _cli_main()
