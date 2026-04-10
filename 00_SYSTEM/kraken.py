#!/usr/bin/env python3
r"""
PROJECT KRAKEN -- Multi-Tentacle Autonomous Evidence Hunting System

Architecture:
  fd (Rust) -> file discovery (C:\, D:\)
  file_inventory DB -> slow drives (F:\, G:\, I:\, J:\)
  Content-first analysis -> PDF/DOCX/TXT/CSV/HTML/JSON/MD
  22 adversary patterns + 6 legal authority types + 8 evidence categories
  Auto-persist HIGH findings -> evidence_quotes + dossier files
  Dedup tracking -> kraken_processed table
  Focus modes -> adversary, legal, housing, judicial, custody, ppo, all

Usage:
  python kraken.py                    # 1 round, 10 files, all focus
  python kraken.py --rounds 5         # 5 rounds of 10
  python kraken.py --count 20         # 20 files per round
  python kraken.py --focus judicial   # focus on judicial patterns
  python kraken.py --drives C,D       # specific drives only
  python kraken.py --status           # show KRAKEN stats
  python kraken.py --resume           # continue from last session
"""

import sqlite3, random, json, os, sys, re, hashlib, traceback, argparse, subprocess
from pathlib import Path
from collections import defaultdict
from datetime import datetime, date

# ===================================================================
# CONFIGURATION
# ===================================================================

DB_PATH = r"C:\Users\andre\LitigationOS\litigation_context.db"
DOSSIER_DIR = r"C:\Users\andre\LitigationOS\04_ANALYSIS\ADVERSARY_TRACKS"
RESULTS_DIR = r"D:\LitigationOS_tmp\kraken_results"
EXTS = {'.pdf', '.txt', '.csv', '.html', '.json', '.docx', '.md'}
MAX_CONTENT_BYTES = 500_000  # 500KB content cap per file
MAX_PDF_PAGES = 30

# Separation counter
SEP_ANCHOR = date(2025, 7, 29)
SEP_DAYS = (date.today() - SEP_ANCHOR).days

def _safe(text):
    """Strip non-ASCII chars to prevent cp1252 encoding crashes on Windows."""
    if not isinstance(text, str):
        return str(text)
    return text.encode('ascii', 'replace').decode('ascii')

# ===================================================================
# ADVERSARY PATTERNS (22 targets)
# ===================================================================

ADVERSARIES = {
    'Emily Watson': r'(?i)\bemily\b.*\bwatson\b|\bwatson\b.*\bemily\b|\bemily\s+a\.?\s+watson\b',
    'Judge McNeill': r'(?i)\bmcneill\b|\bmcneil\b',
    'Pamela Rusco': r'(?i)\brusco\b|\bpamela\b.*\brusco\b',
    'Albert Watson': r'(?i)\balbert\b.*\bwatson\b|\bwatson\b.*\balbert\b',
    'Lori Watson': r'(?i)\blori\b.*\bwatson\b',
    'Ronald Berry': r'(?i)\bronald\b.*\bberry\b|\bron\b.*\bberry\b',
    'Cavan Berry': r'(?i)\bcavan\b.*\bberry\b',
    'Jennifer Barnes': r'(?i)\bbarnes\b.*\bjennifer\b|\bjennifer\b.*\bbarnes\b|\bP55406\b',
    'Mandi Martini': r'(?i)\bmartini\b|\bmandi\b.*\bmartini\b',
    'Kenneth Hoopes': r'(?i)\bhoopes\b|\bkenneth\b.*\bhoopes\b',
    'Maria Ladas-Hoopes': r'(?i)\bladas[\s-]*hoopes\b|\bmaria\b.*\bladas\b',
    'Kim Davis': r'(?i)\bkim\b.*\bdavis\b',
    'Cassandra VanDam': r'(?i)\bvandam\b|\bvan\s*dam\b|\bcassandra\b.*\bvan|\bcasey\b.*\bvan',
    'Nicole Browley': r'(?i)\bbrowley\b|\bnicole\b.*\bbrowley\b',
    'Henry Brandell': r'(?i)\bbrandell\b|\bbrandel\b|\bhenry\b.*\bbrand',
    'Jeremy Brown': r'(?i)\bjeremy\b.*\bbrown\b',
    'Aaron Cox': r'(?i)\baaron\b.*\bcox\b|\bcox\b.*\bP69346\b',
    'DJ Hilson': r'(?i)\bhilson\b|\bd\.?j\.?\b.*\bhilson\b',
    'Lauren Duguid': r'(?i)\bduguid\b|\blauren\b.*\bduguid\b',
    'Shady Oaks': r'(?i)\bshady\s*oaks\b',
    'FOC': r'(?i)\bfriend\s+of\s+(the\s+)?court\b|\bFOC\b',
    'EGLE': r'(?i)\bEGLE\b|\benvironmental\s+quality\b|\bsewage\b.*\bviolat',
}

# Compiled for speed
ADVERSARY_RE = {k: re.compile(v) for k, v in ADVERSARIES.items()}

# ===================================================================
# LEGAL AUTHORITY PATTERNS
# ===================================================================

LEGAL_PATTERNS = {
    'MCR': re.compile(r'\bMCR\s+\d+\.\d+\w*'),
    'MCL': re.compile(r'\bMCL\s+\d+\.\d+\w*'),
    'MRE': re.compile(r'\bMRE\s+\d+\w*'),
    'USC': re.compile(r'\b\d+\s+U\.?S\.?C\.?\s*[S\s]*\d+'),
    'FRCP': re.compile(r'\bF\.?R\.?C\.?P\.?\s*\d+|\bFed\.?\s*R\.?\s*Civ\.?\s*P\.?\s*\d+'),
    'Case_Law': re.compile(r'\b\d+\s+Mich\.?\s+(App\.?\s+)?\d+|\b\d+\s+F\.\d[a-z]*\s+\d+|\b\d+\s+S\.?\s*Ct\.?\s+\d+'),
}

# ===================================================================
# EVIDENCE CATEGORY PATTERNS
# ===================================================================

EVIDENCE_CATEGORIES = {
    'custody': re.compile(r'(?i)\bcustody\b|\bparenting\s*time\b|\bvisitation\b|\bbest\s+interest\b'),
    'PPO': re.compile(r'(?i)\bprotection\s*order\b|\bPPO\b|\brestraining\b|\bstalking\b'),
    'judicial': re.compile(r'(?i)\bjudicial\b|\bbias\b|\bex\s*parte\b|\brecusal\b|\bdisqualif'),
    'housing': re.compile(r'(?i)\beviction\b|\btenant\b|\blandlord\b|\bhousing\b|\bmobile\s*home\b|\btrailer\b'),
    'criminal': re.compile(r'(?i)\bcontempt\b|\bjail\b|\bincarcerat\b|\barrest\b|\bsentenc'),
    'financial': re.compile(r'(?i)\bchild\s*support\b|\bfiling\s*fee\b|\bdamages\b|\bgarnish'),
    'police': re.compile(r'(?i)\bpolice\b|\bNSPD\b|\bofficer\b|\bincident\s*report\b'),
    'medical': re.compile(r'(?i)\bhealthwest\b|\bmental\s*health\b|\bpsych\b|\bmedication\b|\bLOCUS\b'),
}

# ===================================================================
# FOCUS MODE BOOSTERS
# ===================================================================

FOCUS_BOOSTS = {
    'adversary': ['Emily Watson', 'Albert Watson', 'Lori Watson', 'Ronald Berry', 'Cavan Berry'],
    'judicial': ['Judge McNeill', 'Kenneth Hoopes', 'Maria Ladas-Hoopes', 'Cavan Berry', 'Pamela Rusco'],
    'housing': ['Shady Oaks', 'Cassandra VanDam', 'Nicole Browley', 'Henry Brandell', 'Jeremy Brown', 'Aaron Cox', 'EGLE'],
    'custody': ['Emily Watson', 'FOC', 'Pamela Rusco', 'Albert Watson'],
    'ppo': ['Emily Watson', 'Judge McNeill', 'Ronald Berry'],
    'legal': [],  # boosts legal authority matches instead
    'all': [],
}


# ===================================================================
# DATABASE HELPERS
# ===================================================================

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA busy_timeout=60000")
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA cache_size=-32000")
    conn.execute("PRAGMA temp_store=MEMORY")
    conn.execute("PRAGMA synchronous=NORMAL")
    return conn

def ensure_kraken_table(conn):
    """Create tracking table for processed files."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_processed (
            file_hash TEXT PRIMARY KEY,
            file_path TEXT NOT NULL,
            file_size INTEGER,
            processed_at TEXT DEFAULT (datetime('now')),
            round_id TEXT,
            value_score INTEGER DEFAULT 0,
            value_label TEXT DEFAULT 'LOW',
            adversaries_found TEXT,
            legal_found TEXT,
            categories TEXT,
            persisted_to_db INTEGER DEFAULT 0,
            focus_mode TEXT DEFAULT 'all'
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_rounds (
            round_id TEXT PRIMARY KEY,
            started_at TEXT DEFAULT (datetime('now')),
            completed_at TEXT,
            files_scanned INTEGER DEFAULT 0,
            high_count INTEGER DEFAULT 0,
            medium_count INTEGER DEFAULT 0,
            low_count INTEGER DEFAULT 0,
            focus_mode TEXT DEFAULT 'all',
            drives TEXT,
            new_evidence_rows INTEGER DEFAULT 0
        )
    """)
    conn.commit()


def file_hash(path):
    """Quick hash: path + size + mtime for dedup."""
    try:
        st = os.stat(path)
        return hashlib.md5(f"{path}|{st.st_size}|{int(st.st_mtime)}".encode()).hexdigest()
    except:
        return hashlib.md5(path.encode()).hexdigest()


def is_already_processed(conn, fhash):
    row = conn.execute("SELECT 1 FROM kraken_processed WHERE file_hash = ?", (fhash,)).fetchone()
    return row is not None


# ===================================================================
# FILE DISCOVERY (fd + DB inventory -- zero os.walk)
# ===================================================================

def discover_files(drives=None, conn=None):
    """Discover files using fd (fast drives) + file_inventory DB (slow drives)."""
    all_files = []
    drives = drives or ['C', 'D', 'F', 'G', 'I', 'J']

    # Source 1: file_inventory DB (pre-indexed, instant)
    if conn:
        try:
            cols = {r[1] for r in conn.execute("PRAGMA table_info(file_inventory)")}
            path_col = 'file_path' if 'file_path' in cols else 'path' if 'path' in cols else 'full_path'
            
            drive_filter = " OR ".join(f"{path_col} LIKE '{d}:%'" for d in drives)
            rows = conn.execute(f"""
                SELECT {path_col} FROM file_inventory 
                WHERE ({drive_filter}) AND {path_col} IS NOT NULL
                LIMIT 100000
            """).fetchall()
            
            db_files = [fp for (fp,) in rows if fp and any(fp.lower().endswith(e) for e in EXTS)]
            all_files.extend(db_files)
            print(f"  [DB] {len(db_files)} files from file_inventory")
        except Exception as e:
            print(f"  [DB] Error: {e}")

    # Source 2: fd (Rust) for fast drives
    for drive in drives:
        if drive.upper() in ('F', 'G', 'I', 'J'):
            continue  # slow USB/SD -- rely on DB inventory
        
        scan_path = f"{drive.upper()}:\\"
        if drive.upper() == 'C':
            scan_path = r"C:\Users\andre"
        
        try:
            ext_args = []
            for ext in ['pdf', 'txt', 'csv', 'html', 'json', 'docx', 'md']:
                ext_args.extend(['-e', ext])
            
            result = subprocess.run(
                ['fd', '--type', 'f', '--no-ignore', '--hidden'] + ext_args +
                ['--exclude', 'node_modules', '--exclude', '.git', '--exclude', 'AppData',
                 '--exclude', '__pycache__', '--exclude', '.cache', '--exclude', 'pytools_venv',
                 '--exclude', '.mcp_venv',
                 '.', scan_path],
                capture_output=True, text=True, timeout=45,
                encoding='utf-8', errors='replace'
            )
            fd_files = [f.strip() for f in result.stdout.strip().split('\n') if f.strip()]
            print(f"  [fd] {len(fd_files)} files on {drive.upper()}:\\")
            all_files.extend(fd_files)
        except subprocess.TimeoutExpired:
            print(f"  [fd] {drive.upper()}:\\ timed out (45s)")
        except FileNotFoundError:
            print(f"  [fd] fd not found -- install with: cargo install fd-find")
        except Exception as e:
            print(f"  [fd] {drive.upper()}:\\ error: {e}")

    # Deduplicate
    unique = list(set(all_files))
    print(f"  [TOTAL] {len(unique)} unique files discovered")
    return unique


# ===================================================================
# CONTENT EXTRACTION
# ===================================================================

def extract_content(filepath):
    """Read actual file content. Returns (text, method)."""
    ext = Path(filepath).suffix.lower()
    
    if ext == '.pdf':
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(filepath)
            pages = []
            for idx in range(min(len(pdf), MAX_PDF_PAGES)):
                page = pdf[idx]
                text = page.get_textpage().get_text_range()
                pages.append(text)
            content = '\n'.join(pages)
            return content[:MAX_CONTENT_BYTES], f"PDF({len(pdf)}pp)"
        except Exception as e:
            return "", f"PDF_ERR({e})"

    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            content = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            return content[:MAX_CONTENT_BYTES], f"DOCX({len(doc.paragraphs)}P)"
        except Exception as e:
            return "", f"DOCX_ERR({e})"

    elif ext in {'.txt', '.csv', '.html', '.json', '.md'}:
        for enc in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
            try:
                with open(filepath, 'r', encoding=enc, errors='replace') as f:
                    content = f.read(MAX_CONTENT_BYTES)
                return content, f"TEXT({enc})"
            except:
                continue
        return "", "TEXT_ERR"

    return "", "UNSUPPORTED"


# ===================================================================
# CONTENT ANALYSIS ENGINE
# ===================================================================

def analyze_content(content, filepath, focus='all'):
    """Analyze file content for adversary mentions, legal authorities, evidence categories."""
    result = {
        'adversaries': {},
        'legal': {},
        'categories': [],
        'key_quotes': [],
        'value_score': 0,
        'value_label': 'LOW',
    }

    if not content or len(content) < 50:
        result['value_label'] = 'EMPTY'
        return result

    # Adversary detection
    for name, pattern in ADVERSARY_RE.items():
        matches = pattern.findall(content)
        if matches:
            result['adversaries'][name] = len(matches)

    # Legal authority detection
    for cat, pattern in LEGAL_PATTERNS.items():
        matches = pattern.findall(content)
        if matches:
            result['legal'][cat] = list(set(matches))[:15]

    # Evidence category detection
    for cat, pattern in EVIDENCE_CATEGORIES.items():
        if pattern.search(content):
            result['categories'].append(cat)

    # Key quotes extraction (sentences mentioning adversaries)
    sentences = re.split(r'(?<=[.!?])\s+|\n+', content)
    for sent in sentences:
        sent = sent.strip()
        if 30 < len(sent) < 500:
            for name, pattern in ADVERSARY_RE.items():
                if pattern.search(sent):
                    result['key_quotes'].append(sent[:300])
                    break
        if len(result['key_quotes']) >= 8:
            break

    # Scoring with focus boost
    adv_score = len(result['adversaries']) * 3
    legal_score = len(result['legal']) * 2
    cat_score = len(result['categories'])
    quote_score = len(result['key_quotes'])

    # Focus mode boost
    focus_targets = FOCUS_BOOSTS.get(focus, [])
    if focus_targets:
        for target in focus_targets:
            if target in result['adversaries']:
                adv_score += result['adversaries'][target]  # extra boost for focus targets
    elif focus == 'legal':
        legal_score *= 2  # double legal authority score in legal focus

    result['value_score'] = adv_score + legal_score + cat_score + quote_score
    
    if result['value_score'] >= 10:
        result['value_label'] = 'HIGH'
    elif result['value_score'] >= 4:
        result['value_label'] = 'MEDIUM'
    else:
        result['value_label'] = 'LOW'

    return result


# ===================================================================
# AUTO-PERSIST TO DB
# ===================================================================

def persist_high_findings(conn, filepath, analysis, round_id):
    """Persist HIGH-value findings to evidence_quotes table."""
    persisted = 0
    
    if analysis['value_label'] != 'HIGH':
        return 0

    fname = Path(filepath).name
    categories = ', '.join(analysis['categories']) if analysis['categories'] else 'general'
    adversaries_str = ', '.join(sorted(analysis['adversaries'].keys()))

    for quote in analysis['key_quotes'][:5]:
        # Check if this exact quote already exists (dedup)
        existing = conn.execute(
            "SELECT id FROM evidence_quotes WHERE quote_text = ? AND source_file = ? LIMIT 1",
            (quote, fname)
        ).fetchone()
        
        if existing:
            continue

        # Determine lane
        lane = 'A'  # default custody
        if any(cat in analysis['categories'] for cat in ['housing']):
            lane = 'B'
        elif any(cat in analysis['categories'] for cat in ['judicial']):
            lane = 'E'
        elif any(cat in analysis['categories'] for cat in ['PPO']):
            lane = 'D'
        elif any(cat in analysis['categories'] for cat in ['criminal']):
            lane = 'CRIMINAL'

        try:
            conn.execute("""
                INSERT INTO evidence_quotes (source_file, quote_text, category, lane, relevance_score, tags, created_at)
                VALUES (?, ?, ?, ?, ?, ?, datetime('now'))
            """, (fname, quote, categories, lane, 0.7, f"KRAKEN:{round_id}:{adversaries_str}", ))
            persisted += 1
        except Exception as e:
            print(f"      DB insert error: {e}")

    if persisted:
        conn.commit()
    
    return persisted


# ===================================================================
# AUTO-EXPAND DOSSIERS
# ===================================================================

ADVERSARY_TO_DOSSIER = {
    'Emily Watson': 'WATSON_EMILY_DOSSIER.md',
    'Judge McNeill': 'MCNEILL_JENNY_DOSSIER.md',
    'Pamela Rusco': 'RUSCO_PAMELA_DOSSIER.md',
    'Albert Watson': 'WATSON_ALBERT_DOSSIER.md',
    'Lori Watson': 'WATSON_FAMILY_DOSSIER.md',
    'Ronald Berry': 'BERRY_RONALD_DOSSIER.md',
    'Cavan Berry': 'BERRY_CAVAN_DOSSIER.md',
    'Jennifer Barnes': 'BARNES_JENNIFER_DOSSIER.md',
    'Mandi Martini': 'MARTINI_MANDI_DOSSIER.md',
    'Kenneth Hoopes': 'HOOPES_KENNETH_DOSSIER.md',
    'Maria Ladas-Hoopes': 'LADAS_HOOPES_MARIA_DOSSIER.md',
    'Kim Davis': 'KIM_DAVIS_DOSSIER.md',
    'Cassandra VanDam': 'VANDAM_CASSANDRA_DOSSIER.md',
    'Nicole Browley': 'BROWLEY_NICOLE_DOSSIER.md',
    'Henry Brandell': 'BRANDELL_HENRY_DOSSIER.md',
    'Jeremy Brown': 'BROWN_JEREMY_DOSSIER.md',
    'Aaron Cox': 'COX_AARON_DOSSIER.md',
    'DJ Hilson': 'HILSON_DJ_DOSSIER.md',
    'Lauren Duguid': 'DUGUID_LAUREN_DOSSIER.md',
    'Shady Oaks': 'SHADY_OAKS_CORPORATE_DOSSIER.md',
}

def expand_dossiers(filepath, analysis, round_id):
    """Append KRAKEN findings to relevant dossier files."""
    if analysis['value_label'] not in ('HIGH', 'MEDIUM'):
        return 0
    
    expanded = 0
    fname = Path(filepath).name
    
    for adversary in analysis['adversaries']:
        dossier_file = ADVERSARY_TO_DOSSIER.get(adversary)
        if not dossier_file:
            continue
        
        dossier_path = os.path.join(DOSSIER_DIR, dossier_file)
        if not os.path.isfile(dossier_path):
            continue

        # Check if this file was already referenced in the dossier
        try:
            with open(dossier_path, 'r', encoding='utf-8', errors='replace') as f:
                existing = f.read()
            if fname in existing:
                continue  # already referenced
        except:
            continue

        # Get relevant quotes for this adversary
        adv_quotes = []
        for q in analysis['key_quotes']:
            if ADVERSARY_RE[adversary].search(q):
                adv_quotes.append(q)

        if not adv_quotes:
            continue

        # Append to dossier
        try:
            section = f"\n\n### KRAKEN Hunt ({round_id}) -- {datetime.now().strftime('%Y-%m-%d')}\n"
            section += f"**Source:** `{fname}` ({analysis['value_label']})\n"
            section += f"**Mentions:** {analysis['adversaries'][adversary]} references\n"
            if analysis['legal']:
                legal_str = ', '.join(f"{k}: {len(v)}" for k, v in analysis['legal'].items())
                section += f"**Legal refs:** {legal_str}\n"
            for q in adv_quotes[:3]:
                section += f'- "{q[:200]}"\n'

            with open(dossier_path, 'a', encoding='utf-8') as f:
                f.write(section)
            expanded += 1
        except Exception as e:
            print(f"      Dossier append error: {e}")

    return expanded


# ===================================================================
# KRAKEN ROUND EXECUTION
# ===================================================================

def run_round(conn, all_files, count=10, focus='all', round_num=1, total_rounds=1):
    """Execute one KRAKEN hunting round."""
    round_id = f"KRK-{datetime.now().strftime('%Y%m%d-%H%M%S')}-R{round_num}"
    
    print(f"\n{'=' * 70}")
    print(f"[KRAKEN] KRAKEN ROUND {round_num}/{total_rounds} -- {round_id}")
    print(f"   Focus: {focus.upper()} | Count: {count} | Pool: {len(all_files)}")
    print(f"   Separation: {SEP_DAYS} days since Jul 29, 2025")
    print(f"{'=' * 70}")

    # Record round start
    conn.execute("""
        INSERT INTO kraken_rounds (round_id, focus_mode, drives, files_scanned)
        VALUES (?, ?, ?, 0)
    """, (round_id, focus, 'all'))
    conn.commit()

    # Filter out already-processed files
    candidates = []
    for fp in all_files:
        fh = file_hash(fp)
        if not is_already_processed(conn, fh):
            try:
                sz = os.path.getsize(fp)
                if 100 < sz < 50_000_000:
                    candidates.append((fp, fh, sz))
            except:
                pass

    if not candidates:
        print("  [!] No unprocessed files remain!")
        return {'round_id': round_id, 'files': 0, 'high': 0, 'medium': 0, 'low': 0}

    # Evolved lottery draw (weighted by learned yields)
    draw = weighted_sample(candidates, conn, count)
    
    results = []
    high_count = 0
    med_count = 0
    low_count = 0
    total_persisted = 0
    total_dossiers = 0

    for idx, (fp, fh, sz) in enumerate(draw, 1):
        ext = Path(fp).suffix.lower()
        fname = Path(fp).name
        
        print(f"\n  [{idx:02d}/{count}] {_safe(fname)} ({sz:,}B)")
        
        # Extract content
        content, method = extract_content(fp)
        if not content:
            print(f"         ? Empty ({method})")
            conn.execute("""
                INSERT OR IGNORE INTO kraken_processed (file_hash, file_path, file_size, round_id, value_label, focus_mode)
                VALUES (?, ?, ?, ?, 'EMPTY', ?)
            """, (fh, fp, sz, round_id, focus))
            continue

        # Analyze
        analysis = analyze_content(content, fp, focus)
        label = analysis['value_label']
        score = analysis['value_score']
        
        # Color-coded output
        icon = '[!]' if label == 'HIGH' else '[~]' if label == 'MEDIUM' else '[ ]'
        print(f"         {icon} {label} (score={score}) via {method}")
        
        if analysis['adversaries']:
            top = sorted(analysis['adversaries'].items(), key=lambda x: -x[1])[:4]
            print(f"         [P] {_safe(', '.join(f'{n}({c})' for n,c in top))}")
        if analysis['legal']:
            print(f"         [L]  {', '.join(f'{k}:{len(v)}' for k,v in analysis['legal'].items())}")
        if analysis['categories']:
            print(f"         [F] {', '.join(analysis['categories'])}")
        if analysis['key_quotes']:
            print(f'         [Q] "{_safe(analysis["key_quotes"][0][:100])}..."')

        # Auto-persist HIGH findings to DB
        if label == 'HIGH':
            persisted = persist_high_findings(conn, fp, analysis, round_id)
            if persisted:
                print(f"         [DB] Persisted {persisted} quotes to evidence_quotes")
                total_persisted += persisted

        # Auto-expand dossiers for HIGH/MEDIUM
        if label in ('HIGH', 'MEDIUM'):
            expanded = expand_dossiers(fp, analysis, round_id)
            if expanded:
                print(f"         [+] Expanded {expanded} dossier(s)")
                total_dossiers += expanded

        # Track
        if label == 'HIGH': high_count += 1
        elif label == 'MEDIUM': med_count += 1
        else: low_count += 1

        # Record in tracking table
        conn.execute("""
            INSERT OR IGNORE INTO kraken_processed 
            (file_hash, file_path, file_size, round_id, value_score, value_label, 
             adversaries_found, legal_found, categories, persisted_to_db, focus_mode)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (fh, fp, sz, round_id, score, label,
              json.dumps(analysis['adversaries']),
              json.dumps({k: v[:5] for k, v in analysis['legal'].items()}),
              ','.join(analysis['categories']),
              1 if label == 'HIGH' else 0,
              focus))

        results.append({
            'file': fp, 'ext': ext, 'size': sz,
            'method': method, 'content_len': len(content),
            **analysis
        })

    conn.execute("""
        UPDATE kraken_rounds SET 
            completed_at = datetime('now'),
            files_scanned = ?,
            high_count = ?,
            medium_count = ?,
            low_count = ?,
            new_evidence_rows = ?
        WHERE round_id = ?
    """, (len(draw), high_count, med_count, low_count, total_persisted, round_id))
    conn.commit()

    # Round summary
    print(f"\n  {'-' * 50}")
    print(f"  ROUND {round_num} SUMMARY: {len(draw)} files scanned")
    print(f"  [!] HIGH: {high_count} | [~] MEDIUM: {med_count} | [ ] LOW: {low_count}")
    print(f"  [DB] DB inserts: {total_persisted} | [+] Dossiers expanded: {total_dossiers}")

    return {
        'round_id': round_id,
        'files': len(draw),
        'high': high_count,
        'medium': med_count,
        'low': low_count,
        'persisted': total_persisted,
        'dossiers_expanded': total_dossiers,
        'results': results,
    }


# ===================================================================
# STATUS REPORT
# ===================================================================

def show_status(conn):
    """Show KRAKEN operational status."""
    ensure_kraken_table(conn)
    
    print("\n" + "=" * 70)
    print("[KRAKEN] PROJECT KRAKEN -- STATUS REPORT")
    print("=" * 70)

    # Rounds
    rounds = conn.execute("""
        SELECT round_id, files_scanned, high_count, medium_count, low_count, 
               new_evidence_rows, focus_mode, completed_at
        FROM kraken_rounds ORDER BY started_at DESC LIMIT 20
    """).fetchall()

    print(f"\n  Rounds completed: {len(rounds)}")
    total_files = sum(r[1] for r in rounds)
    total_high = sum(r[2] for r in rounds)
    total_med = sum(r[3] for r in rounds)
    total_low = sum(r[4] for r in rounds)
    total_ev = sum(r[5] for r in rounds)
    
    print(f"  Total files scanned: {total_files}")
    print(f"  [!] HIGH: {total_high} | [~] MEDIUM: {total_med} | [ ] LOW: {total_low}")
    print(f"  [DB] Evidence rows persisted: {total_ev}")

    # Processed file stats
    processed = conn.execute("SELECT COUNT(*) FROM kraken_processed").fetchone()[0]
    by_label = conn.execute("""
        SELECT value_label, COUNT(*) FROM kraken_processed GROUP BY value_label ORDER BY COUNT(*) DESC
    """).fetchall()
    
    print(f"\n  Unique files processed: {processed}")
    for label, cnt in by_label:
        print(f"    {label}: {cnt}")

    # Top adversaries found
    print(f"\n  Top adversaries discovered:")
    all_advs = defaultdict(int)
    adv_rows = conn.execute("SELECT adversaries_found FROM kraken_processed WHERE adversaries_found IS NOT NULL").fetchall()
    for (raw,) in adv_rows:
        try:
            d = json.loads(raw)
            for k, v in d.items():
                all_advs[k] += v
        except:
            pass
    for name, cnt in sorted(all_advs.items(), key=lambda x: -x[1])[:10]:
        print(f"    {name}: {cnt} mentions")

    # Recent rounds
    if rounds:
        print(f"\n  Recent rounds:")
        for r in rounds[:5]:
            print(f"    {r[0]} | {r[1]} files | H:{r[2]} M:{r[3]} L:{r[4]} | focus:{r[6]} | {r[7] or 'running'}")

    print(f"\n  Separation: {SEP_DAYS} days since Jul 29, 2025")
    print("=" * 70)


# ===================================================================
# EVOLUTION ENGINE -- Self-tuning yield optimization
# ===================================================================

def ensure_evolution_tables(conn):
    """Create evolution tracking tables for self-tuning."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_evolution (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            round_id TEXT NOT NULL,
            metric_type TEXT NOT NULL,
            metric_key TEXT NOT NULL,
            metric_value REAL DEFAULT 0,
            sample_count INTEGER DEFAULT 0,
            recorded_at TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_weights (
            key TEXT PRIMARY KEY,
            category TEXT NOT NULL,
            weight REAL DEFAULT 1.0,
            total_scanned INTEGER DEFAULT 0,
            total_high INTEGER DEFAULT 0,
            yield_rate REAL DEFAULT 0,
            last_updated TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_cooccurrence (
            adversary_a TEXT NOT NULL,
            adversary_b TEXT NOT NULL,
            count INTEGER DEFAULT 1,
            last_seen TEXT DEFAULT (datetime('now')),
            PRIMARY KEY (adversary_a, adversary_b)
        )
    """)
    conn.commit()


def record_evolution_metrics(conn, round_id, results):
    """Compute and record evolution metrics from a round's results."""
    ensure_evolution_tables(conn)
    if not results:
        return

    items = results.get('results', [])
    if not items:
        return

    # Yield by file extension
    ext_stats = defaultdict(lambda: {'scanned': 0, 'high': 0})
    drive_stats = defaultdict(lambda: {'scanned': 0, 'high': 0})
    cooccur = defaultdict(int)

    for item in items:
        ext = item.get('ext', '.unknown')
        fp = item.get('file', '')
        label = item.get('value_label', 'LOW')
        advs = list(item.get('adversaries', {}).keys())

        ext_stats[ext]['scanned'] += 1
        if label == 'HIGH':
            ext_stats[ext]['high'] += 1

        drive = fp[0].upper() if fp and len(fp) > 1 and fp[1] == ':' else '?'
        drive_stats[drive]['scanned'] += 1
        if label == 'HIGH':
            drive_stats[drive]['high'] += 1

        # Track adversary co-occurrence (which adversaries appear together)
        for i, a1 in enumerate(advs):
            for a2 in advs[i+1:]:
                pair = tuple(sorted([a1, a2]))
                cooccur[pair] += 1

    # Persist extension yield metrics
    for ext, stats in ext_stats.items():
        yld = stats['high'] / max(stats['scanned'], 1)
        conn.execute("""
            INSERT INTO kraken_evolution (round_id, metric_type, metric_key, metric_value, sample_count)
            VALUES (?, 'ext_yield', ?, ?, ?)
        """, (round_id, ext, yld, stats['scanned']))
        conn.execute("""
            INSERT INTO kraken_weights (key, category, weight, total_scanned, total_high, yield_rate, last_updated)
            VALUES (?, 'ext', 1.0, ?, ?, ?, datetime('now'))
            ON CONFLICT(key) DO UPDATE SET
                total_scanned = total_scanned + excluded.total_scanned,
                total_high = total_high + excluded.total_high,
                yield_rate = CAST(total_high + excluded.total_high AS REAL) / 
                             MAX(total_scanned + excluded.total_scanned, 1),
                last_updated = datetime('now')
        """, (ext, stats['scanned'], stats['high'], yld))

    # Persist drive yield metrics
    for drive, stats in drive_stats.items():
        yld = stats['high'] / max(stats['scanned'], 1)
        conn.execute("""
            INSERT INTO kraken_evolution (round_id, metric_type, metric_key, metric_value, sample_count)
            VALUES (?, 'drive_yield', ?, ?, ?)
        """, (round_id, drive, yld, stats['scanned']))
        conn.execute("""
            INSERT INTO kraken_weights (key, category, weight, total_scanned, total_high, yield_rate, last_updated)
            VALUES (?, 'drive', 1.0, ?, ?, ?, datetime('now'))
            ON CONFLICT(key) DO UPDATE SET
                total_scanned = total_scanned + excluded.total_scanned,
                total_high = total_high + excluded.total_high,
                yield_rate = CAST(total_high + excluded.total_high AS REAL) / 
                             MAX(total_scanned + excluded.total_scanned, 1),
                last_updated = datetime('now')
        """, (drive, stats['scanned'], stats['high'], yld))

    # Persist co-occurrence
    for (a1, a2), cnt in cooccur.items():
        conn.execute("""
            INSERT INTO kraken_cooccurrence (adversary_a, adversary_b, count, last_seen)
            VALUES (?, ?, ?, datetime('now'))
            ON CONFLICT(adversary_a, adversary_b) DO UPDATE SET
                count = count + excluded.count,
                last_seen = datetime('now')
        """, (a1, a2, cnt))

    conn.commit()


def get_evolved_weights(conn):
    """Get evolved sampling weights for intelligent file selection."""
    ensure_evolution_tables(conn)
    weights = {'ext': {}, 'drive': {}}
    try:
        rows = conn.execute("""
            SELECT key, category, yield_rate, total_scanned 
            FROM kraken_weights WHERE total_scanned >= 3
            ORDER BY yield_rate DESC
        """).fetchall()
        for key, cat, yld, scanned in rows:
            if cat in weights:
                # Weight = base(1.0) + yield_bonus(0-3.0) scaled by confidence
                confidence = min(scanned / 20.0, 1.0)  # full confidence at 20+ samples
                weights[cat][key] = 1.0 + (yld * 3.0 * confidence)
    except:
        pass
    return weights


def weighted_sample(candidates, conn, count):
    """Sample files using evolved weights instead of uniform random."""
    weights = get_evolved_weights(conn)
    if not weights['ext'] and not weights['drive']:
        return random.sample(candidates, min(count, len(candidates)))

    scored = []
    for fp, fh, sz in candidates:
        ext = Path(fp).suffix.lower()
        drive = fp[0].upper() if fp and len(fp) > 1 and fp[1] == ':' else '?'
        w = weights['ext'].get(ext, 1.0) * weights['drive'].get(drive, 1.0)
        scored.append((fp, fh, sz, w))

    # Weighted random selection
    total_w = sum(s[3] for s in scored)
    if total_w <= 0:
        return random.sample(candidates, min(count, len(candidates)))

    selected = []
    pool = list(scored)
    for _ in range(min(count, len(pool))):
        r = random.uniform(0, sum(s[3] for s in pool))
        cumulative = 0
        for idx, item in enumerate(pool):
            cumulative += item[3]
            if cumulative >= r:
                selected.append((item[0], item[1], item[2]))
                pool.pop(idx)
                break

    return selected


def show_evolution(conn):
    """Display evolution metrics and learned weights."""
    ensure_evolution_tables(conn)

    print("\n" + "=" * 70)
    print("[KRAKEN] EVOLUTION ENGINE -- Self-Tuning Intelligence")
    print("=" * 70)

    # Extension yields
    rows = conn.execute("""
        SELECT key, total_scanned, total_high, yield_rate, weight
        FROM kraken_weights WHERE category = 'ext' AND total_scanned > 0
        ORDER BY yield_rate DESC
    """).fetchall()

    if rows:
        print("\n  File Type Yield Rates (learned from all rounds):")
        print(f"  {'Ext':<8} {'Scanned':>8} {'HIGH':>6} {'Yield':>8} {'Weight':>8}")
        print(f"  {'-'*8} {'-'*8} {'-'*6} {'-'*8} {'-'*8}")
        for key, scanned, high, yld, weight in rows:
            bar = '#' * int(yld * 20)
            print(f"  {key:<8} {scanned:>8} {high:>6} {yld:>7.1%} {weight:>7.2f}  {bar}")
    else:
        print("\n  No extension yield data yet. Run more rounds to train.")

    # Drive yields
    rows = conn.execute("""
        SELECT key, total_scanned, total_high, yield_rate, weight
        FROM kraken_weights WHERE category = 'drive' AND total_scanned > 0
        ORDER BY yield_rate DESC
    """).fetchall()

    if rows:
        print(f"\n  Drive Yield Rates:")
        print(f"  {'Drive':<8} {'Scanned':>8} {'HIGH':>6} {'Yield':>8} {'Weight':>8}")
        print(f"  {'-'*8} {'-'*8} {'-'*6} {'-'*8} {'-'*8}")
        for key, scanned, high, yld, weight in rows:
            bar = '#' * int(yld * 20)
            print(f"  {key:<8} {scanned:>8} {high:>6} {yld:>7.1%} {weight:>7.2f}  {bar}")

    # Co-occurrence network (adversary relationships discovered)
    cooccur = conn.execute("""
        SELECT adversary_a, adversary_b, count 
        FROM kraken_cooccurrence ORDER BY count DESC LIMIT 15
    """).fetchall()

    if cooccur:
        print(f"\n  Adversary Co-Occurrence Network (discovered relationships):")
        for a1, a2, cnt in cooccur:
            strength = '#' * min(cnt, 20)
            print(f"    {a1:<20} <-> {a2:<20} [{cnt:>3}] {strength}")

    # Evolution trajectory
    rounds = conn.execute("""
        SELECT round_id, COUNT(*), 
               SUM(CASE WHEN metric_type='ext_yield' THEN metric_value ELSE 0 END),
               SUM(sample_count)
        FROM kraken_evolution GROUP BY round_id ORDER BY recorded_at DESC LIMIT 10
    """).fetchall()

    if rounds:
        print(f"\n  Evolution Trajectory (last 10 rounds):")
        for rid, metrics, total_yield, samples in rounds:
            short_id = rid.split('-')[-1] if '-' in rid else rid[:12]
            print(f"    {short_id}: {metrics} metrics, avg_yield={total_yield/max(metrics,1):.2%}, samples={samples}")

    print("=" * 70)


# ===================================================================
# CONDENSATION PIPELINE -- JUDICIAL + NARRATIVE intelligence
# ===================================================================

JUDICIAL_PATTERNS = {
    'ex_parte': re.compile(r'(?i)\bex\s*parte\b'),
    'bias': re.compile(r'(?i)\bbias\b|\bprejudice\b|\bimpartial\b'),
    'due_process': re.compile(r'(?i)\bdue\s*process\b|\bnotice\b.*\bhearing\b|\b14th\s*amendment\b'),
    'contempt_abuse': re.compile(r'(?i)\bcontempt\b|\bjail\b|\bincarcerat\b'),
    'evidence_exclusion': re.compile(r'(?i)\bexclud\w+\s+evidence\b|\brefus\w+\s+to\s+consider\b'),
    'recusal_refusal': re.compile(r'(?i)\brecus\w+\b|\bdisqualif\w+\b|\bMCR\s*2\.003\b'),
    'cartel_connection': re.compile(r'(?i)\bformer\s+partner\b|\bsame\s+address\b|\bmarried\b|\bfamily\b.*\bconnect\b'),
    'rights_denial': re.compile(r'(?i)\bparent\w+\s+right\b|\bfundamental\s+right\b|\baccess\s+to\s+court\b'),
}

NARRATIVE_PATTERNS = {
    'date_event': re.compile(r'(?:(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}|\d{1,2}/\d{1,2}/\d{2,4}|\d{4}-\d{2}-\d{2})'),
    'court_action': re.compile(r'(?i)\bordered\b|\bgranted\b|\bdenied\b|\bfiled\b|\bruled\b|\bsustained\b|\boverruled\b'),
    'witness_statement': re.compile(r'(?i)\bstated\b|\btestified\b|\badmitted\b|\bclaimed\b|\balleged\b|\bdeclared\b'),
    'separation_event': re.compile(r'(?i)\bwithh\w+\b.*\bchild\b|\bno\s+contact\b|\bsuspend\w+\b.*\bparent\b'),
    'false_allegation': re.compile(r'(?i)\bfalse\b.*\ballegat\b|\bfabricat\b|\brecant\b'),
    'harm_indicator': re.compile(r'(?i)\blost\s+(?:home|job|employ)\b|\bincarcerat\b|\bemotional\s+distress\b'),
}


def condense_findings(conn, since_round=None):
    """Condense KRAKEN findings into JUDICIAL + NARRATIVE intelligence structures."""
    ensure_evolution_tables(conn)

    print("\n" + "=" * 70)
    print("[KRAKEN] CONDENSATION PIPELINE -- Judicial + Narrative Intelligence")
    print("=" * 70)

    # Gather all HIGH/MEDIUM findings
    query = """
        SELECT kp.file_path, kp.value_score, kp.value_label, kp.adversaries_found,
               kp.legal_found, kp.categories, kp.round_id
        FROM kraken_processed kp
        WHERE kp.value_label IN ('HIGH', 'MEDIUM')
    """
    params = []
    if since_round:
        query += " AND kp.round_id >= ?"
        params.append(since_round)
    query += " ORDER BY kp.value_score DESC"

    findings = conn.execute(query, params).fetchall()
    print(f"\n  Processing {len(findings)} HIGH/MEDIUM findings...")

    if not findings:
        print("  No findings to condense. Run more KRAKEN rounds first.")
        return None

    # Build intelligence structures
    judicial_intel = {
        'violation_types': defaultdict(int),
        'violation_evidence': defaultdict(list),
        'cartel_connections': [],
        'rights_denials': [],
        'timeline_events': [],
    }
    narrative_intel = {
        'dated_events': [],
        'court_actions': [],
        'witness_statements': [],
        'false_allegations': [],
        'harm_indicators': [],
        'separation_events': [],
    }
    adversary_summary = defaultdict(lambda: {'mentions': 0, 'high_count': 0, 'evidence': []})

    for fp, score, label, advs_json, legal_json, categories, round_id in findings:
        fname = Path(fp).name if fp else 'unknown'
        advs = {}
        legal = {}
        try:
            advs = json.loads(advs_json) if advs_json else {}
        except:
            pass
        try:
            legal = json.loads(legal_json) if legal_json else {}
        except:
            pass

        # Re-read content for deep analysis
        content = ""
        if fp and os.path.isfile(fp):
            text, _ = extract_content(fp)
            content = text[:100000]  # cap at 100KB for condensation

        # Judicial classification
        for vtype, pattern in JUDICIAL_PATTERNS.items():
            matches = pattern.findall(content)
            if matches:
                judicial_intel['violation_types'][vtype] += len(matches)
                judicial_intel['violation_evidence'][vtype].append({
                    'file': fname, 'score': score, 'match_count': len(matches),
                    'adversaries': list(advs.keys()), 'round': round_id,
                })

        # Narrative extraction
        dates = NARRATIVE_PATTERNS['date_event'].findall(content)
        for d in dates[:10]:
            # Find the sentence containing this date
            idx = content.find(d)
            if idx >= 0:
                start = max(0, content.rfind('.', max(0, idx-200), idx) + 1)
                end = content.find('.', idx, idx + 300)
                if end < 0:
                    end = min(len(content), idx + 200)
                sentence = content[start:end+1].strip()
                if len(sentence) > 20:
                    judicial_intel['timeline_events'].append({
                        'date_str': d, 'text': sentence[:300], 'file': fname, 'score': score
                    })
                    narrative_intel['dated_events'].append({
                        'date_str': d, 'text': sentence[:300], 'file': fname
                    })

        for ntype, pattern in NARRATIVE_PATTERNS.items():
            if ntype == 'date_event':
                continue  # already handled
            matches = pattern.findall(content)
            if matches and ntype in narrative_intel:
                for m in matches[:5]:
                    idx = content.find(m)
                    if idx >= 0:
                        start = max(0, content.rfind('.', max(0, idx-150), idx) + 1)
                        end = content.find('.', idx, idx + 250)
                        if end < 0:
                            end = min(len(content), idx + 200)
                        sentence = content[start:end+1].strip()
                        if len(sentence) > 20:
                            narrative_intel[ntype].append({
                                'text': sentence[:300], 'file': fname, 'type': ntype
                            })

        # Adversary aggregation
        for adv, cnt in advs.items():
            adversary_summary[adv]['mentions'] += cnt
            if label == 'HIGH':
                adversary_summary[adv]['high_count'] += 1
            adversary_summary[adv]['evidence'].append(fname)

    # Deduplicate narrative events by text similarity (simple prefix match)
    seen_prefixes = set()
    for ntype in narrative_intel:
        deduped = []
        for item in narrative_intel[ntype]:
            prefix = item.get('text', '')[:60].lower()
            if prefix not in seen_prefixes:
                seen_prefixes.add(prefix)
                deduped.append(item)
        narrative_intel[ntype] = deduped

    # Print condensation report
    print(f"\n  JUDICIAL INTELLIGENCE CONDENSED:")
    total_jud = sum(judicial_intel['violation_types'].values())
    for vtype, cnt in sorted(judicial_intel['violation_types'].items(), key=lambda x: -x[1]):
        bar = '#' * min(cnt, 30)
        print(f"    {vtype:<24} {cnt:>4} {bar}")
    print(f"    {'TOTAL':<24} {total_jud:>4}")

    print(f"\n  NARRATIVE INTELLIGENCE CONDENSED:")
    for ntype, items in narrative_intel.items():
        print(f"    {ntype:<24} {len(items):>4} events")

    print(f"\n  ADVERSARY INTELLIGENCE:")
    sorted_advs = sorted(adversary_summary.items(), key=lambda x: -x[1]['mentions'])[:15]
    for adv, data in sorted_advs:
        files_str = f"({len(set(data['evidence']))} files)"
        print(f"    {adv:<24} {data['mentions']:>5} mentions, {data['high_count']:>3} HIGH {files_str}")

    condensed = {
        'judicial': judicial_intel,
        'narrative': narrative_intel,
        'adversaries': dict(adversary_summary),
        'stats': {
            'total_findings': len(findings),
            'judicial_violations': total_jud,
            'narrative_events': sum(len(v) for v in narrative_intel.values()),
            'adversaries_profiled': len(adversary_summary),
            'condensed_at': datetime.now().isoformat(),
        }
    }

    # Save condensation to disk
    cond_path = os.path.join(RESULTS_DIR, f"condensed_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(cond_path, 'w', encoding='utf-8') as f:
        json.dump(condensed, f, indent=2, default=str)
    print(f"\n  [>] Condensation saved: {cond_path}")

    return condensed


# ===================================================================
# MBP BRIDGE -- Transform to THEMANBEARPIG graph data
# ===================================================================

MBP_GRAPH_DATA = os.path.join(
    r"C:\Users\andre\LitigationOS\12_WORKSPACE\THEMANBEARPIG_v7", "graph_data_v7.json"
)

# Colors matching MBP layer aesthetic
KRAKEN_COLORS = {
    'discovery': '#00ffcc',      # cyan-green for new discoveries
    'judicial_violation': '#cc44ff',  # purple for judicial
    'narrative_event': '#ffaa00',     # amber for timeline
    'adversary_link': '#ff2244',      # red for adversary connections
    'cooccurrence': '#44ccff',        # blue for co-occurrence
}


def export_to_mbp(conn, condensed=None):
    """Export KRAKEN discoveries as MBP-compatible graph nodes and links."""
    ensure_evolution_tables(conn)

    print("\n" + "=" * 70)
    print("[KRAKEN] MBP BRIDGE -- Transforming to THEMANBEARPIG graph data")
    print("=" * 70)

    if condensed is None:
        condensed = condense_findings(conn)
    if condensed is None:
        print("  [!] No condensed data available. Run --condense first.")
        return

    new_nodes = []
    new_links = []
    nid_set = set()

    def add_node(nid, label, layer, color, r, threat=0, group="", desc="", data=None):
        if nid in nid_set:
            return
        nid_set.add(nid)
        lod = 1 if r >= 16 else (2 if r >= 8 else (3 if r >= 4 else 4))
        new_nodes.append({
            "id": nid, "label": label, "layer": layer, "color": color,
            "r": r, "threat": threat, "group": group, "desc": desc[:200],
            "lod": lod, "data": data or {}
        })

    def add_link(src, tgt, typ, color="#1a3a5a", w=0.5, strength=0.3):
        new_links.append({
            "source": src, "target": tgt, "type": typ,
            "color": color, "width": w, "strength": strength
        })

    # === JUDICIAL VIOLATION NODES ===
    jud = condensed.get('judicial', {})
    for vtype, cnt in jud.get('violation_types', {}).items():
        if cnt >= 2:  # only significant violations
            nid = f"krk_jud_{vtype}"
            r = max(4, min(14, cnt * 0.5))
            threat = min(10, cnt // 2 + 3)
            add_node(nid, f"JUD: {vtype.replace('_', ' ').title()}", "JUDICIAL_CARTEL",
                     KRAKEN_COLORS['judicial_violation'], r, threat, "Judicial Violation",
                     f"KRAKEN found {cnt} instances of {vtype}",
                     data={"violation_type": vtype, "count": cnt, "source": "KRAKEN"})
            # Link to McNeill (primary judicial target)
            add_link(nid, "adv_judge_mcneill", "VIOLATION",
                     KRAKEN_COLORS['judicial_violation'], max(0.3, cnt/20), 0.3)

    # === NARRATIVE TIMELINE NODES (top 20 dated events) ===
    narr = condensed.get('narrative', {})
    dated = narr.get('dated_events', [])
    for i, evt in enumerate(dated[:20]):
        nid = f"krk_evt_{i}"
        add_node(nid, f"EVT: {evt['date_str']}", "TIMELINE",
                 KRAKEN_COLORS['narrative_event'], 5, 0, "Timeline Event",
                 evt.get('text', '')[:200],
                 data={"date": evt['date_str'], "file": evt.get('file', ''), "source": "KRAKEN"})

    # === ADVERSARY CO-OCCURRENCE LINKS ===
    cooccur = conn.execute("""
        SELECT adversary_a, adversary_b, count 
        FROM kraken_cooccurrence WHERE count >= 2 ORDER BY count DESC LIMIT 30
    """).fetchall()

    for a1, a2, cnt in cooccur:
        # Map adversary names to MBP node IDs
        a1_id = f"adv_{a1.replace(' ', '_').lower()}"
        a2_id = f"adv_{a2.replace(' ', '_').lower()}"
        add_link(a1_id, a2_id, "KRAKEN_COOCCURRENCE",
                 KRAKEN_COLORS['cooccurrence'], max(0.3, cnt / 10), 0.4)

    # === FALSE ALLEGATION NODES ===
    false_allgs = narr.get('false_allegations', [])
    for i, fa in enumerate(false_allgs[:10]):
        nid = f"krk_fa_{i}"
        add_node(nid, f"FALSE: {fa.get('text', '')[:40]}", "WEAPON_CHAINS",
                 '#ff4444', 5, 6, "False Allegation", fa.get('text', '')[:200],
                 data={"type": "false_allegation", "file": fa.get('file', ''), "source": "KRAKEN"})
        add_link(nid, "adv_emily_watson", "WEAPON_USE", '#ff4444', 0.5, 0.3)
        add_link(nid, "andrew", "HARM", '#ff2244', 0.3, 0.2)

    # === HARM INDICATOR NODES ===
    harms = narr.get('harm_indicators', [])
    for i, h in enumerate(harms[:10]):
        nid = f"krk_harm_{i}"
        add_node(nid, f"HARM: {h.get('text', '')[:40]}", "DAMAGES",
                 '#ff8800', 4, 5, "Harm", h.get('text', '')[:200],
                 data={"type": "harm_indicator", "file": h.get('file', ''), "source": "KRAKEN"})

    # === WITNESS STATEMENT NODES ===
    witnesses = narr.get('witness_statements', [])
    for i, ws in enumerate(witnesses[:10]):
        nid = f"krk_wit_{i}"
        add_node(nid, f"WIT: {ws.get('text', '')[:40]}", "IMPEACHMENT",
                 '#ff3366', 4, 0, "Witness Statement", ws.get('text', '')[:200],
                 data={"type": "witness_statement", "file": ws.get('file', ''), "source": "KRAKEN"})

    print(f"\n  Generated {len(new_nodes)} new nodes and {len(new_links)} new links")

    # --- SAVE KRAKEN OVERLAY (for MBP builder to merge during HTML rebuild) ---
    overlay_path = os.path.join(os.path.dirname(MBP_GRAPH_DATA), "kraken_overlay.json")
    kraken_graph = {
        'nodes': new_nodes, 'links': new_links,
        'meta': {
            'source': 'KRAKEN', 'version': '2.0',
            'exported': datetime.now().isoformat(),
            'node_count': len(new_nodes),
            'link_count': len(new_links),
            'stats': condensed.get('stats', {}),
        }
    }
    os.makedirs(os.path.dirname(overlay_path), exist_ok=True)
    with open(overlay_path, 'w', encoding='utf-8') as f:
        json.dump(kraken_graph, f, indent=2, default=str)
    print(f"  [+] KRAKEN overlay saved: {overlay_path}")
    print(f"      {len(new_nodes)} nodes, {len(new_links)} links ready for MBP merge")

    # Also save timestamped standalone export in results dir
    export_path = os.path.join(RESULTS_DIR, f"kraken_graph_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(export_path, 'w', encoding='utf-8') as f:
        json.dump(kraken_graph, f, indent=2, default=str)
    print(f"  [>] Timestamped export: {export_path}")

    # Update graph_data_v7.json metadata (counts only -- node arrays live in HTML)
    if os.path.isfile(MBP_GRAPH_DATA):
        try:
            with open(MBP_GRAPH_DATA, 'r', encoding='utf-8') as f:
                meta = json.load(f)
            meta['kraken_nodes'] = len(new_nodes)
            meta['kraken_links'] = len(new_links)
            meta['kraken_updated'] = datetime.now().isoformat()
            with open(MBP_GRAPH_DATA, 'w', encoding='utf-8') as f:
                json.dump(meta, f, indent=2, default=str)
            print(f"  [+] Updated MBP metadata: {MBP_GRAPH_DATA}")
        except Exception as e:
            print(f"  [~] Could not update MBP metadata: {e}")
    else:
        print(f"  [~] MBP metadata not found (run mbp_build.py to create base graph)")

    print("=" * 70)
    return kraken_graph


# ===================================================================
# CONTINUOUS DAEMON MODE
# ===================================================================

def run_continuous(conn, all_files, interval_seconds=300, max_cycles=50,
                   count=10, focus='all', condense_every=3, mbp_export_every=5):
    """Run KRAKEN in continuous self-improving daemon mode."""
    print("\n" + "=" * 70)
    print("[KRAKEN] CONTINUOUS MODE -- Self-Evolving Daemon")
    print(f"   Interval: {interval_seconds}s | Max cycles: {max_cycles}")
    print(f"   Condense every {condense_every} rounds | MBP export every {mbp_export_every} rounds")
    print("=" * 70)

    import time

    cycle = 0
    total_persisted = 0
    total_high = 0

    while cycle < max_cycles:
        cycle += 1
        print(f"\n{'#' * 70}")
        print(f"[KRAKEN] DAEMON CYCLE {cycle}/{max_cycles}")
        print(f"{'#' * 70}")

        # Run round with evolved sampling
        result = run_round(conn, all_files, count=count, focus=focus,
                          round_num=cycle, total_rounds=max_cycles)

        # Record evolution metrics
        record_evolution_metrics(conn, result.get('round_id', f'cycle_{cycle}'), result)

        total_persisted += result.get('persisted', 0)
        total_high += result.get('high', 0)

        # Periodic condensation
        if cycle % condense_every == 0:
            print(f"\n  [*] Condensation triggered (every {condense_every} rounds)...")
            condense_findings(conn)

        # Periodic MBP export
        if cycle % mbp_export_every == 0:
            print(f"\n  [*] MBP export triggered (every {mbp_export_every} rounds)...")
            condensed = condense_findings(conn)
            export_to_mbp(conn, condensed)

        # Show evolution state periodically
        if cycle % 5 == 0:
            show_evolution(conn)

        # Daemon summary
        print(f"\n  [DAEMON] Cycle {cycle} complete. Cumulative: {total_high} HIGH, {total_persisted} persisted")
        print(f"  [DAEMON] Separation: {SEP_DAYS} days. Next cycle in {interval_seconds}s...")

        if cycle < max_cycles:
            try:
                time.sleep(interval_seconds)
            except KeyboardInterrupt:
                print("\n  [DAEMON] Interrupted by user. Running final condensation...")
                condensed = condense_findings(conn)
                export_to_mbp(conn, condensed)
                break

    print(f"\n{'=' * 70}")
    print(f"[KRAKEN] DAEMON COMPLETE -- {cycle} cycles")
    print(f"   Total HIGH: {total_high} | Total persisted: {total_persisted}")
    print(f"{'=' * 70}")


# ===================================================================
# MAIN
# ===================================================================

def main():
    parser = argparse.ArgumentParser(
        description="PROJECT KRAKEN -- Multi-Tentacle Autonomous Evidence Hunting",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python kraken.py                       # 1 round, 10 files
  python kraken.py --rounds 5            # 5 rounds of 10
  python kraken.py --count 20            # 20 files per round
  python kraken.py --focus judicial      # judicial focus mode
  python kraken.py --drives C,D          # specific drives
  python kraken.py --status              # show stats
  python kraken.py --evolve              # show evolution metrics
  python kraken.py --condense            # run condensation pipeline
  python kraken.py --mbp-export          # export to MBP graph
  python kraken.py --continuous          # daemon mode (self-evolving)
  python kraken.py --continuous --interval 60 --max-cycles 20
        """
    )
    parser.add_argument('--rounds', type=int, default=1, help='Number of hunting rounds (default: 1)')
    parser.add_argument('--count', type=int, default=10, help='Files per round (default: 10)')
    parser.add_argument('--focus', default='all', choices=list(FOCUS_BOOSTS.keys()), help='Focus mode')
    parser.add_argument('--drives', default=None, help='Comma-separated drive letters (default: all)')
    parser.add_argument('--status', action='store_true', help='Show KRAKEN status report')
    parser.add_argument('--resume', action='store_true', help='Resume from last session')
    parser.add_argument('--evolve', action='store_true', help='Show evolution metrics')
    parser.add_argument('--condense', action='store_true', help='Run condensation pipeline')
    parser.add_argument('--mbp-export', action='store_true', help='Export KRAKEN data to MBP graph')
    parser.add_argument('--continuous', action='store_true', help='Daemon mode: continuous self-evolving hunting')
    parser.add_argument('--interval', type=int, default=300, help='Daemon cycle interval in seconds (default: 300)')
    parser.add_argument('--max-cycles', type=int, default=50, help='Max daemon cycles (default: 50)')
    parser.add_argument('--condense-every', type=int, default=3, help='Condense every N rounds (default: 3)')
    parser.add_argument('--mbp-every', type=int, default=5, help='MBP export every N rounds (default: 5)')
    args = parser.parse_args()

    print("=" * 70)
    print("[KRAKEN] PROJECT KRAKEN v2.0 -- Self-Evolving Evidence Hunting")
    print(f"   Version 2.0 | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"   Separation: {SEP_DAYS} days since Jul 29, 2025")
    print("=" * 70)

    conn = get_db()
    ensure_kraken_table(conn)
    ensure_evolution_tables(conn)

    # Quick-exit commands
    if args.status:
        show_status(conn)
        conn.close()
        return

    if args.evolve and not args.continuous:
        show_evolution(conn)
        conn.close()
        return

    if args.condense and not args.continuous:
        condense_findings(conn)
        conn.close()
        return

    if args.mbp_export and not args.continuous:
        condensed = condense_findings(conn)
        export_to_mbp(conn, condensed)
        conn.close()
        return

    # Parse drives
    drives = args.drives.split(',') if args.drives else None

    # Discover files
    print("\n[PHASE 1] File Discovery...")
    all_files = discover_files(drives=drives, conn=conn)

    if not all_files:
        print("  [!] No files discovered!")
        conn.close()
        return

    # Ensure results directory
    os.makedirs(RESULTS_DIR, exist_ok=True)

    # Continuous daemon mode
    if args.continuous:
        if args.evolve:
            show_evolution(conn)
        run_continuous(conn, all_files, interval_seconds=args.interval,
                      max_cycles=args.max_cycles, count=args.count, focus=args.focus,
                      condense_every=args.condense_every, mbp_export_every=args.mbp_every)
        conn.close()
        return

    # Execute rounds (with evolution tracking)
    all_round_results = []
    for rnd in range(1, args.rounds + 1):
        result = run_round(conn, all_files, count=args.count, focus=args.focus,
                          round_num=rnd, total_rounds=args.rounds)
        all_round_results.append(result)
        # Record evolution metrics after each round
        record_evolution_metrics(conn, result.get('round_id', f'R{rnd}'), result)

    # Save combined results
    out_file = os.path.join(RESULTS_DIR, f"kraken_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(out_file, 'w', encoding='utf-8') as f:
        json.dump(all_round_results, f, indent=2, default=str)
    print(f"\n  [>] Results saved: {out_file}")

    # Final summary
    total_files = sum(r['files'] for r in all_round_results)
    total_high = sum(r['high'] for r in all_round_results)
    total_med = sum(r['medium'] for r in all_round_results)
    total_persisted = sum(r.get('persisted', 0) for r in all_round_results)
    total_dossiers = sum(r.get('dossiers_expanded', 0) for r in all_round_results)

    print(f"\n{'=' * 70}")
    print(f"[KRAKEN] KRAKEN MISSION COMPLETE -- {args.rounds} round(s)")
    print(f"   Files scanned: {total_files}")
    print(f"   [!] HIGH: {total_high} | [~] MEDIUM: {total_med}")
    print(f"   [DB] Evidence persisted: {total_persisted} rows")
    print(f"   [+] Dossiers expanded: {total_dossiers}")
    print(f"   Separation: {SEP_DAYS} days -- EVERY DAY COUNTS")
    print(f"{'=' * 70}")

    # Auto-condense if 3+ rounds
    if args.rounds >= 3:
        print("\n  [*] Auto-condensing (3+ rounds completed)...")
        condensed = condense_findings(conn)
        if condensed and args.rounds >= 5:
            print("\n  [*] Auto-exporting to MBP (5+ rounds)...")
            export_to_mbp(conn, condensed)

    # Show cumulative status
    show_status(conn)
    show_evolution(conn)
    conn.close()


if __name__ == '__main__':
    main()
