"""
LOTTERY HARVEST — Pick 10 random files from all drives, read content, extract intel.
No filename filtering. Content-first. Read everything, decide what's useful.
"""
import os, random, json, sys, hashlib, time

DRIVES = [
    r"C:\Users\andre",
    r"D:\\",
    r"F:\\",
    r"G:\\",
    r"I:\\",
    r"J:\\",
]

EXTENSIONS = {'.pdf', '.txt', '.csv', '.html', '.json', '.docx'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB cap per file

# Skip directories that are definitely not useful
SKIP_DIRS = {
    'node_modules', '.git', '__pycache__', 'site-packages', '.cache',
    'AppData', '.cargo', '.rustup', '.ollama', '.EasyOCR', 'pytools_venv',
    '.mcp_venv', '.ml_venv', 'go', '.vscode', '.vscode-server',
    'Music', 'Videos', '3D Objects', 'Saved Games', 'Favorites',
    'Contacts', 'Searches', 'Pictures',
}

def discover_files():
    """Walk all drives, collect every matching file path."""
    all_files = []
    for drive in DRIVES:
        if not os.path.exists(drive):
            print(f"  SKIP {drive} (not accessible)")
            continue
        print(f"  Scanning {drive}...", end=" ", flush=True)
        count = 0
        try:
            for root, dirs, files in os.walk(drive, topdown=True):
                dirs[:] = [d for d in dirs if d not in SKIP_DIRS]
                for f in files:
                    ext = os.path.splitext(f)[1].lower()
                    if ext in EXTENSIONS:
                        fp = os.path.join(root, f)
                        try:
                            sz = os.path.getsize(fp)
                            if 100 < sz < MAX_FILE_SIZE:  # skip empty/tiny and huge
                                all_files.append((fp, sz, ext))
                                count += 1
                        except:
                            pass
        except Exception as e:
            print(f"ERROR: {e}")
        print(f"{count} files found")
    return all_files

def read_file_content(filepath, ext):
    """Read actual content from any supported file type."""
    content = ""
    try:
        if ext == '.pdf':
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(filepath)
                pages = []
                for i in range(min(len(pdf), 30)):  # cap at 30 pages
                    page = pdf[i]
                    tp = page.get_textpage()
                    pages.append(tp.get_text_bounded())
                content = "\n--- PAGE BREAK ---\n".join(pages)
            except:
                try:
                    import fitz
                    doc = fitz.open(filepath)
                    pages = []
                    for i in range(min(len(doc), 30)):
                        pages.append(doc[i].get_text())
                    content = "\n".join(pages)
                except:
                    content = "[PDF extraction failed]"
        
        elif ext == '.docx':
            try:
                from docx import Document
                doc = Document(filepath)
                content = "\n".join(p.text for p in doc.paragraphs)
            except:
                content = "[DOCX extraction failed]"
        
        elif ext in ('.txt', '.csv', '.html', '.json'):
            for enc in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
                try:
                    with open(filepath, 'r', encoding=enc, errors='replace') as f:
                        content = f.read(500_000)  # 500KB cap for text
                    break
                except:
                    continue
    except Exception as e:
        content = f"[Read error: {e}]"
    
    return content

def analyze_content(filepath, content):
    """Analyze content for litigation intelligence. Pure content analysis, no filename games."""
    if not content or len(content) < 50:
        return None
    
    content_lower = content.lower()
    
    # Adversary detection (names + variations)
    adversary_hits = {}
    adversary_patterns = {
        'Emily Watson': ['emily watson', 'emily a. watson', 'emily pigors', 'emily a watson'],
        'Jenny McNeill': ['mcneill', 'mc neill', 'jenny l. mcneill', 'judge mcneill'],
        'Pamela Rusco': ['rusco', 'pamela rusco', 'pam rusco'],
        'Albert Watson': ['albert watson', 'albert w'],
        'Lori Watson': ['lori watson'],
        'Ronald Berry': ['ronald berry', 'ron berry', 'ronald t berry'],
        'Cavan Berry': ['cavan berry', 'cavan b'],
        'Mandi Martini': ['mandi martini', 'martini'],
        'Jennifer Barnes': ['jennifer barnes', 'barnes p55406', 'barnes attorney'],
        'Kenneth Hoopes': ['hoopes', 'kenneth hoopes', 'judge hoopes', 'chief judge hoopes'],
        'Maria Ladas-Hoopes': ['ladas-hoopes', 'ladas hoopes', 'maria ladas'],
        'Kim Davis': ['kim davis'],
        'Nicole Browley': ['browley', 'nicole browley'],
        'Cassandra VanDam': ['vandam', 'van dam', 'cassandra vandam', 'casey vandam'],
        'Jeremy Brown': ['jeremy brown', 'attorney brown'],
        'Aaron Cox': ['aaron cox', 'aaron d. cox', 'cox pllc'],
        'Henry Brandell': ['brandell', 'brandel', 'henry brandell'],
        'Lauren Duguid': ['duguid', 'lauren duguid'],
        'DJ Hilson': ['hilson', 'dj hilson'],
        'Kostrzewa': ['kostrzewa'],
        'Shady Oaks': ['shady oaks', 'homes of america', 'cricklewood'],
        'EGLE': ['egle', 'environmental quality', 'wastewater violation'],
    }
    
    for name, patterns in adversary_patterns.items():
        hits = sum(content_lower.count(p) for p in patterns)
        if hits > 0:
            adversary_hits[name] = hits
    
    # Legal authority detection
    legal_hits = {}
    import re
    mcr_matches = re.findall(r'MCR\s+\d+\.\d+', content, re.IGNORECASE)
    mcl_matches = re.findall(r'MCL\s+\d+\.\d+', content, re.IGNORECASE)
    mre_matches = re.findall(r'MRE\s+\d+', content, re.IGNORECASE)
    usc_matches = re.findall(r'\d+\s+(?:U\.?S\.?C\.?|USC)\s+[§]?\s*\d+', content)
    frcp_matches = re.findall(r'(?:FRCP|Fed\.?\s*R\.?\s*Civ\.?\s*P\.?)\s*\d+', content, re.IGNORECASE)
    case_law = re.findall(r'\d+\s+Mich\s+(?:App\s+)?\d+', content)
    
    if mcr_matches: legal_hits['MCR'] = list(set(mcr_matches))[:20]
    if mcl_matches: legal_hits['MCL'] = list(set(mcl_matches))[:20]
    if mre_matches: legal_hits['MRE'] = list(set(mre_matches))[:20]
    if usc_matches: legal_hits['USC'] = list(set(usc_matches))[:20]
    if frcp_matches: legal_hits['FRCP'] = list(set(frcp_matches))[:20]
    if case_law: legal_hits['CASE_LAW'] = list(set(case_law))[:20]
    
    # Evidence markers
    evidence_markers = {
        'custody': ['custody', 'parenting time', 'visitation', 'best interest'],
        'ppo': ['protection order', 'ppo', 'personal protection', 'stalking'],
        'contempt': ['contempt', 'show cause', 'jail', 'incarcerat'],
        'abuse_allegations': ['abuse', 'assault', 'arsenic', 'poison', 'meth', 'drug'],
        'police': ['police report', 'nspd', 'officer', 'incident report', 'ns250'],
        'housing': ['eviction', 'lease', 'mobile home', 'trailer', 'rent', 'sewage'],
        'judicial_misconduct': ['ex parte', 'bias', 'recusal', 'disqualif', 'jtc', 'judicial tenure'],
        'financial': ['child support', 'arrears', 'garnish', 'income', 'employment'],
        'medical': ['healthwest', 'mental health', 'locus', 'psycho', 'medication'],
        'constitutional': ['due process', 'first amendment', 'fourteenth', '1983', 'civil rights'],
        'foc': ['friend of court', 'foc', 'referee'],
        'appellate': ['court of appeals', 'coa', 'appeal', 'msc', 'supreme court'],
    }
    
    evidence_found = {}
    for category, keywords in evidence_markers.items():
        hits = sum(1 for kw in keywords if kw in content_lower)
        if hits > 0:
            evidence_found[category] = hits
    
    # Court forms detection
    form_patterns = re.findall(r'(?:MC|FOC|DC|CC|COA|AO|JS)\s*-?\s*\d+', content, re.IGNORECASE)
    scao_forms = list(set(form_patterns))[:15] if form_patterns else []
    
    # Extract top quotes (sentences containing adversary names or key phrases)
    sentences = re.split(r'[.!?\n]+', content)
    top_quotes = []
    key_phrases = ['ex parte', 'denied', 'violated', 'contempt', 'custody', 'false', 
                   'fraud', 'retaliat', 'withhold', 'refused', 'obstruct', 'conspir',
                   'threat', 'harass', 'stalk', 'poison', 'arsenic', 'meth', 'jail',
                   'incarcerat', 'evict', 'sewage', 'brandell', 'browley', 'mcneill']
    
    for s in sentences:
        s = s.strip()
        if 30 < len(s) < 500:
            s_lower = s.lower()
            if any(kp in s_lower for kp in key_phrases):
                top_quotes.append(s)
                if len(top_quotes) >= 15:
                    break
    
    total_signal = len(adversary_hits) + len(legal_hits) + len(evidence_found) + len(scao_forms)
    
    if total_signal == 0:
        return None
    
    return {
        'adversaries': adversary_hits,
        'legal_authorities': legal_hits,
        'evidence_categories': evidence_found,
        'court_forms': scao_forms,
        'top_quotes': top_quotes,
        'content_length': len(content),
        'signal_strength': total_signal,
    }

def main():
    print("=" * 70)
    print("LOTTERY HARVEST — 10 Random Files, Content-First Analysis")
    print("=" * 70)
    
    # Step 1: Discover ALL files
    print("\n[PHASE 1] Discovering files across all drives...")
    t0 = time.time()
    all_files = discover_files()
    print(f"\n  TOTAL: {len(all_files):,} files discovered in {time.time()-t0:.1f}s")
    
    # Step 2: Pick 10 truly random files
    print("\n[PHASE 2] Lottery draw — 10 random files...")
    random.seed()  # true random
    if len(all_files) < 10:
        picks = all_files
    else:
        picks = random.sample(all_files, 10)
    
    for i, (fp, sz, ext) in enumerate(picks, 1):
        print(f"  #{i:2d}: [{ext}] {sz:>10,} bytes — {fp}")
    
    # Step 3: Read and analyze each
    print("\n[PHASE 3] Reading content & analyzing intelligence...")
    results = []
    for i, (fp, sz, ext) in enumerate(picks, 1):
        print(f"\n{'='*60}")
        print(f"FILE #{i}: {fp}")
        print(f"  Size: {sz:,} bytes | Type: {ext}")
        
        content = read_file_content(fp, ext)
        if not content:
            print("  ⚠ No content extracted")
            continue
        
        print(f"  Content: {len(content):,} chars extracted")
        print(f"  Preview: {content[:200].replace(chr(10), ' ')[:200]}...")
        
        analysis = analyze_content(fp, content)
        if analysis:
            print(f"\n  🎯 SIGNAL STRENGTH: {analysis['signal_strength']}")
            if analysis['adversaries']:
                print(f"  👤 Adversaries: {analysis['adversaries']}")
            if analysis['legal_authorities']:
                auth_summary = {k: len(v) for k,v in analysis['legal_authorities'].items()}
                print(f"  ⚖️  Legal authorities: {auth_summary}")
            if analysis['evidence_categories']:
                print(f"  📋 Evidence categories: {analysis['evidence_categories']}")
            if analysis['court_forms']:
                print(f"  📄 Court forms: {analysis['court_forms']}")
            if analysis['top_quotes']:
                print(f"  💬 Top quotes ({len(analysis['top_quotes'])}):")
                for q in analysis['top_quotes'][:5]:
                    print(f"     \"{q[:120]}...\"" if len(q) > 120 else f"     \"{q}\"")
            
            results.append({
                'file': fp,
                'size': sz,
                'ext': ext,
                **analysis
            })
        else:
            print("  ❌ No litigation signal detected")
    
    # Step 4: Summary
    print("\n" + "=" * 70)
    print("LOTTERY RESULTS SUMMARY")
    print("=" * 70)
    
    hits = len(results)
    misses = len(picks) - hits
    print(f"\n  🎰 {hits}/10 files contained litigation intelligence")
    print(f"  ❌ {misses}/10 files had no signal")
    
    if results:
        all_adversaries = {}
        all_authorities = {}
        all_categories = {}
        all_quotes = []
        
        for r in results:
            for name, count in r['adversaries'].items():
                all_adversaries[name] = all_adversaries.get(name, 0) + count
            for atype, cites in r['legal_authorities'].items():
                if atype not in all_authorities:
                    all_authorities[atype] = set()
                all_authorities[atype].update(cites)
            for cat, count in r['evidence_categories'].items():
                all_categories[cat] = all_categories.get(cat, 0) + count
            all_quotes.extend(r['top_quotes'])
        
        if all_adversaries:
            print(f"\n  ADVERSARIES DETECTED:")
            for name, count in sorted(all_adversaries.items(), key=lambda x: -x[1]):
                print(f"    {name}: {count} mentions")
        
        if all_authorities:
            print(f"\n  LEGAL AUTHORITIES:")
            for atype, cites in sorted(all_authorities.items()):
                print(f"    {atype}: {len(cites)} unique ({', '.join(list(cites)[:5])}...)")
        
        if all_categories:
            print(f"\n  EVIDENCE CATEGORIES:")
            for cat, count in sorted(all_categories.items(), key=lambda x: -x[1]):
                print(f"    {cat}: {count} signals")
        
        print(f"\n  TOTAL EXTRACTABLE QUOTES: {len(all_quotes)}")
    
    # Save full results
    outpath = r"D:\LitigationOS_tmp\lottery_results.json"
    with open(outpath, 'w', encoding='utf-8') as f:
        json.dump({
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'total_files_discovered': len(all_files),
            'files_sampled': len(picks),
            'files_with_signal': hits,
            'results': results,
            'file_universe_by_ext': {},
        }, f, indent=2, default=str)
    print(f"\n  Results saved to {outpath}")
    
    # Also save the full file universe for future waves
    universe_path = r"D:\LitigationOS_tmp\file_universe.json"
    ext_counts = {}
    drive_counts = {}
    for fp, sz, ext in all_files:
        ext_counts[ext] = ext_counts.get(ext, 0) + 1
        drive = fp[:3]
        drive_counts[drive] = drive_counts.get(drive, 0) + 1
    
    with open(universe_path, 'w', encoding='utf-8') as f:
        json.dump({
            'total_files': len(all_files),
            'by_extension': ext_counts,
            'by_drive': drive_counts,
            'sample_paths': [fp for fp, _, _ in random.sample(all_files, min(100, len(all_files)))],
        }, f, indent=2)
    print(f"  File universe ({len(all_files):,} files) saved to {universe_path}")

if __name__ == "__main__":
    main()
