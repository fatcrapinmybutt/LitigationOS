"""
Extract key SHADY OAKS ledger and evidence files.
Target: CSV ledger + COOKING THE BOOKS PDF + Motion for Sanctions DOCX + Gmail false ledger PDF
"""
import os, sys, sqlite3
sys.path.insert(0, r'C:\Users\andre\LitigationOS\00_SYSTEM\shared')

OUT = r"C:\Users\andre\temp\ledger_extraction.txt"
lines = []

# 1. Read ledger CSV
csv_path = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\shady_oaks_park_mhp_llc_LEDGER.csv"
if os.path.exists(csv_path):
    with open(csv_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read()
    lines.append("=== LEDGER CSV ===")
    # Check if null bytes
    null_count = content.count('\x00')
    lines.append(f"File size: {len(content)} chars, null bytes: {null_count}")
    if null_count > 100:
        lines.append("SPOLIATION CONFIRMED: >100 null bytes in CSV")
        # Try reading as binary to find any real data
        with open(csv_path, 'rb') as f:
            raw = f.read()
        non_null = bytes(b for b in raw if b != 0)
        lines.append(f"Non-null bytes: {len(non_null)}")
        lines.append(f"Sample non-null: {non_null[:500]}")
    else:
        lines.append(content[:3000])

# 2. Read COOKING THE BOOKS PDF
import subprocess
ctb_pdf = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\ledger COOKING THE BOOKS shady oaks.pdf"
if os.path.exists(ctb_pdf):
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(ctb_pdf)
        lines.append(f"\n=== COOKING THE BOOKS PDF ({len(doc)} pages) ===")
        for i in range(min(len(doc), 10)):
            page = doc[i]
            text = page.get_textpage().get_text_range()
            if text.strip():
                lines.append(f"--- Page {i+1} ---")
                lines.append(text[:2000])
        doc.close()
    except ImportError:
        try:
            import fitz
            doc = fitz.open(ctb_pdf)
            lines.append(f"\n=== COOKING THE BOOKS PDF ({doc.page_count} pages) ===")
            for i in range(min(doc.page_count, 10)):
                text = doc[i].get_text()
                if text.strip():
                    lines.append(f"--- Page {i+1} ---")
                    lines.append(text[:2000])
            doc.close()
        except Exception as e:
            lines.append(f"PDF read failed: {e}")

# 3. Read Motion for Sanctions DOCX
ms_docx = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\04_MOTION_FOR_SANCTIONS_LEDGER_FRAUD.docx"
if os.path.exists(ms_docx):
    try:
        import docx
        doc = docx.Document(ms_docx)
        lines.append(f"\n=== MOTION FOR SANCTIONS DOCX ===")
        for p in doc.paragraphs[:80]:
            if p.text.strip():
                lines.append(p.text)
    except Exception as e:
        lines.append(f"DOCX read failed: {e}")

# 4. Read Gmail false ledger email PDF
gmail_pdf = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\Gmail -SHADY OAKS CASSANDRA UNPROFESSIONAL FALSE LEDGER Andrew Pigors Current Ledger.pdf"
if os.path.exists(gmail_pdf):
    try:
        import fitz
        doc = fitz.open(gmail_pdf)
        lines.append(f"\n=== GMAIL FALSE LEDGER EMAIL ({doc.page_count} pages) ===")
        for i in range(min(doc.page_count, 5)):
            text = doc[i].get_text()
            if text.strip():
                lines.append(f"--- Page {i+1} ---")
                lines.append(text[:2000])
        doc.close()
    except Exception as e:
        lines.append(f"PDF read failed: {e}")

# 5. Read "Andrew Pigors Ledger 8-22" PDF
ap_pdf = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\Andrew Pigors Ledger 8-22.pdf"
if os.path.exists(ap_pdf):
    try:
        import fitz
        doc = fitz.open(ap_pdf)
        lines.append(f"\n=== ANDREW PIGORS LEDGER 8-22 PDF ({doc.page_count} pages) ===")
        for i in range(min(doc.page_count, 5)):
            text = doc[i].get_text()
            if text.strip():
                lines.append(f"--- Page {i+1} ---")
                lines.append(text[:2000])
        doc.close()
    except Exception as e:
        lines.append(f"PDF read failed: {e}")

# Write output
with open(OUT, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print('\n'.join(lines[:200]))
print(f"\nFull output: {OUT}")
