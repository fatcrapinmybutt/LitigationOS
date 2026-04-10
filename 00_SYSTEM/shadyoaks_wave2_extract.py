"""
SHADYOAKS-DESTRUCTION Wave 2 — Document Extraction Fleet
Reads: XLS ledger, NoReply PDFs, shady oaks court PDFs, LARA DOCX files
Writes: shadyoaks_brain.db + litigation_context.db
"""
import sys, os, re, json, sqlite3
from pathlib import Path
from datetime import datetime

sys.path.insert(0, r"C:\Users\andre\LitigationOS\00_SYSTEM\shared")

BRAIN_DB = r"C:\Users\andre\LitigationOS\00_SYSTEM\brains\shadyoaks_brain.db"
MAIN_DB  = r"C:\Users\andre\LitigationOS\litigation_context.db"
SHADY    = r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY"
LOG      = []

def db_conn(path):
    c = sqlite3.connect(path)
    c.execute("PRAGMA busy_timeout=60000")
    c.execute("PRAGMA journal_mode=WAL")
    c.execute("PRAGMA cache_size=-32000")
    return c

def log(msg):
    print(msg)
    LOG.append(msg)

# ─────────────────────────────────────────────
# 1. READ XLS LEDGER
# ─────────────────────────────────────────────
def read_xls_ledger():
    xls = os.path.join(SHADY, "shady enhanced redacted ledger$$$ conv_10aa496d.xls")
    if not os.path.exists(xls):
        log(f"XLS NOT FOUND: {xls}")
        return []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(xls, read_only=True, data_only=True)
        rows = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            log(f"XLS SHEET: {sheet}")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if any(v is not None for v in row):
                    rows.append((sheet, i+1, str(row)))
                    if i < 5:
                        log(f"  ROW {i+1}: {row}")
        log(f"XLS TOTAL NON-EMPTY ROWS: {len(rows)}")
        return rows
    except Exception as e:
        log(f"XLS ERROR: {e}")
        # try xlrd as fallback
        try:
            import xlrd
            wb = xlrd.open_workbook(xls)
            rows = []
            for sheet in wb.sheets():
                log(f"XLRD SHEET: {sheet.name}")
                for r in range(sheet.nrows):
                    row_data = sheet.row_values(r)
                    rows.append((sheet.name, r+1, str(row_data)))
                    if r < 5:
                        log(f"  ROW {r+1}: {row_data}")
            log(f"XLRD TOTAL ROWS: {len(rows)}")
            return rows
        except Exception as e2:
            log(f"XLRD ERROR: {e2}")
            return []

# ─────────────────────────────────────────────
# 2. READ PDFs WITH pypdfium2
# ─────────────────────────────────────────────
def read_pdf(path, max_pages=10):
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(path)
        text = []
        for i in range(min(len(doc), max_pages)):
            page = doc[i]
            tp = page.get_textpage()
            text.append(tp.get_text_range())
        return "\n".join(text)
    except Exception as e:
        return f"PDF_ERROR: {e}"

def read_noreply_pdfs():
    """Read first 5 + last 5 of the July 21 NoReply batch"""
    results = {}
    for i in [1,2,3,4,5,28,29,30,31,32]:
        fn = f"NoReply_20250721_102658_{i:04d}.pdf"
        fp = os.path.join(SHADY, fn)
        if os.path.exists(fp):
            txt = read_pdf(fp, max_pages=3)
            results[fn] = txt[:1500]
            log(f"NOREPLY {fn}: {txt[:300]}")
        else:
            log(f"NOREPLY MISSING: {fn}")
    # Also check Aug 27 batch
    for i in [1,2,3]:
        fn = f"NoReply_20250827_142908_{i:04d}.pdf"
        fp = os.path.join(SHADY, fn)
        if os.path.exists(fp):
            txt = read_pdf(fp, max_pages=3)
            results[fn] = txt[:1500]
            log(f"NOREPLY AUG {fn}: {txt[:300]}")
    return results

def read_shady_court_pdfs():
    """Read shady oaks 2_0001 through 0016"""
    results = {}
    for i in range(1, 17):
        fn = f"shady oaks 2_{i:04d}.pdf"
        fp = os.path.join(SHADY, fn)
        if os.path.exists(fp):
            txt = read_pdf(fp, max_pages=5)
            results[fn] = txt[:2000]
            log(f"COURT DOC {fn}: {txt[:400]}")
        else:
            log(f"COURT DOC MISSING: {fn}")
    return results

# ─────────────────────────────────────────────
# 3. READ LARA DOCX FILES
# ─────────────────────────────────────────────
def read_lara_docx():
    results = {}
    try:
        from docx import Document
    except ImportError:
        log("python-docx NOT INSTALLED — skipping DOCX")
        return results
    for fn in os.listdir(SHADY):
        if fn.lower().endswith(".docx") and any(k in fn.lower() for k in ["lara","license","1201891","mhp","egle","sewage"]):
            fp = os.path.join(SHADY, fn)
            try:
                doc = Document(fp)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                results[fn] = text[:3000]
                log(f"DOCX {fn}: {text[:400]}")
            except Exception as e:
                log(f"DOCX ERROR {fn}: {e}")
    return results

# ─────────────────────────────────────────────
# 4. READ REMAINING MEEK1 CSVs
# ─────────────────────────────────────────────
def read_meek1_csvs():
    results = {}
    targets = [
        "MEEK1_ShadyOaks_Facts (2).csv",
        "MEEK1_ShadyOaks_Entity_Map.csv",
        "MEEK1_ShadyOaks_Damages_Leverage_Grid.csv",
        "MEEK1_ShadyOaks_CivilClaim_Architecture.csv",
        "MEEK1_ShadyOaks_Claim_Element_Fact_Grid.csv",
    ]
    for fn in targets:
        fp = os.path.join(SHADY, fn)
        if os.path.exists(fp):
            try:
                with open(fp, "r", encoding="utf-8-sig", errors="replace") as f:
                    content = f.read()
                results[fn] = content[:4000]
                log(f"CSV {fn}: {len(content)} chars, first row: {content.split(chr(10))[1] if chr(10) in content else content[:200]}")
            except Exception as e:
                log(f"CSV ERROR {fn}: {e}")
    return results

# ─────────────────────────────────────────────
# 5. READ TXT FILES (TL_BITEMPORAL, EXHIBIT_MATRIX)
# ─────────────────────────────────────────────
def read_key_txts():
    results = {}
    # Search SHADY folder and subdirs for these
    targets = ["TL_BITEMPORAL", "EXHIBIT_MATRIX", "LT_RFJ_ATTACK", "REAL_RENT_LEDGER", "RUN_LEDGER"]
    for root, dirs, files in os.walk(SHADY):
        for fn in files:
            if any(t.lower() in fn.lower() for t in targets):
                fp = os.path.join(root, fn)
                try:
                    with open(fp, "r", encoding="utf-8", errors="replace") as f:
                        content = f.read(5000)
                    results[fn] = content
                    log(f"TXT {fn} ({fp}): {content[:300]}")
                except Exception as e:
                    log(f"TXT ERROR {fn}: {e}")
    return results

# ─────────────────────────────────────────────
# 6. CHECK for VanDam screenshot metadata
# ─────────────────────────────────────────────
def read_vandm_screenshot_metadata():
    fp = os.path.join(SHADY, "Screenshot_20260218_212004_One UI Home.jpg")
    if not os.path.exists(fp):
        log(f"VANDM SCREENSHOT NOT FOUND: {fp}")
        return None
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS
        img = Image.open(fp)
        log(f"VANDM SCREENSHOT: size={img.size}, mode={img.mode}, format={img.format}")
        exif = img._getexif()
        if exif:
            for tag_id, value in exif.items():
                tag = TAGS.get(tag_id, tag_id)
                log(f"  EXIF {tag}: {value}")
        # Try to get file size and basic info
        sz = os.path.getsize(fp)
        log(f"  File size: {sz} bytes")
        return {"path": fp, "size": sz, "img_size": img.size}
    except Exception as e:
        log(f"VANDM SCREENSHOT ERROR: {e}")
        return {"path": fp, "size": os.path.getsize(fp)}

# ─────────────────────────────────────────────
# 7. WRITE RESULTS TO BRAIN DB
# ─────────────────────────────────────────────
def write_to_brain(conn, ledger_rows, noreply_results, court_results):
    cur = conn.cursor()
    ts = datetime.now().isoformat()
    
    # Write ledger rows summary
    if ledger_rows:
        cur.execute("INSERT OR IGNORE INTO evidence_registry (item_id, description, file_path, category, status, created_at) VALUES (?,?,?,?,?,?)",
            ("LEDGER_XLS_001", f"XLS Ledger extracted: {len(ledger_rows)} rows across sheets", 
             r"C:\Users\andre\Desktop\COURT_FILING_PACKETS\SHADY\shady enhanced redacted ledger$$$ conv_10aa496d.xls",
             "financial", "EXTRACTED", ts))
        # Add sample rows as evidence
        for sheet, rownum, rowdata in ledger_rows[:20]:
            cur.execute("INSERT OR IGNORE INTO evidence_registry (item_id, description, file_path, category, status, created_at) VALUES (?,?,?,?,?,?)",
                (f"LEDGER_XLS_{sheet}_{rownum:04d}", f"Sheet:{sheet} Row:{rownum}: {rowdata[:200]}",
                 r"shady enhanced redacted ledger$$$ conv_10aa496d.xls", "financial", "RAW", ts))

    # Write NoReply PDF extracts
    for fn, text in noreply_results.items():
        eid = f"NOREPLY_{fn[:30].replace(' ','_')}"
        cur.execute("INSERT OR IGNORE INTO evidence_registry (item_id, description, file_path, category, status, created_at) VALUES (?,?,?,?,?,?)",
            (eid, f"MiFILE notification PDF: {text[:300]}", os.path.join(SHADY, fn), "court_filing", "EXTRACTED", ts))

    # Write court doc extracts
    for fn, text in court_results.items():
        eid = f"COURT_{fn[:30].replace(' ','_')}"
        cur.execute("INSERT OR IGNORE INTO evidence_registry (item_id, description, file_path, category, status, created_at) VALUES (?,?,?,?,?,?)",
            (eid, f"Court document: {text[:300]}", os.path.join(SHADY, fn), "court_filing", "EXTRACTED", ts))

    conn.commit()
    log(f"BRAIN DB WRITE COMPLETE")

# ─────────────────────────────────────────────
# 8. WRITE VERBATIM QUOTES TO litigation_context.db
# ─────────────────────────────────────────────
def write_quotes_to_main(noreply_results, court_results):
    conn = db_conn(MAIN_DB)
    cur = conn.cursor()
    # Check if evidence_quotes has expected columns
    cols = {r[1] for r in cur.execute("PRAGMA table_info(evidence_quotes)")}
    log(f"evidence_quotes columns: {cols}")
    ts = datetime.now().isoformat()
    
    inserted = 0
    for fn, text in {**noreply_results, **court_results}.items():
        if len(text) < 50:
            continue
        # Try insert with available columns
        try:
            if "quote_text" in cols and "source_file" in cols and "lane" in cols:
                cur.execute("INSERT OR IGNORE INTO evidence_quotes (quote_text, source_file, lane, category, created_at) VALUES (?,?,?,?,?)",
                    (text[:2000], fn, "B", "housing", ts))
                inserted += 1
        except Exception as e:
            log(f"QUOTE INSERT ERROR: {e}")
    conn.commit()
    conn.close()
    log(f"MAIN DB: {inserted} quotes inserted")
    return inserted

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == "__main__":
    log("="*60)
    log("SHADYOAKS-DESTRUCTION WAVE 2 EXTRACTION")
    log("="*60)

    # XLS Ledger
    log("\n--- XLS LEDGER ---")
    ledger_rows = read_xls_ledger()

    # NoReply PDFs
    log("\n--- NOREPLY PDFs ---")
    noreply_results = read_noreply_pdfs()

    # Court PDFs
    log("\n--- COURT DOCUMENTS (shady oaks 2_XXXX.pdf) ---")
    court_results = read_shady_court_pdfs()

    # LARA DOCX
    log("\n--- LARA DOCX ---")
    lara_results = read_lara_docx()

    # Remaining MEEK1 CSVs
    log("\n--- MEEK1 CSVs ---")
    meek_results = read_meek1_csvs()

    # Key TXTs
    log("\n--- KEY TXT FILES ---")
    txt_results = read_key_txts()

    # VanDam screenshot
    log("\n--- VANDM SCREENSHOT ---")
    vandm = read_vandm_screenshot_metadata()

    # Write to brain DB
    log("\n--- DB WRITES ---")
    brain_conn = db_conn(BRAIN_DB)
    write_to_brain(brain_conn, ledger_rows, noreply_results, court_results)
    brain_conn.close()

    # Write to main DB
    write_quotes_to_main(noreply_results, court_results)

    # Summary
    log("\n" + "="*60)
    log("EXTRACTION COMPLETE")
    log(f"  XLS rows: {len(ledger_rows)}")
    log(f"  NoReply PDFs read: {len(noreply_results)}")
    log(f"  Court docs read: {len(court_results)}")
    log(f"  LARA DOCX: {len(lara_results)}")
    log(f"  MEEK1 CSVs: {len(meek_results)}")
    log(f"  Key TXTs: {len(txt_results)}")
    log(f"  VanDam screenshot: {vandm}")
    log("="*60)

    # Save full log
    outpath = r"C:\Users\andre\temp\wave2_extraction_log.txt"
    with open(outpath, "w", encoding="utf-8") as f:
        f.write("\n".join(LOG))
    log(f"LOG SAVED: {outpath}")
