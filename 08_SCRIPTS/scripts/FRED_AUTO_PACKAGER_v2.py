import os
from pathlib import Path
from datetime import datetime
from docx import Document
import shutil

# === CONFIGURATION ===
SCAN_DIR = Path("F:/")  # Change if needed
OUTPUT_DIR = Path("F:/FRED_PACKAGER_OUTPUT")
TODAY = datetime.today().strftime('%B %d, %Y')

# === CATEGORY KEYWORDS & FOLDER MAPPING ===
CATEGORIES = {
    "lease": "01_Required_Forms",
    "complaint": "01_Required_Forms",
    "summons": "01_Required_Forms",
    "motion": "02_Motions_Orders_Declarations",
    "order": "02_Motions_Orders_Declarations",
    "affidavit": "02_Motions_Orders_Declarations",
    "declaration": "02_Motions_Orders_Declarations",
    "exhibit": "03_Attached_Exhibits",
    "photo": "03_Attached_Exhibits",
    "evidence": "03_Attached_Exhibits",
    "service": "04_Service_Proofs",
    "mailing": "04_Service_Proofs"
}

# === ENSURE OUTPUT DIRECTORIES EXIST ===
for subfolder in set(CATEGORIES.values()):
    (OUTPUT_DIR / subfolder).mkdir(parents=True, exist_ok=True)

# === CLASSIFY FILES BASED ON NAME ===
classified_files = []
for root, dirs, files in os.walk(SCAN_DIR):
    for file in files:
        lower_file = file.lower()
        for keyword, folder in CATEGORIES.items():
            if keyword in lower_file:
                source_path = Path(root) / file
                dest_path = OUTPUT_DIR / folder / file
                try:
                    shutil.copy2(source_path, dest_path)
                    classified_files.append((file, folder))
                except Exception as e:
                    print(f"Failed to copy {file}: {e}")
                break

# === AUTO-GENERATE COURT DOCUMENTS ===
doc_templates = {
    "Declaration": "This is a declaration submitted by Andrew Pigors regarding the ongoing civil matter involving Shady Oaks. Dated: ",
    "Affidavit": "Affidavit of Andrew Pigors, affirming the facts stated in the attached complaint as true. Executed on ",
    "Proposed_Order": "IT IS HEREBY ORDERED that Defendant cease all retaliatory actions, and appear to show cause. Dated: ",
    "Certificate_of_Service": "Certificate of Service – I, Andrew Pigors, certify that all documents were served on all parties on "
}

for name, content in doc_templates.items():
    doc = Document()
    doc.add_heading(name.replace("_", " "), level=1)
    doc.add_paragraph(content + TODAY)
    filename = name.replace(" ", "_") + ".docx"
    save_path = OUTPUT_DIR / "02_Motions_Orders_Declarations" / filename
    doc.save(save_path)

# === GENERATE MANIFEST FILE ===
manifest_path = OUTPUT_DIR / "Filing_Manifest.txt"
with open(manifest_path, "w") as f:
    f.write(f"Filing Manifest - Generated {TODAY}\n\n")
    for file, folder in classified_files:
        f.write(f"{file} => {folder}\n")

print(f"Auto-packaging complete. Output folder: {OUTPUT_DIR}")
