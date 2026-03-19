import os
import shutil
from pathlib import Path
from datetime import datetime
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext

# === CONFIGURATION ===
TODAY = datetime.today().strftime('%B %d, %Y')
CATEGORIES = {
    "lease": "01_Required_Forms",
    "complaint": "01_Required_Forms",
    "summons": "01_Required_Forms",
    "motion": "02_Motions_Orders_Declarations",
    "order": "02_Motions_Orders_Declarations",
    "affidavit": "02_Motions_Orders_Declarations",
    "declaration": "02_Motions_Orders_Declarations",
    "proposed": "02_Motions_Orders_Declarations",
    "exhibit": "03_Attached_Exhibits",
    "photo": "03_Attached_Exhibits",
    "evidence": "03_Attached_Exhibits",
    "service": "04_Service_Proofs",
    "mailing": "04_Service_Proofs",
    "certificate": "04_Service_Proofs"
}

class FredAutoPackagerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("FRED AUTO PACKAGER GUI")
        self.root.geometry("700x500")

        self.src_dir = None
        self.output_dir = None

        self.setup_ui()

    def setup_ui(self):
        frame = tk.Frame(self.root)
        frame.pack(pady=10)

        tk.Button(frame, text="Select Source Folder", command=self.select_source).grid(row=0, column=0, padx=10)
        tk.Button(frame, text="Select Output Folder", command=self.select_output).grid(row=0, column=1, padx=10)
        tk.Button(frame, text="Run Packaging", command=self.run_packager).grid(row=0, column=2, padx=10)

        self.log = scrolledtext.ScrolledText(self.root, width=80, height=25)
        self.log.pack(padx=10, pady=10)

    def select_source(self):
        self.src_dir = Path(filedialog.askdirectory())
        self.log.insert(tk.END, f"Selected source: {self.src_dir}\n")

    def select_output(self):
        self.output_dir = Path(filedialog.askdirectory())
        self.log.insert(tk.END, f"Selected output: {self.output_dir}\n")

    def run_packager(self):
        if not self.src_dir or not self.output_dir:
            messagebox.showwarning("Missing Path", "Please select both source and output folders.")
            return

        try:
            self.classify_files()
            self.generate_documents()
            self.build_manifest()
            messagebox.showinfo("Done", "Packaging complete!")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def classify_files(self):
        self.log.insert(tk.END, "📂 Classifying files...\n")
        for folder in set(CATEGORIES.values()):
            (self.output_dir / folder).mkdir(parents=True, exist_ok=True)

        self.classified_files = []
        for root, _, files in os.walk(self.src_dir):
            for file in files:
                for keyword, folder in CATEGORIES.items():
                    if keyword in file.lower():
                        src_path = Path(root) / file
                        dst_path = self.output_dir / folder / file
                        shutil.copy2(src_path, dst_path)
                        self.classified_files.append((file, folder))
                        self.log.insert(tk.END, f"✅ {file} => {folder}\n")
                        break

    def generate_documents(self):
        self.log.insert(tk.END, "📝 Generating legal documents...\n")
        templates = {
            "Declaration": "This is a declaration submitted by Andrew Pigors regarding the ongoing civil matter involving Shady Oaks. Dated: ",
            "Affidavit": "Affidavit of Andrew Pigors, affirming the facts stated in the attached complaint as true. Executed on ",
            "Proposed_Order": "IT IS HEREBY ORDERED that Defendant cease all retaliatory actions, and appear to show cause. Dated: ",
            "Certificate_of_Service": "Certificate of Service – I, Andrew Pigors, certify that all documents were served on all parties on "
        }

        doc_path = self.output_dir / "02_Motions_Orders_Declarations"
        doc_path.mkdir(parents=True, exist_ok=True)

        for name, text in templates.items():
            doc = Document()
            doc.add_heading(name.replace("_", " "), level=1)
            doc.add_paragraph(text + TODAY)
            file_path = doc_path / f"{name}.docx"
            doc.save(file_path)
            self.log.insert(tk.END, f"📄 Created {file_path}\n")

    def build_manifest(self):
        manifest_path = self.output_dir / "Filing_Manifest.txt"
        with open(manifest_path, "w") as f:
            f.write(f"Filing Manifest - Generated {TODAY}\n\n")
            for file, folder in self.classified_files:
                f.write(f"{file} => {folder}\n")
        self.log.insert(tk.END, f"📜 Manifest created at {manifest_path}\n")

if __name__ == "__main__":
    root = tk.Tk()
    app = FredAutoPackagerApp(root)
    root.mainloop()
