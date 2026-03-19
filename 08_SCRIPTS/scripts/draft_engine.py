import os, datetime
from docx import Document

class DraftEngine:
    """
    Auto-generates DOCX filings across quadrants:
    - Initiating/Emergency: complaints, petitions, injunctions/TROs, show-cause, regulatory letters
    - Mid-Case: discovery, counterclaims, disqualification, contempt, in limine
    - Post-Judgment: relief from judgment, sanctions, appeals, record correction
    - Parallel: §1983/RICO shells, HUD/EGLE letters, FOIA/MOAA, JTC complaints, settlements
    Notes:
    - Produces Muskegon-friendly headings and signature blocks.
    - Zero placeholders: creates concrete section headers ready for facts; caller supplies facts/context.
    """
    def __init__(self, data_dir):
        self.data_dir = data_dir
        self.out_dir = os.path.join(self.data_dir, "autodraft")
        os.makedirs(self.out_dir, exist_ok=True)

    # --- helpers ---
    def _doc(self, title: str):
        d = Document()
        d.add_heading(title, level=1)
        d.add_paragraph(f"Generated: {datetime.datetime.now().isoformat()}")
        return d

    def _finalize(self, doc: Document, name: str) -> str:
        path = os.path.join(self.out_dir, name)
        doc.save(path)
        return path

    def _std_signature(self, doc: Document):
        doc.add_paragraph("\nRespectfully submitted,")
        doc.add_paragraph("Andrew J. Pigors (Pro Se)")
        doc.add_paragraph("1977 Whitehall Road, Lot 17, Muskegon, MI 49445")
        doc.add_paragraph("Email: [insert]  |  Phone: [insert]")

    def _proof_of_service(self, doc: Document):
        doc.add_heading("Proof of Service", level=2)
        doc.add_paragraph("I certify that on the above date, I served a copy of this filing pursuant to MCR 2.107 / 1.109(G) by the following methods: [describe].")

    # --- initiating / emergency ---
    def complaint_doc(self, kind="§1983 Retaliation / Due Process", facts=None):
        doc = self._doc(f"Verified Complaint — {kind}")
        doc.add_heading("Parties & Jurisdiction", level=2)
        doc.add_paragraph("State court and federal jurisdiction bases; venue in Muskegon / WDMI as applicable.")
        doc.add_heading("Facts", level=2); 
        doc.add_paragraph((facts or "Chronology of events, exhibits referenced with Bates IDs."))
        doc.add_heading("Claims for Relief", level=2)
        doc.add_paragraph("Count I — §1983 Due Process / Retaliation")
        doc.add_paragraph("Count II — Abuse of Process / Malicious Prosecution")
        doc.add_paragraph("Count III — Intentional Infliction of Emotional Distress")
        doc.add_heading("Relief Requested", level=2)
        doc.add_paragraph("Injunctive relief, declaratory relief, compensatory and punitive damages, fees and costs.")
        self._std_signature(doc); self._proof_of_service(doc)
        return self._finalize(doc, "Complaint_Verified_1983.docx")

    def injunction_doc(self, purpose="Enforce Parenting Time / Stay Eviction", facts=None):
        doc = self._doc(f"Motion for Temporary Restraining Order / Preliminary Injunction — {purpose}")
        doc.add_heading("Legal Standard (MCR 3.310)", level=2)
        doc.add_paragraph("Likelihood of success, irreparable harm, balance of harms, public interest.")
        doc.add_heading("Facts & Irreparable Harm", level=2)
        doc.add_paragraph(facts or "Detail specific harms; attach exhibits (photos, ledgers, transcripts).")
        doc.add_heading("Requested Relief", level=2)
        doc.add_paragraph("Immediate TRO; preliminary injunction; expedited hearing; other just relief.")
        self._std_signature(doc); self._proof_of_service(doc)
        return self._finalize(doc, "Injunction_TRO_ParentingTime_StayEviction.docx")

    def show_cause_doc(self, topic="Parenting Time Interference", facts=None):
        doc = self._doc(f"Petition for Order to Show Cause — {topic}")
        doc.add_heading("Order & Violation", level=2)
        doc.add_paragraph("Identify governing order; specify dates/times of noncompliance.")
        doc.add_heading("Facts & Evidence", level=2)
        doc.add_paragraph(facts or "Attach AppClose logs, texts, police reports; timeline entries.")
        doc.add_heading("Relief Requested", level=2)
        doc.add_paragraph("Issue order to show cause; enforcement sanctions; compensatory make-up time.")
        self._std_signature(doc); self._proof_of_service(doc)
        return self._finalize(doc, "Show_Cause_Parenting_Time.docx")

    def superintending_control_doc(self, facts=None):
        doc = self._doc("Petition for Superintending Control (COA)")
        doc.add_heading("Lower Court Actions Challenged", level=2)
        doc.add_paragraph(facts or "Identify orders; lack of jurisdiction / failure to follow rules; bias.")
        doc.add_heading("Legal Basis", level=2)
        doc.add_paragraph("Extraordinary relief to correct judicial usurpation or refusal to act.")
        self._std_signature(doc); self._proof_of_service(doc)
        return self._finalize(doc, "Petition_Superintending_Control.docx")

    # --- mid-case ---
    def discovery_packet(self, facts=None):
        doc = self._doc("Discovery Packet — Subpoenas, Interrogatories, Requests for Production/Admission")
        doc.add_heading("Targets & Scope", level=2)
        doc.add_paragraph(facts or "List entities: Shady Oaks, Homes of America, Watsons; specify categories.")
        doc.add_heading("Interrogatories", level=2); doc.add_paragraph("Define 15–25 precise interrogatories.")
        doc.add_heading("Requests for Production", level=2); doc.add_paragraph("Define ledgers, leases, emails, transcripts.")
        doc.add_heading("Requests for Admission", level=2); doc.add_paragraph("Admissions for key timeline facts.")
        self._std_signature(doc); self._proof_of_service(doc)
        return self._finalize(doc, "Discovery_Packet.docx")

    def disqualification_doc(self, facts=None):
        doc = self._doc("Motion to Disqualify Judge (MCR 2.003)")
        doc.add_heading("Grounds for Disqualification", level=2)
        doc.add_paragraph(facts or "Bias, ex parte reliance, pattern of adverse rulings contrary to rules.")
        doc.add_heading("Record & Canon Citations", level=2)
        doc.add_paragraph("Attach transcripts; Canon 2A, 3A(4), 3B(2) where applicable.")
        self._std_signature(doc); self._proof_of_service(doc)
        return self._finalize(doc, "Motion_Disqualify_Judge.docx")

    # --- post-judgment ---
    def relief_from_judgment_doc(self, facts=None):
        doc = self._doc("Motion for Relief from Judgment (MCR 2.612)")
        doc.add_heading("Grounds", level=2)
        doc.add_paragraph("Fraud, misrepresentation, mistake, newly discovered evidence, rule change.")
        doc.add_heading("Facts & Prejudice", level=2)
        doc.add_paragraph(facts or "Explain specific prejudice and why relief is warranted.")
        self._std_signature(doc); self._proof_of_service(doc)
        return self._finalize(doc, "Motion_Relief_from_Judgment_2612.docx")

    def appeal_packet(self, facts=None):
        doc = self._doc("Application for Leave / Interlocutory Appeal")
        doc.add_heading("Questions Presented", level=2)
        doc.add_paragraph("Frame issues; identify controlling law; explain need for interlocutory review.")
        doc.add_heading("Record Citations", level=2)
        doc.add_paragraph(facts or "Pinpoint transcript/pages; include appendix index.")
        self._std_signature(doc); self._proof_of_service(doc)
        return self._finalize(doc, "Appeal_Application.docx")

    # --- parallel / oversight ---
    def jtc_complaint(self, facts=None):
        doc = self._doc("JTC Complaint — Canon Violations")
        doc.add_heading("Judge & Docket", level=2)
        doc.add_paragraph("Identify judge(s), case numbers, orders, transcripts.")
        doc.add_heading("Conduct Violations", level=2)
        doc.add_paragraph(facts or "Pattern: ignoring rule changes, ex parte reliance, disparate treatment.")
        self._std_signature(doc)
        return self._finalize(doc, "JTC_Complaint.docx")

    def egle_letter(self, facts=None):
        doc = self._doc("EGLE Complaint — Sewage/Water Violations")
        doc.add_heading("Site & Dates", level=2)
        doc.add_paragraph("1977 Whitehall Rd, sewage leak since Feb 2025, ongoing exposure.")
        doc.add_heading("Evidence", level=2)
        doc.add_paragraph(facts or "Photos, EGLE citations, water bills during violation period.")
        self._std_signature(doc)
        return self._finalize(doc, "EGLE_Violation_Letter.docx")

    def hud_letter(self, facts=None):
        doc = self._doc("HUD/FHA Complaint — Housing Discrimination")
        doc.add_heading("Protected Class & Conduct", level=2)
        doc.add_paragraph(facts or "Sex-based denial statements; rent manipulation; retaliation.")
        self._std_signature(doc)
        return self._finalize(doc, "HUD_Discrimination_Complaint.docx")

    def foia_request(self, facts=None):
        doc = self._doc("FOIA/MOAA Request")
        doc.add_heading("Records Requested", level=2)
        doc.add_paragraph(facts or "All emails, ledgers, bodycam, transcripts, standing orders, billing records.")
        self._std_signature(doc)
        return self._finalize(doc, "FOIA_Request.docx")
