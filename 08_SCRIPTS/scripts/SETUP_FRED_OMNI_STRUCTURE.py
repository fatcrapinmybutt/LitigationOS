
import os
from pathlib import Path
from docx import Document

base_path = Path("F:/OMNILITIGATION_SYSTEM")

# Required directories
directories = [
    base_path / "VerifiedComplaints",
    base_path / "ProposedOrders",
    base_path / "MotionsToShowCause",
    base_path / "AffidavitDrafts",
    base_path / "MiFILE_READY"
]

# Create folders
for folder in directories:
    folder.mkdir(parents=True, exist_ok=True)

# Create a dummy affidavit
affidavit_path = base_path / "AffidavitDrafts" / "Affidavit_Dummy.docx"
if not affidavit_path.exists():
    doc = Document()
    doc.add_heading("Affidavit of Truth", level=1)
    doc.add_paragraph("This is a test affidavit used to validate the signature page injector.")
    doc.save(affidavit_path)

# Create dummy files for complaints, orders, and motions
(Path(base_path / "VerifiedComplaints" / "Complaint_Dummy.docx")).write_text("Verified Complaint - Test")
(Path(base_path / "ProposedOrders" / "Order_Dummy.docx")).write_text("Proposed Order - Test")
(Path(base_path / "MotionsToShowCause" / "Motion_Dummy.docx")).write_text("Motion to Show Cause - Test")

print("✅ Folder structure and dummy files created.")
