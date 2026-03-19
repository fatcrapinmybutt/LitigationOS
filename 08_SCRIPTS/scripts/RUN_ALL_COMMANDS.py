
import json
import subprocess
import os
from datetime import datetime
from docx import Document

MODULES = {
    "Motion Packet Generator": "./_ENGINES/Motion_Packet_Generator_Engine/module_task.py",
    "Violation Response Engine": "./_ENGINES/Violation_Response_Engine/module_task.py",
    "Evidence Weighing Core": "./_CORES/Evidence_Weighing_Core/module_task.py",
    "Contempt Trigger Workflow": "./_WORKFLOWS/Contempt_Trigger_Workflow/module_task.py",
    "AppClose Analyzer Plugin": "./_PLUGINS/AppClose_Analyzer_Plugin/module_task.py",
    "Trigger Heatmap Dashboard": "./_DASHBOARDS/Trigger_Heatmap_Dashboard/module_task.py",
    "Statute Breach Pipeline": "./_PIPELINES/Statute_Breach_Pipeline/module_task.py",
    "Case Masterlist Core": "./_MASTERLISTS/Case_Masterlist/module_task.py",
    "Test Simulation System": "./_SYSTEMS/Test_Simulation_System/module_task.py",
    "Automation Trigger Daemon": "./_SYSTEMS/Automation_Trigger_Daemon/module_task.py",
    "Master Loader": "./FRED_CASE_MASTER_LOADER.py",
    "Evidence Report Generator": "./Generate_Evidence_Report.py",
    "Plugin Signature Validator": "./_PLUGINS/Plugin_Signature_Validator.py"
}

def run_modules_from_preset(preset_path="default_preset.json"):
    if not os.path.exists(preset_path):
        print("⚠ No preset found.")
        return []
    
    with open(preset_path, 'r') as f:
        selected = json.load(f)

    execution_log = []
    for name in selected:
        module_path = MODULES.get(name)
        if not module_path:
            execution_log.append(f"⚠ {name} not found")
            continue
        try:
            subprocess.run(["python", module_path], check=True)
            execution_log.append(f"✓ Executed: {name}")
        except Exception as e:
            execution_log.append(f"⚠ Failed: {name}")
    return execution_log

def generate_exhibit_fusion_doc(log, output_path="FUSION_OUTPUT.docx"):
    doc = Document()
    doc.add_heading("FRED-PRIME Exhibit Fusion Report", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("This document was created using Exhibit Fusion Mode. The following modules were executed:")

    for entry in log:
        doc.add_paragraph(f"• {entry}")

    doc.save(output_path)
    print(f"📄 Fusion document saved to {output_path}")

if __name__ == "__main__":
    log = run_modules_from_preset()
    if log:
        generate_exhibit_fusion_doc(log)
