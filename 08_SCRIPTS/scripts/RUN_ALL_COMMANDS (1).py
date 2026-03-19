
import json
import subprocess
import os
import hashlib
from datetime import datetime
from docx import Document

HASHMAP_PATH = "Module_HashMap.json"
LOG_PATH = "Execution_Audit_Log.txt"

with open(HASHMAP_PATH, 'r') as f:
    HASHMAP = json.load(f)

def hash_file(filepath):
    with open(filepath, 'rb') as f:
        return hashlib.sha256(f.read()).hexdigest()

def verify_integrity(filepath, expected_hash):
    return hash_file(filepath) == expected_hash

def log_event(entry):
    with open(LOG_PATH, 'a') as log:
        log.write(entry + "\n")

def run_modules_from_preset(preset_path="default_preset.json"):
    if not os.path.exists(preset_path):
        print("⚠ No preset found.")
        return []

    with open(preset_path, 'r') as f:
        selected = json.load(f)

    log = []
    for module_key, meta in HASHMAP.items():
        module_name = os.path.basename(os.path.dirname(module_key))
        if module_name not in selected:
            continue

        file_path = os.path.join(".", module_key)
        status = "✓ Verified" if verify_integrity(file_path, meta['sha256']) else "⚠ HASH MISMATCH"

        try:
            subprocess.run(["python", file_path], check=True)
            result = f"{datetime.now()} | {module_name} | {status} | EXECUTED"
        except Exception as e:
            result = f"{datetime.now()} | {module_name} | {status} | FAILED"
        log.append(result)
        log_event(result)

    return log

def generate_audit_doc(log, output_path="Audit_Report.docx"):
    doc = Document()
    doc.add_heading('ULTRA-FRED-PRIME Execution & Integrity Audit', 0)
    doc.add_paragraph(f"Audit Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Log Entries:")

    for entry in log:
        doc.add_paragraph(f"• {entry}")

    doc.save(output_path)
    print(f"📄 Audit report saved to: {output_path}")

if __name__ == "__main__":
    results = run_modules_from_preset()
    if results:
        generate_audit_doc(results)
