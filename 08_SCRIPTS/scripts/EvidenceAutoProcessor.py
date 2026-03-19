# EvidenceAutoProcessor: Classifies, Builds Affidavits, and Generates Court-Ready ZIPs

import os
import json
import hashlib
import shutil
from datetime import datetime
from docx import Document
from zipfile import ZipFile

# === CONFIG PATHS ===
BASE_DIR = "F:/OMNILITIGATION_SYSTEM"
INPUT_DIR = f"{BASE_DIR}/EvidenceRepository/NotYetReviewed"
CLASSIFIED_DIR = f"{BASE_DIR}/EvidenceRepository/ByLegalType"
AFFIDAVIT_DIR = f"{BASE_DIR}/FilingGenerators/Affidavits_Declarations"
ZIP_DIR = f"{BASE_DIR}/FilingBundles/ZIP_CourtBundles"
MATRIX_PATH = f"{BASE_DIR}/MasterEvidenceMatrix/evidence_matrix.json"

# === HELPER FUNCTIONS ===
def sha256_of_file(path):
    h = hashlib.sha256()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b""):
            h.update(chunk)
    return h.hexdigest()

def classify_file(file_path):
    fname = os.path.basename(file_path).lower()
    if "lease" in fname:
        return "Leases_Contracts"
    elif "bill" in fname or "utility" in fname:
        return "Utility_Ledgers"
    elif "order" in fname or "judgment" in fname:
        return "CourtOrders_Judgments"
    elif fname.endswith(".jpg") or fname.endswith(".png"):
        return "Photos_MetadataAnchored"
    elif "appclose" in fname or "text" in fname:
        return "Communications_Text_Email_AppClose"
    elif "affidavit" in fname or "sworn" in fname:
        return "Affidavits_SwornStatements"
    else:
        return "EvidenceUnclassified_AuditQueue"

def build_affidavit(file_entry):
    doc = Document()
    doc.add_heading("Affidavit of Authenticity", 0)
    doc.add_paragraph(f"I, Andrew J Pigors, attest that the attached document "
                      f"was created or received on {file_entry['timestamp']} and "
                      f"is a true and accurate record titled '{file_entry['file']}'.")
    doc.add_paragraph(f"SHA-256 Hash: {file_entry['sha256']}")
    doc.add_paragraph("Signed under penalty of perjury. Date: " + datetime.today().strftime('%Y-%m-%d'))
    affidavit_path = os.path.join(AFFIDAVIT_DIR, os.path.basename(file_entry['file']) + ".affidavit.docx")
    doc.save(affidavit_path)
    return affidavit_path

def zip_court_bundle(bundle_name, file_paths):
    zip_path = os.path.join(ZIP_DIR, f"{bundle_name}.zip")
    with ZipFile(zip_path, 'w') as zipf:
        for path in file_paths:
            zipf.write(path, arcname=os.path.basename(path))
    return zip_path

# === MAIN ===
if not os.path.exists(MATRIX_PATH):
    evidence_matrix = []
else:
    with open(MATRIX_PATH, 'r') as f:
        evidence_matrix = json.load(f)

file_list = [os.path.join(dp, f) for dp, dn, filenames in os.walk(INPUT_DIR) for f in filenames]
bundle_files = []

for fpath in file_list:
    sha = sha256_of_file(fpath)
    timestamp = datetime.fromtimestamp(os.path.getmtime(fpath)).isoformat()
    category = classify_file(fpath)

    dest_folder = os.path.join(CLASSIFIED_DIR, category)
    os.makedirs(dest_folder, exist_ok=True)
    dest_path = os.path.join(dest_folder, os.path.basename(fpath))
    shutil.copy2(fpath, dest_path)

    matrix_entry = {"file": dest_path, "sha256": sha, "timestamp": timestamp, "category": category}
    evidence_matrix.append(matrix_entry)

    aff_path = build_affidavit(matrix_entry)
    bundle_files.extend([dest_path, aff_path])

# Save updated matrix
with open(MATRIX_PATH, 'w') as mf:
    json.dump(evidence_matrix, mf, indent=2)

# Generate the ZIP
bundle_path = zip_court_bundle("AutoLitigationBundle", bundle_files)
print(f"✅ Court-ready ZIP generated: {bundle_path}")
