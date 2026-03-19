import os
import shutil
from pathlib import Path
from datetime import datetime
from docx import Document

# =============================
# 🧠 FRED_AUTO_PACKAGER v3.0
# Final Audited Version
# =============================

# === USER CONFIGURATION ===
SCAN_DIR = Path("F:/")  # Change if needed for production
OUTPUT_DIR = Path("F:/FRED_PACKAGER_OUTPUT")
TODAY = datetime.today().strftime('%B %d, %Y')

# === FILE CLASSIFICATION MAP ===
CATEGORIES = {
    "lease": "01_Required_Forms",
    "complaint": "01_Required_Forms",
    "summons": "01_Required_Forms",
    "motion": "02_Motions_Orders_Declarations",
    "order": "02_Motions_Orders_Declarations",
    "affidavit": "02_Motions_Orders_Declarations",
    "declaration": "02_Motions_Orders_Declarations",
    "proposed": "02_Motions_Orders_Declarations",
    "exhibit": "03_Attached_Exhibits",
    "photo": "03_Attached_Exhibits",
    "evidence": "03_Attached_Exhibits",
    "service": "04_Service_Proofs",
    "mailing": "04_Service_Proofs",
    "certificate": "04_Service_Proofs"
}

# === ENSURE OUTPUT DIRECTORIES EXIST ===
print("🛠️ Creating folder structure...")
for subfolder in set(CATEGORIES.values()):
    path = OUTPUT_DIR / subfolder
    path.mkdir(parents=True, exist_ok=True)

# === SCAN AND CLASSIFY FILES BASED ON NAME ===
classified_files = []
print("📂 Scanning files and classifying based on keywords...")
for root, _, files in os.walk(SCAN_DIR):
    for file in files:
        file_path = Path(root) / file
        lower_file = file.lower()

        for keyword, folder in CATEGORIES.items():
            if keyword in lower_file:
                dest_path = OUTPUT_DIR / folder / file
                try:
                    shutil.copy2(file_path, dest_path)
                    classified_files.append((file, folder))
                    print("✅ Classified: {} => {}".format(file, folder))
                except Exception as error:
                    print("⚠️ Failed to copy {}: {}".format(file, error))
                break  # Only match first keyword

# === AUTO-GENERATE LEGAL DOCUMENTS ===
print("📄 Generating pre-filled legal documents...")

doc_templates = {
    "Declaration": "This is a declaration submitted by Andrew Pigors regarding the ongoing civil matter involving Shady Oaks. Dated: ",
    "Affidavit": "Affidavit of Andrew Pigors, affirming the facts stated in the attached complaint as true. Executed on ",
    "Proposed_Order": "IT IS HEREBY ORDERED that Defendant cease all retaliatory actions, and appear to show cause. Dated: ",
    "Certificate_of_Service": "Certificate of Service – I, Andrew Pigors, certify that all documents were served on all parties on "
}

for name, text in doc_templates.items():
    try:
        doc = Document()
        doc.add_heading(name.replace("_", " "), level=1)
        doc.add_paragraph(text + TODAY)
        filename = name + ".docx"
        save_path = OUTPUT_DIR / "02_Motions_Orders_Declarations" / filename
        doc.save(save_path)
        print("📝 Created: {}".format(filename))
    except Exception as error:
        print("⚠️ Error generating {}: {}".format(name, error))

# === CREATE MANIFEST ===
manifest_path = OUTPUT_DIR / "Filing_Manifest.txt"
with open(manifest_path, "w") as f:
    f.write("Filing Manifest - Generated {}

".format(TODAY))
    for file, folder in classified_files:
        f.write("{} => {}
".format(file, folder))

print("📄 Filing Manifest created at {}".format(manifest_path))
print("✅ FRED_AUTO_PACKAGER v3.0 complete. Output is ready at: {}".format(OUTPUT_DIR))
