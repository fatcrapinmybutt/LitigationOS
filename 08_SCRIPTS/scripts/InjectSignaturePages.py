
from pathlib import Path
from docx import Document

# Folder containing affidavit drafts
affidavit_folder = Path("F:/OMNILITIGATION_SYSTEM/AffidavitDrafts")
output_folder = Path("F:/OMNILITIGATION_SYSTEM/MiFILE_READY/SignedAffidavits")
output_folder.mkdir(parents=True, exist_ok=True)

signature_block = '''
_______________________________
Signature of Affiant

Subscribed and sworn to before me this _____ day of ____________, 20____.

_______________________________
Notary Public

My Commission Expires: _______________
'''

for file in affidavit_folder.glob("*.docx"):
    doc = Document(file)
    doc.add_page_break()
    doc.add_paragraph("Signature and Notarization", style="Heading 2")
    doc.add_paragraph(signature_block)
    output_path = output_folder / file.name
    doc.save(output_path)

print(f"🧾 Signature pages injected and saved to: {output_folder}")
