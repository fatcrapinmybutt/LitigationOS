import os, json, datetime
from typing import Dict, Any, List
try:
    from docx import Document
except Exception:
    Document = None

class AutoDraftEngine:
    def __init__(self):
        self.root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.data_dir = os.path.join(self.root, "..", "..", "data", "drafts")
        os.makedirs(self.data_dir, exist_ok=True)

    def _safe(self, s: str) -> str:
        return "".join(c for c in s if c.isalnum() or c in ("_", "-", " ")).strip().replace(" ", "_")

    def _write_docx(self, path: str, title: str, sections: List[Dict[str, str]]):
        if Document is None:
            with open(path.replace(".docx",".txt"), "w", encoding="utf-8") as f:
                f.write("# " + title + "\n\n")
                for sec in sections:
                    f.write(sec["heading"] + "\n")
                    f.write(sec["body"] + "\n\n")
            return path.replace(".docx", ".txt")
        doc = Document()
        doc.add_heading(title, level=1)
        for sec in sections:
            doc.add_heading(sec["heading"], level=2)
            doc.add_paragraph(sec["body"])
        doc.save(path)
        return path

    def build(self, action: str, context: Dict[str, Any] | None = None) -> Dict[str, Any]:
        context = context or {}
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = self._safe(action or "Action")
        out_dir = os.path.join(self.data_dir, f"{safe}_{ts}")
        os.makedirs(out_dir, exist_ok=True)

        facts = context.get("facts","Evidence index cited; attach ledger and exhibits.")
        relief = context.get("relief","Relief requested consistent with Michigan law and local practice.")
        authority = context.get("authority","Filed under applicable Michigan Court Rules and Benchbooks.")

        draft_sections = [
            {"heading":"Allegations", "body": "\n".join(context.get("allegations", [])) or "Specific allegations derived from graph and records are incorporated herein."},
            {"heading":"Exhibits Index", "body": "\n".join([f"{i+1}. {x}" for i,x in enumerate(context.get("exhibits", []))]) or "Exhibits attached as labeled."},
            {"heading":"Clause — Service", "body": context.get("clauses", {}).get("service","")},
            {"heading":"Clause — Parenting Time", "body": context.get("clauses", {}).get("parenting_time","")},
            {"heading":"Clause — Motion Standard", "body": context.get("clauses", {}).get("motion_standard","")},
        
            {"heading":"Caption & Parties", "body": context.get("caption","Muskegon County — form-conforming caption.")},
            {"heading":"Introduction", "body": context.get("intro", f"This filing arises from documented harms and seeks lawful relief ({action}).")},
            {"heading":"Factual Background", "body": facts},
            {"heading":"Legal Grounds", "body": authority},
            {"heading":"Requested Relief", "body": relief},
            {"heading":"Verification/Affidavit", "body":"Verification to be signed with correct venue and date; attach exhibits by index and bates."}
        ]

        po_sections = [{"heading":"Order","body": f"The Court having considered the {action}, IT IS ORDERED that appropriate relief is granted according to law."}]
        pos_sections = [{"heading":"Service","body":"List parties/counsel and method/time of service."}]
        forms_list = context.get("forms", [])
        forms_body = "Required forms:\n" + "\n".join(f"- {f}" for f in forms_list) if forms_list else "Required forms: (none)"
        forms_sections = [{"heading":"Forms Index", "body": forms_body}]

        draft_path = os.path.join(out_dir, "Draft.docx")
        order_path = os.path.join(out_dir, "ProposedOrder.docx")
        pos_path = os.path.join(out_dir, "ProofOfService.docx")
        forms_path = os.path.join(out_dir, "Forms.docx")

        draft_written = self._write_docx(draft_path, f"{action} — Draft", draft_sections)
        order_written = self._write_docx(order_path, f"{action} — Proposed Order", po_sections)
        pos_written = self._write_docx(pos_path, f"{action} — Proof of Service", pos_sections)
        forms_written = self._write_docx(forms_path, f"{action} — Forms", forms_sections)

        import zipfile
        zip_path = os.path.join(out_dir, f"{safe}_bundle.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
            for p in [draft_written, order_written, pos_written, forms_written]:
                z.write(p, arcname=os.path.basename(p))

        return {"action": action, "dir": out_dir, "files": [draft_written, order_written, pos_written, forms_written], "zip": zip_path}
