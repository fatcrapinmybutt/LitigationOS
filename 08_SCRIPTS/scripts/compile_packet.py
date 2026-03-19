import os
from docx import Document
from datetime import datetime

base_dir = "F:/FRED-PRIME"
motion_path = os.path.join(base_dir, "MotionTemplates", "motion.docx")
order_path = os.path.join(base_dir, "ProposedOrders", "order.docx")
certificate_path = os.path.join(base_dir, "CertificatesOfService", "certificate.docx")
signature_path = os.path.join(base_dir, "MotionTemplates", "signature.docx")

output_dir = os.path.join(base_dir, "Compiled_Motions")
os.makedirs(output_dir, exist_ok=True)

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
output_file = os.path.join(output_dir, f"Filing_Packet_{timestamp}.docx")

def append_doc(target_doc, path):
    src = Document(path)
    for element in src.element.body:
        target_doc.element.body.append(element)

try:
    final_doc = Document()

    # Cover Page
    final_doc.add_heading("FRED-PRIME Filing Packet", level=1)
    final_doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    final_doc.add_page_break()

    # Append motion, order, certificate, signature
    for section in [motion_path, order_path, certificate_path, signature_path]:
        if os.path.exists(section):
            append_doc(final_doc, section)
            final_doc.add_page_break()
        else:
            final_doc.add_paragraph(f"WARNING: Missing file – {section}")
            final_doc.add_page_break()

    final_doc.save(output_file)
    print(f"Packet created successfully: {output_file}")
except Exception as e:
    print(f"ERROR: Could not compile packet. Details: {str(e)}")
