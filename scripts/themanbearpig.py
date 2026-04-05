#!/usr/bin/env python3
"""THEMANBEARPIG v15.0 -- Self-Organizing Intelligence Graph.

Merges PROJECT KRAKEN evidence hunting + MBP Brain v5.0 graph intelligence
+ filing generation + adversary analytics into a single pywebview desktop app.

Usage:
    python -I scripts/themanbearpig.py
    python -I scripts/themanbearpig.py --debug --hunt --export
"""

import argparse
import hashlib
import gzip
import http.server
import io
import json as _json
import mimetypes
import os
import random
import re
import shutil
import socket
import sqlite3
import subprocess
import sys
import threading
import time
from datetime import date, datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# Paths -- PyInstaller-aware (sys._MEIPASS for bundled exe, normal for dev)
# ---------------------------------------------------------------------------
if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
    _BUNDLE = Path(sys._MEIPASS)
    REPO_ROOT = Path(sys.executable).resolve().parent
else:
    _BUNDLE = Path(__file__).resolve().parent.parent
    REPO_ROOT = _BUNDLE
BRAIN_DB = REPO_ROOT / "mbp_brain.db"
LIT_DB = REPO_ROOT / "litigation_context.db"
VIS_DIR = _BUNDLE / "08_MEDIA" / "MANBEARPIG_V15"
VIS_DIR_V9 = _BUNDLE / "08_MEDIA" / "MANBEARPIG_V9"
VIS_DIR_V5 = _BUNDLE / "08_MEDIA" / "MANBEARPIG_V5"
GRAPH_JSON = VIS_DIR / "graph_clusters.json"
if not GRAPH_JSON.exists():
    GRAPH_JSON = VIS_DIR_V9 / "graph_data.json"
if not GRAPH_JSON.exists():
    GRAPH_JSON = VIS_DIR_V5 / "graph_data.json"
EXPORT_SCRIPT = REPO_ROOT / "scripts" / "export_brain_d3.py"
EVOLVE_SCRIPT = REPO_ROOT / "scripts" / "brain_evolution.py"
KRAKEN_SCRIPT = REPO_ROOT / "07_CODE" / "PROJECT_KRAKEN" / "kraken.py"
FILING_SCRIPT = REPO_ROOT / "scripts" / "generate_filing.py"
COURT_FEED_SCRIPT = REPO_ROOT / "scripts" / "court_feed.py"
DOSSIER_DIR = REPO_ROOT / "04_ANALYSIS" / "ADVERSARY_TRACKS"

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
VERSION = "15.0.0"
SEPARATION_DATE = date(2025, 7, 29)
APP_BG = "#0a0a0f"
APP_WIDTH, APP_HEIGHT = 1920, 1080
APP_MIN_W, APP_MIN_H = 1280, 720

_PRAGMAS = (
    "PRAGMA busy_timeout = 60000",
    "PRAGMA journal_mode = WAL",
    "PRAGMA cache_size = -32000",
    "PRAGMA synchronous = NORMAL",
    "PRAGMA temp_store = MEMORY",
    "PRAGMA mmap_size = 268435456",  # 256 MB — brain DB is 310 MB on NVMe
)

# FTS5 sanitizer
_FTS_CLEAN = re.compile(r'[^\w\s*"]')

# Read-only guard
_WRITE_RE = re.compile(
    r"\b(INSERT|UPDATE|DELETE|DROP|CREATE|ALTER|REPLACE|ATTACH|DETACH|VACUUM|REINDEX)\b",
    re.IGNORECASE,
)


def _find_python():
    """Return path to the real Python interpreter, not the frozen exe.

    When running as a PyInstaller bundle, sys.executable points to the
    .exe itself.  Subprocess calls that need to run .py scripts must
    use the system Python instead.
    """
    if getattr(sys, "frozen", False):
        found = shutil.which("python")
        if found:
            return found
        # Fallback: common install location on this machine
        fallback = Path(r"C:\Users\andre\AppData\Local\Programs\Python\Python312\python.exe")
        if fallback.exists():
            return str(fallback)
        return "python"
    return sys.executable


_PYTHON = _find_python()

# Adversary patterns (from kraken.py)
ADVERSARIES = {
    "Emily Watson": r"(?i)\bemily\b.*\bwatson\b|\bwatson\b.*\bemily\b|\bemily\s+a\.?\s+watson\b",
    "Judge McNeill": r"(?i)\bmcneill\b|\bmcneil\b",
    "Pamela Rusco": r"(?i)\brusco\b|\bpamela\b.*\brusco\b",
    "Albert Watson": r"(?i)\balbert\b.*\bwatson\b|\bwatson\b.*\balbert\b",
    "Lori Watson": r"(?i)\blori\b.*\bwatson\b",
    "Ronald Berry": r"(?i)\bronald\b.*\bberry\b|\bron\b.*\bberry\b",
    "Cavan Berry": r"(?i)\bcavan\b.*\bberry\b",
    "Jennifer Barnes": r"(?i)\bbarnes\b.*\bjennifer\b|\bjennifer\b.*\bbarnes\b|\bP55406\b",
    "Kenneth Hoopes": r"(?i)\bhoopes\b|\bkenneth\b.*\bhoopes\b",
    "Maria Ladas-Hoopes": r"(?i)\bladas[\\s-]*hoopes\b|\bmaria\b.*\bladas\b",
    "FOC": r"(?i)\bfriend\s+of\s+(the\s+)?court\b|\bFOC\b",
    "Shady Oaks": r"(?i)\bshady\s*oaks\b",
}
ADVERSARY_RE = {k: re.compile(v) for k, v in ADVERSARIES.items()}

# Evidence categories
EVIDENCE_CATEGORIES = {
    "custody": re.compile(r"(?i)\bcustody\b|\bparenting\s*time\b|\bvisitation\b|\bbest\s+interest\b"),
    "PPO": re.compile(r"(?i)\bprotection\s*order\b|\bPPO\b|\brestraining\b|\bstalking\b"),
    "judicial": re.compile(r"(?i)\bjudicial\b|\bbias\b|\bex\s*parte\b|\brecusal\b|\bdisqualif"),
    "housing": re.compile(r"(?i)\beviction\b|\btenant\b|\blandlord\b|\bhousing\b|\bmobile\s*home\b"),
    "criminal": re.compile(r"(?i)\bcontempt\b|\bjail\b|\bincarcerat\b|\barrest\b|\bsentenc"),
    "financial": re.compile(r"(?i)\bchild\s*support\b|\bfiling\s*fee\b|\bdamages\b|\bgarnish"),
    "police": re.compile(r"(?i)\bpolice\b|\bNSPD\b|\bofficer\b|\bincident\s*report\b"),
    "medical": re.compile(r"(?i)\bhealthwest\b|\bmental\s*health\b|\bpsych\b|\bmedication\b|\bLOCUS\b"),
}

# Legal authority patterns
LEGAL_PATTERNS = {
    "MCR": re.compile(r"\bMCR\s+\d+\.\d+\w*"),
    "MCL": re.compile(r"\bMCL\s+\d+\.\d+\w*"),
    "MRE": re.compile(r"\bMRE\s+\d+\w*"),
    "USC": re.compile(r"\b\d+\s+U\.?S\.?C\.?\s*[S\s]*\d+"),
    "Case_Law": re.compile(
        r"\b\d+\s+Mich\.?\s+(App\.?\s+)?\d+|\b\d+\s+F\.\d[a-z]*\s+\d+"
    ),
}

# Focus mode boosters
FOCUS_BOOSTS = {
    "adversary": ["Emily Watson", "Albert Watson", "Lori Watson", "Ronald Berry", "Cavan Berry"],
    "judicial": ["Judge McNeill", "Kenneth Hoopes", "Maria Ladas-Hoopes", "Cavan Berry", "Pamela Rusco"],
    "housing": ["Shady Oaks"],
    "custody": ["Emily Watson", "FOC", "Pamela Rusco", "Albert Watson"],
    "ppo": ["Emily Watson", "Judge McNeill", "Ronald Berry"],
    "legal": [],
    "all": [],
}

EXTS = {".pdf", ".txt", ".csv", ".html", ".json", ".docx", ".md"}
MAX_CONTENT_BYTES = 500_000


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _sanitize_fts(query: str) -> str:
    """Strip dangerous FTS5 metacharacters."""
    return _FTS_CLEAN.sub(" ", query).strip()


def _safe_float(val, default=0.0):
    """Coerce a value to float, returning *default* on any failure."""
    if val is None:
        return default
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


def _sep_days() -> int:
    return (date.today() - SEPARATION_DATE).days


def _find_free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("127.0.0.1", 0))
        return s.getsockname()[1]


def _rows_to_dicts(rows):
    """Convert sqlite3.Row list to plain dicts for JSON serialization."""
    return [dict(r) for r in rows] if rows else []


def _connect_brain():
    """Open WAL-mode connection to mbp_brain.db."""
    if not BRAIN_DB.exists():
        return None
    conn = sqlite3.connect(str(BRAIN_DB))
    conn.row_factory = sqlite3.Row
    conn.text_factory = lambda b: b.decode("utf-8", errors="replace")
    for p in _PRAGMAS:
        conn.execute(p)
    return conn


def _connect_lit():
    """Open WAL-mode connection to litigation_context.db."""
    if not LIT_DB.exists():
        return None
    conn = sqlite3.connect(str(LIT_DB))
    conn.row_factory = sqlite3.Row
    conn.text_factory = lambda b: b.decode("utf-8", errors="replace")
    for p in _PRAGMAS:
        conn.execute(p)
    return conn


def _safe_query(conn, sql, params=(), limit=200):
    """Execute read-only SQL with automatic limit. Returns list of dicts."""
    if conn is None:
        return []
    try:
        rows = conn.execute(sql, params).fetchmany(limit)
        return _rows_to_dicts(rows)
    except Exception:
        return []


def _table_exists(conn, table_name):
    """Check if a table exists in the database."""
    if conn is None:
        return False
    try:
        row = conn.execute(
            "SELECT count(*) FROM sqlite_master WHERE type='table' AND name=?",
            (table_name,),
        ).fetchone()
        return row[0] > 0 if row else False
    except Exception:
        return False


def _fts_search(conn, fts_table, base_table, query, columns="*", limit=200):
    """FTS5 search with sanitization and LIKE fallback."""
    safe_q = _sanitize_fts(query)
    if not safe_q:
        return [], "empty"

    if _table_exists(conn, fts_table):
        try:
            sql = (
                f"SELECT {columns} FROM {fts_table} f "
                f"JOIN {base_table} n ON f.rowid = n.rowid "
                f"WHERE {fts_table} MATCH ? LIMIT ?"
            )
            rows = conn.execute(sql, (safe_q, limit)).fetchall()
            if rows:
                return _rows_to_dicts(rows), "fts5"
        except Exception:
            pass

    like_pat = f"%{safe_q}%"
    try:
        cols = [
            r[1]
            for r in conn.execute(f"PRAGMA table_info({base_table})").fetchall()
        ]
        text_cols = [c for c in cols if c in ("label", "description", "id", "quote_text", "event_description")]
        if not text_cols:
            text_cols = cols[:3]
        where = " OR ".join(f"{c} LIKE ?" for c in text_cols)
        sql = f"SELECT {columns} FROM {base_table} WHERE {where} LIMIT ?"
        params = tuple([like_pat] * len(text_cols)) + (limit,)
        rows = conn.execute(sql, params).fetchall()
        return _rows_to_dicts(rows), "like"
    except Exception:
        return [], "error"


# ---------------------------------------------------------------------------
# KRAKEN mini-engine (inline for thread-safe background hunting)
# ---------------------------------------------------------------------------
def _kraken_extract_text(filepath):
    """Extract text from a file (PDF/DOCX/TXT). Returns (text, method)."""
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == ".pdf":
            try:
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(filepath)
                pages = min(len(pdf), 30)
                text_parts = []
                for i in range(pages):
                    page = pdf[i]
                    tp = page.get_textpage()
                    text_parts.append(tp.get_text_range())
                    tp.close()
                    page.close()
                pdf.close()
                return "\n".join(text_parts)[:MAX_CONTENT_BYTES], f"PDF({pages}pp)"
            except ImportError:
                return "", "PDF_NO_LIB"
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(filepath)
                text = "\n".join(p.text for p in doc.paragraphs)
                return text[:MAX_CONTENT_BYTES], f"DOCX({len(doc.paragraphs)}P)"
            except ImportError:
                return "", "DOCX_NO_LIB"
        elif ext in (".txt", ".csv", ".html", ".json", ".md"):
            for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
                try:
                    with open(filepath, "r", encoding=enc, errors="replace") as f:
                        return f.read(MAX_CONTENT_BYTES), f"TEXT({enc})"
                except Exception:
                    continue
            return "", "TEXT_ERR"
    except Exception as exc:
        return "", f"ERR({exc.__class__.__name__})"
    return "", "UNSUPPORTED"


def _kraken_analyze(content, filepath, focus="all"):
    """Analyze content for adversaries, legal authorities, categories, quotes."""
    if len(content) < 50:
        return {
            "adversaries": {},
            "legal": {},
            "categories": [],
            "key_quotes": [],
            "value_score": 0,
            "value_label": "EMPTY",
        }

    adversaries = {}
    for name, pat in ADVERSARY_RE.items():
        matches = pat.findall(content)
        if matches:
            adversaries[name] = len(matches)

    legal = {}
    for ltype, pat in LEGAL_PATTERNS.items():
        found = pat.findall(content)
        if found:
            legal[ltype] = list(set(found))

    categories = []
    for cat, pat in EVIDENCE_CATEGORIES.items():
        if pat.search(content):
            categories.append(cat)

    sentences = re.split(r"[.!?\n]+", content)
    key_quotes = []
    for sent in sentences:
        sent = sent.strip()
        if len(sent) < 30 or len(sent) > 500:
            continue
        for name, pat in ADVERSARY_RE.items():
            if name in adversaries and pat.search(sent):
                key_quotes.append(sent)
                break
        if len(key_quotes) >= 8:
            break

    adv_score = len(adversaries) * 3
    legal_score = len(legal) * 2
    cat_score = len(categories)
    quote_score = len(key_quotes)

    if focus in FOCUS_BOOSTS and FOCUS_BOOSTS[focus]:
        for name in FOCUS_BOOSTS[focus]:
            if name in adversaries:
                adv_score += adversaries[name]
    if focus == "legal":
        legal_score *= 2

    total = adv_score + legal_score + cat_score + quote_score
    if total >= 10:
        label = "HIGH"
    elif total >= 4:
        label = "MEDIUM"
    else:
        label = "LOW"

    return {
        "adversaries": adversaries,
        "legal": legal,
        "categories": categories,
        "key_quotes": key_quotes,
        "value_score": total,
        "value_label": label,
    }


def _kraken_discover_files(conn):
    """Discover candidate files from file_inventory DB table + local dirs."""
    files = []
    if conn and _table_exists(conn, "file_inventory"):
        try:
            rows = conn.execute(
                "SELECT file_path FROM file_inventory WHERE extension IN "
                "('.pdf','.txt','.csv','.html','.json','.docx','.md') "
                "ORDER BY RANDOM() LIMIT 5000"
            ).fetchall()
            files.extend(r[0] for r in rows if r[0])
        except Exception:
            pass

    scan_dirs = [
        REPO_ROOT / "01_EVIDENCE",
        REPO_ROOT / "02_AUTHORITY",
        REPO_ROOT / "04_ANALYSIS",
        REPO_ROOT / "05_FILINGS",
        REPO_ROOT / "09_REFERENCE",
    ]
    for d in scan_dirs:
        if d.exists():
            try:
                for root, _, fnames in os.walk(str(d)):
                    for fn in fnames:
                        if os.path.splitext(fn)[1].lower() in EXTS:
                            files.append(os.path.join(root, fn))
            except Exception:
                pass
    return list(set(files))


def _file_hash(fp):
    """Quick hash for dedup tracking."""
    try:
        st = os.stat(fp)
        raw = f"{fp}|{st.st_size}|{st.st_mtime}".encode()
        return hashlib.md5(raw).hexdigest()
    except Exception:
        return hashlib.md5(fp.encode()).hexdigest()


# ---------------------------------------------------------------------------
# UnifiedAPI -- JS bridge
# ---------------------------------------------------------------------------
class UnifiedAPI:
    """Exposes ALL methods to JavaScript via window.pywebview.api.*"""

    def __init__(self):
        self._brain_conn = None
        self._lit_conn = None
        self._kraken_thread = None
        self._kraken_status = {
            "running": False,
            "rounds_done": 0,
            "rounds_total": 0,
            "files_scanned": 0,
            "findings_high": 0,
            "findings_total": 0,
            "recent": [],
        }
        self._kraken_lock = threading.Lock()

    # -- Connection helpers --

    def _brain(self):
        if self._brain_conn is None:
            self._brain_conn = _connect_brain()
        return self._brain_conn

    def _lit(self):
        if self._lit_conn is None:
            self._lit_conn = _connect_lit()
        return self._lit_conn

    # ===================================================================
    # BRAIN API (from mbp_app.py BrainAPI)
    # ===================================================================

    def get_node_details(self, node_id):
        """Node + edges + chains + gaps for a given node ID."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            node = conn.execute("SELECT * FROM nodes WHERE id = ?", (node_id,)).fetchone()
            if not node:
                return {"error": f"Node {node_id} not found"}

            edges = conn.execute(
                "SELECT e.*, 'outgoing' AS direction FROM edges e WHERE e.source_id = ? "
                "UNION ALL "
                "SELECT e.*, 'incoming' AS direction FROM edges e WHERE e.target_id = ?",
                (node_id, node_id),
            ).fetchall()

            chains = conn.execute(
                "SELECT id, chain_path, strength_score, lane, filing_id "
                "FROM chains WHERE chain_path LIKE ? LIMIT 100",
                (f"%{node_id}%",),
            ).fetchall()

            gaps = []
            if _table_exists(conn, "gaps"):
                gaps = conn.execute(
                    "SELECT * FROM gaps WHERE node_id = ? AND resolved = 0",
                    (node_id,),
                ).fetchall()

            return {
                "node": dict(node),
                "edges": _rows_to_dicts(edges),
                "chains": _rows_to_dicts(chains),
                "gaps": _rows_to_dicts(gaps),
                "edge_count": len(edges),
                "chain_count": len(chains),
                "gap_count": len(gaps),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def trace_chain(self, node_id):
        """All chains passing through a node."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available", "chains": [], "count": 0}
        try:
            rows = conn.execute(
                "SELECT * FROM chains WHERE chain_path LIKE ? "
                "ORDER BY strength_score DESC LIMIT 200",
                (f"%{node_id}%",),
            ).fetchall()
            chains = []
            for r in rows:
                d = dict(r)
                try:
                    d["path_nodes"] = _json.loads(d.get("chain_path", "[]"))
                except Exception:
                    d["path_nodes"] = []
                chains.append(d)
            return {"chains": chains, "count": len(chains)}
        except Exception as exc:
            return {"error": str(exc), "chains": [], "count": 0}

    def search_nodes(self, query):
        """FTS5 search across brain nodes with LIKE fallback."""
        if not query or not query.strip():
            return {"results": [], "count": 0, "method": "empty"}
        conn = self._brain()
        if conn is None:
            return {"results": [], "count": 0, "method": "no_db"}
        results, method = _fts_search(conn, "nodes_fts", "nodes", query, "n.*")
        if method == "fts5":
            return {"results": results, "count": len(results), "method": "fts5"}

        safe_q = _sanitize_fts(query)
        if not safe_q:
            return {"results": [], "count": 0, "method": "sanitized_empty"}
        like_pat = f"%{safe_q}%"
        try:
            rows = conn.execute(
                "SELECT * FROM nodes WHERE label LIKE ? OR description LIKE ? OR id LIKE ? "
                "ORDER BY CASE WHEN label LIKE ? THEN 0 ELSE 1 END LIMIT 200",
                (like_pat, like_pat, like_pat, like_pat),
            ).fetchall()
            return {"results": _rows_to_dicts(rows), "count": len(rows), "method": "like"}
        except Exception as exc:
            return {"results": [], "count": 0, "method": "error", "error": str(exc)}

    def get_stats(self):
        """Brain stats + separation days."""
        conn = self._brain()
        result = {"separation_days": _sep_days(), "brain_available": conn is not None}
        if conn is None:
            return result
        try:
            row = conn.execute(
                "SELECT "
                "(SELECT COUNT(*) FROM nodes) AS node_count, "
                "(SELECT COUNT(*) FROM edges) AS edge_count, "
                "(SELECT COUNT(*) FROM chains) AS chain_count, "
                "(SELECT COUNT(*) FROM gaps) AS gap_total, "
                "(SELECT COUNT(*) FROM gaps WHERE resolved = 0) AS gap_open, "
                "(SELECT COUNT(*) FROM gaps WHERE priority = 'HIGH' AND resolved = 0) AS gap_high, "
                "(SELECT MAX(version) FROM versions) AS brain_version"
            ).fetchone()
            result.update(dict(row))

            layers = conn.execute(
                "SELECT layer, COUNT(*) AS cnt FROM nodes GROUP BY layer ORDER BY cnt DESC"
            ).fetchall()
            result["layers"] = _rows_to_dicts(layers)

            edge_types = conn.execute(
                "SELECT edge_type, COUNT(*) AS cnt FROM edges GROUP BY edge_type ORDER BY cnt DESC"
            ).fetchall()
            result["edge_types"] = _rows_to_dicts(edge_types)

            lanes = conn.execute(
                "SELECT lane, COUNT(*) AS cnt FROM nodes WHERE lane != '' "
                "GROUP BY lane ORDER BY cnt DESC"
            ).fetchall()
            result["lanes"] = _rows_to_dicts(lanes)

            top = conn.execute(
                "SELECT id, strength_score, lane, filing_id FROM chains "
                "ORDER BY strength_score DESC LIMIT 5"
            ).fetchall()
            result["top_chains"] = _rows_to_dicts(top)
        except Exception as exc:
            result["error"] = str(exc)
        return result

    def get_gaps(self, priority=""):
        """Gap listing, optionally filtered by priority."""
        conn = self._brain()
        if conn is None:
            return {"gaps": [], "count": 0, "by_priority": {}}
        try:
            if priority and priority.upper() in ("HIGH", "MEDIUM", "LOW"):
                rows = conn.execute(
                    "SELECT * FROM gaps WHERE resolved = 0 AND priority = ? "
                    "ORDER BY created_at DESC LIMIT 500",
                    (priority.upper(),),
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT * FROM gaps WHERE resolved = 0 "
                    "ORDER BY CASE priority WHEN 'HIGH' THEN 0 WHEN 'MEDIUM' THEN 1 ELSE 2 END, "
                    "created_at DESC LIMIT 500"
                ).fetchall()

            gaps = _rows_to_dicts(rows)
            by_p = {"HIGH": 0, "MEDIUM": 0, "LOW": 0, "UNKNOWN": 0}
            for g in gaps:
                p = g.get("priority", "UNKNOWN")
                by_p[p] = by_p.get(p, 0) + 1
            return {"gaps": gaps, "count": len(gaps), "by_priority": by_p}
        except Exception as exc:
            return {"gaps": [], "count": 0, "error": str(exc), "by_priority": {}}

    def run_evolution(self):
        """Trigger brain_evolution.py subprocess."""
        if not EVOLVE_SCRIPT.exists():
            return {"error": f"Evolution script not found: {EVOLVE_SCRIPT}"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(EVOLVE_SCRIPT), "--stats"],
                capture_output=True, text=True, timeout=120,
                cwd=str(REPO_ROOT),
            )
            return {
                "stdout": result.stdout[-4000:],
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
            }
        except subprocess.TimeoutExpired:
            return {"error": "Evolution script timed out after 120 seconds"}
        except Exception as exc:
            return {"error": str(exc)}

    def refresh_data(self):
        """Re-export graph_data.json from brain DB."""
        if not EXPORT_SCRIPT.exists():
            return {"error": f"Export script not found: {EXPORT_SCRIPT}"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(EXPORT_SCRIPT)],
                capture_output=True, text=True, timeout=120,
                cwd=str(REPO_ROOT),
            )
            return {
                "success": result.returncode == 0,
                "stdout": result.stdout[-4000:],
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
                "message": (
                    "Reload the page to see updated data"
                    if result.returncode == 0
                    else "Export failed -- check stderr"
                ),
            }
        except subprocess.TimeoutExpired:
            return {"error": "Export script timed out after 120 seconds"}
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # LITIGATION API (from adversary_blueprint.py LitigationAPI)
    # ===================================================================

    def get_live_data(self):
        """Separation counter + DB stats."""
        conn = self._lit()
        stats = {"ev": 0, "jv": 0, "im": 0, "tl": 0, "ct": 0}
        if conn:
            try:
                row = conn.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM evidence_quotes) AS ev, "
                    "(SELECT COUNT(*) FROM judicial_violations) AS jv, "
                    "(SELECT COUNT(*) FROM impeachment_matrix) AS im, "
                    "(SELECT COUNT(*) FROM timeline_events) AS tl, "
                    "(SELECT COUNT(*) FROM contradiction_map) AS ct"
                ).fetchone()
                stats = dict(row)
            except Exception:
                pass
        return {
            "sep_days": _sep_days(),
            "sep_date": str(SEPARATION_DATE),
            "today": str(date.today()),
            "db_available": conn is not None,
            "stats": stats,
        }

    def query_evidence(self, query, limit=20):
        """Evidence quote search from litigation_context.db."""
        conn = self._lit()
        if conn is None:
            return []
        safe_q = _sanitize_fts(query)
        if not safe_q:
            return []

        if _table_exists(conn, "evidence_fts"):
            try:
                rows = conn.execute(
                    "SELECT eq.id, eq.quote_text, eq.source_file, eq.page_number, "
                    "eq.category, eq.lane, eq.relevance_score "
                    "FROM evidence_fts f JOIN evidence_quotes eq ON f.rowid = eq.rowid "
                    "WHERE evidence_fts MATCH ? LIMIT ?",
                    (safe_q, limit),
                ).fetchall()
                if rows:
                    return _rows_to_dicts(rows)
            except Exception:
                pass

        like_pat = f"%{safe_q}%"
        return _safe_query(
            conn,
            "SELECT id, quote_text, source_file, page_number, category, lane, relevance_score "
            "FROM evidence_quotes WHERE quote_text LIKE ? OR source_file LIKE ? "
            "ORDER BY relevance_score DESC LIMIT ?",
            (like_pat, like_pat, limit),
            limit,
        )

    def query_adversary(self, name):
        """Adversary intelligence dossier."""
        conn = self._lit()
        if conn is None:
            return {"name": name, "error": "DB not available"}
        like = f"%{name}%"
        result = {"name": name, "evidence_count": 0, "violations": [], "impeachment": [], "contradictions": []}
        try:
            ev = conn.execute(
                "SELECT COUNT(*) FROM evidence_quotes WHERE quote_text LIKE ?",
                (like,),
            ).fetchone()
            result["evidence_count"] = ev[0] if ev else 0

            if _table_exists(conn, "judicial_violations"):
                result["violations"] = _safe_query(
                    conn,
                    "SELECT violation_type, COUNT(*) AS cnt FROM judicial_violations "
                    "WHERE description LIKE ? GROUP BY violation_type ORDER BY cnt DESC LIMIT 20",
                    (like,),
                )

            if _table_exists(conn, "impeachment_matrix"):
                result["impeachment"] = _safe_query(
                    conn,
                    "SELECT category, COUNT(*) AS cnt, AVG(impeachment_value) AS avg_sev "
                    "FROM impeachment_matrix WHERE target LIKE ? "
                    "GROUP BY category ORDER BY cnt DESC LIMIT 20",
                    (like,),
                )

            if _table_exists(conn, "contradiction_map"):
                result["contradictions"] = _safe_query(
                    conn,
                    "SELECT actor, statement_1, statement_2, severity, source_1, source_2 "
                    "FROM contradiction_map WHERE actor LIKE ? "
                    "ORDER BY severity DESC LIMIT 20",
                    (like,),
                )
        except Exception as exc:
            result["error"] = str(exc)
        return result

    def query_timeline(self, start="", end="", limit=50):
        """Timeline events with date range filter."""
        conn = self._lit()
        if conn is None:
            return []
        conditions = []
        params = []
        if start:
            conditions.append("event_date >= ?")
            params.append(start)
        if end:
            conditions.append("event_date <= ?")
            params.append(end)
        where = " AND ".join(conditions) if conditions else "1=1"
        params.append(limit)
        return _safe_query(
            conn,
            f"SELECT event_date, event_description, actors, source_document "
            f"FROM timeline_events WHERE {where} ORDER BY event_date DESC LIMIT ?",
            tuple(params),
            limit,
        )

    def query_judicial(self, jtype=""):
        """Judicial violations by type."""
        conn = self._lit()
        if conn is None:
            return {"violations": [], "cartel": []}
        result = {"violations": [], "cartel": []}
        try:
            if jtype:
                result["violations"] = _safe_query(
                    conn,
                    "SELECT violation_type, description, COUNT(*) AS cnt "
                    "FROM judicial_violations WHERE violation_type LIKE ? "
                    "GROUP BY violation_type ORDER BY cnt DESC LIMIT 30",
                    (f"%{jtype}%",),
                )
            else:
                result["violations"] = _safe_query(
                    conn,
                    "SELECT violation_type, COUNT(*) AS cnt "
                    "FROM judicial_violations GROUP BY violation_type ORDER BY cnt DESC LIMIT 30",
                )

            if _table_exists(conn, "berry_mcneill_intelligence"):
                result["cartel"] = _safe_query(
                    conn,
                    "SELECT person, connection_type, description, evidence_source "
                    "FROM berry_mcneill_intelligence ORDER BY rowid DESC LIMIT 50",
                )
        except Exception as exc:
            result["error"] = str(exc)
        return result

    def query_impeachment(self, target):
        """Impeachment matrix for a target."""
        conn = self._lit()
        if conn is None:
            return []
        return _safe_query(
            conn,
            "SELECT target, category, evidence_summary, impeachment_value, "
            "cross_exam_question, source_file FROM impeachment_matrix "
            "WHERE target LIKE ? ORDER BY impeachment_value DESC LIMIT 50",
            (f"%{target}%",),
        )

    def search_fts5(self, query):
        """Full-text search across multiple tables."""
        conn = self._lit()
        if conn is None:
            return {"results": [], "count": 0}
        safe_q = _sanitize_fts(query)
        if not safe_q:
            return {"results": [], "count": 0}
        all_results = []
        tables = [
            ("evidence_fts", "evidence_quotes", "quote_text", "evidence"),
            ("timeline_fts", "timeline_events", "event_description", "timeline"),
        ]
        for fts_t, base_t, text_col, source_tag in tables:
            if _table_exists(conn, fts_t):
                try:
                    rows = conn.execute(
                        f"SELECT b.rowid, b.{text_col} AS text FROM {fts_t} f "
                        f"JOIN {base_t} b ON f.rowid = b.rowid "
                        f"WHERE {fts_t} MATCH ? LIMIT 50",
                        (safe_q,),
                    ).fetchall()
                    for r in rows:
                        all_results.append({"source": source_tag, "text": r["text"], "rowid": r["rowid"]})
                except Exception:
                    pass

        if not all_results:
            like_pat = f"%{safe_q}%"
            for base_t, text_col, tag in [
                ("evidence_quotes", "quote_text", "evidence"),
                ("timeline_events", "event_description", "timeline"),
            ]:
                if _table_exists(conn, base_t):
                    try:
                        rows = conn.execute(
                            f"SELECT rowid, {text_col} AS text FROM {base_t} "
                            f"WHERE {text_col} LIKE ? LIMIT 30",
                            (like_pat,),
                        ).fetchall()
                        for r in rows:
                            all_results.append({"source": tag, "text": r["text"], "rowid": r["rowid"]})
                    except Exception:
                        pass

        return {"results": all_results, "count": len(all_results)}

    # ===================================================================
    # KRAKEN HUNTING (background thread)
    # ===================================================================

    def start_kraken(self, rounds=3, count=10, focus="all"):
        """Launch KRAKEN hunting in background thread."""
        with self._kraken_lock:
            if self._kraken_status["running"]:
                return {"status": "already_running", "progress": self._kraken_status}

        self._kraken_status = {
            "running": True,
            "rounds_done": 0,
            "rounds_total": int(rounds),
            "files_scanned": 0,
            "findings_high": 0,
            "findings_total": 0,
            "recent": [],
        }

        def _hunt():
            lit_conn = _connect_lit()
            all_files = _kraken_discover_files(lit_conn)
            processed_hashes = set()

            if lit_conn and _table_exists(lit_conn, "kraken_processed"):
                try:
                    rows = lit_conn.execute("SELECT file_hash FROM kraken_processed").fetchall()
                    processed_hashes = {r[0] for r in rows}
                except Exception:
                    pass

            for rnd in range(int(rounds)):
                candidates = []
                for fp in all_files:
                    fh = _file_hash(fp)
                    if fh not in processed_hashes and os.path.isfile(fp):
                        try:
                            sz = os.path.getsize(fp)
                            if 100 <= sz <= 50_000_000:
                                candidates.append((fp, fh, sz))
                        except Exception:
                            pass

                sample = random.sample(candidates, min(int(count), len(candidates))) if candidates else []
                for fp, fh, sz in sample:
                    text, method = _kraken_extract_text(fp)
                    analysis = _kraken_analyze(text, fp, focus)
                    processed_hashes.add(fh)

                    with self._kraken_lock:
                        self._kraken_status["files_scanned"] += 1
                        self._kraken_status["findings_total"] += 1
                        if analysis["value_label"] == "HIGH":
                            self._kraken_status["findings_high"] += 1

                        if analysis["value_score"] >= 4:
                            entry = {
                                "file": os.path.basename(fp),
                                "score": analysis["value_score"],
                                "label": analysis["value_label"],
                                "adversaries": list(analysis["adversaries"].keys()),
                                "categories": analysis["categories"],
                                "quotes": analysis["key_quotes"][:2],
                            }
                            recent = self._kraken_status["recent"]
                            recent.insert(0, entry)
                            self._kraken_status["recent"] = recent[:50]

                    if analysis["value_label"] == "HIGH" and lit_conn:
                        for quote in analysis["key_quotes"][:5]:
                            try:
                                cats = ",".join(analysis["categories"][:3])
                                lane = "A"
                                if "judicial" in analysis["categories"]:
                                    lane = "E"
                                elif "housing" in analysis["categories"]:
                                    lane = "B"
                                elif "PPO" in analysis["categories"]:
                                    lane = "D"
                                elif "criminal" in analysis["categories"]:
                                    lane = "CRIMINAL"
                                lit_conn.execute(
                                    "INSERT OR IGNORE INTO evidence_quotes "
                                    "(source_file, quote_text, category, lane, relevance_score, created_at) "
                                    "VALUES (?, ?, ?, ?, ?, datetime('now'))",
                                    (os.path.basename(fp), quote, cats, lane, analysis["value_score"]),
                                )
                            except Exception:
                                pass
                        try:
                            lit_conn.commit()
                        except Exception:
                            pass

                    if lit_conn:
                        try:
                            if not _table_exists(lit_conn, "kraken_processed"):
                                lit_conn.execute(
                                    "CREATE TABLE IF NOT EXISTS kraken_processed ("
                                    "file_hash TEXT PRIMARY KEY, file_path TEXT, "
                                    "processed_at TEXT DEFAULT (datetime('now')), "
                                    "value_score INTEGER DEFAULT 0, value_label TEXT DEFAULT 'LOW')"
                                )
                            lit_conn.execute(
                                "INSERT OR IGNORE INTO kraken_processed (file_hash, file_path, value_score, value_label) "
                                "VALUES (?, ?, ?, ?)",
                                (fh, fp, analysis["value_score"], analysis["value_label"]),
                            )
                            lit_conn.commit()
                        except Exception:
                            pass

                with self._kraken_lock:
                    self._kraken_status["rounds_done"] = rnd + 1

            if lit_conn:
                try:
                    lit_conn.close()
                except Exception:
                    pass
            with self._kraken_lock:
                self._kraken_status["running"] = False

        t = threading.Thread(target=_hunt, daemon=True, name="kraken-hunter")
        t.start()
        self._kraken_thread = t
        return {"status": "started", "rounds": int(rounds), "count": int(count), "focus": focus}

    def get_kraken_status(self):
        """Current hunting progress."""
        with self._kraken_lock:
            return dict(self._kraken_status)

    # ===================================================================
    # FILING GENERATION (from generate_filing.py)
    # ===================================================================

    def list_filings(self):
        """Available filing IDs with chain statistics."""
        conn = self._brain()
        if conn is None:
            return []
        try:
            rows = conn.execute(
                "SELECT filing_id, COUNT(*) AS chain_count, "
                "AVG(strength_score) AS avg_strength, MAX(strength_score) AS max_strength "
                "FROM chains WHERE filing_id IS NOT NULL AND filing_id != '' "
                "GROUP BY filing_id ORDER BY avg_strength DESC"
            ).fetchall()
            results = []
            for r in rows:
                d = dict(r)
                node = conn.execute(
                    "SELECT label, lane FROM nodes WHERE id = ? LIMIT 1",
                    (d["filing_id"],),
                ).fetchone()
                if node:
                    d["label"] = node["label"]
                    d["lane"] = node["lane"]
                else:
                    d["label"] = d["filing_id"]
                    d["lane"] = ""
                results.append(d)
            return results
        except Exception as exc:
            return [{"error": str(exc)}]

    def generate_brief(self, filing_id):
        """Auto-generate a brief for a filing via generate_filing.py."""
        if not FILING_SCRIPT.exists():
            return {"error": "Filing script not found"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(FILING_SCRIPT), "--filing", filing_id, "--brief"],
                capture_output=True, text=True, timeout=60,
                cwd=str(REPO_ROOT),
            )
            return {
                "text": result.stdout[-16000:] if result.returncode == 0 else "",
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
                "success": result.returncode == 0,
            }
        except subprocess.TimeoutExpired:
            return {"error": "Brief generation timed out (60s)"}
        except Exception as exc:
            return {"error": str(exc)}

    def generate_impeachment(self, actor_id):
        """Cross-examination outline for an actor."""
        if not FILING_SCRIPT.exists():
            return {"error": "Filing script not found"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(FILING_SCRIPT), "--actor", actor_id, "--impeach"],
                capture_output=True, text=True, timeout=60,
                cwd=str(REPO_ROOT),
            )
            return {
                "text": result.stdout[-16000:] if result.returncode == 0 else "",
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
                "success": result.returncode == 0,
            }
        except subprocess.TimeoutExpired:
            return {"error": "Impeachment generation timed out (60s)"}
        except Exception as exc:
            return {"error": str(exc)}

    def get_strongest_filing(self):
        """Strongest chain -> filing info."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            row = conn.execute(
                "SELECT filing_id, MAX(strength_score) AS max_str, lane "
                "FROM chains WHERE filing_id IS NOT NULL AND filing_id != '' "
                "GROUP BY filing_id ORDER BY max_str DESC LIMIT 1"
            ).fetchone()
            if row:
                d = dict(row)
                node = conn.execute(
                    "SELECT label, description FROM nodes WHERE id = ?",
                    (d["filing_id"],),
                ).fetchone()
                if node:
                    d["label"] = node["label"]
                    d["description"] = node["description"]
                return d
            return {"error": "No filings found"}
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # COURT FEED
    # ===================================================================

    def fetch_court_feed(self):
        """Fetch latest Michigan court updates."""
        if not COURT_FEED_SCRIPT.exists():
            return {"error": "court_feed.py not found", "items": []}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(COURT_FEED_SCRIPT), "--json"],
                capture_output=True, text=True, timeout=30,
                cwd=str(REPO_ROOT),
            )
            if result.returncode == 0 and result.stdout.strip():
                try:
                    return _json.loads(result.stdout)
                except Exception:
                    return {"items": [], "raw": result.stdout[:2000]}
            return {"items": [], "stderr": result.stderr[:1000]}
        except subprocess.TimeoutExpired:
            return {"error": "Court feed timed out (30s)", "items": []}
        except Exception as exc:
            return {"error": str(exc), "items": []}

    # ===================================================================
    # ANALYTICS PASSTHROUGH
    # ===================================================================

    def run_analytics(self, command):
        """Run an mbp_intel.py analytics command."""
        intel_script = REPO_ROOT / "07_CODE" / "PROJECT_KRAKEN" / "mbp" / "mbp_intel.py"
        if not intel_script.exists():
            return {"error": "mbp_intel.py not found"}
        if _WRITE_RE.search(command):
            return {"error": "Write operations not allowed"}
        try:
            result = subprocess.run(
                [_PYTHON, "-I", str(intel_script), command],
                capture_output=True, text=True, timeout=30,
                cwd=str(REPO_ROOT),
            )
            return {
                "output": result.stdout[-8000:],
                "stderr": result.stderr[-2000:],
                "returncode": result.returncode,
            }
        except subprocess.TimeoutExpired:
            return {"error": f"Analytics command '{command}' timed out (30s)"}
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # CHRONOLOGY / FILING / ADVERSARY / DASHBOARD / EXPORT / SEARCH
    # ===================================================================

    def get_chronology(self, lane=None, limit=200):
        """Timeline events from litigation_context.db, ordered by date.

        Args:
            lane: Optional lane filter (e.g. "A", "D").
            limit: Max rows to return (default 200).
        Returns:
            dict with ``events`` list and ``count``.
        """
        conn = self._lit()
        if conn is None:
            return {"events": [], "count": 0, "error": "Litigation DB not available"}
        try:
            if lane:
                sql = (
                    "SELECT id, event_date, event_description, actors, lane, "
                    "category, source_table, source_id, severity, filing_relevance "
                    "FROM timeline_events WHERE lane = ? "
                    "ORDER BY event_date ASC LIMIT ?"
                )
                rows = conn.execute(sql, (lane, limit)).fetchall()
            else:
                sql = (
                    "SELECT id, event_date, event_description, actors, lane, "
                    "category, source_table, source_id, severity, filing_relevance "
                    "FROM timeline_events ORDER BY event_date ASC LIMIT ?"
                )
                rows = conn.execute(sql, (limit,)).fetchall()

            events = []
            for r in rows:
                d = dict(r)
                d["severity"] = _safe_float(d.get("severity"))
                events.append(d)

            return {"events": events, "count": len(events)}
        except Exception as exc:
            return {"events": [], "count": 0, "error": str(exc)}

    def get_filing_packet(self, filing_id):
        """Comprehensive packet overview for a filing from mbp_brain.db.

        Args:
            filing_id: e.g. ``"FILING_MCR_2003_DISQUALIFICATION"``.
        Returns:
            dict with filing details, chains, evidence/violation/authority
            counts, exhibit candidates, and blocking gaps.
        """
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        if not filing_id:
            return {"error": "filing_id is required"}
        try:
            # Filing node
            node = conn.execute(
                "SELECT * FROM nodes WHERE id = ?", (filing_id,)
            ).fetchone()
            if not node:
                return {"error": f"Filing node '{filing_id}' not found"}
            filing_info = dict(node)
            filing_info["readiness"] = _safe_float(filing_info.get("readiness"))
            filing_info["severity"] = _safe_float(filing_info.get("severity"))
            filing_info["confidence"] = _safe_float(filing_info.get("confidence"))

            # Chains reaching this filing, sorted by strength
            chain_rows = conn.execute(
                "SELECT id, chain_path, chain_type, total_weight, length, "
                "lane, filing_id, evidence_ids, strength_score "
                "FROM chains WHERE filing_id = ? "
                "ORDER BY strength_score DESC LIMIT 200",
                (filing_id,),
            ).fetchall()
            chains = []
            for r in chain_rows:
                d = dict(r)
                d["strength_score"] = _safe_float(d.get("strength_score"))
                d["total_weight"] = _safe_float(d.get("total_weight"))
                chains.append(d)

            # Count supporting nodes by type across all chains
            evidence_count = 0
            violation_count = 0
            authority_count = 0
            exhibit_candidates = []
            seen_evidence = set()

            for ch in chains:
                raw_ids = ch.get("evidence_ids") or "[]"
                try:
                    eid_list = _json.loads(raw_ids) if raw_ids.startswith("[") else []
                except Exception:
                    eid_list = []
                for eid in eid_list:
                    if eid not in seen_evidence:
                        seen_evidence.add(eid)

            if seen_evidence:
                placeholders = ",".join("?" for _ in seen_evidence)
                type_rows = conn.execute(
                    f"SELECT node_type, COUNT(*) AS cnt FROM nodes "
                    f"WHERE id IN ({placeholders}) GROUP BY node_type",
                    list(seen_evidence),
                ).fetchall()
                for tr in type_rows:
                    nt = (tr["node_type"] or "").lower()
                    cnt = tr["cnt"]
                    if "evidence" in nt or "quote" in nt:
                        evidence_count += cnt
                    elif "violation" in nt or "judicial" in nt:
                        violation_count += cnt
                    elif "authority" in nt or "rule" in nt or "statute" in nt:
                        authority_count += cnt

                # Exhibit candidates = evidence-type nodes in chains
                ex_rows = conn.execute(
                    f"SELECT id, label, node_type, lane, confidence FROM nodes "
                    f"WHERE id IN ({placeholders}) "
                    f"AND (node_type LIKE '%evidence%' OR node_type LIKE '%quote%' "
                    f"     OR node_type LIKE '%exhibit%' OR node_type LIKE '%police%')",
                    list(seen_evidence),
                ).fetchall()
                for er in ex_rows:
                    d = dict(er)
                    d["confidence"] = _safe_float(d.get("confidence"))
                    exhibit_candidates.append(d)

            # Gaps blocking this filing
            gaps = []
            if _table_exists(conn, "gaps"):
                gap_rows = conn.execute(
                    "SELECT id, gap_type, description, priority, acquisition_task "
                    "FROM gaps WHERE node_id = ? AND resolved = 0 "
                    "ORDER BY priority DESC",
                    (filing_id,),
                ).fetchall()
                gaps = _rows_to_dicts(gap_rows)

            return {
                "filing": filing_info,
                "chains": chains,
                "chain_count": len(chains),
                "evidence_count": evidence_count,
                "violation_count": violation_count,
                "authority_count": authority_count,
                "exhibit_candidates": exhibit_candidates,
                "gaps": gaps,
                "gap_count": len(gaps),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def get_adversary_overview(self):
        """All adversary analytics in a single batch query.

        Returns a dict keyed by adversary name, each containing
        evidence_count, violation_count, impeachment_count, and
        top_contradiction from litigation_context.db.
        """
        conn = self._lit()
        if conn is None:
            return {"adversaries": {}, "error": "Litigation DB not available"}
        try:
            adversary_names = list(ADVERSARIES.keys())
            result = {}
            for name in adversary_names:
                result[name] = {
                    "evidence_count": 0,
                    "violation_count": 0,
                    "impeachment_count": 0,
                    "top_contradiction": None,
                }

            # Evidence counts -- batch via CASE WHEN
            case_clauses = []
            params = []
            for name in adversary_names:
                case_clauses.append(
                    "SUM(CASE WHEN quote_text LIKE ? OR source_file LIKE ? THEN 1 ELSE 0 END)"
                )
                like = f"%{name.split()[0]}%"
                params.extend([like, like])
            sql_ev = f"SELECT {', '.join(case_clauses)} FROM evidence_quotes WHERE is_duplicate = 0"
            row = conn.execute(sql_ev, params).fetchone()
            if row:
                for i, name in enumerate(adversary_names):
                    result[name]["evidence_count"] = row[i] or 0

            # Judicial violation counts per adversary
            jv_clauses = []
            jv_params = []
            for name in adversary_names:
                jv_clauses.append(
                    "SUM(CASE WHEN description LIKE ? OR source_quote LIKE ? THEN 1 ELSE 0 END)"
                )
                like = f"%{name.split()[0]}%"
                jv_params.extend([like, like])
            sql_jv = f"SELECT {', '.join(jv_clauses)} FROM judicial_violations"
            row = conn.execute(sql_jv, jv_params).fetchone()
            if row:
                for i, name in enumerate(adversary_names):
                    result[name]["violation_count"] = row[i] or 0

            # Impeachment counts per adversary
            imp_clauses = []
            imp_params = []
            for name in adversary_names:
                imp_clauses.append(
                    "SUM(CASE WHEN target LIKE ? OR evidence_summary LIKE ? THEN 1 ELSE 0 END)"
                )
                like = f"%{name.split()[0]}%"
                imp_params.extend([like, like])
            sql_imp = f"SELECT {', '.join(imp_clauses)} FROM impeachment_matrix"
            row = conn.execute(sql_imp, imp_params).fetchone()
            if row:
                for i, name in enumerate(adversary_names):
                    result[name]["impeachment_count"] = row[i] or 0

            # Top contradiction per adversary
            for name in adversary_names:
                like = f"%{name.split()[0]}%"
                ctr = conn.execute(
                    "SELECT contradiction_text, severity FROM contradiction_map "
                    "WHERE contradiction_text LIKE ? "
                    "ORDER BY severity DESC LIMIT 1",
                    (like,),
                ).fetchone()
                if ctr:
                    result[name]["top_contradiction"] = {
                        "text": ctr["contradiction_text"],
                        "severity": _safe_float(ctr["severity"]),
                    }

            return {"adversaries": result}
        except Exception as exc:
            return {"adversaries": {}, "error": str(exc)}

    def get_dashboard_data(self):
        """All analytics dashboard data in a single call.

        Consolidates separation counter, table counts, brain stats,
        top chains, and recent brain operations.
        """
        data = {
            "separation_days": _sep_days(),
            "separation_date": str(SEPARATION_DATE),
            "today": str(date.today()),
            "table_counts": {},
            "brain_stats": {},
            "top_chains": [],
            "recent_operations": [],
        }

        # Litigation DB table counts
        lit = self._lit()
        if lit:
            try:
                row = lit.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM evidence_quotes) AS evidence_quotes, "
                    "(SELECT COUNT(*) FROM judicial_violations) AS judicial_violations, "
                    "(SELECT COUNT(*) FROM impeachment_matrix) AS impeachment_matrix, "
                    "(SELECT COUNT(*) FROM timeline_events) AS timeline_events, "
                    "(SELECT COUNT(*) FROM contradiction_map) AS contradiction_map"
                ).fetchone()
                data["table_counts"] = dict(row)
            except Exception as exc:
                data["table_counts"]["error"] = str(exc)

        # Brain DB stats
        brain = self._brain()
        if brain:
            try:
                row = brain.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM nodes) AS nodes, "
                    "(SELECT COUNT(*) FROM edges) AS edges, "
                    "(SELECT COUNT(*) FROM chains) AS chains, "
                    "(SELECT COUNT(*) FROM gaps) AS gaps_total, "
                    "(SELECT COUNT(*) FROM gaps WHERE resolved = 0) AS gaps_open, "
                    "(SELECT COUNT(*) FROM gaps WHERE resolved = 1) AS gaps_resolved"
                ).fetchone()
                data["brain_stats"] = dict(row)
            except Exception as exc:
                data["brain_stats"]["error"] = str(exc)

            # Top 5 strongest chains
            try:
                top = brain.execute(
                    "SELECT id, chain_path, chain_type, strength_score, "
                    "lane, filing_id, length "
                    "FROM chains ORDER BY strength_score DESC LIMIT 5"
                ).fetchall()
                chains = []
                for r in top:
                    d = dict(r)
                    d["strength_score"] = _safe_float(d.get("strength_score"))
                    chains.append(d)
                data["top_chains"] = chains
            except Exception:
                pass

            # Recent brain operations (brain_log may be empty or absent)
            if _table_exists(brain, "brain_log"):
                try:
                    ops = brain.execute(
                        "SELECT * FROM brain_log ORDER BY rowid DESC LIMIT 10"
                    ).fetchall()
                    data["recent_operations"] = _rows_to_dicts(ops)
                except Exception:
                    pass

        return data

    def export_subgraph(self, node_ids):
        """Export a subgraph as D3.js-compatible JSON.

        Args:
            node_ids: List of node IDs to include.
        Returns:
            dict with ``nodes`` and ``links`` arrays matching D3 format.
        """
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        if not node_ids or not isinstance(node_ids, list):
            return {"error": "node_ids must be a non-empty list"}
        try:
            placeholders = ",".join("?" for _ in node_ids)

            # Fetch nodes
            node_rows = conn.execute(
                f"SELECT id, node_type, layer, label, description, "
                f"date_start, date_end, severity, confidence, readiness, "
                f"lane, metadata "
                f"FROM nodes WHERE id IN ({placeholders})",
                node_ids,
            ).fetchall()

            nodes = []
            id_set = set()
            for r in node_rows:
                d = dict(r)
                d["severity"] = _safe_float(d.get("severity"))
                d["confidence"] = _safe_float(d.get("confidence"))
                d["readiness"] = _safe_float(d.get("readiness"))
                # Parse metadata JSON if present
                raw_meta = d.get("metadata")
                if raw_meta:
                    try:
                        d["metadata"] = _json.loads(raw_meta)
                    except Exception:
                        pass
                nodes.append(d)
                id_set.add(d["id"])

            if not nodes:
                return {"nodes": [], "links": [], "warning": "No matching nodes found"}

            # Fetch edges between these nodes (both directions must be in set)
            edge_rows = conn.execute(
                f"SELECT id, source_id, target_id, edge_type, weight, evidence "
                f"FROM edges "
                f"WHERE source_id IN ({placeholders}) "
                f"AND target_id IN ({placeholders})",
                node_ids + node_ids,
            ).fetchall()

            links = []
            for r in edge_rows:
                d = dict(r)
                d["source"] = d.pop("source_id")
                d["target"] = d.pop("target_id")
                d["weight"] = _safe_float(d.get("weight"), 1.0)
                links.append(d)

            return {
                "nodes": nodes,
                "links": links,
                "node_count": len(nodes),
                "link_count": len(links),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def search_everything(self, query, limit=50):
        """Unified search across both databases.

        Searches evidence_fts in litigation_context.db and nodes in
        mbp_brain.db, merges + deduplicates results, and sorts by
        relevance.

        Args:
            query: Search text.
            limit: Max results per source (total may be up to 2*limit
                   before dedup).
        Returns:
            dict with ``results`` list and ``count``.
        """
        if not query or not query.strip():
            return {"results": [], "count": 0, "method": "empty"}

        safe_q = _sanitize_fts(query)
        if not safe_q:
            return {"results": [], "count": 0, "method": "sanitized_empty"}

        all_results = []
        seen_ids = set()

        # --- litigation_context.db: evidence_fts ---
        lit = self._lit()
        if lit:
            # FTS5 path
            fts_hit = False
            if _table_exists(lit, "evidence_fts"):
                try:
                    rows = lit.execute(
                        "SELECT e.id, e.quote_text, e.source_file, e.category, "
                        "e.lane, e.relevance_score "
                        "FROM evidence_fts f "
                        "JOIN evidence_quotes e ON f.rowid = e.rowid "
                        "WHERE evidence_fts MATCH ? LIMIT ?",
                        (safe_q, limit),
                    ).fetchall()
                    for r in rows:
                        d = dict(r)
                        uid = f"lit_{d['id']}"
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            all_results.append({
                                "source": "litigation",
                                "id": d["id"],
                                "text": d.get("quote_text", ""),
                                "type": d.get("category", "evidence"),
                                "lane": d.get("lane", ""),
                                "score": _safe_float(d.get("relevance_score"), 0.5),
                                "source_file": d.get("source_file", ""),
                            })
                    fts_hit = len(rows) > 0
                except Exception:
                    pass

            # LIKE fallback if FTS5 missed
            if not fts_hit:
                like_pat = f"%{safe_q}%"
                try:
                    rows = lit.execute(
                        "SELECT id, quote_text, source_file, category, lane, "
                        "relevance_score FROM evidence_quotes "
                        "WHERE quote_text LIKE ? OR source_file LIKE ? "
                        "ORDER BY relevance_score DESC LIMIT ?",
                        (like_pat, like_pat, limit),
                    ).fetchall()
                    for r in rows:
                        d = dict(r)
                        uid = f"lit_{d['id']}"
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            all_results.append({
                                "source": "litigation",
                                "id": d["id"],
                                "text": d.get("quote_text", ""),
                                "type": d.get("category", "evidence"),
                                "lane": d.get("lane", ""),
                                "score": _safe_float(d.get("relevance_score"), 0.3),
                                "source_file": d.get("source_file", ""),
                            })
                except Exception:
                    pass

        # --- mbp_brain.db: nodes ---
        brain = self._brain()
        if brain:
            # FTS5 path
            brain_fts_hit = False
            if _table_exists(brain, "nodes_fts"):
                try:
                    rows = brain.execute(
                        "SELECT n.id, n.label, n.node_type, n.layer, n.lane, "
                        "n.confidence "
                        "FROM nodes_fts f "
                        "JOIN nodes n ON f.rowid = n.rowid "
                        "WHERE nodes_fts MATCH ? LIMIT ?",
                        (safe_q, limit),
                    ).fetchall()
                    for r in rows:
                        d = dict(r)
                        uid = f"brain_{d['id']}"
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            all_results.append({
                                "source": "brain",
                                "id": d["id"],
                                "text": d.get("label", ""),
                                "type": d.get("node_type", ""),
                                "lane": d.get("lane", ""),
                                "score": _safe_float(d.get("confidence"), 0.5),
                                "layer": d.get("layer", ""),
                            })
                    brain_fts_hit = len(rows) > 0
                except Exception:
                    pass

            # LIKE fallback
            if not brain_fts_hit:
                like_pat = f"%{safe_q}%"
                try:
                    rows = brain.execute(
                        "SELECT id, label, node_type, layer, lane, confidence "
                        "FROM nodes WHERE label LIKE ? OR description LIKE ? "
                        "ORDER BY confidence DESC LIMIT ?",
                        (like_pat, like_pat, limit),
                    ).fetchall()
                    for r in rows:
                        d = dict(r)
                        uid = f"brain_{d['id']}"
                        if uid not in seen_ids:
                            seen_ids.add(uid)
                            all_results.append({
                                "source": "brain",
                                "id": d["id"],
                                "text": d.get("label", ""),
                                "type": d.get("node_type", ""),
                                "lane": d.get("lane", ""),
                                "score": _safe_float(d.get("confidence"), 0.3),
                                "layer": d.get("layer", ""),
                            })
                except Exception:
                    pass

        # Sort by score descending, then trim
        all_results.sort(key=lambda x: x.get("score", 0), reverse=True)
        all_results = all_results[: limit * 2]

        return {
            "results": all_results,
            "count": len(all_results),
        }

    # ===================================================================
    # T1: BRAIN REBUILD (build_mbp_brain.py)
    # ===================================================================

    def rebuild_brain(self):
        """Trigger a full brain rebuild (T1 -- build_mbp_brain.py).

        Runs asynchronously in a background thread.  Returns immediately
        with status ``started``.  Poll ``get_health()`` to see when
        new node/edge counts appear.
        """
        script = REPO_ROOT / "scripts" / "build_mbp_brain.py"
        if not script.exists():
            return {"status": "error", "message": f"Script not found: {script}"}

        def _run():
            try:
                subprocess.run(
                    [_PYTHON, "-I", str(script)],
                    cwd=str(REPO_ROOT),
                    timeout=600,
                    capture_output=True,
                )
            except Exception:
                pass

        t = threading.Thread(target=_run, daemon=True, name="brain-rebuild")
        t.start()
        return {"status": "started", "script": str(script)}

    # ===================================================================
    # T2: CHAIN RECOMPUTE (compute_chains.py)
    # ===================================================================

    def recompute_chains(self):
        """Trigger chain recomputation (T2 -- compute_chains.py)."""
        script = REPO_ROOT / "scripts" / "compute_chains.py"
        if not script.exists():
            return {"status": "error", "message": f"Script not found: {script}"}

        def _run():
            try:
                subprocess.run(
                    [_PYTHON, "-I", str(script)],
                    cwd=str(REPO_ROOT),
                    timeout=600,
                    capture_output=True,
                )
            except Exception:
                pass

        t = threading.Thread(target=_run, daemon=True, name="chain-recompute")
        t.start()
        return {"status": "started", "script": str(script)}

    # ===================================================================
    # T3: FILE WATCHER (brain_watcher.py)
    # ===================================================================

    def start_watcher(self):
        """Start the MEEK file watcher daemon (T3 -- brain_watcher.py).

        Launches brain_watcher.py as a background subprocess with
        auto-routing enabled.  Returns immediately.
        """
        script = REPO_ROOT / "scripts" / "brain_watcher.py"
        if not script.exists():
            return {"status": "error", "message": f"Script not found: {script}"}

        def _run():
            try:
                subprocess.Popen(
                    [_PYTHON, "-I", str(script)],
                    cwd=str(REPO_ROOT),
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
            except Exception:
                pass

        t = threading.Thread(target=_run, daemon=True, name="watcher-start")
        t.start()
        return {"status": "started", "script": str(script)}

    # ===================================================================
    # CLOSED LOOP ORCHESTRATOR
    # ===================================================================

    def run_full_cycle(self):
        """Execute the complete closed-loop pipeline:

        Evidence -> Violation -> Authority -> Remedy -> Filing -> back to Evidence

        Runs T1 (brain build) -> T2 (chains) -> T5 (evolution) -> T6 (export)
        sequentially in a background thread.  Returns immediately.
        """
        def _cycle():
            scripts = [
                ("T1-Brain", REPO_ROOT / "scripts" / "build_mbp_brain.py"),
                ("T2-Chains", REPO_ROOT / "scripts" / "compute_chains.py"),
                ("T5-Evolution", REPO_ROOT / "scripts" / "brain_evolution.py"),
                ("T6-Export", REPO_ROOT / "scripts" / "export_brain_d3.py"),
            ]
            for label, script in scripts:
                if script.exists():
                    try:
                        subprocess.run(
                            [_PYTHON, "-I", str(script)],
                            cwd=str(REPO_ROOT),
                            timeout=600,
                            capture_output=True,
                        )
                    except Exception:
                        pass

        t = threading.Thread(target=_cycle, daemon=True, name="full-cycle")
        t.start()
        return {
            "status": "started",
            "pipeline": "T1-Brain -> T2-Chains -> T5-Evolution -> T6-Export",
        }

    # ===================================================================
    # COMMUNITY INTELLIGENCE (v15.0 — Leiden Clusters)
    # ===================================================================

    def get_community_intel(self, community_id):
        """Full intelligence for a community: metadata, top nodes, patterns."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            comm = conn.execute(
                "SELECT * FROM communities WHERE id = ?", (community_id,)
            ).fetchone()
            if not comm:
                return {"error": f"Community {community_id} not found"}

            members = conn.execute(
                "SELECT n.id, n.label, n.node_type, n.lane, na.pagerank, "
                "na.hub_score, na.authority_score "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "LEFT JOIN node_analytics na ON n.id = na.node_id "
                "WHERE cm.community_id = ? "
                "ORDER BY COALESCE(na.pagerank, 0) DESC LIMIT 50",
                (community_id,),
            ).fetchall()

            children = conn.execute(
                "SELECT id, label, member_count, level FROM communities "
                "WHERE parent_id = ? ORDER BY member_count DESC",
                (community_id,),
            ).fetchall()

            edges = conn.execute(
                "SELECT ce.*, c2.label AS target_label "
                "FROM community_edges ce "
                "LEFT JOIN communities c2 ON ce.target_id = c2.id "
                "WHERE ce.source_id = ? OR ce.target_id = ? "
                "ORDER BY ce.total_weight DESC LIMIT 30",
                (community_id, community_id),
            ).fetchall()

            return {
                "community": dict(comm),
                "top_nodes": _rows_to_dicts(members),
                "children": _rows_to_dicts(children),
                "inter_edges": _rows_to_dicts(edges),
                "node_count": len(members),
            }
        except Exception as exc:
            return {"error": str(exc)}

    def search_communities(self, query, limit=30):
        """Search communities by label, narrative, or key actors."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            safe_q = _sanitize_fts(query)
            like_pat = f"%{safe_q}%"
            rows = conn.execute(
                "SELECT id, label, level, lane, member_count, "
                "evidence_strength, authority_completeness, impeachment_density "
                "FROM communities "
                "WHERE label LIKE ? OR narrative LIKE ? OR key_actors LIKE ? "
                "ORDER BY member_count DESC LIMIT ?",
                (like_pat, like_pat, like_pat, limit),
            ).fetchall()
            return {"communities": _rows_to_dicts(rows), "count": len(rows)}
        except Exception as exc:
            return {"error": str(exc)}

    def get_community_timeline(self, community_id, limit=100):
        """Chronological events within a community."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            rows = conn.execute(
                "SELECT n.id, n.label, n.node_type, n.date_start, n.lane "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "WHERE cm.community_id = ? AND n.date_start IS NOT NULL "
                "ORDER BY n.date_start ASC LIMIT ?",
                (community_id, limit),
            ).fetchall()
            return {"events": _rows_to_dicts(rows), "count": len(rows)}
        except Exception as exc:
            return {"error": str(exc)}

    def expand_community(self, community_id, limit=200):
        """Return individual nodes for JS-side expansion of a community."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            nodes = conn.execute(
                "SELECT n.id, n.label, n.node_type, n.lane, n.date_start, "
                "na.pagerank, na.hub_score, na.authority_score "
                "FROM community_members cm "
                "JOIN nodes n ON cm.node_id = n.id "
                "LEFT JOIN node_analytics na ON n.id = na.node_id "
                "WHERE cm.community_id = ? "
                "ORDER BY COALESCE(na.pagerank, 0) DESC LIMIT ?",
                (community_id, limit),
            ).fetchall()

            node_ids = [r["id"] for r in nodes]
            edges = []
            if node_ids:
                placeholders = ",".join("?" * len(node_ids))
                edges = conn.execute(
                    f"SELECT source_id, target_id, edge_type, weight "
                    f"FROM edges "
                    f"WHERE source_id IN ({placeholders}) "
                    f"AND target_id IN ({placeholders}) LIMIT 1000",
                    node_ids + node_ids,
                ).fetchall()

            return {
                "nodes": _rows_to_dicts(nodes),
                "edges": _rows_to_dicts(edges),
                "community_id": community_id,
            }
        except Exception as exc:
            return {"error": str(exc)}

    def get_community_stats(self):
        """Summary stats for the community system."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            row = conn.execute(
                "SELECT "
                "(SELECT COUNT(*) FROM communities WHERE level = 0) AS lanes, "
                "(SELECT COUNT(*) FROM communities WHERE level = 1) AS epochs, "
                "(SELECT COUNT(*) FROM communities WHERE level = 2) AS communities, "
                "(SELECT COUNT(*) FROM community_members) AS memberships, "
                "(SELECT COUNT(*) FROM community_edges) AS inter_edges, "
                "(SELECT COUNT(*) FROM node_analytics WHERE pagerank > 0) AS ranked_nodes"
            ).fetchone()
            return dict(row) if row else {}
        except Exception as exc:
            return {"error": str(exc)}

    def get_detected_patterns(self, limit=50):
        """Autonomously detected patterns across communities."""
        conn = self._brain()
        if conn is None:
            return {"error": "Brain DB not available"}
        try:
            if not _table_exists(conn, "detected_patterns"):
                return {"patterns": [], "count": 0}
            rows = conn.execute(
                "SELECT * FROM detected_patterns "
                "ORDER BY confidence DESC, detected_at DESC LIMIT ?",
                (limit,),
            ).fetchall()
            return {"patterns": _rows_to_dicts(rows), "count": len(rows)}
        except Exception as exc:
            return {"error": str(exc)}

    # ===================================================================
    # HEALTH CHECK
    # ===================================================================

    def get_health(self):
        """Brain health check with comprehensive counts."""
        brain = self._brain()
        lit = self._lit()
        health = {
            "version": VERSION,
            "separation_days": _sep_days(),
            "brain_db": {"available": brain is not None, "path": str(BRAIN_DB)},
            "lit_db": {"available": lit is not None, "path": str(LIT_DB)},
        }

        if brain:
            try:
                row = brain.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM nodes) AS nodes, "
                    "(SELECT COUNT(*) FROM edges) AS edges, "
                    "(SELECT COUNT(*) FROM chains) AS chains, "
                    "(SELECT COUNT(*) FROM gaps WHERE resolved = 0) AS open_gaps, "
                    "(SELECT MAX(version) FROM versions) AS version"
                ).fetchone()
                health["brain_db"].update(dict(row))

                # Community stats (v15.0)
                if _table_exists(brain, "communities"):
                    crow = brain.execute(
                        "SELECT "
                        "(SELECT COUNT(*) FROM communities) AS communities, "
                        "(SELECT COUNT(*) FROM community_members) AS memberships, "
                        "(SELECT COUNT(*) FROM node_analytics WHERE pagerank > 0) AS ranked_nodes"
                    ).fetchone()
                    health["brain_db"]["communities"] = dict(crow) if crow else {}
            except Exception as exc:
                health["brain_db"]["error"] = str(exc)

        if lit:
            try:
                row = lit.execute(
                    "SELECT "
                    "(SELECT COUNT(*) FROM evidence_quotes) AS evidence, "
                    "(SELECT COUNT(*) FROM judicial_violations) AS violations, "
                    "(SELECT COUNT(*) FROM timeline_events) AS timeline, "
                    "(SELECT COUNT(*) FROM impeachment_matrix) AS impeachment"
                ).fetchone()
                health["lit_db"].update(dict(row))
            except Exception as exc:
                health["lit_db"]["error"] = str(exc)

        health["graph_json"] = {
            "exists": GRAPH_JSON.exists(),
            "size_mb": round(GRAPH_JSON.stat().st_size / 1048576, 1) if GRAPH_JSON.exists() else 0,
        }
        return health


# ---------------------------------------------------------------------------
# HTTP Server
# ---------------------------------------------------------------------------
class _QuietHandler(http.server.SimpleHTTPRequestHandler):
    """Serve V15 HTML assets with gzip compression and smart caching."""

    _gzip_cache: dict = {}  # path → (mtime, compressed_bytes)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(VIS_DIR), **kwargs)

    def log_message(self, fmt, *args):
        pass

    def end_headers(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        super().end_headers()

    def guess_type(self, path):
        ext = os.path.splitext(path)[1].lower()
        overrides = {".json": "application/json", ".js": "application/javascript",
                     ".mjs": "application/javascript", ".woff2": "font/woff2"}
        return overrides.get(ext, super().guess_type(path))

    def _serve_gzipped(self, file_path, content_type):
        """Serve a file with gzip compression and aggressive caching."""
        mtime = file_path.stat().st_mtime
        cache_key = str(file_path)

        # Check gzip cache (avoids re-compressing on every request)
        cached = self._gzip_cache.get(cache_key)
        if cached and cached[0] == mtime:
            compressed = cached[1]
        else:
            with open(str(file_path), "rb") as f:
                raw = f.read()
            compressed = gzip.compress(raw, compresslevel=6)
            self._gzip_cache[cache_key] = (mtime, compressed)

        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Encoding", "gzip")
        self.send_header("Content-Length", str(len(compressed)))
        # Cache static assets aggressively (reload via Ctrl+Shift+R if needed)
        self.send_header("Cache-Control", "public, max-age=3600, stale-while-revalidate=86400")
        self.end_headers()
        self.wfile.write(compressed)

    def do_GET(self):
        """Serve graph_data.json with gzip; other files with caching."""
        clean_path = self.path.split("?")[0]

        if clean_path in ("/graph_data.json", "/graph_clusters.json"):
            candidates = [
                VIS_DIR / "graph_clusters.json",
                VIS_DIR / "graph_data.json",
                VIS_DIR_V9 / "graph_data.json",
                VIS_DIR_V5 / "graph_data.json",
            ]
            gj = next((c for c in candidates if c.exists()), None)
            if not gj:
                self.send_error(404, "graph data not found in V15/V9/V5")
                return
            # Gzip compress: 4.3 MB → ~400 KB
            accept = self.headers.get("Accept-Encoding", "")
            if "gzip" in accept:
                self._serve_gzipped(gj, "application/json")
            else:
                # Fallback: uncompressed chunked
                self.send_response(200)
                self.send_header("Content-Type", "application/json")
                sz = gj.stat().st_size
                self.send_header("Content-Length", str(sz))
                self.send_header("Cache-Control", "public, max-age=3600")
                self.end_headers()
                with open(str(gj), "rb") as f:
                    while True:
                        chunk = f.read(65536)
                        if not chunk:
                            break
                        self.wfile.write(chunk)
            return

        # For HTML/JS/CSS: add caching headers
        ext = os.path.splitext(clean_path)[1].lower()
        if ext in (".html", ".js", ".css") and "gzip" in self.headers.get("Accept-Encoding", ""):
            fpath = VIS_DIR / clean_path.lstrip("/")
            if fpath.exists():
                ct = self.guess_type(clean_path)
                self._serve_gzipped(fpath, ct)
                return

        super().do_GET()


class MBPServer(threading.Thread):
    """HTTP server thread for static assets."""

    def __init__(self, port=0):
        super().__init__(daemon=True, name="mbp-http")
        self.port = port or _find_free_port()
        self.server = http.server.HTTPServer(("127.0.0.1", self.port), _QuietHandler)

    def run(self):
        self.server.serve_forever()

    def shutdown(self):
        self.server.shutdown()


# ---------------------------------------------------------------------------
# Banner
# ---------------------------------------------------------------------------
def _print_banner(port, stats):
    sep = _sep_days()
    nodes = stats.get("node_count", "?")
    edges = stats.get("edge_count", "?")
    chains = stats.get("chain_count", "?")
    gaps = stats.get("gap_open", "?")
    ver = stats.get("brain_version", "?")
    banner = (
        "\n"
        "+==============================================================+\n"
        f"|       THEMANBEARPIG v{VERSION} -- Unified Legal Brain            |\n"
        "+==============================================================+\n"
        f"|  HTTP  : http://127.0.0.1:{port:<5}                             |\n"
        f"|  Brain : {str(BRAIN_DB):<50} |\n"
        f"|  Lit DB: {str(LIT_DB):<50} |\n"
        f"|  Version: {ver:<4}  Nodes: {nodes:<8}  Edges: {edges:<8}         |\n"
        f"|  Chains: {chains:<6}  Open Gaps: {gaps:<6}                        |\n"
        f"|  SEPARATION: {sep} DAYS since July 29, 2025              |\n"
        "+==============================================================+\n"
    )
    print(banner)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="THEMANBEARPIG v15.0 -- Self-Organizing Intelligence Graph"
    )
    parser.add_argument("--debug", action="store_true", help="Enable developer tools")
    parser.add_argument("--port", type=int, default=0, help="HTTP server port (0=auto)")
    parser.add_argument("--hunt", action="store_true", help="Launch KRAKEN hunting on startup")
    parser.add_argument("--export", action="store_true", help="Re-export graph data before launch")
    parser.add_argument("--rounds", type=int, default=3, help="KRAKEN hunting rounds (with --hunt)")
    parser.add_argument("--count", type=int, default=10, help="Files per KRAKEN round (with --hunt)")
    parser.add_argument("--focus", default="all", help="KRAKEN focus mode (with --hunt)")
    args = parser.parse_args()

    try:
        import webview
    except ImportError:
        print("ERROR: pywebview not installed. Run: pip install pywebview")
        sys.exit(1)

    if not BRAIN_DB.exists() and not LIT_DB.exists():
        print("WARNING: Neither mbp_brain.db nor litigation_context.db found.")
        print("The application will run in limited mode.")

    if not VIS_DIR.exists():
        VIS_DIR.mkdir(parents=True, exist_ok=True)

    index_html = VIS_DIR / "index.html"
    if not index_html.exists():
        print(f"ERROR: {index_html} not found. Build the V15 visualization first.")
        sys.exit(1)

    if args.export and EXPORT_SCRIPT.exists():
        print("Re-exporting graph data...")
        subprocess.run(
            [_PYTHON, "-I", str(EXPORT_SCRIPT)],
            cwd=str(REPO_ROOT), timeout=300,
        )

    api = UnifiedAPI()
    stats = api.get_stats()

    server = MBPServer(port=args.port)
    server.start()

    _print_banner(server.port, stats)

    if args.hunt:
        print(f"Starting KRAKEN hunter: {args.rounds} rounds, {args.count} files, focus={args.focus}")
        api.start_kraken(rounds=args.rounds, count=args.count, focus=args.focus)

    sep = _sep_days()
    title = f"THEMANBEARPIG v{VERSION} -- {sep} Days Separated"

    window = webview.create_window(
        title,
        url=f"http://127.0.0.1:{server.port}/index.html",
        js_api=api,
        width=APP_WIDTH,
        height=APP_HEIGHT,
        min_size=(APP_MIN_W, APP_MIN_H),
        background_color=APP_BG,
        text_select=True,
    )

    webview.start(debug=args.debug)
    server.shutdown()
    print("THEMANBEARPIG exited.")


if __name__ == "__main__":
    import io as _io
    try:
        sys.stdout = _io.TextIOWrapper(
            sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True
        )
    except Exception:
        pass
    main()
