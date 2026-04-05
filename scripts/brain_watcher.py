#!/usr/bin/env python3
"""brain_watcher.py — File watcher daemon for mbp_brain.db auto-ingestion.

Monitors C:\\Users\\andre\\LitigationOS for new/modified files and auto-ingests
them as Document/Recording/Image nodes into mbp_brain.db with MEEK lane routing.

Usage:
    python -I scripts/brain_watcher.py                   # Watch (foreground)
    python -I scripts/brain_watcher.py --daemon           # Watch (background)
    python -I scripts/brain_watcher.py --status           # Show queue stats
    python -I scripts/brain_watcher.py --process-backlog  # Process pending queue
    python -I scripts/brain_watcher.py --watch-dir D:\\evidence --interval 5
"""

from __future__ import annotations

import argparse
import hashlib
import logging
import os
import re
import signal
import sqlite3
import sys
import threading
import time
from datetime import datetime, timezone
from pathlib import Path

# ---------------------------------------------------------------------------
# watchdog imports
# ---------------------------------------------------------------------------
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler, FileSystemEvent

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MBP_BRAIN_DB = Path(r"C:\Users\andre\LitigationOS\mbp_brain.db")
DEFAULT_WATCH_DIR = Path(r"C:\Users\andre\LitigationOS")
LOG_DIR = DEFAULT_WATCH_DIR / "logs"

WATCHED_EXTENSIONS: frozenset[str] = frozenset({
    ".pdf", ".docx", ".doc", ".txt", ".md", ".json", ".csv",
    ".eml", ".msg", ".html", ".htm",
    ".mp3", ".mp4", ".wav",
    ".png", ".jpg", ".jpeg",
})

AUDIO_VIDEO_EXTENSIONS: frozenset[str] = frozenset({".mp3", ".mp4", ".wav"})
IMAGE_EXTENSIONS: frozenset[str] = frozenset({".png", ".jpg", ".jpeg"})

IGNORE_DIRS: frozenset[str] = frozenset({
    "__pycache__", ".git", "node_modules", ".mcp_venv",
    "pytools_venv", "D_tmp", "logs",
})

IGNORE_SUFFIXES: frozenset[str] = frozenset({".pyc", ".pyo"})

MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
DESCRIPTION_MAX_CHARS = 500

# ---------------------------------------------------------------------------
# MEEK Lane Routing — priority order: E → D → F → C → A → B
# ---------------------------------------------------------------------------

_LANE_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("E", re.compile(
        r"mcneill|judicial|bias|jtc|canon|misconduct|benchbook|ex\s*parte",
        re.IGNORECASE,
    )),
    ("D", re.compile(
        r"ppo|protection\s+order|5907|contempt|stalking|harassment",
        re.IGNORECASE,
    )),
    ("F", re.compile(
        r"coa|366810|appeal|appellant|appellee|brief|appendix",
        re.IGNORECASE,
    )),
    ("C", re.compile(
        r"federal|§\s*1983|1983|conspiracy|civil\s+rights|42\s+usc",
        re.IGNORECASE,
    )),
    ("A", re.compile(
        r"custody|parenting|001507|watson|child|visitation|foc",
        re.IGNORECASE,
    )),
    ("B", re.compile(
        r"shady\s*oaks|eviction|housing|trailer|002760|habitability",
        re.IGNORECASE,
    )),
]

LANE_NODE_IDS: dict[str, str] = {
    "A": "lane-A",
    "B": "lane-B",
    "C": "lane-C",
    "D": "lane-D",
    "E": "lane-E",
    "F": "lane-F",
}

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

logger = logging.getLogger("brain_watcher")


def setup_logging(console: bool = True) -> None:
    """Configure logging to console + file."""
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    log_file = LOG_DIR / "brain_watcher.log"

    formatter = logging.Formatter(
        "%(asctime)s [%(levelname)s] %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    file_handler = logging.FileHandler(log_file, encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)

    if console:
        console_handler = logging.StreamHandler(sys.stderr)
        console_handler.setLevel(logging.INFO)
        console_handler.setFormatter(formatter)
        logger.addHandler(console_handler)

    logger.setLevel(logging.DEBUG)


# ---------------------------------------------------------------------------
# Database helpers
# ---------------------------------------------------------------------------

def get_connection() -> sqlite3.Connection:
    """Open mbp_brain.db with mandatory PRAGMAs."""
    conn = sqlite3.connect(str(MBP_BRAIN_DB), timeout=120)
    conn.execute("PRAGMA busy_timeout = 60000")
    conn.execute("PRAGMA journal_mode = WAL")
    conn.execute("PRAGMA cache_size = -32000")
    conn.execute("PRAGMA synchronous = NORMAL")
    conn.execute("PRAGMA temp_store = MEMORY")
    conn.row_factory = sqlite3.Row
    return conn


def log_brain_op(
    conn: sqlite3.Connection,
    operation: str,
    input_params: str | None = None,
    output_summary: str | None = None,
    nodes_touched: int = 0,
    edges_traversed: int = 0,
    duration_ms: float = 0.0,
) -> None:
    """Record an operation in brain_ops."""
    conn.execute(
        """INSERT INTO brain_ops
           (operation, input_params, output_summary, nodes_touched, edges_traversed, duration_ms)
           VALUES (?, ?, ?, ?, ?, ?)""",
        (operation, input_params, output_summary, nodes_touched, edges_traversed, duration_ms),
    )


def record_version(conn: sqlite3.Connection, mutations: str | None = None) -> None:
    """Snapshot current counts into versions table."""
    row = conn.execute(
        """SELECT
             (SELECT COUNT(*) FROM nodes) AS nc,
             (SELECT COUNT(*) FROM edges) AS ec,
             (SELECT COUNT(*) FROM chains) AS cc,
             (SELECT COUNT(*) FROM gaps) AS gc"""
    ).fetchone()
    conn.execute(
        """INSERT INTO versions (node_count, edge_count, chain_count, gap_count, mutations)
           VALUES (?, ?, ?, ?, ?)""",
        (row["nc"], row["ec"], row["cc"], row["gc"], mutations),
    )


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_pdf(file_path: Path) -> str:
    """Extract text from PDF via pypdfium2, fallback to raw bytes."""
    try:
        import pypdfium2 as pdfium  # type: ignore[import-untyped]
        pdf = pdfium.PdfDocument(str(file_path))
        pages_text: list[str] = []
        for page_idx in range(min(len(pdf), 20)):  # cap at 20 pages for speed
            page = pdf[page_idx]
            textpage = page.get_textpage()
            pages_text.append(textpage.get_text_bounded())
            textpage.close()
            page.close()
        pdf.close()
        return "\n".join(pages_text)
    except Exception:
        logger.debug("pypdfium2 failed for %s, falling back to raw read", file_path)
        try:
            raw = file_path.read_bytes()[:10240]
            return raw.decode("utf-8", errors="replace")
        except Exception:
            return ""


def extract_text_docx(file_path: Path) -> str:
    """Extract text from DOCX via python-docx."""
    try:
        import docx  # type: ignore[import-untyped]
        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception:
        logger.debug("python-docx failed for %s, falling back to raw read", file_path)
        try:
            raw = file_path.read_bytes()[:10240]
            return raw.decode("utf-8", errors="replace")
        except Exception:
            return ""


def extract_text_plain(file_path: Path) -> str:
    """Read text files directly (txt, md, csv, json, html, eml, msg)."""
    try:
        return file_path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        try:
            raw = file_path.read_bytes()[:10240]
            return raw.decode("utf-8", errors="replace")
        except Exception:
            return ""


def extract_text(file_path: Path) -> str:
    """Route extraction by extension. Returns extracted text (may be empty)."""
    ext = file_path.suffix.lower()
    if ext in AUDIO_VIDEO_EXTENSIONS or ext in IMAGE_EXTENSIONS:
        return ""  # metadata-only nodes
    if ext == ".pdf":
        return extract_text_pdf(file_path)
    if ext in {".docx", ".doc"}:
        return extract_text_docx(file_path)
    # All other text-based formats
    return extract_text_plain(file_path)


# ---------------------------------------------------------------------------
# MEEK lane detection
# ---------------------------------------------------------------------------

def detect_lane(file_path: Path, text: str) -> str | None:
    """Detect case lane from filename + extracted text. Priority: E→D→F→C→A→B."""
    combined = f"{file_path.name} {text[:3000]}"
    for lane, pattern in _LANE_PATTERNS:
        if pattern.search(combined):
            return lane
    return None


# ---------------------------------------------------------------------------
# Node type classification
# ---------------------------------------------------------------------------

def classify_node_type(ext: str) -> str:
    """Return node_type based on file extension."""
    if ext in AUDIO_VIDEO_EXTENSIONS:
        return "Recording"
    if ext in IMAGE_EXTENSIONS:
        return "Image"
    return "Document"


# ---------------------------------------------------------------------------
# Actor lookup for edge wiring
# ---------------------------------------------------------------------------

_KNOWN_ACTORS: list[tuple[re.Pattern[str], str]] | None = None


def _load_actors(conn: sqlite3.Connection) -> list[tuple[re.Pattern[str], str]]:
    """Load Person/Actor nodes from brain for text matching. Cached."""
    global _KNOWN_ACTORS
    if _KNOWN_ACTORS is not None:
        return _KNOWN_ACTORS

    actors: list[tuple[re.Pattern[str], str]] = []
    try:
        rows = conn.execute(
            """SELECT id, label FROM nodes
               WHERE node_type IN ('Person', 'Actor', 'Judge', 'Party', 'Witness')
               AND length(label) > 2
               LIMIT 500"""
        ).fetchall()
        for row in rows:
            label = row["label"]
            try:
                pat = re.compile(re.escape(label), re.IGNORECASE)
                actors.append((pat, row["id"]))
            except re.error:
                continue
    except Exception as e:
        logger.debug("Actor load failed: %s", e)
    _KNOWN_ACTORS = actors
    return actors


def find_actor_mentions(conn: sqlite3.Connection, text: str) -> list[str]:
    """Return node IDs of actors mentioned in text."""
    if not text:
        return []
    actors = _load_actors(conn)
    mentioned: list[str] = []
    search_text = text[:5000]  # limit search scope
    for pattern, node_id in actors:
        if pattern.search(search_text):
            mentioned.append(node_id)
    return mentioned


# ---------------------------------------------------------------------------
# Core ingestion logic
# ---------------------------------------------------------------------------

def make_node_id(file_path: str) -> str:
    """Deterministic node ID from file path."""
    h = hashlib.md5(file_path.encode("utf-8")).hexdigest()
    return f"ingest-{h}"


def _file_metadata(fp: Path) -> dict:
    """Gather file metadata as a JSON-serializable dict."""
    try:
        stat = fp.stat()
        return {
            "file_path": str(fp),
            "file_size": stat.st_size,
            "modified_time": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
            "file_type": fp.suffix.lower(),
        }
    except OSError:
        return {
            "file_path": str(fp),
            "file_size": 0,
            "modified_time": None,
            "file_type": fp.suffix.lower(),
        }


def ingest_file(conn: sqlite3.Connection, file_path: Path) -> tuple[int, int]:
    """Ingest a single file into mbp_brain.db. Returns (nodes_created, edges_created)."""
    import json as _json  # local import avoids shadow module at module level

    fp_str = str(file_path)
    ext = file_path.suffix.lower()
    node_id = make_node_id(fp_str)
    nodes_created = 0
    edges_created = 0

    # Extract text
    text = extract_text(file_path)
    description = text[:DESCRIPTION_MAX_CHARS].strip() if text else None

    # Detect lane
    lane = detect_lane(file_path, text)

    # Classify node type
    node_type = classify_node_type(ext)

    # Build metadata JSON
    meta = _file_metadata(file_path)
    meta_json = _json.dumps(meta, ensure_ascii=False)

    # Insert node (idempotent — ON CONFLICT updates description/metadata)
    result = conn.execute(
        """INSERT INTO nodes (id, node_type, layer, label, description, lane, source_table, metadata, confidence)
           VALUES (?, ?, 'EVIDENCE', ?, ?, ?, 'ingest_queue', ?, 0.5)
           ON CONFLICT(id) DO UPDATE SET
             description = excluded.description,
             metadata = excluded.metadata,
             lane = COALESCE(excluded.lane, nodes.lane)""",
        (node_id, node_type, file_path.name, description, lane, meta_json),
    )
    nodes_created += result.rowcount

    # Edge: ASSIGNED_TO lane node
    if lane and lane in LANE_NODE_IDS:
        lane_target = LANE_NODE_IDS[lane]
        result = conn.execute(
            """INSERT INTO edges (source_id, target_id, edge_type, weight, evidence, source_table)
               VALUES (?, ?, 'ASSIGNED_TO', 0.8, ?, 'ingest_queue')
               ON CONFLICT(source_id, target_id, edge_type) DO NOTHING""",
            (node_id, lane_target, f"MEEK lane detection from file: {file_path.name}"),
        )
        edges_created += result.rowcount

    # Edges: RELATED to mentioned actors
    actor_ids = find_actor_mentions(conn, text)
    for actor_id in actor_ids[:10]:  # cap at 10 actor edges per file
        result = conn.execute(
            """INSERT INTO edges (source_id, target_id, edge_type, weight, evidence, source_table)
               VALUES (?, ?, 'RELATED', 0.4, ?, 'ingest_queue')
               ON CONFLICT(source_id, target_id, edge_type) DO NOTHING""",
            (node_id, actor_id, f"Actor mentioned in {file_path.name}"),
        )
        edges_created += result.rowcount

    return nodes_created, edges_created


def process_queue_item(conn: sqlite3.Connection, queue_id: int, file_path_str: str) -> bool:
    """Process a single ingest_queue entry. Returns True on success."""
    t0 = time.perf_counter()
    file_path = Path(file_path_str)

    # Mark processing
    conn.execute(
        "UPDATE ingest_queue SET status = 'processing' WHERE id = ?",
        (queue_id,),
    )
    conn.commit()

    try:
        if not file_path.exists():
            conn.execute(
                "UPDATE ingest_queue SET status = 'error', processed_at = datetime('now') WHERE id = ?",
                (queue_id,),
            )
            conn.commit()
            logger.warning("File not found, skipping: %s", file_path)
            return False

        if file_path.stat().st_size > MAX_FILE_SIZE:
            conn.execute(
                "UPDATE ingest_queue SET status = 'skipped', processed_at = datetime('now') WHERE id = ?",
                (queue_id,),
            )
            conn.commit()
            logger.info("File too large (>100MB), skipping: %s", file_path)
            return False

        nodes_created, edges_created = ingest_file(conn, file_path)

        elapsed_ms = (time.perf_counter() - t0) * 1000
        conn.execute(
            """UPDATE ingest_queue
               SET status = 'ingested', nodes_created = ?, edges_created = ?, processed_at = datetime('now')
               WHERE id = ?""",
            (nodes_created, edges_created, queue_id),
        )
        log_brain_op(
            conn,
            operation="ingest_file",
            input_params=file_path_str,
            output_summary=f"nodes={nodes_created} edges={edges_created}",
            nodes_touched=nodes_created,
            edges_traversed=edges_created,
            duration_ms=elapsed_ms,
        )
        conn.commit()
        logger.info(
            "Ingested: %s → %d nodes, %d edges (%.0fms)",
            file_path.name, nodes_created, edges_created, elapsed_ms,
        )
        return True

    except Exception as e:
        import json as _json
        error_meta = _json.dumps({"error": str(e)}, ensure_ascii=False)
        conn.execute(
            """UPDATE ingest_queue
               SET status = 'error', processed_at = datetime('now')
               WHERE id = ?""",
            (queue_id,),
        )
        log_brain_op(conn, operation="ingest_file_error", input_params=file_path_str, output_summary=str(e))
        conn.commit()
        logger.error("Failed to ingest %s: %s", file_path, e)
        return False


def enqueue_file(conn: sqlite3.Connection, file_path: Path) -> int | None:
    """Add a file to the ingest_queue. Returns queue id or None if duplicate."""
    ext = file_path.suffix.lower()
    try:
        cursor = conn.execute(
            """INSERT INTO ingest_queue (file_path, file_type, status)
               VALUES (?, ?, 'pending')
               ON CONFLICT(file_path) DO UPDATE SET
                 status = CASE
                   WHEN ingest_queue.status IN ('ingested', 'error', 'skipped')
                   THEN 'pending'
                   ELSE ingest_queue.status
                 END,
                 detected_at = datetime('now')
               RETURNING id""",
            (str(file_path), ext),
        )
        row = cursor.fetchone()
        conn.commit()
        return row[0] if row else None
    except sqlite3.Error as e:
        logger.debug("Enqueue failed for %s: %s", file_path, e)
        conn.rollback()
        return None


def process_pending(conn: sqlite3.Connection, limit: int = 100) -> int:
    """Process up to `limit` pending items from the queue. Returns count processed."""
    rows = conn.execute(
        "SELECT id, file_path FROM ingest_queue WHERE status = 'pending' ORDER BY detected_at LIMIT ?",
        (limit,),
    ).fetchall()

    if not rows:
        return 0

    success_count = 0
    for row in rows:
        if process_queue_item(conn, row["id"], row["file_path"]):
            success_count += 1

    # Record version after batch
    if success_count > 0:
        import json as _json
        mutations = _json.dumps({"batch_ingested": success_count, "total_queued": len(rows)})
        record_version(conn, mutations)
        conn.commit()

    return success_count


# ---------------------------------------------------------------------------
# Filesystem event handler with debounce
# ---------------------------------------------------------------------------

class BrainWatcherHandler(FileSystemEventHandler):
    """Debounced file event handler that enqueues files for ingestion."""

    def __init__(self, debounce_seconds: float = 2.0) -> None:
        super().__init__()
        self._debounce = debounce_seconds
        self._pending: dict[str, float] = {}  # file_path -> last_event_time
        self._lock = threading.Lock()
        self._timer: threading.Timer | None = None
        self._conn: sqlite3.Connection | None = None

    def _get_conn(self) -> sqlite3.Connection:
        if self._conn is None:
            self._conn = get_connection()
        return self._conn

    def _should_watch(self, path: str) -> bool:
        """Check if a file event should be processed."""
        p = Path(path)

        # Must have a watched extension
        if p.suffix.lower() not in WATCHED_EXTENSIONS:
            return False

        # Must not be in ignored suffix
        if p.suffix.lower() in IGNORE_SUFFIXES:
            return False

        # Must not be in an ignored directory
        parts = p.parts
        for part in parts:
            if part in IGNORE_DIRS:
                return False

        # Must not be the brain DB itself
        try:
            if p.resolve() == MBP_BRAIN_DB.resolve():
                return False
        except (OSError, ValueError):
            pass

        return True

    def _schedule_flush(self) -> None:
        """Schedule a debounced flush of pending files."""
        if self._timer is not None:
            self._timer.cancel()
        self._timer = threading.Timer(self._debounce, self._flush_pending)
        self._timer.daemon = True
        self._timer.start()

    def _flush_pending(self) -> None:
        """Flush all debounced pending files into the queue and process them."""
        with self._lock:
            now = time.time()
            ready: list[str] = []
            still_pending: dict[str, float] = {}

            for fpath, last_time in self._pending.items():
                if now - last_time >= self._debounce:
                    ready.append(fpath)
                else:
                    still_pending[fpath] = last_time

            self._pending = still_pending

        if not ready:
            if still_pending:
                self._schedule_flush()
            return

        conn = self._get_conn()
        enqueued = 0
        for fpath in ready:
            fp = Path(fpath)
            if fp.exists() and fp.stat().st_size > 0:
                qid = enqueue_file(conn, fp)
                if qid is not None:
                    enqueued += 1

        if enqueued > 0:
            logger.info("Enqueued %d files, processing...", enqueued)
            processed = process_pending(conn)
            logger.info("Processed %d/%d files", processed, enqueued)

        # If there are still pending items, reschedule
        with self._lock:
            if self._pending:
                self._schedule_flush()

    def _record_event(self, path: str) -> None:
        """Record a file event for debounced processing."""
        if not self._should_watch(path):
            return
        with self._lock:
            self._pending[path] = time.time()
        self._schedule_flush()

    def on_created(self, event: FileSystemEvent) -> None:
        if not event.is_directory:
            self._record_event(event.src_path)

    def on_modified(self, event: FileSystemEvent) -> None:
        if not event.is_directory:
            self._record_event(event.src_path)

    def on_moved(self, event: FileSystemEvent) -> None:
        if not event.is_directory:
            self._record_event(event.dest_path)

    def close(self) -> None:
        """Clean up resources."""
        if self._timer is not None:
            self._timer.cancel()
        if self._conn is not None:
            try:
                self._conn.close()
            except Exception:
                pass


# ---------------------------------------------------------------------------
# Status display
# ---------------------------------------------------------------------------

def show_status() -> None:
    """Print ingest queue and brain stats."""
    conn = get_connection()
    try:
        print("=" * 60)
        print("  BRAIN WATCHER STATUS")
        print("=" * 60)

        # Queue stats
        rows = conn.execute(
            """SELECT status, COUNT(*) as cnt
               FROM ingest_queue GROUP BY status ORDER BY cnt DESC"""
        ).fetchall()
        print("\n  Ingest Queue:")
        if rows:
            for r in rows:
                print(f"    {r['status']:>12s}: {r['cnt']:>6d}")
        else:
            print("    (empty)")

        # Brain stats
        stats = conn.execute(
            """SELECT
                 (SELECT COUNT(*) FROM nodes) AS nodes,
                 (SELECT COUNT(*) FROM edges) AS edges,
                 (SELECT COUNT(*) FROM nodes WHERE source_table = 'ingest_queue') AS ingested_nodes,
                 (SELECT COUNT(*) FROM edges WHERE source_table = 'ingest_queue') AS ingested_edges"""
        ).fetchone()
        print(f"\n  Brain Totals:")
        print(f"    Total nodes:         {stats['nodes']:>8d}")
        print(f"    Total edges:         {stats['edges']:>8d}")
        print(f"    Ingested nodes:      {stats['ingested_nodes']:>8d}")
        print(f"    Ingested edges:      {stats['ingested_edges']:>8d}")

        # Recent activity
        recent = conn.execute(
            """SELECT file_path, status, nodes_created, edges_created, processed_at
               FROM ingest_queue ORDER BY detected_at DESC LIMIT 10"""
        ).fetchall()
        if recent:
            print(f"\n  Recent Queue Items (last 10):")
            for r in recent:
                fname = Path(r["file_path"]).name if r["file_path"] else "?"
                ts = r["processed_at"] or "pending"
                print(f"    [{r['status']:>10s}] {fname:<40s} +{r['nodes_created']}n/{r['edges_created']}e  {ts}")

        # Latest version
        ver = conn.execute(
            "SELECT * FROM versions ORDER BY version DESC LIMIT 1"
        ).fetchone()
        if ver:
            print(f"\n  Latest Version: v{ver['version']}")
            print(f"    Nodes: {ver['node_count']}  Edges: {ver['edge_count']}  "
                  f"Chains: {ver['chain_count']}  Gaps: {ver['gap_count']}")

        print("\n" + "=" * 60)
    finally:
        conn.close()


# ---------------------------------------------------------------------------
# Daemon / main watch loop
# ---------------------------------------------------------------------------

_shutdown = threading.Event()


def _signal_handler(signum: int, frame) -> None:  # noqa: ANN001
    """Handle shutdown signals gracefully."""
    logger.info("Shutdown signal received (signal %d), stopping...", signum)
    _shutdown.set()


def run_watcher(watch_dir: Path, debounce: float) -> None:
    """Run the file watcher loop until shutdown."""
    signal.signal(signal.SIGINT, _signal_handler)
    signal.signal(signal.SIGTERM, _signal_handler)

    handler = BrainWatcherHandler(debounce_seconds=debounce)
    observer = Observer()
    observer.schedule(handler, str(watch_dir), recursive=True)

    logger.info("Starting brain watcher on: %s (debounce=%.1fs)", watch_dir, debounce)
    logger.info("Watching extensions: %s", ", ".join(sorted(WATCHED_EXTENSIONS)))
    logger.info("Brain DB: %s", MBP_BRAIN_DB)

    # Log startup op
    conn = get_connection()
    log_brain_op(conn, "watcher_start", input_params=str(watch_dir), output_summary="daemon started")
    conn.commit()
    conn.close()

    observer.start()
    try:
        while not _shutdown.is_set():
            _shutdown.wait(timeout=1.0)
    except KeyboardInterrupt:
        logger.info("KeyboardInterrupt received, stopping...")
    finally:
        observer.stop()
        handler.close()
        observer.join(timeout=5)

        # Log shutdown op
        try:
            conn = get_connection()
            log_brain_op(conn, "watcher_stop", output_summary="daemon stopped")
            conn.commit()
            conn.close()
        except Exception:
            pass

        logger.info("Brain watcher stopped.")


def run_daemon(watch_dir: Path, debounce: float) -> None:
    """Fork into background (Windows: run in a detached thread)."""
    logger.info("Starting brain watcher in daemon mode...")
    daemon_thread = threading.Thread(
        target=run_watcher, args=(watch_dir, debounce), daemon=False
    )
    daemon_thread.start()
    print(f"Brain watcher daemon started. Monitoring: {watch_dir}")
    print(f"Log: {LOG_DIR / 'brain_watcher.log'}")
    print("Press Ctrl+C to stop.")
    try:
        daemon_thread.join()
    except KeyboardInterrupt:
        _shutdown.set()
        daemon_thread.join(timeout=10)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Brain Watcher — auto-ingest files into mbp_brain.db",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""Examples:
  python -I scripts/brain_watcher.py                    # Watch (foreground)
  python -I scripts/brain_watcher.py --daemon           # Watch (background thread)
  python -I scripts/brain_watcher.py --status           # Queue & brain stats
  python -I scripts/brain_watcher.py --process-backlog  # Process pending items
""",
    )
    parser.add_argument(
        "--watch-dir",
        type=Path,
        default=DEFAULT_WATCH_DIR,
        help=f"Directory to watch (default: {DEFAULT_WATCH_DIR})",
    )
    parser.add_argument(
        "--interval",
        type=float,
        default=2.0,
        help="Debounce interval in seconds (default: 2)",
    )
    parser.add_argument(
        "--daemon",
        action="store_true",
        help="Run in daemon mode (background thread)",
    )
    parser.add_argument(
        "--process-backlog",
        action="store_true",
        help="Process all pending items in the ingest queue then exit",
    )
    parser.add_argument(
        "--status",
        action="store_true",
        help="Show ingest queue and brain statistics",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()

    # --status is read-only, no logging setup needed to console
    if args.status:
        show_status()
        return 0

    setup_logging(console=True)

    # Verify brain DB exists
    if not MBP_BRAIN_DB.exists():
        logger.error("Brain DB not found: %s", MBP_BRAIN_DB)
        return 1

    # Verify watch dir exists
    if not args.watch_dir.is_dir():
        logger.error("Watch directory not found: %s", args.watch_dir)
        return 1

    if args.process_backlog:
        logger.info("Processing backlog...")
        conn = get_connection()
        try:
            count = process_pending(conn, limit=10000)
            logger.info("Backlog processing complete: %d items processed", count)
        finally:
            conn.close()
        return 0

    if args.daemon:
        run_daemon(args.watch_dir, args.interval)
    else:
        run_watcher(args.watch_dir, args.interval)

    return 0


if __name__ == "__main__":
    sys.exit(main())
