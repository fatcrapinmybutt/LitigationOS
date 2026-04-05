#!/usr/bin/env python3
r"""
PROJECT KRAKEN -- Multi-Tentacle Autonomous Evidence Hunting System
Standalone Desktop Executable Edition

Architecture:
  fd (Rust) -> file discovery (C:\, D:\)
  file_inventory DB -> slow drives (F:\, G:\, I:\, J:\)
  Content-first analysis -> PDF/DOCX/TXT/CSV/HTML/JSON/MD
  22 adversary patterns + 6 legal authority types + 8 evidence categories
  Auto-persist HIGH findings -> evidence_quotes + dossier files
  Dedup tracking -> kraken_processed table
  Focus modes -> adversary, legal, housing, judicial, custody, ppo, all

Usage (CLI):
  PROJECT_KRAKEN.exe                    # Interactive menu
  PROJECT_KRAKEN.exe --rounds 5         # 5 rounds of 10
  PROJECT_KRAKEN.exe --count 20         # 20 files per round
  PROJECT_KRAKEN.exe --focus judicial   # focus on judicial patterns
  PROJECT_KRAKEN.exe --drives C,D       # specific drives only
  PROJECT_KRAKEN.exe --status           # show KRAKEN stats
  PROJECT_KRAKEN.exe --continuous       # daemon mode (self-evolving)
"""

import sqlite3, random, json, os, sys, re, hashlib, traceback, argparse, subprocess, shutil, time
from pathlib import Path
from collections import defaultdict
from datetime import datetime, date

# ===================================================================
# FROZEN-MODE PATH HANDLING
# ===================================================================

_FROZEN = getattr(sys, 'frozen', False)

def _find_python():
    """Find real python.exe -- in frozen mode sys.executable is the .exe itself."""
    if _FROZEN:
        p = shutil.which("python")
        if p:
            return p
        for candidate in [
            r"C:\Python312\python.exe",
            r"C:\Users\andre\LitigationOS\pytools_venv\Scripts\python.exe",
            r"C:\Users\andre\LitigationOS\.mcp_venv\Scripts\python.exe",
        ]:
            if os.path.isfile(candidate):
                return candidate
        return "python"
    return sys.executable

_PYTHON = _find_python()

# ===================================================================
# CONFIGURATION
# ===================================================================

DB_PATH = r"C:\Users\andre\LitigationOS\litigation_context.db"
DOSSIER_DIR = r"C:\Users\andre\LitigationOS\04_ANALYSIS\ADVERSARY_TRACKS"
RESULTS_DIR = r"D:\LitigationOS_tmp\kraken_results"
EXTS = {'.pdf', '.txt', '.csv', '.html', '.json', '.docx', '.md'}
MAX_CONTENT_BYTES = 500_000  # 500KB content cap per file
MAX_PDF_PAGES = 30

VERSION = "2.3.0"

# Separation counter
SEP_ANCHOR = date(2025, 7, 29)
SEP_DAYS = (date.today() - SEP_ANCHOR).days

# ===================================================================
# ADVERSARY PATTERNS (22 targets) -- PRE-COMPILED for O(1) per match
# ===================================================================

_ADV_RAW = {
    'Emily Watson': r'(?i)\bemily\b.*\bwatson\b|\bwatson\b.*\bemily\b|\bemily\s+a\.?\s+watson\b',
    'Judge McNeill': r'(?i)\bmcneill\b|\bmcneil\b',
    'Pamela Rusco': r'(?i)\brusco\b|\bpamela\s+rusco\b',
    'Ronald Berry': r'(?i)\bronald\s+berry\b|\bron\s+berry\b|\bberry\b.*\bronald\b',
    'Cavan Berry': r'(?i)\bcavan\s+berry\b|\bcavan\b.*\bberry\b',
    'Albert Watson': r'(?i)\balbert\s+watson\b|\balbert\b.*\bwatson\b',
    'Lori Watson': r'(?i)\blori\s+watson\b|\blori\b.*\bwatson\b',
    'Kenneth Hoopes': r'(?i)\bhoopes\b|\bkenneth\s+hoopes\b',
    'Maria Ladas-Hoopes': r'(?i)\bladas[\s-]*hoopes\b|\bmaria\s+ladas\b',
    'Jennifer Barnes': r'(?i)\bbarnes\b.*\bjennifer\b|\bjennifer\s+barnes\b|\bP55406\b',
    'Michael Martini': r'(?i)\bmartini\b|\bmichael\s+martini\b',
    'Shady Oaks': r'(?i)\bshady\s+oaks\b|\bgreenridge\b|\bmobile\s+home\s+park\b',
    'Kim Davis': r'(?i)\bkim\s+davis\b',
    'Lauren Duguid': r'(?i)\bduguid\b|\blauren\s+duguid\b',
    'DJ Hilson': r'(?i)\bdj\s+hilson\b|\bhilson\b',
    'Cassandra VanDam': r'(?i)\bvandam\b|\bcassandra\b.*\bvandam\b|\bcasey\b.*\bvandam\b',
    'Nicole Browley': r'(?i)\bbrowley\b|\bnicole\s+browley\b',
    'Jeremy Brown': r'(?i)\bjeremy\s+brown\b',
    'Henry Brandell': r'(?i)\bbrandell?\b|\bhenry\s+brandell?\b',
    'Aaron Cox': r'(?i)\baaron\s+cox\b|\bcox\b.*\baaron\b|\bP69346\b',
    'FOC Office': r'(?i)\bfriend\s+of\s+(the\s+)?court\b|\bFOC\b|\b990\s+terrace\b',
    'DHHS/CPS': r'(?i)\bDHHS\b|\bCPS\b|\bchild\s+protective\b|\bchildren.?s?\s+services\b',
}
ADVERSARIES = {name: re.compile(pat) for name, pat in _ADV_RAW.items()}

# Pre-compiled mega-regex for fast initial screening (skip files with zero adversary hits)
_SCREEN_PATTERN = re.compile(
    r'(?i)\b(emily|watson|mcneill|mcneil|rusco|berry|hoopes|ladas|barnes|martini|'
    r'shady\s+oaks|greenridge|davis|duguid|hilson|vandam|browley|brandell|cox|'
    r'FOC|DHHS|CPS|friend\s+of.*court)\b'
)

# ===================================================================
# LEGAL AUTHORITY PATTERNS -- PRE-COMPILED
# ===================================================================

_LEGAL_RAW = {
    'MCR': r'\bMCR\s+\d+\.\d+',
    'MCL': r'\bMCL\s+\d+\.\d+',
    'MRE': r'\bMRE\s+\d+',
    'USC': r'\b\d+\s+U\.?S\.?C\.?\s+',
    'FRCP': r'\bFRCP\b|\bFed\.?\s*R\.?\s*Civ\.?\s*P\b',
    'Case Law': r'\b\d+\s+Mich\.?\s*(App\.?)?\s+\d+|\b\d+\s+F\.\s*\d+d\s+\d+|\b\d+\s+S\.?\s*Ct\.?\s+\d+',
}
LEGAL_PATTERNS = {name: re.compile(pat) for name, pat in _LEGAL_RAW.items()}

# ===================================================================
# EVIDENCE CATEGORIES -- PRE-COMPILED
# ===================================================================

_CAT_RAW = {
    'custody': r'(?i)\bcustod(y|ial)\b|\bparenting\s+time\b|\bvisit(ation)?\b|\bbest\s+interest\b|\bchild\b.*\b(welfare|safety)\b',
    'judicial_misconduct': r'(?i)\bex\s+parte\b|\bbias\b|\bprejudice\b|\bmisconduct\b|\bcanon\b|\bJTC\b|\brecus(al|e)\b|\bdisqualif',
    'ppo_abuse': r'(?i)\bPPO\b|\bprotection\s+order\b|\bstalking\b|\bharass(ment)?\b|\bcontempt\b|\bviolat(e|ion)\b',
    'false_allegations': r'(?i)\bfalse\b.*\b(alleg|report|claim|accus)\b|\bfabricat\b|\blie[sd]?\b|\brecant\b|\bperjur',
    'housing': r'(?i)\bevict(ion|ed)?\b|\blockout\b|\bhabitability\b|\butility\b|\bwater\s+shut\b|\bsewage\b|\btitle\b|\bmobile\s+home',
    'financial': r'(?i)\bchild\s+support\b|\balimony\b|\bincome\b|\btax\b|\bfinancial\b|\bemployment\b|\bwages?\b',
    'criminal': r'(?i)\barrest\b|\bcharge[sd]?\b|\bplea\b|\bsentenc\b|\bprobation\b|\bjail\b|\bincarcerat',
    'medical': r'(?i)\bmental\s+health\b|\bpsych(olog|iatr)\b|\bmedication\b|\btherapy\b|\bsubstance\b|\bmeth\b|\bdrug\b',
}
EVIDENCE_CATEGORIES = {name: re.compile(pat) for name, pat in _CAT_RAW.items()}

# Pre-compiled high-value quote patterns
_QUOTE_PATTERNS = [re.compile(p) for p in [
    r'(?i)\bex\s+parte\b', r'(?i)\bwithout\s+notice\b', r'(?i)\bdenied\b.*\bhearing\b',
    r'(?i)\bfalse\b.*\b(alleg|report)\b', r'(?i)\brecant\b', r'(?i)\badmit\b',
    r'(?i)\bthreat\b', r'(?i)\bintimid\b', r'(?i)\bcoer\b', r'(?i)\bretaliat\b',
    r'(?i)\bviolat\b.*\b(order|right|due\s+process)\b', r'(?i)\bmeth\b',
    r'(?i)\bweaponiz\b', r'(?i)\bperjur\b', r'(?i)\bfraud\b',
    r'(?i)\bpremeditat\b', r'(?i)\bconspirac?\b', r'(?i)\bcollu\b',
]]

# ===================================================================
# FOCUS MODE BOOSTS
# ===================================================================

FOCUS_BOOSTS = {
    'all': {},
    'adversary': {'Emily Watson': 3, 'Judge McNeill': 3, 'Pamela Rusco': 2, 'Albert Watson': 2},
    'judicial': {'Judge McNeill': 5, 'Kenneth Hoopes': 4, 'Maria Ladas-Hoopes': 4, 'Cavan Berry': 3, 'judicial_misconduct': 5},
    'housing': {'Shady Oaks': 5, 'housing': 5, 'Cassandra VanDam': 3, 'Nicole Browley': 3},
    'custody': {'custody': 5, 'Emily Watson': 3, 'false_allegations': 3, 'ppo_abuse': 2},
    'ppo': {'ppo_abuse': 5, 'Emily Watson': 3, 'Judge McNeill': 2, 'contempt': 3},
    'legal': {'MCR': 3, 'MCL': 3, 'Case Law': 3, 'MRE': 2},
}


# ===================================================================
# DATABASE HELPERS
# ===================================================================

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA busy_timeout=60000")
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA cache_size=-32000")
    conn.execute("PRAGMA temp_store=MEMORY")
    conn.execute("PRAGMA synchronous=NORMAL")
    return conn

def ensure_kraken_table(conn):
    """Create tracking table for processed files."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_processed (
            file_hash TEXT PRIMARY KEY,
            file_path TEXT NOT NULL,
            file_size INTEGER,
            processed_at TEXT DEFAULT (datetime('now')),
            round_id TEXT,
            value_score INTEGER DEFAULT 0,
            value_label TEXT DEFAULT 'LOW',
            adversaries_found TEXT,
            legal_found TEXT,
            categories TEXT,
            persisted_to_db INTEGER DEFAULT 0,
            focus_mode TEXT DEFAULT 'all'
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_rounds (
            round_id TEXT PRIMARY KEY,
            started_at TEXT DEFAULT (datetime('now')),
            completed_at TEXT,
            files_scanned INTEGER DEFAULT 0,
            high_count INTEGER DEFAULT 0,
            medium_count INTEGER DEFAULT 0,
            low_count INTEGER DEFAULT 0,
            focus_mode TEXT DEFAULT 'all',
            drives TEXT,
            new_evidence_rows INTEGER DEFAULT 0
        )
    """)
    conn.commit()


def file_hash(path):
    """Quick hash: path + size + mtime for dedup."""
    try:
        st = os.stat(path)
        return hashlib.md5(f"{path}|{st.st_size}|{int(st.st_mtime)}".encode()).hexdigest()
    except Exception:
        return hashlib.md5(path.encode()).hexdigest()


def is_already_processed(conn, fhash):
    row = conn.execute("SELECT 1 FROM kraken_processed WHERE file_hash = ?", (fhash,)).fetchone()
    return row is not None


# ===================================================================
# FILE DISCOVERY (fd + DB inventory -- zero os.walk)
# ===================================================================

def discover_files(drives=None, conn=None):
    """Discover files using fd (fast drives) + file_inventory DB (slow drives)."""
    all_files = []
    drives = drives or ['C', 'D', 'F', 'G', 'I', 'J']

    # Source 1: file_inventory DB (pre-indexed, instant)
    if conn:
        try:
            cols = {r[1] for r in conn.execute("PRAGMA table_info(file_inventory)")}
            path_col = 'file_path' if 'file_path' in cols else 'path' if 'path' in cols else 'full_path'
            
            drive_filter = " OR ".join(f"{path_col} LIKE '{d}:%'" for d in drives)
            rows = conn.execute(f"""
                SELECT {path_col} FROM file_inventory 
                WHERE ({drive_filter}) AND {path_col} IS NOT NULL
                LIMIT 100000
            """).fetchall()
            
            db_files = [fp for (fp,) in rows if fp and any(fp.lower().endswith(e) for e in EXTS)]
            all_files.extend(db_files)
            print(f"  [DB] {len(db_files)} files from file_inventory")
        except Exception as e:
            print(f"  [DB] Error: {e}")

    # Source 2: fd (Rust) for fast drives
    for drive in drives:
        if drive.upper() in ('F', 'G', 'I', 'J'):
            continue  # slow USB/SD -- rely on DB inventory
        
        scan_path = f"{drive.upper()}:\\"
        if drive.upper() == 'C':
            scan_path = r"C:\Users\andre"
        
        try:
            ext_args = []
            for ext in ['pdf', 'txt', 'csv', 'html', 'json', 'docx', 'md']:
                ext_args.extend(['-e', ext])
            
            result = subprocess.run(
                ['fd', '--type', 'f', '--no-ignore', '--hidden'] + ext_args +
                ['--exclude', 'node_modules', '--exclude', '.git', '--exclude', 'AppData',
                 '--exclude', '__pycache__', '--exclude', '.cache', '--exclude', 'pytools_venv',
                 '--exclude', '.mcp_venv',
                 '.', scan_path],
                capture_output=True, text=True, timeout=45,
                encoding='utf-8', errors='replace'
            )
            fd_files = [f.strip() for f in result.stdout.strip().split('\n') if f.strip()]
            if fd_files and fd_files != ['']:
                print(f"  [fd] {len(fd_files)} files on {drive.upper()}:\\")
                all_files.extend(fd_files)
            else:
                print(f"  [fd] {drive.upper()}:\\ returned 0 files (empty or not accessible)")
        except subprocess.TimeoutExpired:
            print(f"  [fd] {drive.upper()}:\\ timed out (45s)")
        except FileNotFoundError:
            print(f"  [fd] fd not found -- install with: cargo install fd-find")
        except Exception as e:
            print(f"  [fd] {drive.upper()}:\\ error: {e}")

    # Deduplicate
    unique = list(set(all_files))
    print(f"  [TOTAL] {len(unique)} unique files discovered")
    return unique


# ===================================================================
# CONTENT EXTRACTION
# ===================================================================

def extract_content(file_path):
    """Extract text from file. Supports PDF, DOCX, TXT, CSV, HTML, JSON, MD."""
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == '.pdf':
            return _extract_pdf(file_path)
        elif ext == '.docx':
            return _extract_docx(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read(MAX_CONTENT_BYTES)
    except Exception as e:
        return f"[EXTRACTION ERROR: {e}]"


def _extract_pdf(path):
    """Extract PDF text using pypdfium2 (fast, C-based)."""
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(path)
        pages = min(len(doc), MAX_PDF_PAGES)
        text_parts = []
        for i in range(pages):
            page = doc[i]
            textpage = page.get_textpage()
            text_parts.append(textpage.get_text_range())
            textpage.close()
            page.close()
        doc.close()
        return '\n'.join(text_parts)[:MAX_CONTENT_BYTES]
    except ImportError:
        return "[pypdfium2 not installed -- PDF extraction unavailable]"
    except Exception as e:
        return f"[PDF ERROR: {e}]"


def _extract_docx(path):
    """Extract DOCX text using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        text = '\n'.join(p.text for p in doc.paragraphs)
        return text[:MAX_CONTENT_BYTES]
    except ImportError:
        return "[python-docx not installed -- DOCX extraction unavailable]"
    except Exception as e:
        return f"[DOCX ERROR: {e}]"


# ===================================================================
# ANALYSIS ENGINE
# ===================================================================

def analyze_content(content, file_path, focus='all'):
    """Analyze content against all PRE-COMPILED pattern sets. Returns scored result."""
    result = {
        'file': file_path,
        'adversaries': {},
        'legal': {},
        'categories': {},
        'quotes': [],
        'score': 0,
    }
    
    if not content or content.startswith('['):
        return result
    
    # O1: Fast screen -- skip files with zero adversary/legal signal
    if not _SCREEN_PATTERN.search(content[:50000]):
        # Still check for legal patterns (files can be purely legal authority)
        has_legal = False
        for auth_type, rx in LEGAL_PATTERNS.items():
            m = rx.findall(content[:50000])
            if m:
                result['legal'][auth_type] = m[:10]
                has_legal = True
        if not has_legal:
            return result  # No signal at all -- fast exit
    
    # Adversary detection (compiled regex -- ~3x faster than re.findall with string)
    for name, rx in ADVERSARIES.items():
        matches = rx.findall(content)
        if matches:
            result['adversaries'][name] = len(matches)
    
    # Legal authority detection (compiled)
    if not result['legal']:  # skip if already filled by screen pass
        for auth_type, rx in LEGAL_PATTERNS.items():
            matches = rx.findall(content)
            if matches:
                result['legal'][auth_type] = matches[:10]
    
    # Evidence category detection (compiled)
    for cat, rx in EVIDENCE_CATEGORIES.items():
        matches = rx.findall(content)
        if matches:
            result['categories'][cat] = len(matches)
    
    # O4: Enhanced quote extraction -- context windows + more quotes + priority scoring
    sentences = re.split(r'(?<=[.!?])\s+', content[:150000])
    quote_scores = []  # (score, text) tuples for priority ranking
    for idx, sent in enumerate(sentences):
        slen = len(sent)
        if slen < 25 or slen > 600:
            continue
        qscore = 0
        for qp in _QUOTE_PATTERNS:
            if qp.search(sent):
                qscore += 10
        if qscore > 0:
            # Add context: previous sentence as lead-in if short enough
            context = sent.strip()
            if idx > 0 and len(sentences[idx - 1]) < 200:
                context = sentences[idx - 1].strip() + '. ' + context
            quote_scores.append((qscore, context[:500]))
    
    # Take top 8 quotes by score (was 5, now more comprehensive)
    quote_scores.sort(key=lambda x: -x[0])
    result['quotes'] = [q[1] for q in quote_scores[:8]]
    
    # Score calculation (tuned weights)
    score = 0
    score += len(result['adversaries']) * 5
    score += min(sum(result['adversaries'].values()), 50)  # cap raw mention bonus
    score += len(result['legal']) * 8
    score += sum(min(len(v), 5) for v in result['legal'].values()) * 2  # depth bonus
    score += len(result['categories']) * 3
    score += sum(min(v, 10) for v in result['categories'].values())  # category depth
    score += len(result['quotes']) * 10
    
    # Quote quality bonus (multi-pattern quotes score higher)
    if quote_scores:
        avg_qscore = sum(q[0] for q in quote_scores[:8]) / len(quote_scores[:8])
        score += int(avg_qscore * 0.5)
    
    # Focus mode boost
    boosts = FOCUS_BOOSTS.get(focus, {})
    for name, multiplier in boosts.items():
        if name in result['adversaries']:
            score += result['adversaries'][name] * multiplier
        if name in result['categories']:
            score += result['categories'][name] * multiplier
    
    result['score'] = score
    return result


# ===================================================================
# DB PERSISTENCE
# ===================================================================

def persist_high_value(conn, result, round_id):
    """Persist HIGH-value findings to evidence_quotes with dedup."""
    persisted = 0
    
    # Persist quotes (with dedup check)
    for quote in result.get('quotes', []):
        if len(quote) < 30:
            continue
        # Quote dedup: skip if substantially similar quote already exists from same source
        try:
            existing = conn.execute(
                "SELECT 1 FROM evidence_quotes WHERE source_file = ? AND quote_text = ? LIMIT 1",
                (result['file'], quote[:2000])
            ).fetchone()
            if existing:
                continue
        except Exception:
            pass
        
        advs = ', '.join(result.get('adversaries', {}).keys())
        cats = ', '.join(result.get('categories', {}).keys())
        lane = _detect_lane(result)
        
        try:
            conn.execute("""
                INSERT OR IGNORE INTO evidence_quotes 
                (source_file, quote_text, category, lane, relevance_score, tags, created_at)
                VALUES (?, ?, ?, ?, ?, ?, datetime('now'))
            """, (
                result['file'],
                quote[:2000],
                cats[:200] if cats else 'uncategorized',
                lane,
                min(result['score'] / 100.0, 1.0),
                f"KRAKEN:{round_id}:{advs[:100]}",
            ))
            persisted += 1
        except Exception:
            pass
    
    if persisted > 0:
        conn.commit()
    return persisted


def _detect_lane(result):
    """Detect primary lane from evidence categories."""
    cats = result.get('categories', {})
    if cats.get('judicial_misconduct', 0) > 2:
        return 'E'
    if cats.get('ppo_abuse', 0) > 2:
        return 'D'
    if cats.get('housing', 0) > 2:
        return 'B'
    if cats.get('custody', 0) > 0:
        return 'A'
    if cats.get('criminal', 0) > 0:
        return 'CRIMINAL'
    return 'A'


# ===================================================================
# DOSSIER EXPANSION
# ===================================================================

def expand_dossier(result):
    """Append HIGH-value findings to matching adversary dossier files."""
    expanded = 0
    if not os.path.isdir(DOSSIER_DIR):
        return 0
    
    name_to_file = {
        'Emily Watson': 'WATSON_EMILY_DOSSIER.md',
        'Judge McNeill': 'MCNEILL_JENNY_DOSSIER.md',
        'Pamela Rusco': 'RUSCO_PAMELA_DOSSIER.md',
        'Ronald Berry': 'BERRY_RONALD_DOSSIER.md',
        'Cavan Berry': 'BERRY_CAVAN_DOSSIER.md',
        'Albert Watson': 'WATSON_ALBERT_DOSSIER.md',
        'Lori Watson': 'WATSON_LORI_DOSSIER.md',
        'Kenneth Hoopes': 'HOOPES_KENNETH_DOSSIER.md',
        'Maria Ladas-Hoopes': 'LADAS_HOOPES_MARIA_DOSSIER.md',
        'Jennifer Barnes': 'BARNES_JENNIFER_DOSSIER.md',
        'Michael Martini': 'MARTINI_MICHAEL_DOSSIER.md',
        'Shady Oaks': 'SHADY_OAKS_CORPORATE_DOSSIER.md',
        'Kim Davis': 'KIM_DAVIS_DOSSIER.md',
        'Lauren Duguid': 'DUGUID_LAUREN_DOSSIER.md',
        'DJ Hilson': 'HILSON_DJ_DOSSIER.md',
        'Cassandra VanDam': 'VANDAM_CASSANDRA_DOSSIER.md',
        'Nicole Browley': 'BROWLEY_NICOLE_DOSSIER.md',
        'Jeremy Brown': 'BROWN_JEREMY_DOSSIER.md',
        'Henry Brandell': 'BRANDELL_HENRY_DOSSIER.md',
        'Aaron Cox': 'COX_AARON_DOSSIER.md',
    }
    
    for adv_name in result.get('adversaries', {}):
        dossier_file = name_to_file.get(adv_name)
        if not dossier_file:
            continue
        
        dossier_path = os.path.join(DOSSIER_DIR, dossier_file)
        if not os.path.isfile(dossier_path):
            continue
        
        quotes = result.get('quotes', [])
        if not quotes:
            continue
        
        try:
            with open(dossier_path, 'a', encoding='utf-8') as f:
                f.write(f"\n\n### KRAKEN Auto-Discovery ({datetime.now().strftime('%Y-%m-%d %H:%M')})\n")
                f.write(f"- **Source**: `{result['file']}`\n")
                f.write(f"- **Score**: {result['score']} | **Adversaries**: {', '.join(result['adversaries'].keys())}\n")
                for q in quotes[:3]:
                    f.write(f"- > \"{q[:300]}\"\n")
            expanded += 1
        except Exception:
            pass
    
    return expanded


# ===================================================================
# ROUND EXECUTION
# ===================================================================

def run_round(conn, all_files, count=10, focus='all', round_num=1, total_rounds=1):
    """Execute one hunting round with batch dedup + smart selection + adaptive scoring."""
    round_id = f"KR-{datetime.now().strftime('%Y%m%d-%H%M%S')}-R{round_num}"

    # O5: Get adaptive thresholds from historical yield data
    high_thresh, med_thresh = _get_adaptive_thresholds(conn)

    print(f"\n{'=' * 60}")
    print(f"  [ROUND {round_num}/{total_rounds}] {round_id}")
    print(f"  Focus: {focus} | Count: {count} | Thresholds: H>={high_thresh} M>={med_thresh}")
    print(f"{'=' * 60}")
    
    # Record round start
    conn.execute("""
        INSERT OR REPLACE INTO kraken_rounds (round_id, focus_mode)
        VALUES (?, ?)
    """, (round_id, focus))
    conn.commit()
    
    # O2: Batch-load ALL processed hashes into memory set (1 query vs N queries)
    processed_hashes = set()
    try:
        rows = conn.execute("SELECT file_hash FROM kraken_processed").fetchall()
        processed_hashes = {r[0] for r in rows}
        print(f"  [*] {len(processed_hashes)} files previously processed (in-memory dedup)")
    except Exception:
        pass
    
    # O3: Smart file selection -- weight toward HIGH-yield extensions
    ext_weights = _get_extension_weights(conn)
    
    selected = []
    candidates = list(all_files)  # copy to avoid mutating original
    random.shuffle(candidates)
    
    # O3: If evolution data exists, sort candidates to prioritize high-yield extensions
    if ext_weights:
        def _ext_weight(fp):
            ext = os.path.splitext(fp)[1].lower()
            return ext_weights.get(ext, 1.0)
        # Weighted shuffle: partition into tiers, shuffle within each tier
        tier1 = [f for f in candidates if _ext_weight(f) >= 3.0]  # docx, etc.
        tier2 = [f for f in candidates if 1.5 <= _ext_weight(f) < 3.0]  # pdf
        tier3 = [f for f in candidates if _ext_weight(f) < 1.5]  # txt, csv, etc.
        random.shuffle(tier1)
        random.shuffle(tier2)
        random.shuffle(tier3)
        # Mix: 40% tier1, 35% tier2, 25% tier3 (ensures exploration of all tiers)
        mixed = []
        iters = [iter(tier1), iter(tier2), iter(tier3)]
        ratios = [0.4, 0.35, 0.25]
        target = count * 5  # need enough to fill 'count' after dedup filtering
        for tier_iter, ratio in zip(iters, ratios):
            n = max(1, int(target * ratio))
            for _ in range(n):
                try:
                    mixed.append(next(tier_iter))
                except StopIteration:
                    break
        # Fill remainder from any tier
        for tier_iter in iters:
            for f in tier_iter:
                mixed.append(f)
                if len(mixed) >= target:
                    break
        candidates = mixed if mixed else candidates
    
    for fp in candidates:
        if len(selected) >= count:
            break
        fh = file_hash(fp)
        if fh not in processed_hashes:  # O(1) set lookup vs O(n) DB query
            selected.append((fp, fh))
            processed_hashes.add(fh)  # prevent selecting same file twice in round
    
    if not selected:
        print("  [!] No unprocessed files found!")
        return {'round_id': round_id, 'files': 0, 'high': 0, 'medium': 0, 'low': 0, 'persisted': 0, 'dossiers_expanded': 0}
    
    high_count = 0
    med_count = 0
    low_count = 0
    total_persisted = 0
    total_dossiers = 0
    t_start = time.time()
    
    for i, (fp, fh) in enumerate(selected, 1):
        # O6: Per-file error resilience
        try:
            # Extract
            content = extract_content(fp)
            
            # Analyze
            result = analyze_content(content, fp, focus=focus)
        except Exception as e:
            print(f"  [X] {i:>2}/{count} ERROR: {os.path.basename(fp)[:40]} -- {e}")
            continue
        
        # O5: Adaptive thresholds from historical yield data
        score = result['score']
        if score >= high_thresh:
            label = 'HIGH'
            high_count += 1
            marker = '[!]'
        elif score >= med_thresh:
            label = 'MEDIUM'
            med_count += 1
            marker = '[~]'
        else:
            label = 'LOW'
            low_count += 1
            marker = '[ ]'
        
        # O6: Progress indicator with ETA
        elapsed = time.time() - t_start
        rate = i / elapsed if elapsed > 0 else 0
        eta = (count - i) / rate if rate > 0 else 0
        eta_str = f"ETA {int(eta)}s" if eta > 0 else ""
        
        # Display
        basename = os.path.basename(fp)
        advs = ', '.join(result['adversaries'].keys())[:40] if result['adversaries'] else '-'
        print(f"  {marker} {i:>2}/{count} [{score:>3}] {basename[:42]:42s} | {advs} {eta_str}")
        
        # Record processed
        try:
            conn.execute("""
                INSERT OR REPLACE INTO kraken_processed 
                (file_hash, file_path, file_size, round_id, value_score, value_label,
                 adversaries_found, legal_found, categories, focus_mode)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                fh, fp, os.path.getsize(fp) if os.path.exists(fp) else 0,
                round_id, score, label,
                json.dumps(result['adversaries']),
                json.dumps({k: len(v) if isinstance(v, list) else v for k, v in result['legal'].items()}),
                json.dumps(result['categories']),
                focus,
            ))
        except Exception as e:
            print(f"    [warn] DB record: {e}")
        
        # Persist HIGH to DB + expand dossiers
        if label == 'HIGH':
            persisted = persist_high_value(conn, result, round_id)
            total_persisted += persisted
            
            dossiers = expand_dossier(result)
            total_dossiers += dossiers
            
            if persisted > 0:
                conn.execute("UPDATE kraken_processed SET persisted_to_db = ? WHERE file_hash = ?",
                           (persisted, fh))
    
    # Update round record
    conn.execute("""
        UPDATE kraken_rounds SET completed_at = datetime('now'),
        files_scanned = ?, high_count = ?, medium_count = ?, low_count = ?,
        new_evidence_rows = ?
        WHERE round_id = ?
    """, (len(selected), high_count, med_count, low_count, total_persisted, round_id))
    conn.commit()
    
    # Round summary with timing
    elapsed_total = time.time() - t_start
    print(f"\n  --- Round {round_num} Summary ---")
    print(f"  Files: {len(selected)} | HIGH: {high_count} | MED: {med_count} | LOW: {low_count}")
    print(f"  Persisted: {total_persisted} | Dossiers: {total_dossiers}")
    print(f"  Time: {elapsed_total:.1f}s ({elapsed_total/len(selected):.2f}s/file)")
    
    return {
        'round_id': round_id,
        'files': len(selected),
        'high': high_count,
        'medium': med_count,
        'low': low_count,
        'persisted': total_persisted,
        'dossiers_expanded': total_dossiers,
    }


def _get_extension_weights(conn):
    """O3: Query evolution data for extension yield rates. Returns {'.ext': weight}."""
    try:
        rows = conn.execute("""
            SELECT 
                LOWER(SUBSTR(file_path, INSTR(file_path, '.') - LENGTH(file_path))) as ext,
                COUNT(*) as total,
                SUM(CASE WHEN value_label = 'HIGH' THEN 1 ELSE 0 END) as highs
            FROM kraken_processed 
            WHERE file_path IS NOT NULL
            GROUP BY ext
            HAVING total >= 5
        """).fetchall()
        
        weights = {}
        for ext_raw, total, highs in rows:
            # Extract actual extension (the SQL above is approximate)
            ext = '.' + ext_raw.rsplit('.', 1)[-1] if '.' in ext_raw else ext_raw
            if ext and total > 0:
                yield_rate = highs / total
                # Weight: 1.0 baseline, up to 5.0 for 50%+ HIGH yield
                weights[ext] = max(1.0, 1.0 + yield_rate * 8.0)
        return weights
    except Exception:
        return {}


def _get_adaptive_thresholds(conn):
    """O5: Adaptive scoring thresholds based on historical yield. Target ~25% HIGH rate."""
    try:
        row = conn.execute("""
            SELECT COUNT(*) as total,
                   SUM(CASE WHEN value_label = 'HIGH' THEN 1 ELSE 0 END) as highs,
                   AVG(value_score) as avg_score,
                   ROUND(AVG(CASE WHEN value_label = 'HIGH' THEN value_score ELSE NULL END), 1) as avg_high
            FROM kraken_processed
            WHERE value_score > 0
        """).fetchone()

        if not row or row[0] < 20:
            return 40, 15  # default thresholds until enough data

        total, highs, avg_score, avg_high_score = row
        yield_rate = highs / total if total > 0 else 0

        # Target ~25% HIGH rate: adjust thresholds to converge
        high_thresh = 40
        med_thresh = 15

        if yield_rate > 0.40:
            # Too many HIGHs -- raise threshold (more selective)
            high_thresh = int(avg_high_score * 0.85) if avg_high_score else 50
            med_thresh = int(high_thresh * 0.45)
        elif yield_rate > 0.30:
            # Slightly high -- nudge up
            high_thresh = 45
            med_thresh = 18
        elif yield_rate < 0.10:
            # Too few HIGHs -- lower threshold (more permissive)
            high_thresh = max(25, int(avg_score * 1.2)) if avg_score else 30
            med_thresh = max(8, int(high_thresh * 0.35))
        elif yield_rate < 0.18:
            # Somewhat low -- nudge down
            high_thresh = 35
            med_thresh = 12

        # Clamp to reasonable bounds
        high_thresh = max(20, min(60, high_thresh))
        med_thresh = max(5, min(30, med_thresh))

        return high_thresh, med_thresh
    except Exception:
        return 40, 15


# ===================================================================
# STATUS REPORT
# ===================================================================

def show_status(conn):
    """O7: Enhanced KRAKEN status with top adversaries, directories, yield stats."""
    print(f"\n{'=' * 60}")
    print(f"  [KRAKEN] STATUS REPORT v2.2")
    print(f"{'=' * 60}")
    
    try:
        row = conn.execute("""
            SELECT 
                COUNT(*) as total,
                SUM(CASE WHEN value_label = 'HIGH' THEN 1 ELSE 0 END) as high,
                SUM(CASE WHEN value_label = 'MEDIUM' THEN 1 ELSE 0 END) as med,
                SUM(CASE WHEN value_label = 'LOW' THEN 1 ELSE 0 END) as low,
                SUM(CASE WHEN persisted_to_db > 0 THEN persisted_to_db ELSE 0 END) as persisted
            FROM kraken_processed
        """).fetchone()
        
        if row:
            yield_pct = (row[1] / row[0] * 100) if row[0] > 0 else 0
            print(f"  Total files: {row[0]} | Yield: {yield_pct:.1f}% HIGH")
            print(f"  HIGH: {row[1]} | MEDIUM: {row[2]} | LOW: {row[3]}")
            print(f"  Evidence persisted: {row[4]}")
    except Exception as e:
        print(f"  Error reading stats: {e}")
    
    try:
        rounds = conn.execute("""
            SELECT COUNT(*), SUM(files_scanned), SUM(high_count), SUM(new_evidence_rows)
            FROM kraken_rounds
        """).fetchone()
        if rounds:
            print(f"  Rounds: {rounds[0]} | Scanned: {rounds[1]} | New evidence: {rounds[3]}")
    except Exception:
        pass
    
    # O7: Top adversaries found
    try:
        print(f"\n  -- Top Adversaries Found --")
        rows = conn.execute("""
            SELECT adversaries_found FROM kraken_processed 
            WHERE value_label = 'HIGH' AND adversaries_found IS NOT NULL
        """).fetchall()
        adv_counts = defaultdict(int)
        for (ajson,) in rows:
            try:
                d = json.loads(ajson)
                for name, cnt in d.items():
                    adv_counts[name] += cnt
            except Exception:
                pass
        top_advs = sorted(adv_counts.items(), key=lambda x: -x[1])[:8]
        for name, cnt in top_advs:
            print(f"    {name:30s} {cnt:>5} mentions")
    except Exception:
        pass
    
    # O7: Yield by extension
    try:
        print(f"\n  -- Yield by Extension --")
        rows = conn.execute("""
            SELECT 
                CASE 
                    WHEN file_path LIKE '%.pdf' THEN '.pdf'
                    WHEN file_path LIKE '%.docx' THEN '.docx'
                    WHEN file_path LIKE '%.txt' THEN '.txt'
                    WHEN file_path LIKE '%.csv' THEN '.csv'
                    WHEN file_path LIKE '%.html' THEN '.html'
                    WHEN file_path LIKE '%.json' THEN '.json'
                    WHEN file_path LIKE '%.md' THEN '.md'
                    ELSE 'other'
                END as ext,
                COUNT(*) as total,
                SUM(CASE WHEN value_label = 'HIGH' THEN 1 ELSE 0 END) as highs
            FROM kraken_processed
            GROUP BY ext
            HAVING total >= 3
            ORDER BY highs DESC
        """).fetchall()
        for ext, total, highs in rows:
            pct = (highs / total * 100) if total > 0 else 0
            bar = '#' * int(pct / 5)
            print(f"    {ext:>6s}: {highs:>4}/{total:<5} ({pct:>5.1f}%) {bar}")
    except Exception:
        pass
    
    # O7: Top source directories
    try:
        print(f"\n  -- Top Source Directories (HIGH) --")
        rows = conn.execute("""
            SELECT 
                SUBSTR(file_path, 1, 
                    CASE WHEN INSTR(SUBSTR(file_path, 4), '\\') > 0 
                    THEN INSTR(SUBSTR(file_path, 4), '\\') + 3
                    ELSE LENGTH(file_path) END
                ) as dir_prefix,
                COUNT(*) as cnt
            FROM kraken_processed
            WHERE value_label = 'HIGH'
            GROUP BY dir_prefix
            ORDER BY cnt DESC
            LIMIT 6
        """).fetchall()
        for dir_pre, cnt in rows:
            print(f"    {dir_pre[:50]:50s} {cnt:>4}")
    except Exception:
        pass
    
    print(f"\n  Separation: {SEP_DAYS} days since Jul 29, 2025")
    print(f"{'=' * 60}")


# ===================================================================
# EVOLUTION ENGINE
# ===================================================================

def ensure_evolution_tables(conn):
    """Create tables for evolution/learning metrics."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_evolution (
            metric_id TEXT PRIMARY KEY,
            round_id TEXT,
            metric_type TEXT,
            metric_key TEXT,
            metric_value REAL,
            recorded_at TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS kraken_learning (
            learning_id INTEGER PRIMARY KEY AUTOINCREMENT,
            lesson_type TEXT,
            lesson_key TEXT,
            lesson_value TEXT,
            confidence REAL DEFAULT 0.5,
            applied_count INTEGER DEFAULT 0,
            discovered_at TEXT DEFAULT (datetime('now')),
            last_applied TEXT
        )
    """)
    conn.commit()


def record_evolution_metrics(conn, round_id, round_result):
    """Record evolution metrics from a round."""
    try:
        ts = datetime.now().isoformat()
        metrics = [
            (f"{round_id}_yield", round_id, 'yield_rate',
             'high_pct', (round_result['high'] / max(round_result['files'], 1)) * 100, ts),
            (f"{round_id}_persist", round_id, 'persist_rate',
             'persist_pct', (round_result.get('persisted', 0) / max(round_result['files'], 1)) * 100, ts),
        ]
        conn.executemany("""
            INSERT OR REPLACE INTO kraken_evolution 
            (metric_id, round_id, metric_type, metric_key, metric_value, recorded_at)
            VALUES (?, ?, ?, ?, ?, ?)
        """, metrics)
        conn.commit()
    except Exception:
        pass


def show_evolution(conn):
    """Display evolution/learning metrics."""
    print(f"\n  [EVOLUTION] Learning Metrics")
    
    try:
        rows = conn.execute("""
            SELECT metric_type, metric_key, 
                   ROUND(AVG(metric_value), 1) as avg_val,
                   ROUND(MIN(metric_value), 1) as min_val,
                   ROUND(MAX(metric_value), 1) as max_val,
                   COUNT(*) as samples
            FROM kraken_evolution
            GROUP BY metric_type, metric_key
        """).fetchall()
        
        if rows:
            for r in rows:
                print(f"    {r[0]}/{r[1]}: avg={r[2]}, min={r[3]}, max={r[4]} ({r[5]} samples)")
        else:
            print("    No evolution data yet.")
    except Exception as e:
        print(f"    Error: {e}")
    
    # Extension yield by type
    try:
        rows = conn.execute("""
            SELECT 
                CASE 
                    WHEN file_path LIKE '%.pdf' THEN '.pdf'
                    WHEN file_path LIKE '%.docx' THEN '.docx'
                    WHEN file_path LIKE '%.txt' THEN '.txt'
                    WHEN file_path LIKE '%.csv' THEN '.csv'
                    WHEN file_path LIKE '%.html' THEN '.html'
                    WHEN file_path LIKE '%.json' THEN '.json'
                    WHEN file_path LIKE '%.md' THEN '.md'
                    ELSE 'other'
                END as ext,
                COUNT(*) as total,
                SUM(CASE WHEN value_label = 'HIGH' THEN 1 ELSE 0 END) as high,
                ROUND(SUM(CASE WHEN value_label = 'HIGH' THEN 1.0 ELSE 0 END) / COUNT(*) * 100, 1) as yield_pct
            FROM kraken_processed
            GROUP BY ext
            ORDER BY yield_pct DESC
        """).fetchall()
        
        if rows:
            print("\n    Extension Yield:")
            for r in rows:
                bar = '#' * int(r[3] / 5) if r[3] else ''
                print(f"      {r[0]:6s} {r[1]:>5d} files, {r[2]:>4d} HIGH ({r[3]:>5.1f}%) {bar}")
    except Exception:
        pass


# ===================================================================
# CONDENSATION PIPELINE
# ===================================================================

def condense_findings(conn):
    """Condense KRAKEN findings into structured intelligence."""
    print(f"\n  [CONDENSE] Running condensation pipeline...")
    
    condensed = {
        'judicial': [],
        'adversary': [],
        'evidence': [],
        'legal': [],
        'timestamp': datetime.now().isoformat(),
    }
    
    # Judicial condensation
    try:
        rows = conn.execute("""
            SELECT file_path, adversaries_found, categories, value_score
            FROM kraken_processed
            WHERE value_label = 'HIGH'
            AND (categories LIKE '%judicial%' OR adversaries_found LIKE '%McNeill%' 
                 OR adversaries_found LIKE '%Hoopes%' OR adversaries_found LIKE '%Ladas%')
            ORDER BY value_score DESC LIMIT 200
        """).fetchall()
        
        for r in rows:
            condensed['judicial'].append({
                'source': r[0], 'adversaries': r[1],
                'categories': r[2], 'score': r[3]
            })
        print(f"    Judicial items: {len(condensed['judicial'])}")
    except Exception as e:
        print(f"    Judicial error: {e}")
    
    # Adversary condensation
    try:
        rows = conn.execute("""
            SELECT adversaries_found, COUNT(*) as cnt, 
                   ROUND(AVG(value_score), 1) as avg_score
            FROM kraken_processed
            WHERE value_label IN ('HIGH', 'MEDIUM')
            AND adversaries_found != '{}'
            GROUP BY adversaries_found
            ORDER BY cnt DESC LIMIT 50
        """).fetchall()
        
        for r in rows:
            condensed['adversary'].append({
                'pattern': r[0], 'count': r[1], 'avg_score': r[2]
            })
        print(f"    Adversary patterns: {len(condensed['adversary'])}")
    except Exception as e:
        print(f"    Adversary error: {e}")
    
    # Evidence condensation
    try:
        rows = conn.execute("""
            SELECT quote_text, source_file, category, lane, relevance_score
            FROM evidence_quotes
            WHERE tags LIKE 'KRAKEN:%'
            ORDER BY relevance_score DESC LIMIT 100
        """).fetchall()
        
        for r in rows:
            condensed['evidence'].append({
                'quote': r[0][:200], 'source': r[1],
                'category': r[2], 'lane': r[3], 'score': r[4]
            })
        print(f"    Evidence items: {len(condensed['evidence'])}")
    except Exception as e:
        print(f"    Evidence error: {e}")
    
    # Save condensed data
    os.makedirs(RESULTS_DIR, exist_ok=True)
    out_path = os.path.join(RESULTS_DIR, f"condensed_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(condensed, f, indent=2, default=str)
    print(f"    Saved: {out_path}")
    
    return condensed


# ===================================================================
# MBP BRIDGE (THEMANBEARPIG Integration)
# ===================================================================

def export_to_mbp(conn, condensed=None):
    """Export KRAKEN data as graph nodes/links for THEMANBEARPIG visualization."""
    print(f"\n  [MBP] Exporting to THEMANBEARPIG graph format...")
    
    if not condensed:
        condensed = condense_findings(conn)
    
    nodes = []
    links = []
    node_ids = set()
    
    # Create adversary co-occurrence network
    try:
        rows = conn.execute("""
            SELECT adversaries_found FROM kraken_processed
            WHERE value_label = 'HIGH' AND adversaries_found != '{}'
        """).fetchall()
        
        co_occur = defaultdict(int)
        for (advs_json,) in rows:
            try:
                advs = json.loads(advs_json)
                names = list(advs.keys())
                for i in range(len(names)):
                    nid = f"KRAKEN_{names[i].replace(' ', '_')}"
                    if nid not in node_ids:
                        nodes.append({
                            'id': nid, 'label': names[i],
                            'type': 'ADVERSARY_CORE', 'tier': 0,
                            'layer': 'ADVERSARY', 'source': 'KRAKEN',
                        })
                        node_ids.add(nid)
                    for j in range(i + 1, len(names)):
                        pair = tuple(sorted([names[i], names[j]]))
                        co_occur[pair] += 1
            except Exception:
                pass
        
        for (a, b), weight in sorted(co_occur.items(), key=lambda x: -x[1])[:50]:
            links.append({
                'source': f"KRAKEN_{a.replace(' ', '_')}",
                'target': f"KRAKEN_{b.replace(' ', '_')}",
                'type': 'CO_OCCURRENCE', 'weight': weight,
            })
        
        print(f"    Adversary nodes: {len(nodes)}")
        print(f"    Co-occurrence links: {len(links)}")
    except Exception as e:
        print(f"    Co-occurrence error: {e}")
    
    # Judicial violation nodes
    jud_count = 0
    for item in condensed.get('judicial', [])[:50]:
        nid = f"KRAKEN_JUD_{jud_count}"
        nodes.append({
            'id': nid, 'label': f"Violation #{jud_count+1}",
            'type': 'VIOLATION', 'tier': 2,
            'layer': 'JUDICIAL', 'source': 'KRAKEN',
            'score': item.get('score', 0),
        })
        node_ids.add(nid)
        jud_count += 1
    
    # Save MBP export
    mbp_data = {
        'nodes': nodes, 'links': links,
        'metadata': {
            'source': 'PROJECT_KRAKEN',
            'version': VERSION,
            'exported_at': datetime.now().isoformat(),
            'total_nodes': len(nodes),
            'total_links': len(links),
        }
    }
    
    os.makedirs(RESULTS_DIR, exist_ok=True)
    out_path = os.path.join(RESULTS_DIR, f"kraken_mbp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(mbp_data, f, indent=2, default=str)
    print(f"    MBP export: {out_path}")
    
    return mbp_data


# ===================================================================
# CONTINUOUS DAEMON
# ===================================================================

def run_continuous(conn, all_files, interval_seconds=300, max_cycles=50,
                   count=10, focus='all', condense_every=3, mbp_export_every=5):
    """Run continuous self-evolving hunting cycles."""
    import time
    
    print(f"\n  [DAEMON] Continuous mode: {max_cycles} cycles, {interval_seconds}s interval")
    print(f"    Condense every {condense_every} rounds | MBP export every {mbp_export_every} rounds")
    
    for cycle in range(1, max_cycles + 1):
        print(f"\n{'#' * 60}")
        print(f"  [DAEMON] Cycle {cycle}/{max_cycles}")
        print(f"{'#' * 60}")
        
        result = run_round(conn, all_files, count=count, focus=focus,
                          round_num=cycle, total_rounds=max_cycles)
        record_evolution_metrics(conn, result.get('round_id', f'C{cycle}'), result)
        
        # Condense periodically
        if cycle % condense_every == 0:
            print(f"\n  [DAEMON] Condensation at cycle {cycle}...")
            condensed = condense_findings(conn)
            
            # MBP export on schedule
            if cycle % mbp_export_every == 0:
                print(f"\n  [DAEMON] MBP export at cycle {cycle}...")
                export_to_mbp(conn, condensed)
        
        # Show evolution periodically
        if cycle % 5 == 0:
            show_evolution(conn)
        
        if cycle < max_cycles:
            print(f"\n  [DAEMON] Sleeping {interval_seconds}s until next cycle...")
            time.sleep(interval_seconds)
    
    print(f"\n  [DAEMON] Completed {max_cycles} cycles!")
    show_status(conn)
    show_evolution(conn)


# ===================================================================
# INTERACTIVE MENU (for Desktop double-click)
# ===================================================================

def interactive_menu():
    """Interactive menu for when exe is double-clicked with no args."""
    print("")
    print("  Select an operation:")
    print("")
    print("    1) Quick Hunt    -- 3 rounds, 10 files each")
    print("    2) Deep Hunt     -- 10 rounds, 20 files each")
    print("    3) Judicial Hunt -- 5 rounds, judicial focus")
    print("    4) Custody Hunt  -- 5 rounds, custody focus")
    print("    5) Full Sweep    -- 20 rounds, all focus")
    print("    6) Status Report -- show KRAKEN stats")
    print("    7) Evolution     -- show learning metrics")
    print("    8) Condense      -- run condensation pipeline")
    print("    9) MBP Export    -- export to THEMANBEARPIG graph")
    print("    0) Daemon Mode   -- continuous self-evolving (5 cycles)")
    print("    Q) Quit")
    print("")
    
    choice = input("  >> ").strip().upper()
    return choice


# ===================================================================
# MAIN
# ===================================================================

def main():
    parser = argparse.ArgumentParser(
        description="PROJECT KRAKEN -- Multi-Tentacle Autonomous Evidence Hunting",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  PROJECT_KRAKEN.exe                       # Interactive menu
  PROJECT_KRAKEN.exe --rounds 5            # 5 rounds of 10
  PROJECT_KRAKEN.exe --count 20            # 20 files per round
  PROJECT_KRAKEN.exe --focus judicial      # judicial focus mode
  PROJECT_KRAKEN.exe --drives C,D          # specific drives
  PROJECT_KRAKEN.exe --status              # show stats
  PROJECT_KRAKEN.exe --evolve              # show evolution metrics
  PROJECT_KRAKEN.exe --condense            # run condensation pipeline
  PROJECT_KRAKEN.exe --mbp-export          # export to MBP graph
  PROJECT_KRAKEN.exe --continuous           # daemon mode (self-evolving)
        """
    )
    parser.add_argument('--rounds', type=int, default=0, help='Number of hunting rounds (0 = interactive menu)')
    parser.add_argument('--version', action='version', version='PROJECT KRAKEN v%s' % VERSION)
    parser.add_argument('--count', type=int, default=10, help='Files per round (default: 10)')
    parser.add_argument('--focus', default='all', choices=list(FOCUS_BOOSTS.keys()), help='Focus mode')
    parser.add_argument('--drives', default=None, help='Comma-separated drive letters (default: all)')
    parser.add_argument('--status', action='store_true', help='Show KRAKEN status report')
    parser.add_argument('--resume', action='store_true', help='Resume from last session')
    parser.add_argument('--evolve', action='store_true', help='Show evolution metrics')
    parser.add_argument('--condense', action='store_true', help='Run condensation pipeline')
    parser.add_argument('--mbp-export', action='store_true', help='Export KRAKEN data to MBP graph')
    parser.add_argument('--continuous', action='store_true', help='Daemon mode: continuous self-evolving hunting')
    parser.add_argument('--interval', type=int, default=300, help='Daemon cycle interval in seconds (default: 300)')
    parser.add_argument('--max-cycles', type=int, default=50, help='Max daemon cycles (default: 50)')
    parser.add_argument('--condense-every', type=int, default=3, help='Condense every N rounds (default: 3)')
    parser.add_argument('--mbp-every', type=int, default=5, help='MBP export every N rounds (default: 5)')
    args = parser.parse_args()

    print("=" * 70)
    print("[KRAKEN] PROJECT KRAKEN v%s -- Autonomous Evidence Hunting" % VERSION)
    print("   Standalone Desktop Edition | %s" % datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    print("   Separation: %d days since Jul 29, 2025" % SEP_DAYS)
    if _FROZEN:
        print("   (Icon wrong? Run: ie4uinit.exe -show)")
    print("=" * 70)

    # Verify DB exists
    if not os.path.isfile(DB_PATH):
        print("\n  [ERROR] Database not found: %s" % DB_PATH)
        print("  Make sure LitigationOS is at C:\\Users\\andre\\LitigationOS\\")
        input("\n  Press Enter to exit...")
        return

    conn = get_db()
    ensure_kraken_table(conn)
    ensure_evolution_tables(conn)

    # Quick-exit commands
    if args.status:
        show_status(conn)
        conn.close()
        input("\n  Press Enter to exit...")
        return

    if args.evolve and not args.continuous:
        show_evolution(conn)
        conn.close()
        input("\n  Press Enter to exit...")
        return

    if args.condense and not args.continuous:
        condense_findings(conn)
        conn.close()
        input("\n  Press Enter to exit...")
        return

    if args.mbp_export and not args.continuous:
        condensed = condense_findings(conn)
        export_to_mbp(conn, condensed)
        conn.close()
        input("\n  Press Enter to exit...")
        return

    # Interactive menu if no rounds specified and no flags
    if args.rounds == 0 and not args.continuous:
        while True:
            choice = interactive_menu()
            
            if choice == '1':
                args.rounds = 3; args.count = 10; args.focus = 'all'; break
            elif choice == '2':
                args.rounds = 10; args.count = 20; args.focus = 'all'; break
            elif choice == '3':
                args.rounds = 5; args.count = 15; args.focus = 'judicial'; break
            elif choice == '4':
                args.rounds = 5; args.count = 15; args.focus = 'custody'; break
            elif choice == '5':
                args.rounds = 20; args.count = 20; args.focus = 'all'; break
            elif choice == '6':
                show_status(conn)
                continue
            elif choice == '7':
                show_evolution(conn)
                continue
            elif choice == '8':
                condense_findings(conn)
                continue
            elif choice == '9':
                condensed = condense_findings(conn)
                export_to_mbp(conn, condensed)
                continue
            elif choice == '0':
                args.continuous = True; args.max_cycles = 5; break
            elif choice == 'Q':
                conn.close()
                return
            else:
                print("  Invalid choice. Try again.")
                continue

    # Parse drives
    drives = args.drives.split(',') if args.drives else None

    # Discover files
    print("\n[PHASE 1] File Discovery...")
    all_files = discover_files(drives=drives, conn=conn)

    if not all_files:
        print("  [!] No files discovered!")
        conn.close()
        input("\n  Press Enter to exit...")
        return

    # Ensure results directory
    os.makedirs(RESULTS_DIR, exist_ok=True)

    # Continuous daemon mode
    if args.continuous:
        if args.evolve:
            show_evolution(conn)
        run_continuous(conn, all_files, interval_seconds=args.interval,
                      max_cycles=args.max_cycles, count=args.count, focus=args.focus,
                      condense_every=args.condense_every, mbp_export_every=args.mbp_every)
        conn.close()
        input("\n  Press Enter to exit...")
        return

    # Execute rounds (with evolution tracking)
    all_round_results = []
    for rnd in range(1, args.rounds + 1):
        result = run_round(conn, all_files, count=args.count, focus=args.focus,
                          round_num=rnd, total_rounds=args.rounds)
        all_round_results.append(result)
        record_evolution_metrics(conn, result.get('round_id', 'R%d' % rnd), result)

    # Save combined results
    out_file = os.path.join(RESULTS_DIR, "kraken_%s.json" % datetime.now().strftime('%Y%m%d_%H%M%S'))
    with open(out_file, 'w', encoding='utf-8') as f:
        json.dump(all_round_results, f, indent=2, default=str)
    print("\n  [>] Results saved: %s" % out_file)

    # Final summary
    total_files = sum(r['files'] for r in all_round_results)
    total_high = sum(r['high'] for r in all_round_results)
    total_med = sum(r['medium'] for r in all_round_results)
    total_persisted = sum(r.get('persisted', 0) for r in all_round_results)
    total_dossiers = sum(r.get('dossiers_expanded', 0) for r in all_round_results)

    print("\n" + "=" * 70)
    print("[KRAKEN] KRAKEN MISSION COMPLETE -- %d round(s)" % args.rounds)
    print("   Files scanned: %d" % total_files)
    print("   [!] HIGH: %d | [~] MEDIUM: %d" % (total_high, total_med))
    print("   [DB] Evidence persisted: %d rows" % total_persisted)
    print("   [+] Dossiers expanded: %d" % total_dossiers)
    print("   Separation: %d days -- EVERY DAY COUNTS" % SEP_DAYS)
    print("=" * 70)

    # Auto-condense if 3+ rounds
    if args.rounds >= 3:
        print("\n  [*] Auto-condensing (3+ rounds completed)...")
        condensed = condense_findings(conn)
        if condensed and args.rounds >= 5:
            print("\n  [*] Auto-exporting to MBP (5+ rounds)...")
            export_to_mbp(conn, condensed)

    # Show cumulative status
    show_status(conn)
    show_evolution(conn)
    conn.close()
    
    input("\n  Press Enter to exit...")


if __name__ == '__main__':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except (AttributeError, OSError):
        pass
    main()
